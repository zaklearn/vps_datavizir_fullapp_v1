import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import plotly.figure_factory as ff
import scipy.stats as stats
from docx import Document
from docx.shared import Inches
import tempfile
import os
from config import translations, egra_columns, egma_columns


def show_contextual_factors(df, language):
    """
    Analyze the impact of contextual factors on student performance with tabbed interface.
    
    Args:
        df (pd.DataFrame): Dataset containing contextual variables
        language (str): Selected language for UI text
    """
    t = translations[language]  # Get translations
    
    st.subheader(t.get("title_contextual_factors", "Contextual Factors Analysis"))
    
    # Verify required columns exist
    required_columns = ["ses", "home_support"]
    missing_columns = [col for col in required_columns if col not in df.columns]
    
    if missing_columns:
        st.warning(f"{t.get('missing_columns', 'Missing required columns')}: {', '.join(missing_columns)}")
        st.info(t.get('contextual_columns_info', 'This analysis requires socioeconomic status (ses) and home support information.'))
        return
    
    # Select variables for analysis
    selected_columns = _select_variables(df, t)
    
    if selected_columns:
        try:
            # Prepare analysis data
            df_analysis = _prepare_contextual_data(df, selected_columns, t)
            
            if df_analysis is None:
                st.error(t.get("data_preparation_error", "Error preparing data for analysis"))
                return
            
            # Create tabs for different analyses
            tab1, tab2, tab3, tab4, tab5 = st.tabs([
                t.get("tab_overview", "Overview"),
                t.get("tab_ses_analysis", "SES Analysis"),
                t.get("tab_home_support", "Home Support"),
                t.get("tab_combined_effects", "Combined Effects"),
                t.get("tab_export", "Export Results")
            ])
            
            with tab1:
                _show_overview(df_analysis, selected_columns, t)
            
            with tab2:
                _show_ses_analysis(df_analysis, selected_columns, t)
            
            with tab3:
                _show_home_support_analysis(df_analysis, selected_columns, t)
            
            with tab4:
                _show_combined_factor_analysis(df_analysis, selected_columns, t)
            
            with tab5:
                _provide_export_options(df_analysis, selected_columns, t)
        
        except Exception as e:
            st.error(f"{t.get('error_message', 'An error occurred')}: {str(e)}")
            st.info("Please ensure your data meets the required format for analysis.")
    else:
        st.warning(t.get("warning_select_indicator", "Please select at least one indicator to analyze."))


def _select_variables(df, t):
    """
    Allow user to select variables for analysis.
    
    Args:
        df (pd.DataFrame): Dataset
        t (dict): Translation dictionary
        
    Returns:
        list: Selected column names
    """
    st.subheader(t.get("select_variables", "Select Variables"))
    
    # Get available assessment columns
    available_columns = []
    for col in egra_columns + egma_columns:
        if col in df.columns:
            available_columns.append(col)
    
    # If no assessment columns found, allow selection from all numeric columns
    if not available_columns:
        numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
        available_columns = [col for col in numeric_cols 
                            if col not in ['school', 'stgender', 'ses', 'home_support']]
    
    # Create selection widget
    selected_columns = st.multiselect(
        t.get("variables_all", "Variables:"),
        options=available_columns,
        default=available_columns[:min(3, len(available_columns))],
        format_func=lambda x: t["columns_of_interest"].get(x, x)
    )
    
    return selected_columns


def _prepare_contextual_data(df, selected_columns, t):
    """
    Prepare dataset with additional calculated variables and proper data types.
    
    Args:
        df (pd.DataFrame): Original dataset
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
        
    Returns:
        pd.DataFrame: Enhanced dataset with additional variables
    """
    df_analysis = df.copy()
    
    # Ensure data types are correct
    try:
        # Convert SES to numeric if it's not already
        if not pd.api.types.is_numeric_dtype(df_analysis["ses"]):
            df_analysis["ses"] = pd.to_numeric(df_analysis["ses"], errors='coerce')
            
            # Check if conversion resulted in too many NaN values
            if df_analysis["ses"].isna().sum() > len(df) * 0.5:
                st.warning(t.get("ses_conversion_warning", "Unable to convert SES to numeric values. Using categorical approach."))
                df_analysis["ses"] = df["ses"].rank(pct=True) * 10
        
        # Convert home_support to numeric if necessary
        if not pd.api.types.is_numeric_dtype(df_analysis["home_support"]):
            try:
                df_analysis["home_support"] = pd.to_numeric(df_analysis["home_support"], errors='coerce')
            except:
                # If conversion fails, use rank transformation
                df_analysis["home_support"] = df_analysis["home_support"].rank(pct=True) * 10
        
        # Calculate total score (sum of selected columns)
        df_analysis["total_score"] = df_analysis[selected_columns].sum(axis=1)
        
        # Create SES groups (Low, Medium, High)
        if len(df_analysis["ses"].dropna()) > 2:
            ses_bins = [df_analysis["ses"].min() - 0.001, 
                        df_analysis["ses"].quantile(0.33), 
                        df_analysis["ses"].quantile(0.67), 
                        df_analysis["ses"].max() + 0.001]
            
            df_analysis["ses_group"] = pd.cut(
                df_analysis["ses"], 
                bins=ses_bins, 
                labels=["Low", "Medium", "High"]
            )
        else:
            df_analysis["ses_group"] = "Insufficient Data"
        
        # Create home support groups
        if len(df_analysis["home_support"].dropna()) > 2:
            support_bins = [df_analysis["home_support"].min() - 0.001, 
                          df_analysis["home_support"].quantile(0.33), 
                          df_analysis["home_support"].quantile(0.67), 
                          df_analysis["home_support"].max() + 0.001]
            
            df_analysis["home_support_group"] = pd.cut(
                df_analysis["home_support"], 
                bins=support_bins, 
                labels=["Low", "Medium", "High"]
            )
        else:
            df_analysis["home_support_group"] = "Insufficient Data"
        
        return df_analysis
    
    except Exception as e:
        st.error(f"{t.get('data_preparation_error', 'Error preparing data')}: {str(e)}")
        return None


def _show_overview(df_analysis, selected_columns, t):
    """
    Display overview of contextual factors analysis.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    st.header(t.get("overview_title", "Overview"))
    
    # Display key metrics
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric(
            t.get("total_students", "Total Students"),
            len(df_analysis)
        )
    
    with col2:
        if pd.api.types.is_numeric_dtype(df_analysis["ses"]):
            mean_ses = df_analysis["ses"].mean()
            st.metric(
                t.get("average_ses", "Average SES"),
                f"{mean_ses:.2f}"
            )
        else:
            st.metric(
                t.get("average_ses", "Average SES"),
                "N/A"
            )
    
    with col3:
        if pd.api.types.is_numeric_dtype(df_analysis["home_support"]):
            mean_support = df_analysis["home_support"].mean()
            st.metric(
                t.get("average_home_support", "Average Home Support"),
                f"{mean_support:.2f}"
            )
        else:
            st.metric(
                t.get("average_home_support", "Average Home Support"),
                "N/A"
            )
    
    # Display distribution charts
    st.subheader(t.get("distribution_analysis", "Distribution Analysis"))
    
    col1, col2 = st.columns(2)
    
    with col1:
        # SES distribution
        fig = px.histogram(
            df_analysis,
            x="ses",
            nbins=30,
            title=t.get("ses_distribution", "SES Distribution"),
            labels={"ses": t.get("ses", "Socioeconomic Status")}
        )
        st.plotly_chart(fig, use_container_width=True)
    
    with col2:
        # Home support distribution
        fig = px.histogram(
            df_analysis,
            x="home_support",
            nbins=30,
            title=t.get("home_support_distribution", "Home Support Distribution"),
            labels={"home_support": t.get("home_support", "Home Support")}
        )
        st.plotly_chart(fig, use_container_width=True)
    
    # Quick correlation overview
    st.subheader(t.get("quick_correlations", "Quick Correlations"))
    
    correlation_vars = ["ses", "home_support", "total_score"]
    numeric_corr_vars = [col for col in correlation_vars if pd.api.types.is_numeric_dtype(df_analysis[col])]
    
    if len(numeric_corr_vars) >= 2:
        corr_matrix = df_analysis[numeric_corr_vars].corr().round(2)
        
        fig = px.imshow(
            corr_matrix,
            labels=dict(color="Correlation"),
            x=numeric_corr_vars,
            y=numeric_corr_vars,
            color_continuous_scale="RdBu_r",
            aspect="auto"
        )
        st.plotly_chart(fig, use_container_width=True)


def _show_ses_analysis(df_analysis, selected_columns, t):
    """
    Display analysis of socioeconomic status impact.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    st.header(t.get("ses_analysis_title", "SES Analysis"))
    
    # Check if SES is numeric
    if not pd.api.types.is_numeric_dtype(df_analysis["ses"]):
        st.warning(t.get("ses_not_numeric", "SES variable is not numeric. Limited analysis available."))
        _show_categorical_ses_analysis(df_analysis, selected_columns, t)
        return
    
    # Regression analysis
    st.subheader(t.get("ses_relation", "Relationship Between SES and Scores"))
    
    # Create scatter plots with regression lines for each selected column
    for i in range(0, len(selected_columns), 2):
        col1, col2 = st.columns(2)
        
        with col1:
            if i < len(selected_columns):
                _create_regression_plot(df_analysis, "ses", selected_columns[i], t)
        
        with col2:
            if i + 1 < len(selected_columns):
                _create_regression_plot(df_analysis, "ses", selected_columns[i + 1], t)
    
    # Also show regression plot for total score
    st.subheader(t.get("ses_total_score_relation", "SES vs Total Score"))
    _create_regression_plot(df_analysis, "ses", "total_score", t)
    
    # Add boxplot view of SES groups
    st.subheader(t.get("ses_group_comparison", "Performance by SES Group"))
    
    if "ses_group" in df_analysis.columns and df_analysis["ses_group"].nunique() > 1:
        fig = px.box(
            df_analysis,
            x="ses_group",
            y="total_score",
            color="ses_group",
            title=t.get("total_score_by_ses", "Total Score by SES Group"),
            labels={
                "ses_group": t.get("ses_group", "SES Group"),
                "total_score": t.get("total_score", "Total Score")
            },
            category_orders={"ses_group": ["Low", "Medium", "High"]}
        )
        
        fig.update_layout(showlegend=False)
        st.plotly_chart(fig, use_container_width=True)
        
        # Show insights
        _show_ses_insights(df_analysis, selected_columns, t)


def _show_categorical_ses_analysis(df_analysis, selected_columns, t):
    """
    Display analysis for categorical SES data.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    # Bar charts for categorical SES
    for col in selected_columns:
        fig = px.box(
            df_analysis,
            x="ses",
            y=col,
            title=f"{t.get('ses', 'SES')} vs {t['columns_of_interest'].get(col, col)}",
            labels={
                "ses": t.get("ses", "Socioeconomic Status"),
                col: t["columns_of_interest"].get(col, col)
            }
        )
        st.plotly_chart(fig, use_container_width=True)


def _show_home_support_analysis(df_analysis, selected_columns, t):
    """
    Display analysis of home support impact.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    st.header(t.get("home_support_analysis_title", "Home Support Analysis"))
    
    # Show distribution of home support
    st.subheader(t.get("home_support_distribution", "Home Support Distribution"))
    
    if "home_support_group" in df_analysis.columns:
        support_counts = df_analysis["home_support_group"].value_counts().reset_index()
        support_counts.columns = ["home_support_group", "count"]
        support_counts["percentage"] = (support_counts["count"] / support_counts["count"].sum() * 100).round(1)
        
        if len(support_counts) >= 2:
            fig = px.bar(
                support_counts,
                x="home_support_group",
                y="count",
                text="percentage",
                title=t.get("home_support_distribution_chart", "Home Support Group Distribution"),
                labels={
                    "home_support_group": t.get("home_support", "Home Support"),
                    "count": t.get("student_count", "Number of Students"),
                    "percentage": t.get("percentage", "Percentage")
                },
                color="home_support_group",
                category_orders={"home_support_group": ["Low", "Medium", "High"]}
            )
            
            fig.update_traces(texttemplate='%{text:.1f}%', textposition='outside')
            fig.update_layout(showlegend=False)
            st.plotly_chart(fig, use_container_width=True)
    
    # Impact analysis
    st.subheader(t.get("home_support_impact", "Impact of Home Support on Scores"))
    
    # Display boxplots for each selected variable by home support
    for i in range(0, len(selected_columns), 2):
        col1, col2 = st.columns(2)
        
        with col1:
            if i < len(selected_columns):
                fig = px.box(
                    df_analysis,
                    x="home_support_group",
                    y=selected_columns[i],
                    color="home_support_group",
                    title=t.get("histogram_title", "Distribution of {}").format(
                        t["columns_of_interest"].get(selected_columns[i], selected_columns[i])
                    ),
                    labels={
                        "home_support_group": t.get("home_support", "Home Support"),
                        selected_columns[i]: t["columns_of_interest"].get(selected_columns[i], selected_columns[i])
                    },
                    category_orders={"home_support_group": ["Low", "Medium", "High"]}
                )
                
                fig.update_layout(showlegend=False)
                st.plotly_chart(fig, use_container_width=True)
        
        with col2:
            if i + 1 < len(selected_columns):
                fig = px.box(
                    df_analysis,
                    x="home_support_group",
                    y=selected_columns[i + 1],
                    color="home_support_group",
                    title=t.get("histogram_title", "Distribution of {}").format(
                        t["columns_of_interest"].get(selected_columns[i + 1], selected_columns[i + 1])
                    ),
                    labels={
                        "home_support_group": t.get("home_support", "Home Support"),
                        selected_columns[i + 1]: t["columns_of_interest"].get(selected_columns[i + 1], selected_columns[i + 1])
                    },
                    category_orders={"home_support_group": ["Low", "Medium", "High"]}
                )
                
                fig.update_layout(showlegend=False)
                st.plotly_chart(fig, use_container_width=True)
    
    # Show insights
    _show_home_support_insights(df_analysis, selected_columns, t)


def _show_combined_factor_analysis(df_analysis, selected_columns, t):
    """
    Display analysis of combined contextual factors.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    st.header(t.get("combined_analysis_title", "Combined Factor Analysis"))
    
    # Create correlation matrix
    st.subheader(t.get("correlation_matrix", "Correlation Matrix"))
    
    correlation_vars = ["ses", "home_support", "total_score"] + selected_columns
    
    # Filter to only include numeric columns that actually exist
    numeric_vars = [col for col in correlation_vars 
                   if col in df_analysis.columns and pd.api.types.is_numeric_dtype(df_analysis[col])]
    
    if len(numeric_vars) >= 2:
        # Calculate correlation matrix
        corr_matrix = df_analysis[numeric_vars].corr().round(2)
        
        # Create translation mapping for variable names
        var_translations = {}
        for var in numeric_vars:
            if var == "ses":
                var_translations[var] = t.get("ses", "SES")
            elif var == "home_support":
                var_translations[var] = t.get("home_support", "Home Support")
            elif var == "total_score":
                var_translations[var] = t.get("total_score", "Total Score")
            else:
                var_translations[var] = t["columns_of_interest"].get(var, var)
        
        # Create heatmap
        fig = ff.create_annotated_heatmap(
            z=corr_matrix.values,
            x=[var_translations.get(col, col) for col in corr_matrix.columns],
            y=[var_translations.get(col, col) for col in corr_matrix.index],
            annotation_text=corr_matrix.values.round(2),
            colorscale='RdBu_r',
            showscale=True,
            zmin=-1, zmax=1
        )
        
        # Update layout for better readability
        fig.update_layout(
            title=t.get("correlation_matrix_title", "Correlation Matrix"),
            height=600,
            xaxis_title=t.get("variables", "Variables"),
            yaxis_title=t.get("variables", "Variables")
        )
        
        # Rotate x-axis labels for better readability
        fig.update_xaxes(tickangle=-45)
        
        st.plotly_chart(fig, use_container_width=True)
        
        # Add explanatory text for correlation interpretation
        st.markdown(t.get("correlation_interpretation", """
        **Interpreting the Correlation Matrix:**
        - Values close to 1 indicate a strong positive relationship
        - Values close to -1 indicate a strong negative relationship
        - Values close to 0 indicate a weak or no relationship
        """))
        
        # Show interaction effects
        st.subheader(t.get("interaction_effects", "Interaction Effects"))
        
        # Create heatmap of SES x Home Support interaction for total score
        if all(col in df_analysis.columns for col in ["ses_group", "home_support_group", "total_score"]):
            if df_analysis["ses_group"].nunique() > 1 and df_analysis["home_support_group"].nunique() > 1:
                try:
                    # Calculate mean total score for each SES x Home Support combination
                    heatmap_data = df_analysis.groupby(["ses_group", "home_support_group"])["total_score"].mean().reset_index()
                    
                    # Pivot the data for heatmap format
                    heatmap_pivot = heatmap_data.pivot(
                        index="ses_group", 
                        columns="home_support_group", 
                        values="total_score"
                    )
                    
                    # Ensure correct order of categories
                    if isinstance(heatmap_pivot.index, pd.CategoricalIndex):
                        heatmap_pivot = heatmap_pivot.reindex(["Low", "Medium", "High"])
                    
                    if isinstance(heatmap_pivot.columns, pd.CategoricalIndex):
                        heatmap_pivot = heatmap_pivot.reindex(columns=["Low", "Medium", "High"])
                    
                    # Create heatmap
                    fig = ff.create_annotated_heatmap(
                        z=heatmap_pivot.values,
                        x=list(heatmap_pivot.columns),
                        y=list(heatmap_pivot.index),
                        annotation_text=np.around(heatmap_pivot.values, 1),
                        colorscale='Viridis',
                        showscale=True
                    )
                    
                    # Update layout
                    fig.update_layout(
                        title=t.get("interaction_effect", "Interaction Effect on Total Score"),
                        xaxis_title=t.get("home_support", "Home Support"),
                        yaxis_title=t.get("ses_group", "SES Group")
                    )
                    
                    st.plotly_chart(fig, use_container_width=True)
                    
                    # Add interpretation
                    _interpret_interaction_effects(heatmap_pivot, t)
                except Exception as e:
                    st.warning(f"{t.get('heatmap_error', 'Could not create interaction heatmap')}: {str(e)}")
            else:
                st.info(t.get("insufficient_groups", "Insufficient group variation for interaction analysis"))


def _provide_export_options(df_analysis, selected_columns, t):
    """
    Provide options to export analysis results.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    st.header(t.get("export_title", "Export Results"))
    
    # Create analysis summary
    analysis_summary = _create_analysis_summary(df_analysis, selected_columns, t)
    
    col1, col2 = st.columns(2)
    
    # CSV Export
    with col1:
        # Create CSV content
        csv_data = pd.DataFrame()
        
        # Add summary statistics
        for key, value in analysis_summary.items():
            if isinstance(value, pd.DataFrame):
                csv_data = pd.concat([csv_data, value], axis=0)
        
        csv = csv_data.to_csv(index=True).encode('utf-8-sig')
        
        st.download_button(
            t.get("export_csv", "Export to CSV"),
            csv,
            "contextual_factors_analysis.csv",
            "text/csv",
            key='download-contextual-csv'
        )
    
    # Word Export
    with col2:
        if st.button(t.get("export_word", "Generate Word Report")):
            try:
                doc = _create_contextual_factors_word_report(df_analysis, selected_columns, analysis_summary, t)
                
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    doc.save(tmp.name)
                    with open(tmp.name, 'rb') as f:
                        docx = f.read()
                    st.download_button(
                        t.get("download_word", "Download Word Report"),
                        docx,
                        "contextual_factors_report.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                os.unlink(tmp.name)
            except Exception as e:
                st.error(f"{t.get('error_report_generation', 'Error generating report')}: {str(e)}")


def _create_regression_plot(df, x_var, y_var, t):
    """
    Create a scatter plot with regression line.
    
    Args:
        df (pd.DataFrame): Dataset
        x_var (str): X-axis variable
        y_var (str): Y-axis variable
        t (dict): Translation dictionary
    """
    # Get translated variable names
    x_var_name = t.get("ses", "Socioeconomic Status") if x_var == "ses" else t.get(x_var, x_var)
    y_var_name = t["columns_of_interest"].get(y_var, y_var)
    
    # Create scatter plot with regression line
    fig = px.scatter(
        df,
        x=x_var,
        y=y_var,
        trendline="ols",
        title=f"{x_var_name} vs {y_var_name}",
        labels={
            x_var: x_var_name,
            y_var: y_var_name
        }
    )
    
    # Get trendline info
    try:
        model_results = px.get_trendline_results(fig)
        
        if model_results is not None and not model_results.empty:
            if isinstance(model_results, pd.DataFrame) and len(model_results.px_fit_results) > 0:
                results = model_results.iloc[0]["px_fit_results"]
                alpha = results.params[0]  # Intercept
                beta = results.params[1]   # Slope
                r_squared = results.rsquared
                
                # Add equation annotation
                fig.add_annotation(
                    x=0.05,
                    y=0.95,
                    xref="paper",
                    yref="paper",
                    text=f"y = {alpha:.2f} + {beta:.2f}x<br>R² = {r_squared:.2f}",
                    showarrow=False,
                    font=dict(size=12),
                    bgcolor="rgba(255, 255, 255, 0.8)",
                    bordercolor="black",
                    borderwidth=1
                )
    except Exception as e:
        st.warning(f"Could not calculate regression statistics: {str(e)}")
    
    # Display the plot
    st.plotly_chart(fig, use_container_width=True)


def _show_ses_insights(df_analysis, selected_columns, t):
    """
    Generate and display insights about SES impact.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    st.subheader(t.get("ses_insights", "Key Insights"))
    
    # Calculate correlations between SES and performance measures
    correlations = {}
    if pd.api.types.is_numeric_dtype(df_analysis["ses"]):
        for col in selected_columns:
            if pd.api.types.is_numeric_dtype(df_analysis[col]):
                corr, p_value = stats.pearsonr(
                    df_analysis["ses"].dropna(),
                    df_analysis[col].dropna()
                )
                correlations[col] = {"correlation": corr, "p_value": p_value}
    
    # Calculate group means
    if "ses_group" in df_analysis.columns and df_analysis["ses_group"].nunique() > 1:
        ses_group_means = df_analysis.groupby("ses_group")["total_score"].mean()
        
        # Display performance gap
        if len(ses_group_means) >= 2:
            max_group = ses_group_means.idxmax()
            min_group = ses_group_means.idxmin()
            gap = ses_group_means.max() - ses_group_means.min()
            gap_percentage = (gap / ses_group_means.min()) * 100
            
            st.markdown(f"**{t.get('ses_performance_gap', 'SES Performance Gap')}:**")
            st.markdown(f"- {t.get('highest_ses_group', 'Highest performing SES group')}: **{max_group}** ({ses_group_means[max_group]:.1f})")
            st.markdown(f"- {t.get('lowest_ses_group', 'Lowest performing SES group')}: **{min_group}** ({ses_group_means[min_group]:.1f})")
            st.markdown(f"- {t.get('performance_gap', 'Performance gap')}: **{gap:.1f}** {t.get('points', 'points')} ({gap_percentage:.1f}%)")
    
    # Display significant correlations
    if correlations:
        st.markdown(f"**{t.get('significant_correlations', 'Significant Correlations')}:**")
        for col, stats_data in correlations.items():
            if stats_data['p_value'] < 0.05:
                col_name = t["columns_of_interest"].get(col, col)
                corr = stats_data['correlation']
                
                if abs(corr) >= 0.5:
                    strength = t.get("strong", "strong")
                elif abs(corr) >= 0.3:
                    strength = t.get("moderate", "moderate")
                else:
                    strength = t.get("weak", "weak")
                
                direction = t.get("positive", "positive") if corr > 0 else t.get("negative", "negative")
                
                st.markdown(f"- {col_name}: {corr:.2f} ({strength} {direction})")


def _show_home_support_insights(df_analysis, selected_columns, t):
    """
    Generate and display insights about home support impact.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    st.subheader(t.get("home_support_insights", "Key Insights"))
    
    # Calculate group comparisons
    if "home_support_group" in df_analysis.columns and df_analysis["home_support_group"].nunique() > 1:
        support_group_means = df_analysis.groupby("home_support_group")["total_score"].mean()
        
        # Display performance gap
        if len(support_group_means) >= 2:
            max_group = support_group_means.idxmax()
            min_group = support_group_means.idxmin()
            gap = support_group_means.max() - support_group_means.min()
            gap_percentage = (gap / support_group_means.min()) * 100
            
            st.markdown(f"**{t.get('support_performance_gap', 'Home Support Performance Gap')}:**")
            st.markdown(f"- {t.get('highest_support_group', 'Highest performing support group')}: **{max_group}** ({support_group_means[max_group]:.1f})")
            st.markdown(f"- {t.get('lowest_support_group', 'Lowest performing support group')}: **{min_group}** ({support_group_means[min_group]:.1f})")
            st.markdown(f"- {t.get('performance_gap', 'Performance gap')}: **{gap:.1f}** {t.get('points', 'points')} ({gap_percentage:.1f}%)")


def _interpret_interaction_effects(heatmap_pivot, t):
    """
    Interpret and display interaction effects from heatmap.
    
    Args:
        heatmap_pivot (pd.DataFrame): Pivot table for heatmap
        t (dict): Translation dictionary
    """
    st.markdown(f"**{t.get('interaction_interpretation', 'Interpretation')}:**")
    
    # Find the highest and lowest performing combinations
    if not heatmap_pivot.empty:
        max_val = heatmap_pivot.max().max()
        min_val = heatmap_pivot.min().min()
        
        # Find the groups with max and min values
        max_groups = []
        min_groups = []
        
        for ses in heatmap_pivot.index:
            for support in heatmap_pivot.columns:
                if heatmap_pivot.loc[ses, support] == max_val:
                    max_groups.append((ses, support))
                if heatmap_pivot.loc[ses, support] == min_val:
                    min_groups.append((ses, support))
        
        # Show insights
        for ses, support in max_groups:
            st.markdown(f"- {t.get('highest_combination', 'The highest performing combination is')} **{ses} SES** {t.get('with', 'with')} **{support} {t.get('home_support', 'Home Support')}** ({max_val:.1f})")
        
        for ses, support in min_groups:
            st.markdown(f"- {t.get('lowest_combination', 'The lowest performing combination is')} **{ses} SES** {t.get('with', 'with')} **{support} {t.get('home_support', 'Home Support')}** ({min_val:.1f})")
        
        # Calculate the performance gap
        performance_gap = max_val - min_val
        st.markdown(f"- {t.get('performance_gap', 'Performance gap')}: **{performance_gap:.1f}** {t.get('points', 'points')}")
        
        # Check for compensation effects
        if 'Low' in heatmap_pivot.index and 'High' in heatmap_pivot.columns:
            low_ses_high_support = heatmap_pivot.loc['Low', 'High'] if pd.notna(heatmap_pivot.loc['Low', 'High']) else 0
            
            if 'High' in heatmap_pivot.index and 'Low' in heatmap_pivot.columns:
                high_ses_low_support = heatmap_pivot.loc['High', 'Low'] if pd.notna(heatmap_pivot.loc['High', 'Low']) else 0
                
                if low_ses_high_support > high_ses_low_support:
                    st.markdown(f"- **{t.get('compensation_effect', 'Compensation Effect')}:** {t.get('compensation_explanation', 'High home support appears to compensate for low SES, as this combination outperforms high SES with low home support.')}")


def _create_analysis_summary(df_analysis, selected_columns, t):
    """
    Create a summary of all analysis results.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
        
    Returns:
        dict: Summary of analysis results
    """
    summary = {}
    
    # Basic statistics
    summary['basic_stats'] = {
        'total_students': len(df_analysis),
        'variables_analyzed': selected_columns
    }
    
    # SES statistics
    if pd.api.types.is_numeric_dtype(df_analysis["ses"]):
        summary['ses_stats'] = {
            'mean': df_analysis["ses"].mean(),
            'std': df_analysis["ses"].std(),
            'min': df_analysis["ses"].min(),
            'max': df_analysis["ses"].max()
        }
    
    # Home support statistics
    if pd.api.types.is_numeric_dtype(df_analysis["home_support"]):
        summary['home_support_stats'] = {
            'mean': df_analysis["home_support"].mean(),
            'std': df_analysis["home_support"].std(),
            'min': df_analysis["home_support"].min(),
            'max': df_analysis["home_support"].max()
        }
    
    # Group means
    if "ses_group" in df_analysis.columns:
        summary['ses_group_means'] = df_analysis.groupby("ses_group")[selected_columns + ["total_score"]].mean()
    
    if "home_support_group" in df_analysis.columns:
        summary['home_support_group_means'] = df_analysis.groupby("home_support_group")[selected_columns + ["total_score"]].mean()
    
    # Correlations
    numeric_vars = ["ses", "home_support", "total_score"] + selected_columns
    numeric_vars = [col for col in numeric_vars if col in df_analysis.columns and pd.api.types.is_numeric_dtype(df_analysis[col])]
    
    if len(numeric_vars) >= 2:
        summary['correlations'] = df_analysis[numeric_vars].corr()
    
    return summary


def _create_contextual_factors_word_report(df_analysis, selected_columns, analysis_summary, t):
    """
    Create a Word document with contextual factors analysis.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        analysis_summary (dict): Summary of analysis results
        t (dict): Translation dictionary
        
    Returns:
        Document: Word document with report
    """
    doc = Document()
    
    # Title
    doc.add_heading(t.get("title_contextual_factors", "Contextual Factors Analysis"), level=1)
    
    # Introduction
    doc.add_paragraph(t.get("contextual_report_intro", 
        "This report analyzes the impact of socioeconomic status (SES) and home support on student performance."
    ))
    
    # Basic statistics
    doc.add_heading(t.get("basic_statistics", "Basic Statistics"), level=2)
    
    if 'basic_stats' in analysis_summary:
        doc.add_paragraph(f"{t.get('total_students', 'Total Students')}: {analysis_summary['basic_stats']['total_students']}")
        doc.add_paragraph(f"{t.get('variables_analyzed', 'Variables Analyzed')}: {', '.join(analysis_summary['basic_stats']['variables_analyzed'])}")
    
    # SES Analysis
    doc.add_heading(t.get("ses_analysis", "SES Analysis"), level=2)
    
    if 'ses_stats' in analysis_summary:
        doc.add_paragraph(f"{t.get('mean_ses', 'Mean SES')}: {analysis_summary['ses_stats']['mean']:.2f}")
        doc.add_paragraph(f"{t.get('std_ses', 'Standard Deviation')}: {analysis_summary['ses_stats']['std']:.2f}")
        doc.add_paragraph(f"{t.get('range_ses', 'Range')}: {analysis_summary['ses_stats']['min']:.2f} - {analysis_summary['ses_stats']['max']:.2f}")
    
    # Home Support Analysis
    doc.add_heading(t.get("home_support_analysis", "Home Support Analysis"), level=2)
    
    if 'home_support_stats' in analysis_summary:
        doc.add_paragraph(f"{t.get('mean_support', 'Mean Home Support')}: {analysis_summary['home_support_stats']['mean']:.2f}")
        doc.add_paragraph(f"{t.get('std_support', 'Standard Deviation')}: {analysis_summary['home_support_stats']['std']:.2f}")
        doc.add_paragraph(f"{t.get('range_support', 'Range')}: {analysis_summary['home_support_stats']['min']:.2f} - {analysis_summary['home_support_stats']['max']:.2f}")
    
    # Group Comparisons
    doc.add_heading(t.get("group_comparisons", "Group Comparisons"), level=2)
    
    if 'ses_group_means' in analysis_summary:
        doc.add_heading(t.get("ses_groups", "SES Groups"), level=3)
        table = doc.add_table(rows=1, cols=len(selected_columns) + 2)
        table.style = 'Table Grid'
        
        # Headers
        header_cells = table.rows[0].cells
        header_cells[0].text = t.get("group", "Group")
        header_cells[1].text = t.get("total_score", "Total Score")
        for i, col in enumerate(selected_columns):
            header_cells[i + 2].text = t["columns_of_interest"].get(col, col)
        
        # Data
        for group, means in analysis_summary['ses_group_means'].iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(group)
            row_cells[1].text = f"{means['total_score']:.2f}"
            for i, col in enumerate(selected_columns):
                row_cells[i + 2].text = f"{means[col]:.2f}"
    
    # Add correlation matrix as table
    if 'correlations' in analysis_summary:
        doc.add_heading(t.get("correlation_matrix", "Correlation Matrix"), level=2)
        corr_df = analysis_summary['correlations']
        
        table = doc.add_table(rows=len(corr_df) + 1, cols=len(corr_df.columns) + 1)
        table.style = 'Table Grid'
        
        # Headers
        for i, col in enumerate(corr_df.columns):
            table.cell(0, i + 1).text = str(col)
            table.cell(i + 1, 0).text = str(corr_df.index[i])
        
        # Data
        for i in range(len(corr_df)):
            for j in range(len(corr_df.columns)):
                table.cell(i + 1, j + 1).text = f"{corr_df.iloc[i, j]:.2f}"
    
    # Add recommendations
    doc.add_heading(t.get("recommendations", "Recommendations"), level=2)
    doc.add_paragraph(t.get("recommendation_text", 
        "Based on the analysis of contextual factors, consider the following recommendations:"
    ))
    
    recommendations = [
        t.get("rec1", "Develop targeted support programs for students from low SES backgrounds"),
        t.get("rec2", "Enhance parent engagement initiatives to improve home support"),
        t.get("rec3", "Monitor the interaction between SES and home support for early intervention"),
        t.get("rec4", "Implement differentiated instruction based on contextual factors")
    ]
    
    for rec in recommendations:
        doc.add_paragraph(rec, style='List Bullet')
    
    return doc