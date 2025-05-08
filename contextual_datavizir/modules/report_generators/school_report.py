"""
School comparison report generator for EduInsight.

This module generates comprehensive reports analyzing performance differences
across schools, including visualizations and insights.
"""

import tempfile
import os
import pandas as pd
import numpy as np
import plotly.express as px
from docx import Document
from docx.shared import Inches

def generate_school_report(data, t, doc=None):
    """
    Generate a school comparison analysis report.
    
    Args:
        data (dict): Dictionary containing:
            - df (pd.DataFrame): Original dataset
            - stats_by_school (pd.DataFrame): Statistics by school
            - selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
        doc (Document, optional): Existing Word document to use
        
    Returns:
        Document: Word document with school comparison report
    """
    # Extract data
    df = data.get('df')
    stats_by_school = data.get('stats_by_school')
    selected_columns = data.get('selected_columns')
    
    # Create document if not provided
    if doc is None:
        doc = Document()
    
    # Add title
    doc.add_heading(t.get("title_school_comparison", "School Comparison Analysis"), level=0)
    
    # Add introduction
    doc.add_paragraph(t.get("school_comparison_intro", 
        "This report analyzes performance differences across schools, identifying patterns and areas for improvement."
    ))
    
    # Add school distribution section
    _add_school_distribution(doc, df, t)
    
    # Add statistics table
    _add_statistics_table(doc, stats_by_school, t)
    
    # Add performance visualizations
    _add_performance_visualizations(doc, df, selected_columns, t)
    
    # Add insights
    _add_insights(doc, df, selected_columns, t)
    
    # Add recommendations
    _add_recommendations(doc, df, selected_columns, t)
    
    return doc

def _add_school_distribution(doc, df, t):
    """
    Add school distribution information to the report.
    
    Args:
        doc (Document): Word document
        df (pd.DataFrame): Original dataset
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("school_distribution", "School Distribution"), level=1)
    
    # Count students per school
    school_counts = df["school"].value_counts().reset_index()
    school_counts.columns = ["school", "count"]
    school_counts["percentage"] = (school_counts["count"] / df.shape[0] * 100).round(1)
    
    # Create table for school distribution
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = t.get("school", "School")
    header_cells[1].text = t.get("student_count", "Number of Students")
    header_cells[2].text = t.get("percentage", "Percentage")
    
    # Add data
    for i, row in school_counts.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["school"])
        row_cells[1].text = str(row["count"])
        row_cells[2].text = f"{row['percentage']}%"
    
    # Add space after table
    doc.add_paragraph()
    
    # Add pie chart of school distribution
    with tempfile.TemporaryDirectory() as tmp_dir:
        # Create pie chart
        fig = px.pie(
            school_counts, 
            values="count", 
            names="school",
            title=t.get("school_distribution_chart", "Student Distribution by School"),
            hole=0.4
        )
        
        # Save figure to temporary file
        img_path = os.path.join(tmp_dir, "school_distribution.png")
        fig.write_image(img_path, width=800, height=500)
        
        # Add figure to document
        doc.add_picture(img_path, width=Inches(6))
    
    # Add space after figure
    doc.add_paragraph()

def _add_statistics_table(doc, stats_by_school, t):
    """
    Add statistics table to the report.
    
    Args:
        doc (Document): Word document
        stats_by_school (pd.DataFrame): Statistics by school
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("table_school_comparison", "Descriptive Statistics by School"), level=1)
    
    # Create table for overall stats
    table = doc.add_table(rows=1, cols=len(stats_by_school.columns))
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, col in enumerate(stats_by_school.columns):
        header_cells[i].text = str(col)
    
    # Add data
    for _, row in stats_by_school.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    # Add space after table
    doc.add_paragraph()
    
    # Add description of key statistics
    doc.add_paragraph(t.get("statistics_explanation", 
        """
        Key statistics explanation:
        - count: Number of observations (students)
        - mean: Average score
        - std: Standard deviation (spread of scores)
        - min: Minimum score
        - 25%: First quartile (25% of scores are below this)
        - 50%: Median (middle score)
        - 75%: Third quartile (75% of scores are below this)
        - max: Maximum score
        """
    ))

def _add_performance_visualizations(doc, df, selected_columns, t):
    """
    Add performance visualizations to the report.
    
    Args:
        doc (Document): Word document
        df (pd.DataFrame): Original dataset
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("school_comparison_chart_title", "Score Distribution by School"), level=1)
    
    # Create visualizations for selected columns
    with tempfile.TemporaryDirectory() as tmp_dir:
        for col in selected_columns:
            # Get translated column name
            col_name = t["columns_of_interest"].get(col, col)
            
            # Add subheading for variable
            doc.add_heading(col_name, level=2)
            
            # Create boxplot
            fig = px.box(
                df,
                x="school",
                y=col,
                color="school",
                title=t.get("histogram_title", "Distribution of {}").format(col_name)
            )
            
            # Save figure to temporary file
            img_path = os.path.join(tmp_dir, f"{col}_boxplot.png")
            fig.write_image(img_path, width=800, height=400)
            
            # Add figure to document
            doc.add_picture(img_path, width=Inches(6))
            
            # Add statistical summary for this variable
            summary = []
            mean_values = df.groupby('school')[col].mean().sort_values(ascending=False)
            
            # Best and worst schools for this indicator
            if not mean_values.empty:
                top_school = mean_values.index[0]
                bottom_school = mean_values.index[-1]
                
                summary.append(f"{t.get('highest_school_var', 'Highest performing school')}: {top_school} ({mean_values.iloc[0]:.2f})")
                summary.append(f"{t.get('lowest_school_var', 'Lowest performing school')}: {bottom_school} ({mean_values.iloc[-1]:.2f})")
                
                # Calculate difference between highest and lowest
                max_diff = mean_values.iloc[0] - mean_values.iloc[-1]
                summary.append(f"{t.get('performance_gap', 'Performance gap')}: {max_diff:.2f}")
            
            # Add summary points
            for point in summary:
                doc.add_paragraph(point, style='List Bullet')
            
            doc.add_paragraph()
    
    # Add mean score comparison by school
    doc.add_heading(t.get("mean_score_comparison", "Mean Score Comparison by School"), level=1)
    
    with tempfile.TemporaryDirectory() as tmp_dir:
        # Calculate mean scores by school
        mean_scores = df.groupby("school")[selected_columns].mean().reset_index()
        mean_scores_melted = pd.melt(
            mean_scores, 
            id_vars=["school"], 
            value_vars=selected_columns,
            var_name=t.get("indicator", "Indicator"),
            value_name=t.get("mean_score", "Mean Score")
        )
        
        # Translate column names for display
        mean_scores_melted[t.get("indicator", "Indicator")] = mean_scores_melted[t.get("indicator", "Indicator")].map(
            lambda x: t["columns_of_interest"].get(x, x)
        )
        
        # Create grouped bar chart
        fig = px.bar(
            mean_scores_melted,
            x="school",
            y=t.get("mean_score", "Mean Score"),
            color=t.get("indicator", "Indicator"),
            barmode="group",
            title=t.get("mean_score_by_school", "Mean Scores by School")
        )
        
        # Save figure to temporary file
        img_path = os.path.join(tmp_dir, "mean_scores.png")
        fig.write_image(img_path, width=800, height=500)
        
        # Add figure to document
        doc.add_picture(img_path, width=Inches(6))

def _add_insights(doc, df, selected_columns, t):
    """
    Add insights based on school comparison.
    
    Args:
        doc (Document): Word document
        df (pd.DataFrame): Original dataset
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("school_insights", "Key Insights"), level=1)
    
    # Calculate mean scores by school
    mean_scores = df.groupby("school")[selected_columns].mean()
    
    # Find highest and lowest performing schools overall
    school_overall = mean_scores.mean(axis=1).sort_values(ascending=False)
    
    if not school_overall.empty:
        top_school = school_overall.index[0]
        bottom_school = school_overall.index[-1]
        
        # Add overall insights
        doc.add_paragraph(t.get('overall_performance', 'Overall Performance'), style='Heading 2')
        doc.add_paragraph(f"{t.get('highest_school', 'Highest performing school')}: {top_school}", style='List Bullet')
        doc.add_paragraph(f"{t.get('lowest_school', 'Lowest performing school')}: {bottom_school}", style='List Bullet')
        
        # Calculate performance gap
        performance_gap = school_overall.iloc[0] - school_overall.iloc[-1]
        doc.add_paragraph(f"{t.get('performance_gap', 'Performance gap')}: {performance_gap:.2f}", style='List Bullet')
        
        # Add specific strengths/weaknesses
        doc.add_paragraph(t.get('specific_strengths', 'Specific Strengths/Weaknesses'), style='Heading 2')
        
        # For each school, find their best and worst area
        for school in df['school'].unique():
            if school in mean_scores.index:
                school_scores = mean_scores.loc[school]
                best_area = school_scores.idxmax()
                worst_area = school_scores.idxmin()
                
                best_area_name = t["columns_of_interest"].get(best_area, best_area)
                worst_area_name = t["columns_of_interest"].get(worst_area, worst_area)
                
                doc.add_paragraph(f"{school}: {t.get('strongest_in', 'Strongest in')} {best_area_name}, {t.get('weakest_in', 'weakest in')} {worst_area_name}", style='List Bullet')
    
    # Add pattern analysis
    doc.add_paragraph(t.get('performance_patterns', 'Performance Patterns'), style='Heading 2')
    
    # Look for consistent patterns across schools
    # Check if there's a skill where most schools perform well or poorly
    consistent_strengths = []
    consistent_weaknesses = []
    
    for col in selected_columns:
        # Count how many schools have this as their best/worst area
        best_count = sum(1 for school in mean_scores.index if mean_scores.loc[school].idxmax() == col)
        worst_count = sum(1 for school in mean_scores.index if mean_scores.loc[school].idxmin() == col)
        
        # If more than 1/3 of schools have this as their best/worst, it's a pattern
        if best_count >= len(mean_scores) / 3:
            consistent_strengths.append((col, best_count))
        if worst_count >= len(mean_scores) / 3:
            consistent_weaknesses.append((col, worst_count))
    
    # Add insights about consistent patterns
    if consistent_strengths:
        for col, count in consistent_strengths:
            col_name = t["columns_of_interest"].get(col, col)
            doc.add_paragraph(f"{col_name}: {t.get('common_strength', 'Common strength')} ({count} {t.get('schools', 'schools')})", style='List Bullet')
    
    if consistent_weaknesses:
        for col, count in consistent_weaknesses:
            col_name = t["columns_of_interest"].get(col, col)
            doc.add_paragraph(f"{col_name}: {t.get('common_weakness', 'Common weakness')} ({count} {t.get('schools', 'schools')})", style='List Bullet')
    
    if not consistent_strengths and not consistent_weaknesses:
        doc.add_paragraph(t.get('no_consistent_patterns', 'No consistent patterns across schools. Each school has unique strengths and weaknesses.'), style='List Bullet')

def _add_recommendations(doc, df, selected_columns, t):
    """
    Add recommendations based on school comparison.
    
    Args:
        doc (Document): Word document
        df (pd.DataFrame): Original dataset
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("recommendations", "Recommendations"), level=1)
    
    # Calculate mean scores by school
    mean_scores = df.groupby("school")[selected_columns].mean()
    school_overall = mean_scores.mean(axis=1).sort_values(ascending=False)
    
    # Generic recommendations
    doc.add_paragraph(t.get("recommendation_text", 
        "Based on the analysis of school performance differences, consider the following recommendations:"
    ))
    
    # Generic recommendations
    recommendations = [
        t.get("rec1", "Share best practices from high-performing schools with lower-performing schools"),
        t.get("rec2", "Investigate factors behind performance differences (resources, teaching methods, etc.)"),
        t.get("rec3", "Develop targeted interventions for specific areas of weakness identified in each school"),
        t.get("rec4", "Establish mentoring relationships between schools with complementary strengths")
    ]
    
    for rec in recommendations:
        doc.add_paragraph(rec, style='List Bullet')
    
    # School-specific recommendations
    if not school_overall.empty and len(school_overall) > 1:
        doc.add_heading(t.get("school_specific_recommendations", "School-Specific Recommendations"), level=2)
        
        # For top-performing school
        top_school = school_overall.index[0]
        doc.add_paragraph(f"{top_school}:", style='Heading 3')
        doc.add_paragraph(t.get("top_school_rec1", "Document successful practices to share with other schools"), style='List Bullet')
        doc.add_paragraph(t.get("top_school_rec2", "Focus on addressing any specific areas of weakness"), style='List Bullet')
        
        # For bottom-performing school
        bottom_school = school_overall.index[-1]
        doc.add_paragraph(f"{bottom_school}:", style='Heading 3')
        doc.add_paragraph(t.get("bottom_school_rec1", "Implement comprehensive improvement plan focusing on key performance areas"), style='List Bullet')
        doc.add_paragraph(t.get("bottom_school_rec2", "Consider partnering with higher-performing schools for mentoring and knowledge sharing"), style='List Bullet')
        
        # For schools with specific patterns
        for school in mean_scores.index:
            if school not in [top_school, bottom_school]:
                school_scores = mean_scores.loc[school]
                worst_area = school_scores.idxmin()
                worst_area_name = t["columns_of_interest"].get(worst_area, worst_area)
                
                doc.add_paragraph(f"{school}:", style='Heading 3')
                doc.add_paragraph(f"{t.get('address_weakness', 'Address specific weakness in')} {worst_area_name}", style='List Bullet')
    
    # Add system-level recommendations
    doc.add_heading(t.get("system_recommendations", "System-Level Recommendations"), level=2)
    system_recommendations = [
        t.get("system_rec1", "Establish a school performance monitoring system to track progress over time"),
        t.get("system_rec2", "Create professional learning communities across schools to share expertise"),
        t.get("system_rec3", "Allocate resources based on identified needs and performance gaps"),
        t.get("system_rec4", "Develop leadership capacity in underperforming schools")
    ]
    
    for rec in system_recommendations:
        doc.add_paragraph(rec, style='List Bullet')
