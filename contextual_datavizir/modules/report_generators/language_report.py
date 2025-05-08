"""
Language effect report generator for EduInsight.

This module generates comprehensive reports analyzing how students' language backgrounds
affect their academic performance.
"""

import tempfile
import os
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.figure_factory as ff
from docx import Document
from docx.shared import Inches

def generate_language_report(data, t, doc=None):
    """
    Generate a language effect analysis report.
    
    Args:
        data (dict): Dictionary containing:
            - df_analysis (pd.DataFrame): Dataset with language groups
            - stats_by_language (pd.DataFrame): Statistics by language group
            - selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
        doc (Document, optional): Existing Word document to use
        
    Returns:
        Document: Word document with language effect report
    """
    # Extract data
    df_analysis = data.get('df_analysis')
    stats_by_language = data.get('stats_by_language')
    selected_columns = data.get('selected_columns')
    
    # Create document if not provided
    if doc is None:
        doc = Document()
    
    # Add title
    doc.add_heading(t.get("title_language_effect", "Language Effect Analysis"), level=0)
    
    # Add introduction
    doc.add_paragraph(t.get("language_effect_intro", 
        "This report analyzes how students' language background affects their academic performance, "
        "identifying patterns and areas where language may impact learning outcomes."
    ))
    
    # Add language distribution section
    _add_language_distribution(doc, df_analysis, t)
    
    # Add statistics table
    _add_statistics_table(doc, stats_by_language, t)
    
    # Add performance visualizations
    _add_performance_visualizations(doc, df_analysis, selected_columns, t)
    
    # Add insights
    _add_insights(doc, df_analysis, selected_columns, t)
    
    # Add recommendations
    _add_recommendations(doc, df_analysis, selected_columns, t)
    
    return doc

def _add_language_distribution(doc, df_analysis, t):
    """
    Add language distribution information to the report.
    
    Args:
        doc (Document): Word document
        df_analysis (pd.DataFrame): Dataset with language groups
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("language_distribution", "Language Group Distribution"), level=1)
    
    # Count students per language group
    language_counts = df_analysis["language_group"].value_counts().reset_index()
    language_counts.columns = ["language_group", "count"]
    language_counts["percentage"] = (language_counts["count"] / df_analysis.shape[0] * 100).round(1)
    
    # Create table for language distribution
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = t.get("language_group", "Language Group")
    header_cells[1].text = t.get("student_count", "Number of Students")
    header_cells[2].text = t.get("percentage", "Percentage")
    
    # Add data
    for i, row in language_counts.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["language_group"])
        row_cells[1].text = str(row["count"])
        row_cells[2].text = f"{row['percentage']}%"
    
    # Add space after table
    doc.add_paragraph()
    
    # Add pie chart of language distribution
    with tempfile.TemporaryDirectory() as tmp_dir:
        # Create pie chart
        fig = px.pie(
            language_counts, 
            values="count", 
            names="language_group",
            title=t.get("language_distribution_chart", "Language Background Distribution"),
            hole=0.4
        )
        
        # Save figure to temporary file
        img_path = os.path.join(tmp_dir, "language_distribution.png")
        fig.write_image(img_path, width=800, height=500)
        
        # Add figure to document
        doc.add_picture(img_path, width=Inches(6))
    
    # Add space after figure
    doc.add_paragraph()

def _add_statistics_table(doc, stats_by_language, t):
    """
    Add statistics table to the report.
    
    Args:
        doc (Document): Word document
        stats_by_language (pd.DataFrame): Statistics by language group
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("table_language_effect", "Descriptive Statistics by Language Group"), level=1)
    
    # Create table for overall stats
    table = doc.add_table(rows=1, cols=len(stats_by_language.columns))
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, col in enumerate(stats_by_language.columns):
        header_cells[i].text = str(col)
    
    # Add data
    for _, row in stats_by_language.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    # Add space after table
    doc.add_paragraph()
    
    # Add description of key statistics
    doc.add_paragraph(t.get("statistics_explanation", 
        """
        Key statistics explanation:
        - count: Number of observations (students)
        - mean: Average score
        - std: Standard deviation (spread of scores)
        - min: Minimum score
        - 25%: First quartile (25% of scores are below this)
        - 50%: Median (middle score)
        - 75%: Third quartile (75% of scores are below this)
        - max: Maximum score
        """
    ))

def _add_performance_visualizations(doc, df_analysis, selected_columns, t):
    """
    Add performance visualizations to the report.
    
    Args:
        doc (Document): Word document
        df_analysis (pd.DataFrame): Dataset with language groups
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("language_effect_chart_title", "Score Distribution by Language Group"), level=1)
    
    # Create visualizations for selected columns
    with tempfile.TemporaryDirectory() as tmp_dir:
        for col in selected_columns:
            # Get translated column name
            col_name = t["columns_of_interest"].get(col, col)
            
            # Add subheading for variable
            doc.add_heading(col_name, level=2)
            
            # Create boxplot
            fig = px.box(
                df_analysis,
                x="language_group",
                y=col,
                color="language_group",
                title=t.get("histogram_title", "Distribution of {}").format(col_name)
            )
            
            # Remove legend (redundant with x-axis)
            fig.update_layout(showlegend=False)
            
            # Save figure to temporary file
            img_path = os.path.join(tmp_dir, f"{col}_boxplot.png")
            fig.write_image(img_path, width=800, height=400)
            
            # Add figure to document
            doc.add_picture(img_path, width=Inches(6))
            
            # Add statistical summary for this variable
            summary = []
            mean_values = df_analysis.groupby('language_group')[col].mean().sort_values(ascending=False)
            
            # Best and worst language groups for this indicator
            if not mean_values.empty:
                top_language = mean_values.index[0]
                bottom_language = mean_values.index[-1]
                
                summary.append(f"{t.get('highest_language_var', 'Highest performing language group')}: {top_language} ({mean_values.iloc[0]:.2f})")
                summary.append(f"{t.get('lowest_language_var', 'Lowest performing language group')}: {bottom_language} ({mean_values.iloc[-1]:.2f})")
                
                # Calculate difference between highest and lowest
                max_diff = mean_values.iloc[0] - mean_values.iloc[-1]
                summary.append(f"{t.get('performance_gap', 'Performance gap')}: {max_diff:.2f}")
            
            # Add summary points
            for point in summary:
                doc.add_paragraph(point, style='List Bullet')
            
            doc.add_paragraph()
    
    # Add mean score comparison by language group
    doc.add_heading(t.get("mean_score_comparison", "Mean Score Comparison by Language Group"), level=1)
    
    with tempfile.TemporaryDirectory() as tmp_dir:
        # Calculate mean scores by language group
        mean_scores = df_analysis.groupby("language_group")[selected_columns].mean().reset_index()
        mean_scores_melted = pd.melt(
            mean_scores, 
            id_vars=["language_group"], 
            value_vars=selected_columns,
            var_name=t.get("indicator", "Indicator"),
            value_name=t.get("mean_score", "Mean Score")
        )
        
        # Translate column names for display
        mean_scores_melted[t.get("indicator", "Indicator")] = mean_scores_melted[t.get("indicator", "Indicator")].map(
            lambda x: t["columns_of_interest"].get(x, x)
        )
        
        # Create grouped bar chart
        fig = px.bar(
            mean_scores_melted,
            x="language_group",
            y=t.get("mean_score", "Mean Score"),
            color=t.get("indicator", "Indicator"),
            barmode="group",
            title=t.get("mean_score_by_language", "Mean Scores by Language Group")
        )
        
        # Save figure to temporary file
        img_path = os.path.join(tmp_dir, "mean_scores.png")
        fig.write_image(img_path, width=800, height=500)
        
        # Add figure to document
        doc.add_picture(img_path, width=Inches(6))

def _add_insights(doc, df_analysis, selected_columns, t):
    """
    Add insights based on language group analysis.
    
    Args:
        doc (Document): Word document
        df_analysis (pd.DataFrame): Dataset with language groups
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("language_insights", "Key Insights"), level=1)
    
    # Calculate mean scores by language group
    mean_scores = df_analysis.groupby("language_group")[selected_columns].mean()
    
    # Find highest and lowest performing language groups overall
    language_overall = mean_scores.mean(axis=1).sort_values(ascending=False)
    
    if not language_overall.empty:
        top_language = language_overall.index[0]
        bottom_language = language_overall.index[-1]
        
        # Add overall insights
        doc.add_paragraph(t.get('overall_performance', 'Overall Performance'), style='Heading 2')
        doc.add_paragraph(f"{t.get('highest_language', 'Highest performing language group')}: {top_language}", style='List Bullet')
        doc.add_paragraph(f"{t.get('lowest_language', 'Lowest performing language group')}: {bottom_language}", style='List Bullet')
        
        # Calculate performance gap
        performance_gap = language_overall.iloc[0] - language_overall.iloc[-1]
        gap_percentage = (performance_gap / language_overall.iloc[-1]) * 100
        
        doc.add_paragraph(f"{t.get('language_gap', 'Performance gap')}: {performance_gap:.2f} ({gap_percentage:.1f}%)", style='List Bullet')
        
        # Language effect interpretation
        doc.add_paragraph(t.get('language_interpretation', 'Interpretation'), style='Heading 2')
        
        # English vs. Dutch comparison if available
        if "English Always" in language_overall.index and "Dutch Always" in language_overall.index:
            eng_score = language_overall.loc["English Always"]
            dutch_score = language_overall.loc["Dutch Always"]
            
            if eng_score > dutch_score:
                diff = eng_score - dutch_score
                doc.add_paragraph(f"{t.get('eng_dutch_diff', 'Students who always speak English at home perform better than those who always speak Dutch, with a difference of')} {diff:.2f} {t.get('points', 'points')}.", style='List Bullet')
            elif dutch_score > eng_score:
                diff = dutch_score - eng_score
                doc.add_paragraph(f"{t.get('dutch_eng_diff', 'Students who always speak Dutch at home perform better than those who always speak English, with a difference of')} {diff:.2f} {t.get('points', 'points')}.", style='List Bullet')
        
        # Look at specific strengths/weaknesses
        doc.add_paragraph(t.get('specific_strengths', 'Specific Strengths/Weaknesses by Language'), style='Heading 2')
        
        # For each language group, find their best and worst area
        for lang_group in df_analysis['language_group'].unique():
            if lang_group in mean_scores.index:
                lang_scores = mean_scores.loc[lang_group]
                best_area = lang_scores.idxmax()
                worst_area = lang_scores.idxmin()
                
                best_area_name = t["columns_of_interest"].get(best_area, best_area)
                worst_area_name = t["columns_of_interest"].get(worst_area, worst_area)
                
                doc.add_paragraph(f"{lang_group}: {t.get('strongest_in', 'Strongest in')} {best_area_name}, {t.get('weakest_in', 'weakest in')} {worst_area_name}", style='List Bullet')
        
        # Add implications for instruction
        doc.add_paragraph(t.get('instructional_implications', 'Implications for Instruction'), style='Heading 2')
        
        doc.add_paragraph(t.get('language_match_importance', 'The language match between home and instruction appears to be an important factor in student performance.'), style='List Bullet')
        doc.add_paragraph(t.get('language_support', 'Consider providing additional language support for students with limited proficiency in the language of instruction.'), style='List Bullet')
        doc.add_paragraph(t.get('differentiated_instruction', 'Differentiated instruction based on language background may help address performance gaps.'), style='List Bullet')

def _add_recommendations(doc, df_analysis, selected_columns, t):
    """
    Add recommendations based on language group analysis.
    
    Args:
        doc (Document): Word document
        df_analysis (pd.DataFrame): Dataset with language groups
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("recommendations", "Recommendations"), level=1)
    
    # Introduction to recommendations
    doc.add_paragraph(t.get("recommendation_text", 
        "Based on the analysis of language effects on performance, consider the following recommendations:"
    ))
    
    # General recommendations
    general_recommendations = [
        t.get("lang_rec1", "Provide targeted language support for students who don't speak the language of instruction at home"),
        t.get("lang_rec2", "Develop differentiated instructional strategies based on students' language backgrounds"),
        t.get("lang_rec3", "Create resources in multiple languages to support parent engagement"),
        t.get("lang_rec4", "Consider bilingual approaches for specific content areas where language gaps are largest"),
        t.get("lang_rec5", "Monitor language development alongside academic progress")
    ]
    
    for rec in general_recommendations:
        doc.add_paragraph(rec, style='List Bullet')
    
    # Language group-specific recommendations
    doc.add_heading(t.get("language_specific_recommendations", "Language-Specific Recommendations"), level=2)
    
    # Calculate mean scores to determine recommendations
    language_means = df_analysis.groupby("language_group")[selected_columns].mean()
    language_overall = language_means.mean(axis=1).sort_values()
    
    # Get lowest performing groups (bottom 2 if available)
    if len(language_overall) >= 2:
        bottom_languages = language_overall.index[:2]
        
        for lang in bottom_languages:
            doc.add_paragraph(f"{lang}:", style='Heading 3')
            
            # Add targeted recommendations
            if "English" in lang and "Dutch" in df_analysis["language_teaching"].values:
                doc.add_paragraph(t.get("english_dutch_rec", "Provide additional Dutch language support for English-speaking students"), style='List Bullet')
            elif "Dutch" in lang and "English" in df_analysis["language_teaching"].values:
                doc.add_paragraph(t.get("dutch_english_rec", "Provide additional English language support for Dutch-speaking students"), style='List Bullet')
            elif "Other" in lang:
                doc.add_paragraph(t.get("other_language_rec", "Develop specific support resources for students from other language backgrounds"), style='List Bullet')
            
            # Add general recommendations for lower-performing groups
            doc.add_paragraph(t.get("low_language_rec1", "Implement intensive language development programs"), style='List Bullet')
            doc.add_paragraph(t.get("low_language_rec2", "Create peer language partnerships with native speakers"), style='List Bullet')
    
    # Add instructional recommendations
    doc.add_heading(t.get("instructional_recommendations", "Instructional Approach Recommendations"), level=2)
    
    instructional_recommendations = [
        t.get("inst_rec1", "Incorporate visual supports and graphic organizers in instruction"),
        t.get("inst_rec2", "Provide vocabulary instruction in context, especially for academic language"),
        t.get("inst_rec3", "Use formative assessment to monitor language development"),
        t.get("inst_rec4", "Create language-rich classroom environments"),
        t.get("inst_rec5", "Train teachers in effective strategies for multilingual classrooms")
    ]
    
    for rec in instructional_recommendations:
        doc.add_paragraph(rec, style='List Bullet')
    
    # Add system-level recommendations
    doc.add_heading(t.get("system_recommendations", "System-Level Recommendations"), level=2)
    
    system_recommendations = [
        t.get("sys_rec1", "Develop a comprehensive language support policy"),
        t.get("sys_rec2", "Create professional development focusing on language diversity"),
        t.get("sys_rec3", "Establish community partnerships to support language learning"),
        t.get("sys_rec4", "Evaluate curriculum materials for language accessibility")
    ]
    
    for rec in system_recommendations:
        doc.add_paragraph(rec, style='List Bullet')
