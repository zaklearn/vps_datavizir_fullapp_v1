"""
Gender effect report generator for EduInsight.

This module generates comprehensive reports analyzing how gender affects
student performance across different academic areas.
"""

import tempfile
import os
import pandas as pd
import numpy as np
import plotly.express as px
import scipy.stats as stats
from docx import Document
from docx.shared import Inches

def generate_gender_report(data, t, doc=None):
    """
    Generate a gender effect analysis report.
    
    Args:
        data (dict): Dictionary containing:
            - df_analysis (pd.DataFrame): Dataset with gender mapping
            - mean_scores (pd.DataFrame): Mean scores by gender
            - selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
        doc (Document, optional): Existing Word document to use
        
    Returns:
        Document: Word document with gender effect report
    """
    # Extract data
    df_analysis = data.get('df_analysis')
    mean_scores = data.get('mean_scores', None)
    selected_columns = data.get('selected_columns')
    
    # Create document if not provided
    if doc is None:
        doc = Document()
    
    # Add title
    doc.add_heading(t.get("title_gender_effect", "Gender Effect Analysis"), level=0)
    
    # Add introduction
    doc.add_paragraph(t.get("gender_effect_intro", 
        "This report analyzes how gender affects student performance, identifying patterns and areas "
        "where gender differences may impact learning outcomes."
    ))
    
    # Add gender distribution section
    _add_gender_distribution(doc, df_analysis, t)
    
    # Add performance comparison
    _add_performance_comparison(doc, df_analysis, selected_columns, mean_scores, t)
    
    # Add statistical tests
    _add_statistical_tests(doc, df_analysis, selected_columns, t)
    
    # Add subject-specific patterns
    _add_subject_patterns(doc, df_analysis, selected_columns, t)
    
    # Add recommendations
    _add_recommendations(doc, df_analysis, selected_columns, t)
    
    return doc

def _add_gender_distribution(doc, df_analysis, t):
    """
    Add gender distribution information to the report.
    
    Args:
        doc (Document): Word document
        df_analysis (pd.DataFrame): Dataset with gender mapping
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("gender_distribution", "Gender Distribution"), level=1)
    
    # Count students by gender
    gender_counts = df_analysis["gender"].value_counts().reset_index()
    gender_counts.columns = ["gender", "count"]
    gender_counts["percentage"] = (gender_counts["count"] / df_analysis.shape[0] * 100).round(1)
    
    # Create table for gender distribution
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = t.get("gender", "Gender")
    header_cells[1].text = t.get("student_count", "Number of Students")
    header_cells[2].text = t.get("percentage", "Percentage")
    
    # Add data
    for i, row in gender_counts.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["gender"])
        row_cells[1].text = str(row["count"])
        row_cells[2].text = f"{row['percentage']}%"
    
    # Add space after table
    doc.add_paragraph()
    
    # Add pie chart of gender distribution
    with tempfile.TemporaryDirectory() as tmp_dir:
        # Create pie chart
        fig = px.pie(
            gender_counts, 
            values="count", 
            names="gender",
            title=t.get("gender_distribution_chart", "Gender Distribution"),
            hole=0.4
        )
        
        # Save figure to temporary file
        img_path = os.path.join(tmp_dir, "gender_distribution.png")
        fig.write_image(img_path, width=800, height=500)
        
        # Add figure to document
        doc.add_picture(img_path, width=Inches(6))
    
    # Add space after figure
    doc.add_paragraph()

def _add_performance_comparison(doc, df_analysis, selected_columns, mean_scores, t):
    """
    Add performance comparison by gender to the report.
    
    Args:
        doc (Document): Word document
        df_analysis (pd.DataFrame): Dataset with gender mapping
        selected_columns (list): Selected columns for analysis
        mean_scores (pd.DataFrame, optional): Mean scores by gender if precomputed
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("mean_score_gender", "Average scores by gender and by task"), level=1)
    
    # Calculate mean scores by gender if not provided
    if mean_scores is None:
        mean_scores = df_analysis.groupby("gender")[selected_columns].mean().round(2)
    
    # Create table for mean scores
    table = doc.add_table(rows=1, cols=len(selected_columns) + 1)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = t.get("gender", "Gender")
    
    for i, col in enumerate(selected_columns):
        header_cells[i + 1].text = t["columns_of_interest"].get(col, col)
    
    # Add data
    for gender, row in mean_scores.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(gender)
        
        for i, value in enumerate(row):
            row_cells[i + 1].text = f"{value:.2f}"
    
    # Add space after table
    doc.add_paragraph()
    
    # Add bar chart visualization
    with tempfile.TemporaryDirectory() as tmp_dir:
        # Prepare data for bar chart
        mean_scores_reset = mean_scores.reset_index()
        mean_scores_melted = pd.melt(
            mean_scores_reset,
            id_vars=["gender"],
            value_vars=selected_columns,
            var_name="Indicator",
            value_name="Mean Score"
        )
        
        # Translate column names for display
        mean_scores_melted["Indicator"] = mean_scores_melted["Indicator"].map(
            lambda x: t["columns_of_interest"].get(x, x)
        )
        
        # Create grouped bar chart
        fig = px.bar(
            mean_scores_melted,
            x="Indicator",
            y="Mean Score",
            color="gender",
            barmode="group",
            title=t.get("mean_score_by_gender", "Mean Scores by Gender"),
            labels={
                "Indicator": t.get("indicator", "Indicator"),
                "Mean Score": t.get("mean_score", "Mean Score"),
                "gender": t.get("gender", "Gender")
            }
        )
        
        # Improve layout
        fig.update_layout(xaxis_tickangle=-45)
        
        # Save figure to temporary file
        img_path = os.path.join(tmp_dir, "gender_means.png")
        fig.write_image(img_path, width=800, height=500)
        
        # Add figure to document
        doc.add_picture(img_path, width=Inches(6))
    
    # Add performance distributions
    doc.add_heading(t.get("score_distribution_gender", "Score distribution by gender"), level=1)
    
    # Create visualizations for selected columns
    with tempfile.TemporaryDirectory() as tmp_dir:
        for col in selected_columns:
            # Get translated column name
            col_name = t["columns_of_interest"].get(col, col)
            
            # Add subheading for the variable
            doc.add_heading(col_name, level=2)
            
            # Create boxplot
            fig = px.box(
                df_analysis,
                x="gender",
                y=col,
                color="gender",
                title=t.get("histogram_title", "Distribution of {}").format(col_name),
                labels={
                    "gender": t.get("gender", "Gender"),
                    col: col_name
                }
            )
            
            # Save figure to temporary file
            img_path = os.path.join(tmp_dir, f"{col}_boxplot.png")
            fig.write_image(img_path, width=800, height=400)
            
            # Add figure to document
            doc.add_picture(img_path, width=Inches(6))
            
            # Add statistical summary
            group_stats = df_analysis.groupby("gender")[col].agg(['mean', 'median', 'std']).round(2)
            
            # Create table
            stats_table = doc.add_table(rows=1, cols=4)
            stats_table.style = 'Table Grid'
            
            # Add headers
            header_cells = stats_table.rows[0].cells
            header_cells[0].text = t.get("gender", "Gender")
            header_cells[1].text = t.get("mean", "Mean")
            header_cells[2].text = t.get("median", "Median")
            header_cells[3].text = t.get("std_dev", "Std Dev")
            
            # Add data
            for gender, stats in group_stats.iterrows():
                row_cells = stats_table.add_row().cells
                row_cells[0].text = str(gender)
                row_cells[1].text = f"{stats['mean']:.2f}"
                row_cells[2].text = f"{stats['median']:.2f}"
                row_cells[3].text = f"{stats['std']:.2f}"
            
            # Add space after table
            doc.add_paragraph()
    
    # Add histogram comparison for one selected variable
    if selected_columns:
        # Choose first variable for detailed distribution
        dist_var = selected_columns[0]
        dist_var_name = t["columns_of_interest"].get(dist_var, dist_var)
        
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Create histogram with density curves
            fig = px.histogram(
                df_analysis,
                x=dist_var,
                color="gender",
                barmode="overlay",
                opacity=0.7,
                histnorm="probability density",
                marginal="box",
                title=t.get("distribution_title", "Distribution of {}").format(dist_var_name),
                labels={
                    dist_var: dist_var_name,
                    "gender": t.get("gender", "Gender")
                }
            )
            
            # Save figure to temporary file
            img_path = os.path.join(tmp_dir, "histogram_comparison.png")
            fig.write_image(img_path, width=800, height=500)
            
            # Add figure to document
            doc.add_heading(t.get("detailed_distribution", "Detailed Distribution Comparison"), level=2)
            doc.add_picture(img_path, width=Inches(6))
            
            # Add note about interpretation
            doc.add_paragraph(t.get("distribution_note", 
                "This visualization shows the detailed score distribution for each gender group, "
                "allowing for comparison of not just averages but the entire distribution pattern."
            ))

def _add_statistical_tests(doc, df_analysis, selected_columns, t):
    """
    Add statistical tests for gender differences to the report.
    
    Args:
        doc (Document): Word document
        df_analysis (pd.DataFrame): Dataset with gender mapping
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("stat_tests", "Statistical Tests"), level=1)
    
    # Get gender groups
    gender_groups = df_analysis["gender"].unique()
    
    if len(gender_groups) < 2:
        doc.add_paragraph(t.get("insufficient_groups", "Insufficient gender groups for statistical comparison."))
        return
    
    # Select the first two groups for analysis (typically boy/girl)
    group1, group2 = gender_groups[:2]
    
    # Calculate statistical tests for each selected variable
    results_data = []
    
    for col in selected_columns:
        # Get data for each group
        group1_data = df_analysis[df_analysis["gender"] == group1][col].dropna()
        group2_data = df_analysis[df_analysis["gender"] == group2][col].dropna()
        
        # Only perform test if we have enough data
        if len(group1_data) > 0 and len(group2_data) > 0:
            try:
                # Perform Mann-Whitney U test
                stat, p = stats.mannwhitneyu(group1_data, group2_data, alternative='two-sided')
                
                # Calculate effect size (r = Z / sqrt(N))
                n1, n2 = len(group1_data), len(group2_data)
                Z = stat / np.sqrt((n1 * n2 * (n1 + n2 + 1)) / 12)
                r = Z / np.sqrt(n1 + n2)
                
                # Determine effect size interpretation
                if abs(r) < 0.1:
                    effect_size = t.get("negligible", "Negligible")
                elif abs(r) < 0.3:
                    effect_size = t.get("small", "Small")
                elif abs(r) < 0.5:
                    effect_size = t.get("medium", "Medium")
                else:
                    effect_size = t.get("large", "Large")
                
                # Calculate difference in means
                mean1 = group1_data.mean()
                mean2 = group2_data.mean()
                diff = mean1 - mean2
                
                # Add result to table
                results_data.append({
                    "Variable": t["columns_of_interest"].get(col, col),
                    "Mann-Whitney U": f"{stat:.0f}",
                    "p-value": f"{p:.4f}",
                    "Effect Size": f"{r:.2f} ({effect_size})",
                    "Difference": f"{diff:.2f}",
                    "Significant": p < 0.05,
                    "Higher Group": group1 if mean1 > mean2 else group2
                })
            except:
                # Skip if test fails
                pass
    
    if results_data:
        # Create table in document
        test_table = doc.add_table(rows=1, cols=5)
        test_table.style = 'Table Grid'
        
        # Add headers
        header_cells = test_table.rows[0].cells
        header_cells[0].text = t.get("variable", "Variable")
        header_cells[1].text = "p-value"
        header_cells[2].text = t.get("effect_size", "Effect Size")
        header_cells[3].text = t.get("difference", "Difference")
        header_cells[4].text = t.get("higher_group", "Higher Group")
        
        # Add data
        for result in results_data:
            row_cells = test_table.add_row().cells
            row_cells[0].text = result["Variable"]
            # Add asterisk for significance
            row_cells[1].text = f"{result['p-value']}{'*' if result['Significant'] else ''}"
            row_cells[2].text = result["Effect Size"]
            row_cells[3].text = result["Difference"]
            row_cells[4].text = result["Higher Group"]
        
        # Add significance note
        doc.add_paragraph(t.get("significance_note", "*p < 0.05 indicates statistical significance"))
        
        # Highlight significant findings
        significant_results = [r for r in results_data if r["Significant"]]
        
        if significant_results:
            doc.add_heading(t.get("significant_findings", "Significant Gender Differences"), level=2)
            
            for result in significant_results:
                # Format as "Girls outperform boys in X by Y points (p=0.001, medium effect)"
                higher_gender = result["Higher Group"]
                lower_gender = group2 if higher_gender == group1 else group1
                
                doc.add_paragraph(
                    f"{higher_gender} {t.get('outperform', 'outperform')} {lower_gender} "
                    f"{t.get('in', 'in')} {result['Variable']} "
                    f"{t.get('by', 'by')} {abs(float(result['Difference'])):.2f} "
                    f"{t.get('points', 'points')} "
                    f"(p={result['p-value'].replace('*', '')}, {result['Effect Size'].split('(')[1].replace(')', '')} "
                    f"{t.get('effect', 'effect')})",
                    style='List Bullet'
                )
        else:
            doc.add_paragraph(t.get("no_significant_diff", "No statistically significant gender differences were found."))

def _add_subject_patterns(doc, df_analysis, selected_columns, t):
    """
    Add analysis of gender patterns by subject area to the report.
    
    Args:
        doc (Document): Word document
        df_analysis (pd.DataFrame): Dataset with gender mapping
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    # Check if we have both EGRA and EGMA variables
    from config import egra_columns, egma_columns
    
    egra_selected = [col for col in selected_columns if col in egra_columns]
    egma_selected = [col for col in selected_columns if col in egma_columns]
    
    if egra_selected and egma_selected:
        doc.add_heading(t.get("subject_patterns", "Gender Patterns by Subject Area"), level=1)
        
        # Get gender groups
        gender_groups = df_analysis["gender"].unique()
        
        if len(gender_groups) >= 2:
            group1, group2 = gender_groups[:2]
            
            # Calculate means for each subject area
            group1_egra_mean = df_analysis[df_analysis["gender"] == group1][egra_selected].mean().mean()
            group1_egma_mean = df_analysis[df_analysis["gender"] == group1][egma_selected].mean().mean()
            group2_egra_mean = df_analysis[df_analysis["gender"] == group2][egra_selected].mean().mean()
            group2_egma_mean = df_analysis[df_analysis["gender"] == group2][egma_selected].mean().mean()
            
            # Create data for visualization
            subject_data = pd.DataFrame({
                "Gender": [group1, group1, group2, group2],
                "Subject": [t.get("reading", "Reading"), t.get("math", "Math"), 
                            t.get("reading", "Reading"), t.get("math", "Math")],
                "Mean Score": [group1_egra_mean, group1_egma_mean, group2_egra_mean, group2_egma_mean]
            })
            
            with tempfile.TemporaryDirectory() as tmp_dir:
                # Create bar chart
                fig = px.bar(
                    subject_data,
                    x="Subject",
                    y="Mean Score",
                    color="Gender",
                    barmode="group",
                    title=t.get("subject_comparison", "Gender Comparison by Subject Area")
                )
                
                # Save figure to temporary file
                img_path = os.path.join(tmp_dir, "subject_comparison.png")
                fig.write_image(img_path, width=800, height=500)
                
                # Add figure to document
                doc.add_picture(img_path, width=Inches(6))
            
            # Create table for subject comparison
            subject_table = doc.add_table(rows=1, cols=5)
            subject_table.style = 'Table Grid'
            
            # Add headers
            header_cells = subject_table.rows[0].cells
            header_cells[0].text = t.get("subject", "Subject")
            header_cells[1].text = f"{group1} {t.get('mean', 'Mean')}"
            header_cells[2].text = f"{group2} {t.get('mean', 'Mean')}"
            header_cells[3].text = t.get("difference", "Difference")
            header_cells[4].text = t.get("higher_group", "Higher Group")
            
            # Calculate differences
            reading_diff = group1_egra_mean - group2_egra_mean
            math_diff = group1_egma_mean - group2_egma_mean
            
            # Add reading row
            row_cells = subject_table.add_row().cells
            row_cells[0].text = t.get("reading", "Reading")
            row_cells[1].text = f"{group1_egra_mean:.2f}"
            row_cells[2].text = f"{group2_egra_mean:.2f}"
            row_cells[3].text = f"{abs(reading_diff):.2f}"
            row_cells[4].text = group1 if reading_diff > 0 else group2
            
            # Add math row
            row_cells = subject_table.add_row().cells
            row_cells[0].text = t.get("math", "Math")
            row_cells[1].text = f"{group1_egma_mean:.2f}"
            row_cells[2].text = f"{group2_egma_mean:.2f}"
            row_cells[3].text = f"{abs(math_diff):.2f}"
            row_cells[4].text = group1 if math_diff > 0 else group2
            
            # Add space after table
            doc.add_paragraph()
            
            # Add summary text
            doc.add_paragraph(t.get("subject_insights", "Subject-Specific Insights:"))
            
            # Reading comparison
            if abs(reading_diff) < 0.1:
                doc.add_paragraph(t.get("reading_similar", "Reading performance is similar across genders."), style='List Bullet')
            else:
                higher_gender = group1 if reading_diff > 0 else group2
                doc.add_paragraph(
                    f"{t.get('reading_diff', 'In reading skills,')} {higher_gender} "
                    f"{t.get('higher_by', 'outperform others by')} {abs(reading_diff):.2f} "
                    f"{t.get('points', 'points')}.",
                    style='List Bullet'
                )
            
            # Math comparison
            if abs(math_diff) < 0.1:
                doc.add_paragraph(t.get("math_similar", "Math performance is similar across genders."), style='List Bullet')
            else:
                higher_gender = group1 if math_diff > 0 else group2
                doc.add_paragraph(
                    f"{t.get('math_diff', 'In math skills,')} {higher_gender} "
                    f"{t.get('higher_by', 'outperform others by')} {abs(math_diff):.2f} "
                    f"{t.get('points', 'points')}.",
                    style='List Bullet'
                )
            
            # Compare differences across subjects
            if abs(reading_diff) > abs(math_diff):
                doc.add_paragraph(t.get("greater_reading_gap", "The gender gap is larger in reading than in math."), style='List Bullet')
            elif abs(math_diff) > abs(reading_diff):
                doc.add_paragraph(t.get("greater_math_gap", "The gender gap is larger in math than in reading."), style='List Bullet')
            else:
                doc.add_paragraph(t.get("equal_gaps", "The gender gaps in reading and math are similar in magnitude."), style='List Bullet')

def _add_recommendations(doc, df_analysis, selected_columns, t):
    """
    Add recommendations based on gender analysis to the report.
    
    Args:
        doc (Document): Word document
        df_analysis (pd.DataFrame): Dataset with gender mapping
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("recommendations", "Recommendations"), level=1)
    
    # Introduction to recommendations
    doc.add_paragraph(t.get("recommendation_text", 
        "Based on the analysis of gender effects, consider the following recommendations:"
    ))
    
    # Get gender groups for recommendations
    gender_groups = df_analysis["gender"].unique()
    
    # General recommendations regardless of findings
    general_recommendations = [
        t.get("gender_rec1", "Address gender stereotypes in educational materials and teaching practices"),
        t.get("gender_rec2", "Provide diverse role models that challenge gender-based expectations"),
        t.get("gender_rec3", "Create inclusive learning environments that support all students regardless of gender"),
        t.get("gender_rec4", "Monitor gender gaps in performance and adjust teaching strategies accordingly")
    ]
    
    for rec in general_recommendations:
        doc.add_paragraph(rec, style='List Bullet')
    
    # Subject-specific recommendations if we have both EGRA and EGMA data
    from config import egra_columns, egma_columns
    
    egra_selected = [col for col in selected_columns if col in egra_columns]
    egma_selected = [col for col in selected_columns if col in egma_columns]
    
    if egra_selected and egma_selected and len(gender_groups) >= 2:
        doc.add_heading(t.get("subject_recommendations", "Subject-Specific Recommendations:"), level=2)
        
        group1, group2 = gender_groups[:2]
        
        # Calculate means for each subject area to determine recommendations
        group1_egra_mean = df_analysis[df_analysis["gender"] == group1][egra_selected].mean().mean()
        group1_egma_mean = df_analysis[df_analysis["gender"] == group1][egma_selected].mean().mean()
        group2_egra_mean = df_analysis[df_analysis["gender"] == group2][egra_selected].mean().mean()
        group2_egma_mean = df_analysis[df_analysis["gender"] == group2][egma_selected].mean().mean()
        
        reading_diff = group1_egra_mean - group2_egra_mean
        math_diff = group1_egma_mean - group2_egma_mean
        
        # Reading recommendations
        if abs(reading_diff) >= 0.1:
            lower_gender = group2 if reading_diff > 0 else group1
            doc.add_paragraph(t.get("reading_rec", f"Focus on strengthening reading skills for {lower_gender} through targeted interventions and engaging reading materials"), style='List Bullet')
        
        # Math recommendations
        if abs(math_diff) >= 0.1:
            lower_gender = group2 if math_diff > 0 else group1
            doc.add_paragraph(t.get("math_rec", f"Develop strategies to enhance math confidence and performance for {lower_gender}"), style='List Bullet')
    
    # Add pedagogical recommendations
    doc.add_heading(t.get("pedagogical_recommendations", "Pedagogical Approach Recommendations:"), level=2)
    
    pedagogical_recommendations = [
        t.get("ped_rec1", "Employ a variety of teaching methods to accommodate different learning styles"),
        t.get("ped_rec2", "Create mixed-gender collaborative learning groups"),
        t.get("ped_rec3", "Provide equitable attention and feedback to all students"),
        t.get("ped_rec4", "Use gender-neutral language and examples in teaching materials")
    ]
    
    for rec in pedagogical_recommendations:
        doc.add_paragraph(rec, style='List Bullet')
    
    # Add system-level recommendations
    doc.add_heading(t.get("system_recommendations", "System-Level Recommendations:"), level=2)
    
    system_recommendations = [
        t.get("sys_rec1", "Implement ongoing monitoring of gender-based performance patterns"),
        t.get("sys_rec2", "Provide professional development for teachers on gender-responsive teaching"),
        t.get("sys_rec3", "Review curriculum and assessment methods for potential gender bias"),
        t.get("sys_rec4", "Set targets for reducing any identified gender gaps in academic performance")
    ]
    
    for rec in system_recommendations:
        doc.add_paragraph(rec, style='List Bullet')
