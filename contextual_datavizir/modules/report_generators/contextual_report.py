"""
Contextual factors report generator for EduInsight.

This module generates comprehensive reports analyzing how contextual factors like
socioeconomic status (SES) and home support affect student performance.
"""

import tempfile
import os
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.figure_factory as ff
from docx import Document
from docx.shared import Inches

def generate_contextual_report(data, t, doc=None):
    """
    Generate a contextual factors analysis report.
    
    Args:
        data (dict): Dictionary containing:
            - df_analysis (pd.DataFrame): Dataset with enhanced variables
            - selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
        doc (Document, optional): Existing Word document to use
        
    Returns:
        Document: Word document with contextual factors report
    """
    # Extract data
    df_analysis = data.get('df_analysis')
    selected_columns = data.get('selected_columns')
    
    # Create document if not provided
    if doc is None:
        doc = Document()
    
    # Add title
    doc.add_heading(t.get("title_contextual_factors", "Contextual Factors Analysis"), level=0)
    
    # Add introduction
    doc.add_paragraph(t.get("contextual_factors_intro", 
        "This report analyzes how socioeconomic status (SES) and home support affect student performance, "
        "identifying patterns and areas for targeted intervention."
    ))
    
    # Add SES impact analysis
    _add_ses_analysis(doc, df_analysis, selected_columns, t)
    
    # Add home support impact analysis
    _add_home_support_analysis(doc, df_analysis, selected_columns, t)
    
    # Add combined effect analysis
    _add_combined_effect_analysis(doc, df_analysis, selected_columns, t)
    
    # Add correlation analysis
    _add_correlation_analysis(doc, df_analysis, selected_columns, t)
    
    # Add recommendations
    _add_recommendations(doc, df_analysis, selected_columns, t)
    
    return doc

def _add_ses_analysis(doc, df_analysis, selected_columns, t):
    """
    Add SES impact analysis to the report.
    
    Args:
        doc (Document): Word document
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("ses_relation", "Relationship Between SES and Performance"), level=1)
    
    # Add SES group distribution
    _add_ses_distribution(doc, df_analysis, t)
    
    # Only proceed with regression analysis if SES is numeric
    if pd.api.types.is_numeric_dtype(df_analysis["ses"]):
        # Add regression analysis
        doc.add_heading(t.get("ses_regression", "SES Impact Analysis"), level=2)
        
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Create scatter plot with regression line for total score
            fig = px.scatter(
                df_analysis,
                x="ses",
                y="total_score",
                trendline="ols",
                labels={
                    "ses": t.get("ses", "Socioeconomic Status"),
                    "total_score": t.get("total_score", "Total Score")
                },
                title=t.get("ses_total_score", "SES vs Total Score")
            )
            
            # Save figure to temporary file
            img_path = os.path.join(tmp_dir, "ses_regression.png")
            fig.write_image(img_path, width=800, height=500)
            
            # Add figure to document
            doc.add_picture(img_path, width=Inches(6))
        
        # Calculate correlation and add to the document
        try:
            import scipy.stats as stats
            corr, p = stats.pearsonr(df_analysis["ses"].dropna(), df_analysis["total_score"].dropna())
            
            doc.add_paragraph(f"{t.get('correlation_coefficient', 'Correlation Coefficient')}: {corr:.3f}")
            doc.add_paragraph(f"{t.get('p_value', 'p-value')}: {p:.4f}")
            
            # Add interpretation
            if p < 0.05:
                if corr > 0:
                    doc.add_paragraph(t.get("positive_significant_correlation", 
                        "There is a statistically significant positive relationship between SES and academic performance. "
                        "Higher SES is associated with better performance."
                    ))
                else:
                    doc.add_paragraph(t.get("negative_significant_correlation", 
                        "There is a statistically significant negative relationship between SES and academic performance. "
                        "Lower SES is associated with better performance, which is an unusual pattern."
                    ))
            else:
                doc.add_paragraph(t.get("no_significant_correlation", 
                    "There is no statistically significant relationship between SES and academic performance in this dataset."
                ))
        except:
            # Skip correlation analysis if it fails
            pass
    
    # Add boxplot analysis by SES group
    doc.add_heading(t.get("ses_group_analysis", "Performance by SES Group"), level=2)
    
    with tempfile.TemporaryDirectory() as tmp_dir:
        # Create boxplot
        fig = px.box(
            df_analysis,
            x="ses_group",
            y="total_score",
            color="ses_group",
            labels={
                "ses_group": t.get("ses_group", "SES Group"),
                "total_score": t.get("total_score", "Total Score")
            },
            title=t.get("score_by_ses_group", "Performance by SES Group"),
            category_orders={"ses_group": ["Low", "Medium", "High"]}
        )
        
        # Remove legend (redundant with x-axis)
        fig.update_layout(showlegend=False)
        
        # Save figure to temporary file
        img_path = os.path.join(tmp_dir, "ses_boxplot.png")
        fig.write_image(img_path, width=800, height=500)
        
        # Add figure to document
        doc.add_picture(img_path, width=Inches(6))
    
    # Add SES impact on specific skills
    if len(selected_columns) > 0:
        doc.add_heading(t.get("ses_impact_skills", "SES Impact on Specific Skills"), level=2)
        
        # Calculate correlations for each skill
        if pd.api.types.is_numeric_dtype(df_analysis["ses"]):
            correlations = {}
            p_values = {}
            
            try:
                import scipy.stats as stats
                for col in selected_columns:
                    if pd.api.types.is_numeric_dtype(df_analysis[col]):
                        corr, p = stats.pearsonr(df_analysis["ses"].dropna(), df_analysis[col].dropna())
                        correlations[col] = corr
                        p_values[col] = p
            except:
                # Skip if correlation calculation fails
                pass
            
            if correlations:
                # Create correlation table
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                # Add headers
                header_cells = table.rows[0].cells
                header_cells[0].text = t.get("skill", "Skill")
                header_cells[1].text = t.get("correlation", "Correlation")
                header_cells[2].text = t.get("p_value", "p-value")
                header_cells[3].text = t.get("significance", "Significance")
                
                # Add data
                for col, corr in correlations.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = t["columns_of_interest"].get(col, col)
                    row_cells[1].text = f"{corr:.3f}"
                    row_cells[2].text = f"{p_values[col]:.4f}"
                    row_cells[3].text = t.get("significant", "Significant") if p_values[col] < 0.05 else t.get("not_significant", "Not significant")
                
                # Add space after table
                doc.add_paragraph()
                
                # Add interpretation of most and least affected skills
                if len(correlations) > 1:
                    # Most affected skill
                    most_affected = max(correlations.items(), key=lambda x: abs(x[1]))
                    most_affected_skill = t["columns_of_interest"].get(most_affected[0], most_affected[0])
                    most_affected_corr = most_affected[1]
                    
                    # Least affected skill
                    least_affected = min(correlations.items(), key=lambda x: abs(x[1]))
                    least_affected_skill = t["columns_of_interest"].get(least_affected[0], least_affected[0])
                    least_affected_corr = least_affected[1]
                    
                    doc.add_paragraph(t.get("most_affected_skill", "Most affected skill"), style='Heading 3')
                    doc.add_paragraph(f"{most_affected_skill} ({most_affected_corr:.3f})")
                    
                    if most_affected_corr > 0:
                        doc.add_paragraph(t.get("positive_impact", "Higher SES is associated with better performance in this skill."))
                    else:
                        doc.add_paragraph(t.get("negative_impact", "Lower SES is associated with better performance in this skill."))
                    
                    doc.add_paragraph(t.get("least_affected_skill", "Least affected skill"), style='Heading 3')
                    doc.add_paragraph(f"{least_affected_skill} ({least_affected_corr:.3f})")
                    
                    if abs(least_affected_corr) < 0.1:
                        doc.add_paragraph(t.get("minimal_impact", "This skill shows minimal impact from SES differences."))

def _add_ses_distribution(doc, df_analysis, t):
    """
    Add SES distribution information to the report.
    
    Args:
        doc (Document): Word document
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("ses_distribution", "SES Distribution"), level=2)
    
    # Count students per SES group
    ses_counts = df_analysis["ses_group"].value_counts().reset_index()
    ses_counts.columns = ["ses_group", "count"]
    ses_counts["percentage"] = (ses_counts["count"] / df_analysis.shape[0] * 100).round(1)
    
    # Create table for SES distribution
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = t.get("ses_group", "SES Group")
    header_cells[1].text = t.get("student_count", "Number of Students")
    header_cells[2].text = t.get("percentage", "Percentage")
    
    # Add data
    for i, row in ses_counts.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["ses_group"])
        row_cells[1].text = str(row["count"])
        row_cells[2].text = f"{row['percentage']}%"
    
    # Add space after table
    doc.add_paragraph()
    
    # Add pie chart of SES distribution
    with tempfile.TemporaryDirectory() as tmp_dir:
        # Create pie chart
        fig = px.pie(
            ses_counts, 
            values="count", 
            names="ses_group",
            title=t.get("ses_distribution_chart", "SES Group Distribution"),
            hole=0.4,
            category_orders={"ses_group": ["Low", "Medium", "High"]}
        )
        
        # Save figure to temporary file
        img_path = os.path.join(tmp_dir, "ses_distribution.png")
        fig.write_image(img_path, width=800, height=500)
        
        # Add figure to document
        doc.add_picture(img_path, width=Inches(6))
    
    # Add space after figure
    doc.add_paragraph()

def _add_home_support_analysis(doc, df_analysis, selected_columns, t):
    """
    Add home support impact analysis to the report.
    
    Args:
        doc (Document): Word document
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("home_support_impact", "Impact of Home Support on Performance"), level=1)
    
    # Add home support distribution
    _add_home_support_distribution(doc, df_analysis, t)
    
    # Add home support impact on total score
    doc.add_heading(t.get("home_support_impact_total", "Home Support Impact on Overall Performance"), level=2)
    
    with tempfile.TemporaryDirectory() as tmp_dir:
        # Create boxplot
        fig = px.box(
            df_analysis,
            x="home_support_group",
            y="total_score",
            color="home_support_group",
            labels={
                "home_support_group": t.get("home_support", "Home Support"),
                "total_score": t.get("total_score", "Total Score")
            },
            title=t.get("score_by_support", "Performance by Home Support Level"),
            category_orders={"home_support_group": ["Low", "Medium", "High"]}
        )
        
        # Remove legend (redundant with x-axis)
        fig.update_layout(showlegend=False)
        
        # Save figure to temporary file
        img_path = os.path.join(tmp_dir, "home_support_boxplot.png")
        fig.write_image(img_path, width=800, height=500)
        
        # Add figure to document
        doc.add_picture(img_path, width=Inches(6))
    
    # Calculate statistics by home support group
    support_stats = df_analysis.groupby("home_support_group")["total_score"].agg(['mean', 'std']).round(2)
    
    # Create table for home support group comparison
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = t.get("home_support_group", "Home Support Group")
    header_cells[1].text = t.get("mean", "Mean")
    header_cells[2].text = t.get("std_dev", "Std Dev")
    
    # Add data
    for group, stats in support_stats.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(group)
        row_cells[1].text = f"{stats['mean']:.2f}"
        row_cells[2].text = f"{stats['std']:.2f}"
    
    # Add space after table
    doc.add_paragraph()
    
    # Calculate and show difference between highest and lowest groups
    if len(support_stats) >= 2:
        max_mean = support_stats['mean'].max()
        min_mean = support_stats['mean'].min()
        difference = max_mean - min_mean
        percentage = (difference / min_mean) * 100 if min_mean > 0 else 0
        
        doc.add_paragraph(f"{t.get('high_low_difference', 'Difference between High and Low Support')}: {difference:.2f} ({percentage:.1f}%)")
    
    # Add home support impact on specific skills
    if len(selected_columns) > 0:
        doc.add_heading(t.get("home_support_impact_skills", "Home Support Impact on Specific Skills"), level=2)
        
        # Create visualizations for selected skills
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Select a subset of skills for visualization (max 4)
            visual_columns = selected_columns[:min(4, len(selected_columns))]
            
            for col in visual_columns:
                # Get translated column name
                col_name = t["columns_of_interest"].get(col, col)
                
                # Create boxplot
                fig = px.box(
                    df_analysis,
                    x="home_support_group",
                    y=col,
                    color="home_support_group",
                    labels={
                        "home_support_group": t.get("home_support", "Home Support"),
                        col: col_name
                    },
                    title=f"{t.get('home_support', 'Home Support')} vs {col_name}",
                    category_orders={"home_support_group": ["Low", "Medium", "High"]}
                )
                
                # Remove legend (redundant with x-axis)
                fig.update_layout(showlegend=False)
                
                # Save figure to temporary file
                img_path = os.path.join(tmp_dir, f"home_support_{col}.png")
                fig.write_image(img_path, width=800, height=400)
                
                # Add figure to document
                doc.add_picture(img_path, width=Inches(6))
                
                # Add space after figure
                doc.add_paragraph()

def _add_home_support_distribution(doc, df_analysis, t):
    """
    Add home support distribution information to the report.
    
    Args:
        doc (Document): Word document
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("home_support_distribution", "Home Support Distribution"), level=2)
    
    # Count students per home support group
    support_counts = df_analysis["home_support_group"].value_counts().reset_index()
    support_counts.columns = ["home_support_group", "count"]
    support_counts["percentage"] = (support_counts["count"] / df_analysis.shape[0] * 100).round(1)
    
    # Create table for home support distribution
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = t.get("home_support_group", "Home Support Group")
    header_cells[1].text = t.get("student_count", "Number of Students")
    header_cells[2].text = t.get("percentage", "Percentage")
    
    # Add data
    for i, row in support_counts.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["home_support_group"])
        row_cells[1].text = str(row["count"])
        row_cells[2].text = f"{row['percentage']}%"
    
    # Add space after table
    doc.add_paragraph()
    
    # Add pie chart of home support distribution
    with tempfile.TemporaryDirectory() as tmp_dir:
        # Create pie chart
        fig = px.pie(
            support_counts, 
            values="count", 
            names="home_support_group",
            title=t.get("home_support_distribution_chart", "Home Support Distribution"),
            hole=0.4,
            category_orders={"home_support_group": ["Low", "Medium", "High"]}
        )
        
        # Save figure to temporary file
        img_path = os.path.join(tmp_dir, "home_support_distribution.png")
        fig.write_image(img_path, width=800, height=500)
        
        # Add figure to document
        doc.add_picture(img_path, width=Inches(6))
    
    # Add space after figure
    doc.add_paragraph()

def _add_combined_effect_analysis(doc, df_analysis, selected_columns, t):
    """
    Add combined SES and home support effect analysis to the report.
    
    Args:
        doc (Document): Word document
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("combined_effects", "Combined Effects of SES and Home Support"), level=1)
    
    # Create interaction heatmap for total score
    if all(col in df_analysis.columns for col in ["ses_group", "home_support_group", "total_score"]):
        try:
            # Calculate mean total score for each SES x Home Support combination
            heatmap_data = df_analysis.groupby(["ses_group", "home_support_group"])["total_score"].mean().reset_index()
            
            # Pivot the data for heatmap format
            heatmap_pivot = heatmap_data.pivot(
                index="ses_group", 
                columns="home_support_group", 
                values="total_score"
            )
            
            # Ensure correct order of categories if they're categorical
            if isinstance(heatmap_pivot.index, pd.CategoricalIndex):
                heatmap_pivot = heatmap_pivot.reindex(["Low", "Medium", "High"])
            
            if isinstance(heatmap_pivot.columns, pd.CategoricalIndex):
                heatmap_pivot = heatmap_pivot.reindex(columns=["Low", "Medium", "High"])
            
            with tempfile.TemporaryDirectory() as tmp_dir:
                # Create heatmap
                fig = ff.create_annotated_heatmap(
                    z=heatmap_pivot.values,
                    x=list(heatmap_pivot.columns),
                    y=list(heatmap_pivot.index),
                    annotation_text=np.around(heatmap_pivot.values, 1),
                    colorscale='Viridis',
                    showscale=True
                )
                
                # Update layout
                fig.update_layout(
                    title=t.get("interaction_effect", "Interaction Effect on Total Score"),
                    xaxis_title=t.get("home_support", "Home Support"),
                    yaxis_title=t.get("ses_group", "SES Group")
                )
                
                # Save figure to temporary file
                img_path = os.path.join(tmp_dir, "interaction_heatmap.png")
                fig.write_image(img_path, width=800, height=600)
                
                # Add figure to document
                doc.add_picture(img_path, width=Inches(6))
            
            # Add interpretation
            if not heatmap_pivot.empty:
                doc.add_heading(t.get("heatmap_interpretation", "Interpretation of Interaction Effects"), level=2)
                
                # Find the highest and lowest performing combinations
                max_val = heatmap_pivot.max().max()
                min_val = heatmap_pivot.min().min()
                
                # Find the groups with max and min values
                max_groups = []
                min_groups = []
                
                for ses in heatmap_pivot.index:
                    for support in heatmap_pivot.columns:
                        if heatmap_pivot.at[ses, support] == max_val:
                            max_groups.append((ses, support))
                        if heatmap_pivot.at[ses, support] == min_val:
                            min_groups.append((ses, support))
                
                # Add insights
                for ses, support in max_groups:
                    doc.add_paragraph(f"{t.get('highest_combination', 'The highest performing combination is')} {ses} SES {t.get('with', 'with')} {support} {t.get('home_support', 'Home Support')} ({max_val:.1f})", style='List Bullet')
                
                for ses, support in min_groups:
                    doc.add_paragraph(f"{t.get('lowest_combination', 'The lowest performing combination is')} {ses} SES {t.get('with', 'with')} {support} {t.get('home_support', 'Home Support')} ({min_val:.1f})", style='List Bullet')
                
                # Calculate the performance gap
                performance_gap = max_val - min_val
                doc.add_paragraph(f"{t.get('performance_gap', 'Performance gap')}: {performance_gap:.1f} {t.get('points', 'points')}", style='List Bullet')
                
                # Add highlight if high home support compensates for low SES
                if 'Low' in heatmap_pivot.index and 'High' in heatmap_pivot.columns:
                    low_ses_high_support = heatmap_pivot.at['Low', 'High']
                    
                    if 'High' in heatmap_pivot.index and 'Low' in heatmap_pivot.columns:
                        high_ses_low_support = heatmap_pivot.at['High', 'Low']
                        
                        if low_ses_high_support > high_ses_low_support:
                            doc.add_paragraph(f"{t.get('compensation_effect', 'Compensation Effect')}: {t.get('compensation_explanation', 'High home support appears to compensate for low SES, as this combination outperforms high SES with low home support.')}", style='List Bullet')
        except Exception as e:
            # If heatmap creation fails, add a note
            doc.add_paragraph(t.get("heatmap_error", "Could not create interaction heatmap."))
    
    # Add violin plot comparing distributions
    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fig = px.violin(
                df_analysis,
                x="ses_group",
                y="total_score",
                color="home_support_group",
                box=True,
                points="all",
                labels={
                    "ses_group": t.get("ses_group", "SES Group"),
                    "total_score": t.get("total_score", "Total Score"),
                    "home_support_group": t.get("home_support", "Home Support")
                },
                title=t.get("combined_distribution", "Score Distribution by SES and Home Support"),
                category_orders={
                    "ses_group": ["Low", "Medium", "High"],
                    "home_support_group": ["Low", "Medium", "High"]
                }
            )
            
            # Save figure to temporary file
            img_path = os.path.join(tmp_dir, "combined_violin.png")
            fig.write_image(img_path, width=800, height=600)
            
            # Add figure to document
            doc.add_picture(img_path, width=Inches(6))
    except:
        # If violin plot creation fails, skip it
        pass

def _add_correlation_analysis(doc, df_analysis, selected_columns, t):
    """
    Add correlation analysis to the report.
    
    Args:
        doc (Document): Word document
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("correlation_matrix", "Correlation Matrix"), level=1)
    
    # Select variables for correlation matrix
    correlation_vars = ["ses", "home_support", "total_score"] + selected_columns
    
    # Filter to only include numeric columns that actually exist
    numeric_vars = [col for col in correlation_vars 
                   if col in df_analysis.columns and pd.api.types.is_numeric_dtype(df_analysis[col])]
    
    if len(numeric_vars) >= 2:
        # Calculate correlation matrix
        corr_matrix = df_analysis[numeric_vars].corr().round(2)
        
        # Create translation mapping for variable names
        var_translations = {}
        for var in numeric_vars:
            if var == "ses":
                var_translations[var] = t.get("ses", "SES")
            elif var == "home_support":
                var_translations[var] = t.get("home_support", "Home Support")
            elif var == "total_score":
                var_translations[var] = t.get("total_score", "Total Score")
            else:
                var_translations[var] = t["columns_of_interest"].get(var, var)
        
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Create heatmap
            fig = ff.create_annotated_heatmap(
                z=corr_matrix.values,
                x=[var_translations.get(col, col) for col in corr_matrix.columns],
                y=[var_translations.get(col, col) for col in corr_matrix.index],
                annotation_text=corr_matrix.values.round(2),
                colorscale='Viridis',
                showscale=True,
                zmin=-1, zmax=1
            )
            
            # Update layout for better readability
            fig.update_layout(
                title=t.get("correlation_matrix_title", "Correlation Matrix of Contextual Factors and Performance"),
                height=600,
                xaxis_title=t.get("variables", "Variables"),
                yaxis_title=t.get("variables", "Variables")
            )
            
            # Rotate x-axis labels for better readability
            fig.update_xaxes(tickangle=-45)
            
            # Save figure to temporary file
            img_path = os.path.join(tmp_dir, "correlation_matrix.png")
            fig.write_image(img_path, width=800, height=600)
            
            # Add figure to document
            doc.add_picture(img_path, width=Inches(6))
        
        # Add explanatory text for correlation interpretation
        doc.add_paragraph(t.get("correlation_interpretation", """
        Interpreting the Correlation Matrix:
        - Values close to 1 indicate a strong positive relationship
        - Values close to -1 indicate a strong negative relationship
        - Values close to 0 indicate a weak or no relationship
        """))
        
        # Calculate the correlation between SES and home support
        if "ses" in numeric_vars and "home_support" in numeric_vars:
            ses_support_corr = corr_matrix.loc["ses", "home_support"]
            
            doc.add_paragraph(f"{t.get('ses_support_correlation', 'SES and Home Support Correlation')}: {ses_support_corr:.2f}")
            
            if abs(ses_support_corr) > 0.5:
                doc.add_paragraph(t.get('strong_ses_support_relation', 'There is a strong relationship between socioeconomic status and home support levels.'), style='List Bullet')
            elif abs(ses_support_corr) > 0.3:
                doc.add_paragraph(t.get('moderate_ses_support_relation', 'There is a moderate relationship between socioeconomic status and home support levels.'), style='List Bullet')
            else:
                doc.add_paragraph(t.get('weak_ses_support_relation', 'There is a weak relationship between socioeconomic status and home support levels.'), style='List Bullet')
            
            if ses_support_corr > 0:
                doc.add_paragraph(t.get('positive_ses_support_relation', 'Higher SES tends to be associated with higher levels of home support.'), style='List Bullet')
            else:
                doc.add_paragraph(t.get('negative_ses_support_relation', 'Higher SES tends to be associated with lower levels of home support, which is an unusual pattern.'), style='List Bullet')

def _add_recommendations(doc, df_analysis, selected_columns, t):
    """
    Add recommendations based on contextual factors analysis.
    
    Args:
        doc (Document): Word document
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    doc.add_heading(t.get("recommendations", "Recommendations"), level=1)
    
    # SES-focused recommendations
    doc.add_heading(t.get("ses_recommendations", "SES-Based Interventions"), level=2)
    
    # Check if SES has a strong impact
    if pd.api.types.is_numeric_dtype(df_analysis["ses"]) and pd.api.types.is_numeric_dtype(df_analysis["total_score"]):
        ses_corr = df_analysis[["ses", "total_score"]].corr().iloc[0, 1]
        
        if abs(ses_corr) > 0.3:
            ses_recs = [
                t.get("ses_rec1", "Provide additional resources and support for students from low SES backgrounds"),
                t.get("ses_rec2", "Implement targeted interventions for specific skill areas most affected by SES disparities"),
                t.get("ses_rec3", "Develop school programs that provide access to educational resources typically less available in low-SES homes"),
                t.get("ses_rec4", "Consider community partnerships to address broader SES-related barriers to education")
            ]
        else:
            ses_recs = [
                t.get("ses_rec_weak1", "Continue monitoring SES effects while focusing on other factors that may have stronger impacts"),
                t.get("ses_rec_weak2", "Identify school practices that may be helping mitigate typical SES-related disparities"),
                t.get("ses_rec_weak3", "Share successful approaches with other schools that may be struggling with stronger SES effects")
            ]
        
        for rec in ses_recs:
            doc.add_paragraph(rec, style='List Bullet')
    
    # Home support recommendations
    doc.add_heading(t.get("home_support_recommendations", "Home Support Interventions"), level=2)
    
    # Check support effect status
    if "home_support_group" in df_analysis.columns and len(df_analysis["home_support_group"].unique()) > 1:
        group_means = df_analysis.groupby("home_support_group")["total_score"].mean()
        
        if "High" in group_means.index and "Low" in group_means.index:
            support_diff = group_means["High"] - group_means["Low"]
            
            if support_diff > 0:
                support_recs = [
                    t.get("support_rec1", "Develop parent education programs on supporting learning at home"),
                    t.get("support_rec2", "Create and distribute home learning materials with parent guidance"),
                    t.get("support_rec3", "Implement regular parent-teacher communication about learning support strategies"),
                    t.get("support_rec4", "Consider after-school programs for students with limited home support")
                ]
            else:
                support_recs = [
                    t.get("support_rec_weak1", "Investigate why home support appears to have limited impact in this context"),
                    t.get("support_rec_weak2", "Review how home support is being measured and defined"),
                    t.get("support_rec_weak3", "Identify if school practices may be compensating for home support variations")
                ]
            
            for rec in support_recs:
                doc.add_paragraph(rec, style='List Bullet')
    
    # Combined/targeted recommendations
    doc.add_heading(t.get("targeted_recommendations", "Targeted Recommendations"), level=2)
    
    targeted_recs = [
        t.get("targeted_rec1", "Implement a multi-tiered support system based on both SES and home support factors"),
        t.get("targeted_rec2", "Prioritize resources for students with both low SES and low home support"),
        t.get("targeted_rec3", "Develop differentiated approaches based on each student's contextual profile")
    ]
    
    # Check if we have interaction data for more specific recommendations
    if all(col in df_analysis.columns for col in ["ses_group", "home_support_group", "total_score"]):
        try:
            interaction_data = df_analysis.groupby(["ses_group", "home_support_group"])["total_score"].mean()
            
            # Look for compensation effect (low SES, high support outperforms high SES, low support)
            if ('Low', 'High') in interaction_data.index and ('High', 'Low') in interaction_data.index:
                comp_effect = interaction_data[('Low', 'High')] > interaction_data[('High', 'Low')]
                
                if comp_effect:
                    targeted_recs.append(t.get("compensation_rec", "Focus on improving home support for low-SES families as it shows strong compensatory effects"))
            
            # Find the lowest performing group
            if interaction_data.idxmin():
                low_ses_group, low_support_group = interaction_data.idxmin()
                targeted_recs.append(t.get("lowest_group_rec", f"Implement comprehensive intervention for students with {low_ses_group} SES and {low_support_group} home support, as they show the lowest performance"))
        except:
            # Use default recommendations if interaction analysis fails
            pass
    
    for rec in targeted_recs:
        doc.add_paragraph(rec, style='List Bullet')
    
    # System-level recommendations
    doc.add_heading(t.get("system_recommendations", "System-Level Recommendations"), level=2)
    
    system_recs = [
        t.get("system_rec1", "Develop a comprehensive contextual factors assessment system"),
        t.get("system_rec2", "Create professional development focused on addressing contextual impacts"),
        t.get("system_rec3", "Establish a resource allocation model that accounts for contextual needs"),
        t.get("system_rec4", "Implement regular monitoring of contextual effects on student outcomes")
    ]
    
    for rec in system_recs:
        doc.add_paragraph(rec, style='List Bullet')
