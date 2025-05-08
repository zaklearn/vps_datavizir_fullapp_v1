# modules/language_effect.py
import streamlit as st
import pandas as pd
import plotly.express as px
from docx import Document
from docx.shared import Inches
import tempfile
import os
import logging
from config import translations, egra_columns, egma_columns
from error_handling import (
    error_handler, validate_dataframe, DataValidationError,
    safe_plot_creation, safe_data_operation
)

# Get module logger
logger = logging.getLogger('eduinsight.language_effect')

@error_handler('eduinsight.language_effect')
def show_language_effect(df, language):
    """
    Analyze the effect of home language on performance.
    
    Args:
        df (pd.DataFrame): Dataset containing language variables
        language (str): Selected language for UI text
    """
    logger.info(f"Starting language effect analysis with language: {language}")
    t = translations[language]  # Get translations
    
    st.subheader(t.get("title_language_effect", "Language Effect Analysis"))
    
    # Verify required columns exist
    required_language_columns = ["st_english_home", "st_dutch_home", "st_other_language"]
    if missing_columns := [col for col in required_language_columns if col not in df.columns]:
        logger.warning(f"Missing required columns: {missing_columns}")
        st.warning(f"{t.get('missing_columns', 'Missing required columns')}: {', '.join(missing_columns)}")
        st.info(t.get('language_columns_info', 'This analysis requires language background information (st_english_home, st_dutch_home, st_other_language).'))
        return
    
    # Select variables for analysis
    if selected_columns := _select_variables(df, t):
        try:
            # Prepare analysis data
            df_analysis = _prepare_language_data(df)
            
            if df_analysis is None:
                st.error(t.get('data_preparation_failed', 'Could not prepare data for analysis'))
                return
            
            # Calculate statistics by language group
            stats_by_language = _calculate_language_statistics(df_analysis, selected_columns)
            
            if stats_by_language is not None:
                # Display statistics table
                _display_language_statistics(stats_by_language, t)
                
                # Provide export options
                _provide_export_options(df_analysis, stats_by_language, selected_columns, t)
                
                # Show visualizations
                _show_language_visualizations(df_analysis, selected_columns, t)
            else:
                logger.warning("Failed to calculate language statistics")
                st.warning(t.get("stats_calculation_failed", "Could not calculate statistics by language group."))
        
        except Exception as e:
            logger.exception("Error in language effect analysis")
            st.error(f"{t.get('error_message', 'An error occurred')}: {str(e)}")
    else:
        logger.warning("No columns selected for analysis")
        st.warning(t.get("warning_select_indicator", "Please select at least one indicator to analyze."))

@error_handler('eduinsight.language_effect')
def _select_variables(df, t):
    """
    Allow user to select variables for analysis.
    
    Args:
        df (pd.DataFrame): Dataset
        t (dict): Translation dictionary
        
    Returns:
        list: Selected column names
    """
    logger.debug("Selecting variables for analysis")
    st.subheader(t.get("select_variables", "Select Variables"))
    
    # Get available assessment columns
    available_columns = [col for col in egra_columns + egma_columns if col in df.columns]
    
    # If no assessment columns found, allow selection from all numeric columns
    if not available_columns:
        logger.info("No predefined assessment columns found, using all numeric columns")
        available_columns = [col for col in df.select_dtypes(include=['number']).columns 
                            if col not in {'school', 'stgender', 'ses', 'st_english_home', 'st_dutch_home', 'st_other_language'}]
    
    logger.debug(f"Available columns: {available_columns}")
    
    # Create selection widget
    selected_columns = st.multiselect(
        t.get("variables_all", "Variables:"),
        options=available_columns,
        default=available_columns[:min(3, len(available_columns))],
        format_func=lambda x: t["columns_of_interest"].get(x, x)
    )
    
    logger.info(f"User selected columns: {selected_columns}")
    return selected_columns

@error_handler('eduinsight.language_effect')
def map_language(row):
    """
    Determines the language group based on responses.
    
    Args:
        row (Series): Row of data with language columns
        
    Returns:
        str: Language group classification
    """
    # Handle string values
    if isinstance(row["st_english_home"], str):
        freq_map = {"Always": "Always", "Frequently": "Frequently", "Sometimes": "Sometimes"}
        if row["st_english_home"] == "Always":
            return "English Always"
        elif row["st_dutch_home"] == "Always":
            return "Dutch Always"
        elif row["st_english_home"] in freq_map:
            return f"English {freq_map[row['st_english_home']]}"
        elif row["st_dutch_home"] in freq_map:
            return f"Dutch {freq_map[row['st_dutch_home']]}"
        elif row["st_other_language"] == "Yes":
            return "Other Language"
        else:
            return "Mixed/Other"
    
    # Handle numeric values (assuming 3=Always, 2=Frequently, 1=Sometimes, 0=Never)
    else:
        if row["st_english_home"] == 3:
            return "English Always"
        elif row["st_dutch_home"] == 3:
            return "Dutch Always"
        elif row["st_english_home"] in [2, 1]:
            freq = "Frequently" if row["st_english_home"] == 2 else "Sometimes"
            return f"English {freq}"
        elif row["st_dutch_home"] in [2, 1]:
            freq = "Frequently" if row["st_dutch_home"] == 2 else "Sometimes"
            return f"Dutch {freq}"
        elif row["st_other_language"] == 1:
            return "Other Language"
        else:
            return "Mixed/Other"

@error_handler('eduinsight.language_effect')
def _prepare_language_data(df):
    """
    Prepare dataset with language group classification.
    
    Args:
        df (pd.DataFrame): Original dataset
        
    Returns:
        pd.DataFrame: Dataset with language group column
    """
    logger.debug("Preparing language data")
    
    try:
        df_analysis = df.copy()
        
        # Handle potential data type issues
        for col in ["st_english_home", "st_dutch_home", "st_other_language"]:
            if col in df_analysis.columns and df_analysis[col].dtype == 'object':
                # Try to convert string categories to numeric if they appear to be numeric
                try:
                    if df_analysis[col].str.isnumeric().all():
                        df_analysis[col] = df_analysis[col].astype(float)
                        logger.info(f"Converted {col} from object to numeric")
                except:
                    logger.warning(f"Could not convert {col} to numeric, keeping as is")
        
        # Generate language group classification
        df_analysis["language_group"] = df_analysis.apply(map_language, axis=1)
        
        # Log language group distribution
        language_counts = df_analysis["language_group"].value_counts()
        logger.info(f"Language group distribution: {language_counts.to_dict()}")
        
        return df_analysis
    
    except Exception as e:
        logger.error(f"Error preparing language data: {str(e)}")
        return None

@error_handler('eduinsight.language_effect')
def _calculate_language_statistics(df_analysis, selected_columns):
    """
    Calculate statistics by language group.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with language group column
        selected_columns (list): Columns to analyze
        
    Returns:
        pd.DataFrame: Statistics by language group
    """
    logger.debug("Calculating language statistics")
    
    try:
        # Calculate statistics by language group
        stats_by_language = df_analysis.groupby("language_group")[selected_columns].describe().round(2)
        
        # Create a more usable format with flattened column hierarchy
        stats_by_language.columns = ['_'.join(col).strip() for col in stats_by_language.columns]
        stats_by_language = stats_by_language.reset_index()
        
        logger.info(f"Successfully calculated statistics for {len(stats_by_language)} language groups")
        return stats_by_language
    
    except Exception as e:
        logger.error(f"Error calculating language statistics: {str(e)}")
        return None

@error_handler('eduinsight.language_effect')
def _display_language_statistics(stats_by_language, t):
    """
    Display statistics table in the Streamlit interface.
    
    Args:
        stats_by_language (pd.DataFrame): Statistics by language group
        t (dict): Translation dictionary
    """
    logger.debug("Displaying language statistics")
    
    st.subheader(t.get("table_language_effect", "Descriptive Statistics by Language Group"))
    st.dataframe(stats_by_language, use_container_width=True)
    
    # Add explanation for statistics
    with st.expander(t.get("statistics_explanation", "Understanding these statistics")):
        st.markdown("""
        - **count**: Number of observations (students)
        - **mean**: Average score
        - **std**: Standard deviation (spread of scores)
        - **min**: Minimum score
        - **25%**: First quartile (25% of scores are below this)
        - **50%**: Median (middle score)
        - **75%**: Third quartile (75% of scores are below this)
        - **max**: Maximum score
        """)
    
    # Display language group distribution
    st.subheader(t.get("language_distribution", "Language Group Distribution"))
    
    try:
        # Calculate group counts and percentages
        language_counts = stats_by_language["language_group"].value_counts().reset_index()
        language_counts.columns = ["language_group", "count"]
        language_counts["percentage"] = (language_counts["count"] / language_counts["count"].sum() * 100).round(1)
        
        # Display in a cleaner format
        language_distribution = pd.DataFrame({
            t.get("language_group", "Language Group"): language_counts["language_group"],
            t.get("student_count", "Number of Students"): language_counts["count"],
            t.get("percentage", "Percentage"): language_counts["percentage"].apply(lambda x: f"{x}%")
        })
        
        st.dataframe(language_distribution, use_container_width=True)
        
        # Create pie chart of language distribution
        def create_language_pie_chart():
            fig = px.pie(
                language_counts, 
                values="count", 
                names="language_group",
                title=t.get("language_distribution_chart", "Language Background Distribution"),
                hole=0.4
            )
            return fig
        
        safe_plot_creation(
            create_language_pie_chart,
            "language_pie_chart",
            t.get("visualization_error", "Could not create visualization"),
            "eduinsight.language_effect"
        )
    
    except Exception as e:
        logger.error(f"Error displaying language distribution: {str(e)}")
        st.warning(t.get("distribution_display_error", "Could not display language distribution"))

@error_handler('eduinsight.language_effect')
def _provide_export_options(df_analysis, stats_by_language, selected_columns, t):
    """
    Provide options to export analysis results.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with language groups
        stats_by_language (pd.DataFrame): Statistics by language group
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    logger.debug("Providing export options")
    
    col1, col2 = st.columns(2)
    
    # CSV Export
    with col1:
        try:
            csv = stats_by_language.to_csv(index=False).encode('utf-8-sig')
            st.download_button(
                t.get("export_language_effect_csv", "Download CSV"),
                csv,
                "language_effect_stats.csv",
                "text/csv",
                key='download-language-stats-csv'
            )
            logger.info("CSV export button created successfully")
        except Exception as e:
            logger.error(f"Error creating CSV export: {str(e)}")
            st.error(t.get("export_csv_error", "Error creating CSV export"))
    
    # Word Export
    with col2:
        if st.button(t.get("export_language_effect_word", "Export to Word")):
            try:
                logger.info("Starting Word report generation")
                doc = _create_language_effect_word_report(df_analysis, stats_by_language, selected_columns, t)
                
                if doc is not None:
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                        doc.save(tmp.name)
                        logger.debug(f"Word report saved to temporary file: {tmp.name}")
                        
                        with open(tmp.name, 'rb') as f:
                            docx = f.read()
                        st.download_button(
                            t.get("download_language_effect_word", "Download Word Report"),
                            docx,
                            "language_effect_report.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    os.unlink(tmp.name)
                    logger.info("Word report generated and download button created successfully")
                else:
                    logger.warning("Word report generation returned None")
                    st.warning(t.get("word_report_failed", "Could not generate Word report"))
            except Exception as e:
                logger.error(f"Error generating Word report: {str(e)}")
                st.error(f"{t.get('error_report_generation', 'Error generating report')}: {str(e)}")

@error_handler('eduinsight.language_effect')
def _show_language_visualizations(df_analysis, selected_columns, t):
    """
    Display visualizations of language effect analysis.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with language groups
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    logger.debug("Creating language visualizations")
    
    st.subheader(t.get("language_effect_chart_title", "Score Distribution by Language Group"))
    
    # Show boxplots for each selected column
    for i in range(0, len(selected_columns), 2):
        col1, col2 = st.columns(2)
        
        # Left column visualization
        with col1:
            if i < len(selected_columns):
                safe_plot_creation(
                    lambda: _create_boxplot(df_analysis, "language_group", selected_columns[i], t),
                    f"boxplot_{selected_columns[i]}",
                    t.get("visualization_error", "Could not create visualization"),
                    "eduinsight.language_effect"
                )
        
        # Right column visualization
        with col2:
            if i + 1 < len(selected_columns):
                safe_plot_creation(
                    lambda: _create_boxplot(df_analysis, "language_group", selected_columns[i + 1], t),
                    f"boxplot_{selected_columns[i + 1]}",
                    t.get("visualization_error", "Could not create visualization"),
                    "eduinsight.language_effect"
                )
    
    # Add mean score comparison by language group
    st.subheader(t.get("mean_score_comparison", "Mean Score Comparison by Language Group"))
    
    try:
        mean_scores = df_analysis.groupby("language_group")[selected_columns].mean().reset_index()
        mean_scores_melted = pd.melt(
            mean_scores, 
            id_vars=["language_group"], 
            value_vars=selected_columns,
            var_name=t.get("indicator", "Indicator"),
            value_name=t.get("mean_score", "Mean Score")
        )
        
        # Translate column names for display
        mean_scores_melted[t.get("indicator", "Indicator")] = mean_scores_melted[t.get("indicator", "Indicator")].map(
            lambda x: t["columns_of_interest"].get(x, x)
        )
        
        # Create grouped bar chart
        def create_mean_score_chart():
            fig = px.bar(
                mean_scores_melted,
                x="language_group",
                y=t.get("mean_score", "Mean Score"),
                color=t.get("indicator", "Indicator"),
                barmode="group",
                title=t.get("mean_score_by_language", "Mean Scores by Language Group"),
                height=500
            )
            
            # Improve layout
            fig.update_layout(
                xaxis_title=t.get("language_group", "Language Group"),
                yaxis_title=t.get("mean_score", "Mean Score"),
                legend_title=t.get("indicator", "Indicator")
            )
            
            return fig
        
        safe_plot_creation(
            create_mean_score_chart,
            "mean_score_chart",
            t.get("visualization_error", "Could not create visualization"),
            "eduinsight.language_effect"
        )
    
    except Exception as e:
        logger.error(f"Error creating mean score comparison: {str(e)}")
        st.warning(t.get("mean_score_chart_error", "Could not create mean score comparison chart"))
    
    # Add language effect insights
    st.subheader(t.get("language_insights", "Key Insights"))
    
    safe_data_operation(
        lambda: _generate_language_insights(df_analysis, selected_columns, t),
        "generate_insights",
        t.get("insights_error", "Could not generate insights"),
        "eduinsight.language_effect"
    )

@error_handler('eduinsight.language_effect')
def _create_boxplot(df, x_var, y_var, t):
    """
    Create a boxplot for a specific variable by language group.
    
    Args:
        df (pd.DataFrame): Dataset
        x_var (str): X-axis variable (grouping)
        y_var (str): Y-axis variable (measure)
        t (dict): Translation dictionary
        
    Returns:
        plotly.Figure: Boxplot figure or None if creation fails
    """
    logger.debug(f"Creating boxplot for {y_var} by {x_var}")
    
    # Get translated variable name
    y_var_name = t["columns_of_interest"].get(y_var, y_var)
    
    try:
        # Create boxplot
        fig = px.box(
            df,
            x=x_var,
            y=y_var,
            color=x_var,
            title=t.get("histogram_title", "Distribution of {}").format(y_var_name),
            labels={
                x_var: t.get("language_group", "Language Group"),
                y_var: y_var_name
            }
        )
        
        # Remove legend (redundant with x-axis)
        fig.update_layout(showlegend=False)
        
        # Display the plot
        st.plotly_chart(fig, use_container_width=True)
        logger.debug(f"Boxplot created successfully for {y_var}")
        return fig
    except Exception as e:
        logger.error(f"Error creating boxplot for {y_var}: {str(e)}")
        return None

@error_handler('eduinsight.language_effect')
def _generate_language_insights(df_analysis, selected_columns, t):
    """
    Generate key insights based on language group analysis.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with language groups
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    logger.debug("Generating language insights")
    
    try:
        # Calculate mean scores by language group
        mean_scores = df_analysis.groupby("language_group")[selected_columns].mean()
        
        # Find highest and lowest performing language groups overall
        language_overall = mean_scores.mean(axis=1).sort_values(ascending=False)
        
        if not language_overall.empty:
            top_language = language_overall.index[0]
            bottom_language = language_overall.index[-1]
            
            # Display overall insights
            st.markdown(f"**{t.get('overall_performance', 'Overall Performance')}:**")
            st.markdown(f"- {t.get('highest_language', 'Highest performing language group')}: **{top_language}**")
            st.markdown(f"- {t.get('lowest_language', 'Lowest performing language group')}: **{bottom_language}**")
            
            # Calculate performance gap
            performance_gap = language_overall.iloc[0] - language_overall.iloc[-1]
            gap_percentage = (performance_gap / language_overall.iloc[-1]) * 100 if language_overall.iloc[-1] != 0 else 0
            
            st.markdown(f"- {t.get('language_gap', 'Performance gap')}: **{performance_gap:.2f}** ({gap_percentage:.1f}%)")
            
            # Language effect interpretation
            st.markdown(f"**{t.get('language_interpretation', 'Interpretation')}:**")
            
            if "English Always" in language_overall.index and "Dutch Always" in language_overall.index:
                eng_score = language_overall.loc["English Always"]
                dutch_score = language_overall.loc["Dutch Always"]
                
                if eng_score > dutch_score:
                    diff = eng_score - dutch_score
                    st.markdown(f"- {t.get('eng_dutch_diff', 'Students who always speak English at home perform better than those who always speak Dutch, with a difference of')} **{diff:.2f}** {t.get('points', 'points')}.")
                elif dutch_score > eng_score:
                    diff = dutch_score - eng_score
                    st.markdown(f"- {t.get('dutch_eng_diff', 'Students who always speak Dutch at home perform better than those who always speak English, with a difference of')} **{diff:.2f}** {t.get('points', 'points')}.")
            
            # Look at specific strengths/weaknesses
            st.markdown(f"**{t.get('specific_strengths', 'Specific Strengths/Weaknesses by Language')}:**")
            
            # For each language group, find their best and worst area
            for lang_group in df_analysis['language_group'].unique():
                if lang_group in mean_scores.index:
                    lang_scores = mean_scores.loc[lang_group]
                    if not lang_scores.empty:
                        best_area = lang_scores.idxmax()
                        worst_area = lang_scores.idxmin()
                        
                        best_area_name = t["columns_of_interest"].get(best_area, best_area)
                        worst_area_name = t["columns_of_interest"].get(worst_area, worst_area)
                        
                        st.markdown(f"- **{lang_group}**: {t.get('strongest_in', 'Strongest in')} {best_area_name}, {t.get('weakest_in', 'weakest in')} {worst_area_name}")
            
            # Add implications for instruction
            st.markdown(f"**{t.get('instructional_implications', 'Implications for Instruction')}:**")
            
            st.markdown(f"- {t.get('language_match_importance', 'The language match between home and instruction appears to be an important factor in student performance.')}")
            st.markdown(f"- {t.get('language_support', 'Consider providing additional language support for students with limited proficiency in the language of instruction.')}")
            st.markdown(f"- {t.get('differentiated_instruction', 'Differentiated instruction based on language background may help address performance gaps.')}")
        else:
            logger.warning("No data available for insights generation")
            st.warning(t.get("no_data_for_insights", "Not enough data to generate insights"))
    except Exception as e:
        logger.error(f"Error generating insights: {str(e)}")
        st.warning(t.get("insights_generation_error", "Error generating insights"))

@error_handler('eduinsight.language_effect')
def _create_language_effect_word_report(df_analysis, stats_by_language, selected_columns, t):
    """
    Create a Word document with language effect analysis.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with language groups
        stats_by_language (pd.DataFrame): Statistics by language group
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
        
    Returns:
        Document: Word document with report or None if creation fails
    """
    logger.debug("Creating Word report for language effect analysis")
    
    try:
        # Import language_report to handle the report creation
        from modules.report_generators.language_report import generate_language_report
        
        # Prepare data for report generator
        report_data = {
            "df_analysis": df_analysis,
            "stats_by_language": stats_by_language,
            "selected_columns": selected_columns
        }
        
        # Generate the report using the specialized module
        logger.info("Generating language report using specialized module")
        return generate_language_report(report_data, t)
    except ImportError as e:
        logger.error(f"Error importing language_report module: {str(e)}")
        # Fallback to simple report generation if import fails
        return _create_simple_language_report(df_analysis, stats_by_language, selected_columns, t)
    except Exception as e:
        logger.error(f"Error creating Word report: {str(e)}")
        return None

@error_handler('eduinsight.language_effect')
def _create_simple_language_report(df_analysis, stats_by_language, selected_columns, t):
    """
    Create a simple Word document with language effect analysis.
    Fallback method if specialized report generator is unavailable.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with language groups
        stats_by_language (pd.DataFrame): Statistics by language group
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
        
    Returns:
        Document: Word document with report
    """
    logger.info("Creating simple language report (fallback method)")
    
    try:
        doc = Document()
        
        # Title
        doc.add_heading(t.get("title_language_effect", "Language Effect Analysis"), level=1)
        
        # Introduction
        doc.add_paragraph(t.get("language_effect_intro", 
            "This report analyzes how students' language background affects their academic performance."
        ))
        
        # Language distribution
        doc.add_heading(t.get("language_distribution", "Language Group Distribution"), level=2)
        
        try:
            language_counts = df_analysis["language_group"].value_counts()
            total_students = len(df_analysis)
            
            # Create table for language distribution
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # Add headers
            header_cells = table.rows[0].cells
            header_cells[0].text = t.get("language_group", "Language Group")
            header_cells[1].text = t.get("student_count", "Number of Students")
            header_cells[2].text = t.get("percentage", "Percentage")
            
            # Add data
            for group, count in language_counts.items():
                percentage = (count / total_students) * 100
                row_cells = table.add_row().cells
                row_cells[0].text = str(group)
                row_cells[1].text = str(count)
                row_cells[2].text = f"{percentage:.1f}%"
            
            doc.add_paragraph()
        except Exception as e:
            logger.error(f"Error adding language distribution table: {str(e)}")
            doc.add_paragraph(t.get("table_creation_error", "Error creating distribution table"))
        
        # Add recommendations
        doc.add_heading(t.get("recommendations", "Recommendations"), level=2)
        doc.add_paragraph(t.get("recommendation_text", 
            "Based on the analysis of language effects on performance, consider the following recommendations:"
        ))
        
        # Generic recommendations
        recommendations = [
            t.get("lang_rec1", "Provide targeted language support for students who don't speak the language of instruction at home"),
            t.get("lang_rec2", "Develop differentiated instructional strategies based on students' language backgrounds"),
            t.get("lang_rec3", "Create resources in multiple languages to support parent engagement")
        ]
        
        for rec in recommendations:
            doc.add_paragraph(rec, style='List Bullet')
        
        logger.info("Simple language report created successfully")
        return doc
    
    except Exception as e:
        logger.error(f"Error creating simple report: {str(e)}")
        return None