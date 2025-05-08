import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import plotly.figure_factory as ff
import scipy.stats as stats
from docx import Document
from docx.shared import Inches
import tempfile
import os
from config import translations, egra_columns, egma_columns


def show_ses_home_support(df, language):
    """
    Analyze the impact of SES and home support on student performance with tabbed interface.
    
    Args:
        df (pd.DataFrame): Dataset containing SES and home support variables
        language (str): Selected language for UI text
    """
    t = translations[language]  # Get translations
    
    st.subheader(t.get("title_ses_home_support", "SES & Home Support Analysis"))
    
    # Verify required columns exist
    required_columns = ["ses", "home_support"]
    missing_columns = [col for col in required_columns if col not in df.columns]
    
    if missing_columns:
        st.warning(f"{t.get('missing_columns', 'Missing required columns')}: {', '.join(missing_columns)}")
        st.info(t.get('ses_columns_info', 'This analysis requires socioeconomic status (ses) and home support information.'))
        return
    
    # Select variables for analysis
    selected_columns = _select_variables(df, t)
    
    if selected_columns:
        try:
            # Prepare analysis data
            df_analysis = _prepare_ses_data(df, selected_columns)
            
            if df_analysis is None:
                st.error(t.get("data_preparation_error", "Error preparing data for analysis"))
                return
            
            # Create tabs for different analyses
            tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
                t.get("tab_overview", "Overview"),
                t.get("tab_ses_analysis", "SES Analysis"),
                t.get("tab_home_support", "Home Support"),
                t.get("tab_combined_effects", "Combined Effects"),
                t.get("tab_correlations", "Correlations"),
                t.get("tab_recommendations", "Recommendations & Export")
            ])
            
            with tab1:
                _show_overview_ses(df_analysis, selected_columns, t)
            
            with tab2:
                _show_ses_correlations(df_analysis, selected_columns, t)
            
            with tab3:
                _show_home_support_impact(df_analysis, selected_columns, t)
            
            with tab4:
                _show_combined_effects(df_analysis, selected_columns, t)
            
            with tab5:
                _show_correlation_matrix(df_analysis, selected_columns, t)
            
            with tab6:
                _show_recommendations(df_analysis, selected_columns, t)
                st.divider()
                _provide_export_options(df_analysis, selected_columns, t)
        
        except Exception as e:
            st.error(f"{t.get('error_message', 'An error occurred')}: {str(e)}")
            st.info("Please ensure your data meets the required format for analysis.")
    else:
        st.warning(t.get("warning_select_variable", "Please select at least one variable to analyze."))


def _select_variables(df, t):
    """
    Allow user to select variables for analysis.
    
    Args:
        df (pd.DataFrame): Dataset
        t (dict): Translation dictionary
        
    Returns:
        list: Selected column names
    """
    st.subheader(t.get("select_variables", "Select Variables"))
    
    # Get available assessment columns
    available_columns = []
    for col in egra_columns + egma_columns:
        if col in df.columns:
            available_columns.append(col)
    
    # If no assessment columns found, allow selection from all numeric columns
    if not available_columns:
        numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
        available_columns = [col for col in numeric_cols 
                            if col not in ['school', 'stgender', 'ses', 'home_support']]
    
    # Create selection widget
    selected_columns = st.multiselect(
        t.get("variables_all", "Variables:"),
        options=available_columns,
        default=available_columns[:min(3, len(available_columns))],
        format_func=lambda x: t["columns_of_interest"].get(x, x)
    )
    
    return selected_columns


def _prepare_ses_data(df, selected_columns):
    """
    Prepare dataset with additional calculated variables and proper data types.
    
    Args:
        df (pd.DataFrame): Original dataset
        selected_columns (list): Selected columns for analysis
        
    Returns:
        pd.DataFrame: Enhanced dataset with additional variables
    """
    df_analysis = df.copy()
    
    try:
        # Calculate total score (sum of selected columns)
        df_analysis["total_score"] = df_analysis[selected_columns].sum(axis=1)
        
        # Ensure SES is numeric
        if not pd.api.types.is_numeric_dtype(df_analysis["ses"]):
            try:
                df_analysis["ses"] = pd.to_numeric(df_analysis["ses"], errors='coerce')
            except:
                # If conversion fails, use rank percentiles
                df_analysis["ses"] = df_analysis["ses"].rank(pct=True) * 10
        
        # Create SES groups (Low, Medium, High)
        if df_analysis["ses"].notna().sum() > 2:
            ses_bins = [df_analysis["ses"].min() - 0.001, 
                        df_analysis["ses"].quantile(0.33), 
                        df_analysis["ses"].quantile(0.67), 
                        df_analysis["ses"].max() + 0.001]
            
            df_analysis["ses_group"] = pd.cut(
                df_analysis["ses"], 
                bins=ses_bins, 
                labels=["Low", "Medium", "High"]
            )
        else:
            df_analysis["ses_group"] = "Insufficient Data"
        
        # Ensure home_support is properly formatted
        if not pd.api.types.is_numeric_dtype(df_analysis["home_support"]):
            try:
                df_analysis["home_support"] = pd.to_numeric(df_analysis["home_support"], errors='coerce')
            except:
                # If conversion fails, use rank percentiles
                df_analysis["home_support"] = df_analysis["home_support"].rank(pct=True) * 10
        
        # Create home support groups
        if df_analysis["home_support"].notna().sum() > 2:
            support_bins = [df_analysis["home_support"].min() - 0.001, 
                          df_analysis["home_support"].quantile(0.33), 
                          df_analysis["home_support"].quantile(0.67), 
                          df_analysis["home_support"].max() + 0.001]
            
            df_analysis["home_support_group"] = pd.cut(
                df_analysis["home_support"], 
                bins=support_bins, 
                labels=["Low", "Medium", "High"]
            )
        else:
            df_analysis["home_support_group"] = "Insufficient Data"
        
        return df_analysis
    
    except Exception as e:
        st.error(f"Error in data preparation: {str(e)}")
        return None


def _show_overview_ses(df_analysis, selected_columns, t):
    """
    Display overview of SES and home support analysis.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    st.header(t.get("overview_title", "Overview"))
    
    # Display key metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric(
            t.get("total_students", "Total Students"),
            len(df_analysis)
        )
    
    with col2:
        if pd.api.types.is_numeric_dtype(df_analysis["ses"]):
            mean_ses = df_analysis["ses"].mean()
            st.metric(
                t.get("average_ses", "Average SES"),
                f"{mean_ses:.2f}"
            )
        else:
            st.metric(
                t.get("average_ses", "Average SES"),
                "N/A"
            )
    
    with col3:
        if pd.api.types.is_numeric_dtype(df_analysis["home_support"]):
            mean_support = df_analysis["home_support"].mean()
            st.metric(
                t.get("average_home_support", "Average Home Support"),
                f"{mean_support:.2f}"
            )
        else:
            st.metric(
                t.get("average_home_support", "Average Home Support"),
                "N/A"
            )
    
    with col4:
        if pd.api.types.is_numeric_dtype(df_analysis["ses"]) and pd.api.types.is_numeric_dtype(df_analysis["home_support"]):
            corr = df_analysis[["ses", "home_support"]].corr().iloc[0, 1]
            st.metric(
                t.get("ses_support_correlation", "SES-Support Correlation"),
                f"{corr:.2f}"
            )
        else:
            st.metric(
                t.get("ses_support_correlation", "SES-Support Correlation"),
                "N/A"
            )
    
    # Display distribution charts
    st.subheader(t.get("distribution_analysis", "Distribution Analysis"))
    
    col1, col2 = st.columns(2)
    
    with col1:
        # SES distribution
        if pd.api.types.is_numeric_dtype(df_analysis["ses"]):
            fig = px.histogram(
                df_analysis,
                x="ses",
                nbins=30,
                title=t.get("ses_distribution", "SES Distribution"),
                labels={"ses": t.get("ses", "Socioeconomic Status")},
                color_discrete_sequence=['#1f77b4']
            )
            st.plotly_chart(fig, use_container_width=True)
    
    with col2:
        # Home support distribution
        if pd.api.types.is_numeric_dtype(df_analysis["home_support"]):
            fig = px.histogram(
                df_analysis,
                x="home_support",
                nbins=30,
                title=t.get("home_support_distribution", "Home Support Distribution"),
                labels={"home_support": t.get("home_support", "Home Support")},
                color_discrete_sequence=['#ff7f0e']
            )
            st.plotly_chart(fig, use_container_width=True)
    
    # Add group distribution if groups exist
    if "ses_group" in df_analysis.columns and df_analysis["ses_group"].nunique() > 1:
        st.subheader(t.get("group_distribution", "Group Distribution"))
        
        col1, col2 = st.columns(2)
        
        with col1:
            ses_counts = df_analysis["ses_group"].value_counts()
            fig = px.pie(
                values=ses_counts.values,
                names=ses_counts.index,
                title=t.get("ses_group_distribution", "SES Group Distribution"),
                hole=0.4
            )
            st.plotly_chart(fig, use_container_width=True)
        
        with col2:
            if "home_support_group" in df_analysis.columns and df_analysis["home_support_group"].nunique() > 1:
                support_counts = df_analysis["home_support_group"].value_counts()
                fig = px.pie(
                    values=support_counts.values,
                    names=support_counts.index,
                    title=t.get("support_group_distribution", "Home Support Group Distribution"),
                    hole=0.4
                )
                st.plotly_chart(fig, use_container_width=True)


def _show_ses_correlations(df_analysis, selected_columns, t):
    """
    Display analysis of SES correlations with performance.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    st.header(t.get("ses_relation", "Relationship Between SES and Performance"))
    
    # Only proceed if SES is numeric
    if not pd.api.types.is_numeric_dtype(df_analysis["ses"]):
        st.warning(t.get("ses_not_numeric", "SES variable is not numeric. Some analyses will be limited."))
        _show_categorical_ses_analysis(df_analysis, selected_columns, t)
        return
    
    # Create scatter plots with regression lines for total score
    fig = px.scatter(
        df_analysis,
        x="ses",
        y="total_score",
        trendline="ols",
        labels={
            "ses": t.get("ses", "Socioeconomic Status"),
            "total_score": t.get("total_score", "Total Score")
        },
        title=t.get("ses_total_score", "SES vs Total Score"),
        color_discrete_sequence=['#1f77b4']
    )
    
    # Get trendline info for annotation
    model_results = px.get_trendline_results(fig)
    
    if model_results is not None and len(model_results.px_fit_results) > 0:
        results = model_results.iloc[0]["px_fit_results"]
        alpha = results.params[0]  # Intercept
        beta = results.params[1]   # Slope
        r_squared = results.rsquared
        p_value = results.pvalues[1]  # p-value for slope
        
        # Add equation annotation
        fig.add_annotation(
            x=0.05,
            y=0.95,
            xref="paper",
            yref="paper",
            text=f"y = {alpha:.2f} + {beta:.2f}x<br>R² = {r_squared:.2f} (p = {p_value:.4f})",
            showarrow=False,
            font=dict(size=12),
            bgcolor="rgba(255, 255, 255, 0.8)",
            bordercolor="black",
            borderwidth=1
        )
    
    st.plotly_chart(fig, use_container_width=True)
    
    # Show boxplot of total score by SES group
    if "ses_group" in df_analysis.columns and df_analysis["ses_group"].nunique() > 1:
        st.subheader(t.get("ses_group_performance", "Performance by SES Group"))
        
        fig = px.box(
            df_analysis,
            x="ses_group",
            y="total_score",
            color="ses_group",
            labels={
                "ses_group": t.get("ses_group", "SES Group"),
                "total_score": t.get("total_score", "Total Score")
            },
            title=t.get("score_by_ses_group", "Performance by SES Group"),
            category_orders={"ses_group": ["Low", "Medium", "High"]}
        )
        
        # Remove legend (redundant with x-axis)
        fig.update_layout(showlegend=False)
        
        st.plotly_chart(fig, use_container_width=True)
    
    # Show correlation table for all selected variables
    st.subheader(t.get("ses_correlations", "SES Correlations with Performance Measures"))
    
    # Calculate correlations
    correlations = {}
    p_values = {}
    
    for col in selected_columns + ["total_score"]:
        if pd.api.types.is_numeric_dtype(df_analysis[col]):
            # Calculate Pearson correlation
            corr, p = stats.pearsonr(df_analysis["ses"].dropna(), df_analysis[col].dropna())
            correlations[col] = corr
            p_values[col] = p
    
    # Create correlation table
    corr_data = []
    for col, corr in correlations.items():
        # Get translated variable name
        var_name = t.get("total_score", "Total Score") if col == "total_score" else t["columns_of_interest"].get(col, col)
        
        # Determine correlation strength
        if abs(corr) >= 0.5:
            strength = t.get("strong", "Strong")
        elif abs(corr) >= 0.3:
            strength = t.get("moderate", "Moderate")
        else:
            strength = t.get("weak", "Weak")
        
        # Determine direction
        direction = t.get("positive", "Positive") if corr > 0 else t.get("negative", "Negative")
        
        # Determine significance
        is_significant = p_values[col] < 0.05
        significance = t.get("significant", "Significant") if is_significant else t.get("not_significant", "Not significant")
        
        corr_data.append({
            t.get("variable", "Variable"): var_name,
            t.get("correlation", "Correlation"): f"{corr:.3f}",
            t.get("p_value", "p-value"): f"{p_values[col]:.4f}{'*' if is_significant else ''}",
            t.get("strength", "Strength"): strength,
            t.get("direction", "Direction"): direction
        })
    
    # Sort by absolute correlation
    corr_data = sorted(corr_data, key=lambda x: abs(float(x[t.get("correlation", "Correlation")])), reverse=True)
    
    # Display as dataframe
    st.dataframe(pd.DataFrame(corr_data), use_container_width=True)
    
    # Add significance note
    st.markdown(t.get("significance_note", "*p < 0.05 indicates statistical significance"))
    
    # Show individual scatter plots for selected variables
    st.subheader(t.get("detailed_ses_correlations", "Detailed SES Correlations"))
    
    # Create scatter plots for each selected variable
    for i in range(0, len(selected_columns), 2):
        col1, col2 = st.columns(2)
        
        with col1:
            if i < len(selected_columns):
                _create_scatter_plot(df_analysis, "ses", selected_columns[i], t)
        
        with col2:
            if i + 1 < len(selected_columns):
                _create_scatter_plot(df_analysis, "ses", selected_columns[i + 1], t)


def _show_categorical_ses_analysis(df_analysis, selected_columns, t):
    """
    Display analysis for categorical SES data.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    st.subheader(t.get("categorical_ses_analysis", "Categorical SES Analysis"))
    
    # Show group comparisons for each variable
    for i in range(0, len(selected_columns), 2):
        col1, col2 = st.columns(2)
        
        with col1:
            if i < len(selected_columns):
                fig = px.box(
                    df_analysis,
                    x="ses",
                    y=selected_columns[i],
                    color="ses",
                    title=f"{t.get('ses', 'SES')} vs {t['columns_of_interest'].get(selected_columns[i], selected_columns[i])}",
                    labels={
                        "ses": t.get("ses", "Socioeconomic Status"),
                        selected_columns[i]: t["columns_of_interest"].get(selected_columns[i], selected_columns[i])
                    }
                )
                fig.update_layout(showlegend=False)
                st.plotly_chart(fig, use_container_width=True)
        
        with col2:
            if i + 1 < len(selected_columns):
                fig = px.box(
                    df_analysis,
                    x="ses",
                    y=selected_columns[i + 1],
                    color="ses",
                    title=f"{t.get('ses', 'SES')} vs {t['columns_of_interest'].get(selected_columns[i + 1], selected_columns[i + 1])}",
                    labels={
                        "ses": t.get("ses", "Socioeconomic Status"),
                        selected_columns[i + 1]: t["columns_of_interest"].get(selected_columns[i + 1], selected_columns[i + 1])
                    }
                )
                fig.update_layout(showlegend=False)
                st.plotly_chart(fig, use_container_width=True)


def _show_home_support_impact(df_analysis, selected_columns, t):
    """
    Display analysis of home support impact on performance.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    st.header(t.get("home_support_impact", "Impact of Home Support on Performance"))
    
    # Check if home_support is numeric
    if not pd.api.types.is_numeric_dtype(df_analysis["home_support"]):
        st.warning(t.get("support_not_numeric", "Home support variable is not numeric. Some analyses will be limited."))
    
    # Show boxplot of total score by home support group
    if "home_support_group" in df_analysis.columns and df_analysis["home_support_group"].nunique() > 1:
        fig = px.box(
            df_analysis,
            x="home_support_group",
            y="total_score",
            color="home_support_group",
            labels={
                "home_support_group": t.get("home_support", "Home Support"),
                "total_score": t.get("total_score", "Total Score")
            },
            title=t.get("score_by_support", "Performance by Home Support Level"),
            category_orders={"home_support_group": ["Low", "Medium", "High"]}
        )
        
        # Remove legend (redundant with x-axis)
        fig.update_layout(showlegend=False)
        
        st.plotly_chart(fig, use_container_width=True)
    
    # If home_support is numeric, show scatter plot with regression line
    if pd.api.types.is_numeric_dtype(df_analysis["home_support"]):
        fig = px.scatter(
            df_analysis,
            x="home_support",
            y="total_score",
            trendline="ols",
            labels={
                "home_support": t.get("home_support", "Home Support"),
                "total_score": t.get("total_score", "Total Score")
            },
            title=t.get("support_total_score", "Home Support vs Total Score"),
            color_discrete_sequence=['#ff7f0e']
        )
        
        # Get trendline info for annotation
        model_results = px.get_trendline_results(fig)
        
        if model_results is not None and len(model_results.px_fit_results) > 0:
            results = model_results.iloc[0]["px_fit_results"]
            alpha = results.params[0]  # Intercept
            beta = results.params[1]   # Slope
            r_squared = results.rsquared
            p_value = results.pvalues[1]  # p-value for slope
            
            # Add equation annotation
            fig.add_annotation(
                x=0.05,
                y=0.95,
                xref="paper",
                yref="paper",
                text=f"y = {alpha:.2f} + {beta:.2f}x<br>R² = {r_squared:.2f} (p = {p_value:.4f})",
                showarrow=False,
                font=dict(size=12),
                bgcolor="rgba(255, 255, 255, 0.8)",
                bordercolor="black",
                borderwidth=1
            )
        
        st.plotly_chart(fig, use_container_width=True)
    
    # Calculate comparisons between home support groups
    if "home_support_group" in df_analysis.columns and df_analysis["home_support_group"].nunique() > 1:
        support_group_stats = df_analysis.groupby("home_support_group")["total_score"].agg(['mean', 'std', 'count']).round(2)
        
        # Show group comparison table
        st.markdown(f"**{t.get('home_support_comparison', 'Home Support Group Comparison')}:**")
        
        if not support_group_stats.empty:
            # Format for display
            comparison_df = pd.DataFrame({
                t.get("support_level", "Support Level"): support_group_stats.index,
                t.get("mean_score", "Mean Score"): support_group_stats['mean'],
                t.get("std_dev", "Std Dev"): support_group_stats['std'],
                t.get("n_students", "N Students"): support_group_stats['count']
            })
            
            st.dataframe(comparison_df, use_container_width=True)
            
            # Calculate and show difference between highest and lowest groups
            if support_group_stats.index.to_list() == ["Low", "Medium", "High"]:
                high_mean = support_group_stats.loc["High", "mean"]
                low_mean = support_group_stats.loc["Low", "mean"]
                difference = high_mean - low_mean
                percentage = (difference / low_mean) * 100 if low_mean > 0 else 0
                
                st.markdown(f"**{t.get('high_low_difference', 'Difference between High and Low Support')}:** {difference:.2f} ({percentage:.1f}%)")
    
    # Show home support impact on specific skills
    st.subheader(t.get("home_support_skill_impact", "Home Support Impact on Specific Skills"))
    
    # Create visualizations for each skill
    for i in range(0, len(selected_columns), 2):
        col1, col2 = st.columns(2)
        
        with col1:
            if i < len(selected_columns) and "home_support_group" in df_analysis.columns:
                fig = px.box(
                    df_analysis,
                    x="home_support_group",
                    y=selected_columns[i],
                    color="home_support_group",
                    labels={
                        "home_support_group": t.get("home_support", "Home Support"),
                        selected_columns[i]: t["columns_of_interest"].get(selected_columns[i], selected_columns[i])
                    },
                    title=f"{t.get('home_support', 'Home Support')} vs {t['columns_of_interest'].get(selected_columns[i], selected_columns[i])}",
                    category_orders={"home_support_group": ["Low", "Medium", "High"]}
                )
                fig.update_layout(showlegend=False)
                st.plotly_chart(fig, use_container_width=True)
        
        with col2:
            if i + 1 < len(selected_columns) and "home_support_group" in df_analysis.columns:
                fig = px.box(
                    df_analysis,
                    x="home_support_group",
                    y=selected_columns[i + 1],
                    color="home_support_group",
                    labels={
                        "home_support_group": t.get("home_support", "Home Support"),
                        selected_columns[i + 1]: t["columns_of_interest"].get(selected_columns[i + 1], selected_columns[i + 1])
                    },
                    title=f"{t.get('home_support', 'Home Support')} vs {t['columns_of_interest'].get(selected_columns[i + 1], selected_columns[i + 1])}",
                    category_orders={"home_support_group": ["Low", "Medium", "High"]}
                )
                fig.update_layout(showlegend=False)
                st.plotly_chart(fig, use_container_width=True)


def _show_combined_effects(df_analysis, selected_columns, t):
    """
    Display analysis of combined SES and home support effects.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    st.header(t.get("combined_effects", "Combined Effects of SES and Home Support"))
    
    # Create interaction heatmap for total score
    if all(col in df_analysis.columns for col in ["ses_group", "home_support_group", "total_score"]):
        if df_analysis["ses_group"].nunique() > 1 and df_analysis["home_support_group"].nunique() > 1:
            try:
                # Calculate mean total score for each SES x Home Support combination
                heatmap_data = df_analysis.groupby(["ses_group", "home_support_group"])["total_score"].mean().reset_index()
                
                # Pivot the data for heatmap format
                heatmap_pivot = heatmap_data.pivot(
                    index="ses_group", 
                    columns="home_support_group", 
                    values="total_score"
                )
                
                # Ensure correct order of categories if they're categorical
                if isinstance(heatmap_pivot.index, pd.CategoricalIndex):
                    heatmap_pivot = heatmap_pivot.reindex(["Low", "Medium", "High"])
                
                if isinstance(heatmap_pivot.columns, pd.CategoricalIndex):
                    heatmap_pivot = heatmap_pivot.reindex(columns=["Low", "Medium", "High"])
                
                # Create heatmap
                fig = ff.create_annotated_heatmap(
                    z=heatmap_pivot.values,
                    x=list(heatmap_pivot.columns),
                    y=list(heatmap_pivot.index),
                    annotation_text=np.around(heatmap_pivot.values, 1),
                    colorscale='Viridis',
                    showscale=True
                )
                
                # Update layout
                fig.update_layout(
                    title=t.get("interaction_effect", "Interaction Effect on Total Score"),
                    xaxis_title=t.get("home_support", "Home Support"),
                    yaxis_title=t.get("ses_group", "SES Group")
                )
                
                st.plotly_chart(fig, use_container_width=True)
                
                # Add interpretation
                if not heatmap_pivot.empty:
                    st.markdown(f"**{t.get('interaction_interpretation', 'Interpretation')}:**")
                    
                    # Find the highest and lowest performing combinations
                    max_val = heatmap_pivot.max().max()
                    min_val = heatmap_pivot.min().min()
                    
                    # Find the groups with max and min values
                    max_groups = []
                    min_groups = []
                    
                    for ses in heatmap_pivot.index:
                        for support in heatmap_pivot.columns:
                            if heatmap_pivot.at[ses, support] == max_val:
                                max_groups.append((ses, support))
                            if heatmap_pivot.at[ses, support] == min_val:
                                min_groups.append((ses, support))
                    
                    # Show insights
                    for ses, support in max_groups:
                        st.markdown(f"- {t.get('highest_combination', 'The highest performing combination is')} **{ses} SES** {t.get('with', 'with')} **{support} {t.get('home_support', 'Home Support')}** ({max_val:.1f})")
                    
                    for ses, support in min_groups:
                        st.markdown(f"- {t.get('lowest_combination', 'The lowest performing combination is')} **{ses} SES** {t.get('with', 'with')} **{support} {t.get('home_support', 'Home Support')}** ({min_val:.1f})")
                    
                    # Calculate the performance gap
                    performance_gap = max_val - min_val
                    st.markdown(f"- {t.get('performance_gap', 'Performance gap')}: **{performance_gap:.1f}** {t.get('points', 'points')}")
                    
                    # Add highlight if high home support compensates for low SES
                    if 'Low' in heatmap_pivot.index and 'High' in heatmap_pivot.columns:
                        low_ses_high_support = heatmap_pivot.at['Low', 'High']
                        
                        if 'High' in heatmap_pivot.index and 'Low' in heatmap_pivot.columns:
                            high_ses_low_support = heatmap_pivot.at['High', 'Low']
                            
                            if low_ses_high_support > high_ses_low_support:
                                st.markdown(f"- **{t.get('compensation_effect', 'Compensation Effect')}:** {t.get('compensation_explanation', 'High home support appears to compensate for low SES, as this combination outperforms high SES with low home support.')}")
            except Exception as e:
                st.warning(f"{t.get('heatmap_error', 'Could not create interaction heatmap')}: {str(e)}")
        else:
            st.info(t.get("insufficient_group_variation", "Insufficient group variation to show interaction effects."))
    
    # Add violin plot comparing distributions
    if all(col in df_analysis.columns for col in ["ses_group", "home_support_group", "total_score"]):
        st.subheader(t.get("distribution_comparison", "Distribution Comparison"))
        
        try:
            fig = px.violin(
                df_analysis,
                x="ses_group",
                y="total_score",
                color="home_support_group",
                box=True,
                points="all",
                labels={
                    "ses_group": t.get("ses_group", "SES Group"),
                    "total_score": t.get("total_score", "Total Score"),
                    "home_support_group": t.get("home_support", "Home Support")
                },
                title=t.get("combined_distribution", "Score Distribution by SES and Home Support"),
                category_orders={
                    "ses_group": ["Low", "Medium", "High"],
                    "home_support_group": ["Low", "Medium", "High"]
                }
            )
            
            st.plotly_chart(fig, use_container_width=True)
        except Exception as e:
            st.warning(f"{t.get('violin_error', 'Could not create violin plot')}: {str(e)}")
    
    # Add 3D scatter plot if data is numeric
    if all(pd.api.types.is_numeric_dtype(df_analysis[col]) for col in ["ses", "home_support", "total_score"]):
        st.subheader(t.get("3d_visualization", "3D Visualization"))
        
        fig = go.Figure(data=[go.Scatter3d(
            x=df_analysis["ses"],
            y=df_analysis["home_support"],
            z=df_analysis["total_score"],
            mode='markers',
            marker=dict(
                size=5,
                color=df_analysis["total_score"],
                colorscale='Viridis',
                opacity=0.8
            )
        )])
        
        fig.update_layout(
            scene=dict(
                xaxis_title=t.get("ses", "SES"),
                yaxis_title=t.get("home_support", "Home Support"),
                zaxis_title=t.get("total_score", "Total Score")
            ),
            title=t.get("3d_relationship", "3D Relationship: SES, Home Support, and Performance")
        )
        
        st.plotly_chart(fig, use_container_width=True)


def _show_correlation_matrix(df_analysis, selected_columns, t):
    """
    Display correlation matrix between variables.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    st.header(t.get("correlation_matrix", "Correlation Matrix"))
    
    # Select variables for correlation matrix
    correlation_vars = ["ses", "home_support", "total_score"] + selected_columns
    
    # Filter to only include numeric columns that actually exist
    numeric_vars = [col for col in correlation_vars 
                   if col in df_analysis.columns and pd.api.types.is_numeric_dtype(df_analysis[col])]
    
    if len(numeric_vars) >= 2:
        # Calculate correlation matrix
        corr_matrix = df_analysis[numeric_vars].corr().round(2)
        
        # Create translation mapping for variable names
        var_translations = {}
        for var in numeric_vars:
            if var == "ses":
                var_translations[var] = t.get("ses", "SES")
            elif var == "home_support":
                var_translations[var] = t.get("home_support", "Home Support")
            elif var == "total_score":
                var_translations[var] = t.get("total_score", "Total Score")
            else:
                var_translations[var] = t["columns_of_interest"].get(var, var)
        
        # Create heatmap with absolute values colored
        fig = ff.create_annotated_heatmap(
            z=corr_matrix.values,
            x=[var_translations.get(col, col) for col in corr_matrix.columns],
            y=[var_translations.get(col, col) for col in corr_matrix.index],
            annotation_text=corr_matrix.values.round(2),
            colorscale='RdBu_r',
            showscale=True,
            zmin=-1, zmax=1
        )
        
        # Update layout for better readability
        fig.update_layout(
            height=600,
            xaxis_title=t.get("variables", "Variables"),
            yaxis_title=t.get("variables", "Variables")
        )
        
        # Rotate x-axis labels for better readability
        fig.update_xaxes(tickangle=-45)
        
        st.plotly_chart(fig, use_container_width=True)
        
        # Add explanatory text for correlation interpretation
        st.markdown(t.get("correlation_interpretation", """
        **Interpreting the Correlation Matrix:**
        - Values close to 1 indicate a strong positive relationship
        - Values close to -1 indicate a strong negative relationship
        - Values close to 0 indicate a weak or no relationship
        """))
        
        # Show strong correlations
        st.subheader(t.get("strong_correlations", "Notable Correlations"))
        
        # Extract correlations and filter for strong ones
        strong_correlations = []
        
        for i in range(len(numeric_vars)):
            for j in range(i+1, len(numeric_vars)):
                var1 = numeric_vars[i]
                var2 = numeric_vars[j]
                corr = corr_matrix.iloc[i, j]
                
                if abs(corr) >= 0.3:  # Threshold for notable correlation
                    strong_correlations.append({
                        t.get("variable_1", "Variable 1"): var_translations.get(var1, var1),
                        t.get("variable_2", "Variable 2"): var_translations.get(var2, var2),
                        t.get("correlation", "Correlation"): f"{corr:.2f}",
                        t.get("strength", "Strength"): t.get("strong", "Strong") if abs(corr) >= 0.5 else t.get("moderate", "Moderate"),
                        t.get("direction", "Direction"): t.get("positive", "Positive") if corr > 0 else t.get("negative", "Negative")
                    })
        
        if strong_correlations:
            st.dataframe(pd.DataFrame(strong_correlations), use_container_width=True)
        else:
            st.info(t.get("no_strong_correlations", "No strong correlations found (|r| ≥ 0.3)"))
        
        # Highlight SES-Support correlation if it exists
        if "ses" in numeric_vars and "home_support" in numeric_vars:
            ses_idx = numeric_vars.index("ses")
            support_idx = numeric_vars.index("home_support")
            ses_support_corr = corr_matrix.iloc[ses_idx, support_idx]
            
            st.markdown(f"### {t.get('ses_support_correlation', 'SES and Home Support Correlation')}")
            st.markdown(f"**{ses_support_corr:.2f}**")
            
            if abs(ses_support_corr) > 0.5:
                st.markdown(t.get('strong_ses_support_relation', 'There is a strong relationship between socioeconomic status and home support levels.'))
            elif abs(ses_support_corr) > 0.3:
                st.markdown(t.get('moderate_ses_support_relation', 'There is a moderate relationship between socioeconomic status and home support levels.'))
            else:
                st.markdown(t.get('weak_ses_support_relation', 'There is a weak relationship between socioeconomic status and home support levels.'))
            
            if ses_support_corr > 0:
                st.markdown(t.get('positive_ses_support_relation', 'Higher SES tends to be associated with higher levels of home support.'))
            else:
                st.markdown(t.get('negative_ses_support_relation', 'Higher SES tends to be associated with lower levels of home support, which is an unusual pattern.'))
    else:
        st.info(t.get("insufficient_numeric_vars", "Not enough numeric variables for correlation analysis."))


def _show_recommendations(df_analysis, selected_columns, t):
    """
    Display targeted recommendations based on analysis.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    st.header(t.get("recommendations", "Recommendations"))
    
    # SES-focused recommendations
    st.subheader(t.get("ses_recommendations", "SES-Based Interventions"))
    
    # Check if SES has a strong impact
    ses_impact = _assess_ses_impact(df_analysis, selected_columns)
    
    if ses_impact['strong_impact']:
        ses_recs = [
            t.get("ses_rec1", "Provide additional resources and support for students from low SES backgrounds"),
            t.get("ses_rec2", "Implement targeted interventions for specific skill areas most affected by SES disparities"),
            t.get("ses_rec3", "Develop school programs that provide access to educational resources typically less available in low-SES homes"),
            t.get("ses_rec4", "Consider community partnerships to address broader SES-related barriers to education")
        ]
    else:
        ses_recs = [
            t.get("ses_rec_weak1", "Continue monitoring SES effects while focusing on other factors that may have stronger impacts"),
            t.get("ses_rec_weak2", "Identify school practices that may be helping mitigate typical SES-related disparities"),
            t.get("ses_rec_weak3", "Share successful approaches with other schools that may be struggling with stronger SES effects")
        ]
    
    for rec in ses_recs:
        st.markdown(f"- {rec}")
    
    # Home support recommendations
    st.subheader(t.get("home_support_recommendations", "Home Support Interventions"))
    
    # Check support effect status
    support_impact = _assess_home_support_impact(df_analysis, selected_columns)
    
    if support_impact['strong_impact']:
        support_recs = [
            t.get("support_rec1", "Develop parent education programs on supporting learning at home"),
            t.get("support_rec2", "Create and distribute home learning materials with parent guidance"),
            t.get("support_rec3", "Implement regular parent-teacher communication about learning support strategies"),
            t.get("support_rec4", "Consider after-school programs for students with limited home support")
        ]
    else:
        support_recs = [
            t.get("support_rec_weak1", "Investigate why home support appears to have limited impact in this context"),
            t.get("support_rec_weak2", "Review how home support is being measured and defined"),
            t.get("support_rec_weak3", "Identify if school practices may be compensating for home support variations")
        ]
    
    for rec in support_recs:
        st.markdown(f"- {rec}")
    
    # Combined/targeted recommendations
    st.subheader(t.get("targeted_recommendations", "Targeted Recommendations"))
    
    # Generate targeted recommendations based on analysis
    targeted_recs = _generate_targeted_recommendations(df_analysis, selected_columns, t)
    
    for rec in targeted_recs:
        st.markdown(f"- {rec}")


def _provide_export_options(df_analysis, selected_columns, t):
    """
    Provide options to export analysis results.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    st.header(t.get("export_results", "Export Results"))
    
    col1, col2 = st.columns(2)
    
    # Prepare summary data for export
    export_data = _prepare_export_data(df_analysis, selected_columns, t)
    
    # CSV Export
    with col1:
        csv = export_data.to_csv(index=False).encode('utf-8-sig')
        st.download_button(
            t.get("export_csv", "Download CSV"),
            csv,
            "ses_home_support_analysis.csv",
            "text/csv",
            key='download-ses-csv'
        )
    
    # Word Export
    with col2:
        if st.button(t.get("export_word", "Export to Word")):
            try:
                doc = _create_ses_word_report(df_analysis, selected_columns, t)
                
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    doc.save(tmp.name)
                    with open(tmp.name, 'rb') as f:
                        docx = f.read()
                    st.download_button(
                        t.get("download_word", "Download Word Report"),
                        docx,
                        "ses_home_support_report.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                os.unlink(tmp.name)
            except Exception as e:
                st.error(f"{t.get('error_report_generation', 'Error generating report')}: {str(e)}")


def _create_scatter_plot(df, x_var, y_var, t):
    """
    Create a scatter plot with regression line.
    
    Args:
        df (pd.DataFrame): Dataset
        x_var (str): X-axis variable
        y_var (str): Y-axis variable
        t (dict): Translation dictionary
    """
    # Get translated variable names
    x_var_name = t.get(x_var, x_var) if x_var in ["ses", "home_support"] else x_var
    y_var_name = t["columns_of_interest"].get(y_var, y_var)
    
    # Create scatter plot with regression line
    fig = px.scatter(
        df,
        x=x_var,
        y=y_var,
        trendline="ols",
        title=f"{x_var_name} vs {y_var_name}",
        labels={
            x_var: x_var_name,
            y_var: y_var_name
        }
    )
    
    # Add regression statistics
    try:
        model_results = px.get_trendline_results(fig)
        
        if model_results is not None and len(model_results.px_fit_results) > 0:
            results = model_results.iloc[0]["px_fit_results"]
            r_squared = results.rsquared
            p_value = results.pvalues[1] if len(results.pvalues) > 1 else None
            
            # Add R-squared annotation
            annotation_text = f"R² = {r_squared:.3f}"
            if p_value is not None:
                annotation_text += f"\np = {p_value:.4f}"
            
            fig.add_annotation(
                x=0.05,
                y=0.95,
                xref="paper",
                yref="paper",
                text=annotation_text,
                showarrow=False,
                font=dict(size=12),
                bgcolor="rgba(255, 255, 255, 0.8)",
                bordercolor="black",
                borderwidth=1
            )
    except Exception as e:
        st.warning(f"Could not calculate regression statistics: {str(e)}")
    
    # Display the plot
    st.plotly_chart(fig, use_container_width=True)


def _assess_ses_impact(df_analysis, selected_columns):
    """
    Assess the impact of SES on performance.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        
    Returns:
        dict: Assessment of SES impact
    """
    impact = {'strong_impact': False, 'correlation': 0, 'gap': 0}
    
    # Check correlation if SES is numeric
    if pd.api.types.is_numeric_dtype(df_analysis["ses"]) and pd.api.types.is_numeric_dtype(df_analysis["total_score"]):
        correlation = df_analysis[["ses", "total_score"]].corr().iloc[0, 1]
        impact['correlation'] = correlation
        impact['strong_impact'] = abs(correlation) > 0.3
    
    # Check performance gap between groups
    if "ses_group" in df_analysis.columns and df_analysis["ses_group"].nunique() > 1:
        mean_scores = df_analysis.groupby("ses_group")["total_score"].mean()
        if len(mean_scores) >= 2:
            gap = mean_scores.max() - mean_scores.min()
            impact['gap'] = gap
            impact['strong_impact'] = impact['strong_impact'] or (gap > 10)  # Arbitrary threshold
    
    return impact


def _assess_home_support_impact(df_analysis, selected_columns):
    """
    Assess the impact of home support on performance.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        
    Returns:
        dict: Assessment of home support impact
    """
    impact = {'strong_impact': False, 'correlation': 0, 'gap': 0}
    
    # Check correlation if home_support is numeric
    if pd.api.types.is_numeric_dtype(df_analysis["home_support"]) and pd.api.types.is_numeric_dtype(df_analysis["total_score"]):
        correlation = df_analysis[["home_support", "total_score"]].corr().iloc[0, 1]
        impact['correlation'] = correlation
        impact['strong_impact'] = abs(correlation) > 0.3
    
    # Check performance gap between groups
    if "home_support_group" in df_analysis.columns and df_analysis["home_support_group"].nunique() > 1:
        mean_scores = df_analysis.groupby("home_support_group")["total_score"].mean()
        if len(mean_scores) >= 2:
            gap = mean_scores.max() - mean_scores.min()
            impact['gap'] = gap
            impact['strong_impact'] = impact['strong_impact'] or (gap > 10)  # Arbitrary threshold
    
    return impact


def _generate_targeted_recommendations(df_analysis, selected_columns, t):
    """
    Generate targeted recommendations based on analysis.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
        
    Returns:
        list: List of targeted recommendations
    """
    targeted_recs = [
        t.get("targeted_rec1", "Implement a multi-tiered support system based on both SES and home support factors"),
        t.get("targeted_rec2", "Prioritize resources for students with both low SES and low home support"),
        t.get("targeted_rec3", "Develop differentiated approaches based on each student's contextual profile")
    ]
    
    # Check if we have interaction data for more specific recommendations
    if all(col in df_analysis.columns for col in ["ses_group", "home_support_group", "total_score"]):
        try:
            interaction_data = df_analysis.groupby(["ses_group", "home_support_group"])["total_score"].mean()
            
            # Look for compensation effect (low SES, high support outperforms high SES, low support)
            if ('Low', 'High') in interaction_data.index and ('High', 'Low') in interaction_data.index:
                comp_effect = interaction_data[('Low', 'High')] > interaction_data[('High', 'Low')]
                
                if comp_effect:
                    targeted_recs.append(t.get("compensation_rec", "Focus on improving home support for low-SES families as it shows strong compensatory effects"))
            
            # Find the lowest performing group
            if not interaction_data.empty:
                low_ses_group, low_support_group = interaction_data.idxmin()
                targeted_recs.append(t.get("lowest_group_rec", f"Implement comprehensive intervention for students with {low_ses_group} SES and {low_support_group} home support, as they show the lowest performance"))
        except:
            # Use default recommendations if interaction analysis fails
            pass
    
    return targeted_recs


def _prepare_export_data(df_analysis, selected_columns, t):
    """
    Prepare data for export.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
        
    Returns:
        pd.DataFrame: Data ready for export
    """
    # Create multiple DataFrames to join later
    export_data = {}
    
    # Add basic statistics
    if pd.api.types.is_numeric_dtype(df_analysis["ses"]):
        ses_stats = df_analysis["ses"].describe()
        export_data["SES Statistics"] = pd.DataFrame({"Value": ses_stats})
    
    if pd.api.types.is_numeric_dtype(df_analysis["home_support"]):
        support_stats = df_analysis["home_support"].describe()
        export_data["Home Support Statistics"] = pd.DataFrame({"Value": support_stats})
    
    # Add group means
    if "ses_group" in df_analysis.columns:
        ses_means = df_analysis.groupby("ses_group")[selected_columns + ["total_score"]].mean()
        export_data["SES Group Means"] = ses_means
    
    if "home_support_group" in df_analysis.columns:
        support_means = df_analysis.groupby("home_support_group")[selected_columns + ["total_score"]].mean()
        export_data["Home Support Group Means"] = support_means
    
    # Add correlation matrix if possible
    numeric_vars = ["ses", "home_support", "total_score"] + selected_columns
    numeric_vars = [col for col in numeric_vars if col in df_analysis.columns and pd.api.types.is_numeric_dtype(df_analysis[col])]
    
    if len(numeric_vars) >= 2:
        corr_matrix = df_analysis[numeric_vars].corr()
        export_data["Correlation Matrix"] = corr_matrix
    
    # Combine all data into a single DataFrame
    # First, convert each section into a DataFrame with consistent formatting
    formatted_sections = []
    
    for section_name, section_data in export_data.items():
        # Add section header
        formatted_sections.append(pd.DataFrame({
            "Section": [section_name],
            "Data": ["---" * 20]
        }))
        
        # Add section data
        if isinstance(section_data, pd.DataFrame):
            # Convert DataFrame to a string representation for export
            section_str = section_data.reset_index().to_string()
            formatted_sections.append(pd.DataFrame({
                "Section": [""],
                "Data": [section_str]
            }))
        else:
            formatted_sections.append(pd.DataFrame({
                "Section": [""],
                "Data": [str(section_data)]
            }))
        
        # Add spacing between sections
        formatted_sections.append(pd.DataFrame({
            "Section": [""],
            "Data": [""]
        }))
    
    # Concatenate all sections
    if formatted_sections:
        final_export = pd.concat(formatted_sections, ignore_index=True)
    else:
        final_export = pd.DataFrame({"Message": ["No data available for export"]})
    
    return final_export


def _create_ses_word_report(df_analysis, selected_columns, t):
    """
    Create a Word document with SES and home support analysis.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with enhanced variables
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
        
    Returns:
        Document: Word document with report
    """
    # Import ses_report_generator to handle the report creation
    doc = Document()
    
    # Title
    doc.add_heading(t.get("title_ses_home_support", "SES & Home Support Analysis"), level=1)
    
    # Introduction
    doc.add_paragraph(t.get("ses_report_intro", 
        "This report analyzes how socioeconomic status (SES) and home support affect student performance, "
        "identifying patterns and areas for targeted intervention."
    ))
    
    # Basic statistics
    doc.add_heading(t.get("basic_statistics", "Basic Statistics"), level=2)
    
    doc.add_paragraph(f"{t.get('total_students', 'Total Students')}: {len(df_analysis)}")
    
    if pd.api.types.is_numeric_dtype(df_analysis["ses"]):
        doc.add_paragraph(f"{t.get('average_ses', 'Average SES')}: {df_analysis['ses'].mean():.2f}")
        doc.add_paragraph(f"{t.get('ses_std', 'SES Standard Deviation')}: {df_analysis['ses'].std():.2f}")
    
    if pd.api.types.is_numeric_dtype(df_analysis["home_support"]):
        doc.add_paragraph(f"{t.get('average_home_support', 'Average Home Support')}: {df_analysis['home_support'].mean():.2f}")
        doc.add_paragraph(f"{t.get('support_std', 'Home Support Standard Deviation')}: {df_analysis['home_support'].std():.2f}")
    
    # SES Analysis
    doc.add_heading(t.get("ses_analysis", "SES Analysis"), level=2)
    
    if "ses_group" in df_analysis.columns and df_analysis["ses_group"].nunique() > 1:
        table = doc.add_table(rows=1, cols=len(selected_columns) + 2)
        table.style = 'Table Grid'
        
        # Headers
        header_cells = table.rows[0].cells
        header_cells[0].text = t.get("ses_group", "SES Group")
        header_cells[1].text = t.get("total_score", "Total Score")
        for i, col in enumerate(selected_columns):
            header_cells[i + 2].text = t["columns_of_interest"].get(col, col)
        
        # Data
        group_means = df_analysis.groupby("ses_group")[selected_columns + ["total_score"]].mean()
        for group, means in group_means.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(group)
            row_cells[1].text = f"{means['total_score']:.2f}"
            for i, col in enumerate(selected_columns):
                row_cells[i + 2].text = f"{means[col]:.2f}"
    
    # Home Support Analysis
    doc.add_heading(t.get("home_support_analysis", "Home Support Analysis"), level=2)
    
    if "home_support_group" in df_analysis.columns and df_analysis["home_support_group"].nunique() > 1:
        table = doc.add_table(rows=1, cols=len(selected_columns) + 2)
        table.style = 'Table Grid'
        
        # Headers
        header_cells = table.rows[0].cells
        header_cells[0].text = t.get("home_support_group", "Home Support Group")
        header_cells[1].text = t.get("total_score", "Total Score")
        for i, col in enumerate(selected_columns):
            header_cells[i + 2].text = t["columns_of_interest"].get(col, col)
        
        # Data
        group_means = df_analysis.groupby("home_support_group")[selected_columns + ["total_score"]].mean()
        for group, means in group_means.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(group)
            row_cells[1].text = f"{means['total_score']:.2f}"
            for i, col in enumerate(selected_columns):
                row_cells[i + 2].text = f"{means[col]:.2f}"
    
    # Correlations
    doc.add_heading(t.get("correlation_analysis", "Correlation Analysis"), level=2)
    
    numeric_vars = ["ses", "home_support", "total_score"] + selected_columns
    numeric_vars = [col for col in numeric_vars if col in df_analysis.columns and pd.api.types.is_numeric_dtype(df_analysis[col])]
    
    if len(numeric_vars) >= 2:
        corr_matrix = df_analysis[numeric_vars].corr()
        
        table = doc.add_table(rows=len(numeric_vars) + 1, cols=len(numeric_vars) + 1)
        table.style = 'Table Grid'
        
        # Headers
        for i, var in enumerate(numeric_vars):
            table.cell(0, i + 1).text = var
            table.cell(i + 1, 0).text = var
        
        # Data
        for i in range(len(numeric_vars)):
            for j in range(len(numeric_vars)):
                table.cell(i + 1, j + 1).text = f"{corr_matrix.iloc[i, j]:.2f}"
    
    # Recommendations
    doc.add_heading(t.get("recommendations", "Recommendations"), level=2)
    
    ses_impact = _assess_ses_impact(df_analysis, selected_columns)
    support_impact = _assess_home_support_impact(df_analysis, selected_columns)
    
    # SES recommendations
    doc.add_heading(t.get("ses_recommendations", "SES-Based Interventions"), level=3)
    if ses_impact['strong_impact']:
        doc.add_paragraph(t.get("ses_rec1", "Provide additional resources and support for students from low SES backgrounds"), style='List Bullet')
        doc.add_paragraph(t.get("ses_rec2", "Implement targeted interventions for specific skill areas most affected by SES disparities"), style='List Bullet')
        doc.add_paragraph(t.get("ses_rec3", "Develop school programs that provide access to educational resources typically less available in low-SES homes"), style='List Bullet')
        doc.add_paragraph(t.get("ses_rec4", "Consider community partnerships to address broader SES-related barriers to education"), style='List Bullet')
    else:
        doc.add_paragraph(t.get("ses_rec_weak1", "Continue monitoring SES effects while focusing on other factors that may have stronger impacts"), style='List Bullet')
        doc.add_paragraph(t.get("ses_rec_weak2", "Identify school practices that may be helping mitigate typical SES-related disparities"), style='List Bullet')
        doc.add_paragraph(t.get("ses_rec_weak3", "Share successful approaches with other schools that may be struggling with stronger SES effects"), style='List Bullet')
    
    # Home support recommendations
    doc.add_heading(t.get("home_support_recommendations", "Home Support Interventions"), level=3)
    if support_impact['strong_impact']:
        doc.add_paragraph(t.get("support_rec1", "Develop parent education programs on supporting learning at home"), style='List Bullet')
        doc.add_paragraph(t.get("support_rec2", "Create and distribute home learning materials with parent guidance"), style='List Bullet')
        doc.add_paragraph(t.get("support_rec3", "Implement regular parent-teacher communication about learning support strategies"), style='List Bullet')
        doc.add_paragraph(t.get("support_rec4", "Consider after-school programs for students with limited home support"), style='List Bullet')
    else:
        doc.add_paragraph(t.get("support_rec_weak1", "Investigate why home support appears to have limited impact in this context"), style='List Bullet')
        doc.add_paragraph(t.get("support_rec_weak2", "Review how home support is being measured and defined"), style='List Bullet')
        doc.add_paragraph(t.get("support_rec_weak3", "Identify if school practices may be compensating for home support variations"), style='List Bullet')
    
    # Targeted recommendations
    doc.add_heading(t.get("targeted_recommendations", "Targeted Recommendations"), level=3)
    targeted_recs = _generate_targeted_recommendations(df_analysis, selected_columns, t)
    for rec in targeted_recs:
        doc.add_paragraph(rec, style='List Bullet')
    
    # Add visualizations (if needed)
    try:
        import io
        import matplotlib.pyplot as plt
        
        # Create a correlation heatmap
        if len(numeric_vars) >= 2:
            fig, ax = plt.subplots(figsize=(8, 6))
            corr_matrix = df_analysis[numeric_vars].corr()
            
            im = ax.imshow(corr_matrix, cmap='RdBu_r', vmin=-1, vmax=1)
            
            # Add labels
            ax.set_xticks(range(len(numeric_vars)))
            ax.set_yticks(range(len(numeric_vars)))
            ax.set_xticklabels(numeric_vars, rotation=45, ha='right')
            ax.set_yticklabels(numeric_vars)
            
            # Add colorbar
            cbar = ax.figure.colorbar(im, ax=ax)
            cbar.ax.set_ylabel("Correlation", rotation=-90, va="bottom")
            
            # Add correlation values
            for i in range(len(numeric_vars)):
                for j in range(len(numeric_vars)):
                    text = ax.text(j, i, f"{corr_matrix.iloc[i, j]:.2f}",
                                   ha="center", va="center", color="black" if abs(corr_matrix.iloc[i, j]) < 0.5 else "white")
            
            plt.tight_layout()
            
            # Save to bytes buffer
            img_buffer = io.BytesIO()
            plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
            img_buffer.seek(0)
            
            # Add to document
            doc.add_heading(t.get("correlation_heatmap", "Correlation Heatmap"), level=2)
            doc.add_picture(img_buffer, width=Inches(6))
            plt.close()
            
        # Create SES group comparison chart
        if "ses_group" in df_analysis.columns and df_analysis["ses_group"].nunique() > 1:
            fig, ax = plt.subplots(figsize=(8, 6))
            
            group_means = df_analysis.groupby("ses_group")["total_score"].mean()
            group_std = df_analysis.groupby("ses_group")["total_score"].std()
            
            bars = ax.bar(group_means.index, group_means.values, yerr=group_std.values, capsize=5)
            ax.set_ylabel(t.get("total_score", "Total Score"))
            ax.set_xlabel(t.get("ses_group", "SES Group"))
            ax.set_title(t.get("performance_by_ses", "Performance by SES Group"))
            
            plt.tight_layout()
            
            # Save to bytes buffer
            img_buffer = io.BytesIO()
            plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
            img_buffer.seek(0)
            
            # Add to document
            doc.add_heading(t.get("ses_performance_chart", "SES Performance Comparison"), level=2)
            doc.add_picture(img_buffer, width=Inches(6))
            plt.close()
            
        # Create home support group comparison chart
        if "home_support_group" in df_analysis.columns and df_analysis["home_support_group"].nunique() > 1:
            fig, ax = plt.subplots(figsize=(8, 6))
            
            group_means = df_analysis.groupby("home_support_group")["total_score"].mean()
            group_std = df_analysis.groupby("home_support_group")["total_score"].std()
            
            bars = ax.bar(group_means.index, group_means.values, yerr=group_std.values, capsize=5)
            ax.set_ylabel(t.get("total_score", "Total Score"))
            ax.set_xlabel(t.get("home_support_group", "Home Support Group"))
            ax.set_title(t.get("performance_by_support", "Performance by Home Support Level"))
            
            plt.tight_layout()
            
            # Save to bytes buffer
            img_buffer = io.BytesIO()
            plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
            img_buffer.seek(0)
            
            # Add to document
            doc.add_heading(t.get("support_performance_chart", "Home Support Performance Comparison"), level=2)
            doc.add_picture(img_buffer, width=Inches(6))
            plt.close()
            
    except Exception as e:
        st.warning(f"Could not generate some visualizations for Word report: {str(e)}")
    
    # Conclusions
    doc.add_heading(t.get("conclusions", "Conclusions"), level=2)
    
    # Generate conclusions based on the analysis
    conclusions = []
    
    if pd.api.types.is_numeric_dtype(df_analysis["ses"]) and pd.api.types.is_numeric_dtype(df_analysis["total_score"]):
        ses_corr = df_analysis[["ses", "total_score"]].corr().iloc[0, 1]
        if abs(ses_corr) > 0.3:
            conclusions.append(t.get("ses_significant_impact", "SES has a significant impact on student performance."))
        else:
            conclusions.append(t.get("ses_minimal_impact", "SES shows minimal impact on student performance in this dataset."))
    
    if pd.api.types.is_numeric_dtype(df_analysis["home_support"]) and pd.api.types.is_numeric_dtype(df_analysis["total_score"]):
        support_corr = df_analysis[["home_support", "total_score"]].corr().iloc[0, 1]
        if abs(support_corr) > 0.3:
            conclusions.append(t.get("support_significant_impact", "Home support has a significant impact on student performance."))
        else:
            conclusions.append(t.get("support_minimal_impact", "Home support shows minimal impact on student performance in this dataset."))
    
    if pd.api.types.is_numeric_dtype(df_analysis["ses"]) and pd.api.types.is_numeric_dtype(df_analysis["home_support"]):
        ses_support_corr = df_analysis[["ses", "home_support"]].corr().iloc[0, 1]
        if abs(ses_support_corr) > 0.5:
            conclusions.append(t.get("strong_ses_support_relationship", "There is a strong relationship between SES and home support."))
        elif abs(ses_support_corr) > 0.3:
            conclusions.append(t.get("moderate_ses_support_relationship", "There is a moderate relationship between SES and home support."))
        else:
            conclusions.append(t.get("weak_ses_support_relationship", "There is a weak relationship between SES and home support."))
    
    for conclusion in conclusions:
        doc.add_paragraph(conclusion, style='List Bullet')
    
    # Add timestamp
    from datetime import datetime
    doc.add_paragraph()
    doc.add_paragraph(f"{t.get('report_generated', 'Report generated on')}: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    return doc