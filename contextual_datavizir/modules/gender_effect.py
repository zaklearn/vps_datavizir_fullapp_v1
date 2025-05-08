# modules/gender_effect.py
import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import scipy.stats as stats
from docx import Document
from docx.shared import Inches
import tempfile
import os
import logging
from config import translations, egra_columns, egma_columns
from error_handling import (
    error_handler, validate_dataframe, DataValidationError,
    safe_plot_creation, safe_data_operation
)

# Get module logger
logger = logging.getLogger('eduinsight.gender_effect')

@error_handler('eduinsight.gender_effect')
def show_gender_effect(df, language):
    """
    Analyze the effect of gender on student performance.
    
    Args:
        df (pd.DataFrame): Dataset containing gender variable
        language (str): Selected language for UI text
    """
    logger.info(f"Starting gender effect analysis with language: {language}")
    t = translations[language]  # Get translations
    
    st.subheader(t.get("title_gender_effect", "Gender Effect Analysis"))
    
    # Verify required column exists
    if "stgender" not in df.columns:
        logger.error("Gender column 'stgender' not found in dataset")
        st.warning(f"{t.get('missing_columns', 'Missing required columns')}: stgender")
        st.info(t.get('gender_column_info', 'This analysis requires gender information (stgender column).'))
        return
    
    try:
        # Prepare analysis data with gender mapping
        df_analysis = _prepare_gender_data(df, t)
        
        if df_analysis is None:
            logger.warning("Failed to prepare gender data")
            st.error(t.get('data_preparation_failed', 'Could not prepare data for analysis'))
            return
        
        # Select variables for analysis
        selected_columns = _select_variables(df_analysis, t)
        
        if selected_columns:
            logger.info(f"Selected columns for analysis: {selected_columns}")
            
            # Calculate gender statistics
            mean_scores, mean_scores_melted = _calculate_gender_statistics(df_analysis, selected_columns, t)
            
            if mean_scores is not None and mean_scores_melted is not None:
                # Show overview bar chart
                _show_gender_overview(mean_scores_melted, t)
                
                # Show detailed distributions
                _show_gender_distributions(df_analysis, selected_columns, t)
                
                # Show statistical tests
                _show_gender_statistics(df_analysis, selected_columns, t)
                
                # Show subject-specific gender patterns
                _show_subject_patterns(df_analysis, selected_columns, t)
                
                # Provide export options
                _provide_export_options(df_analysis, mean_scores, selected_columns, t)
            else:
                logger.warning("Failed to calculate gender statistics")
                st.warning(t.get("stats_calculation_failed", "Could not calculate statistics by gender."))
        else:
            logger.warning("No columns selected for analysis")
            st.warning(t.get("warning_select_variable", "Please select at least one variable to analyze."))
    
    except Exception as e:
        logger.exception("Error in gender effect analysis")
        st.error(f"{t.get('error_message', 'An error occurred')}: {str(e)}")

@error_handler('eduinsight.gender_effect')
def _prepare_gender_data(df, t):
    """
    Prepare dataset with gender mapping and additional calculated variables.
    
    Args:
        df (pd.DataFrame): Original dataset
        t (dict): Translation dictionary
        
    Returns:
        pd.DataFrame: Enhanced dataset with gender mapping
    """
    logger.debug("Preparing gender data")
    
    try:
        df_analysis = df.copy()
        
        # Map gender codes to labels
        # Standard coding is 1=male, 0=female
        if pd.api.types.is_numeric_dtype(df_analysis["stgender"]):
            logger.debug("Converting numeric gender codes to labels")
            df_analysis["gender"] = df_analysis["stgender"].map({
                1: t.get("boy", "Boy"),
                0: t.get("girl", "Girl")
            }).fillna(t.get("unknown", "Unknown"))
        elif pd.api.types.is_object_dtype(df_analysis["stgender"]):
            logger.debug("Converting text gender values to labels")
            # Try to map common text values
            gender_map = {
                "boy": t.get("boy", "Boy"),
                "girl": t.get("girl", "Girl"),
                "male": t.get("boy", "Boy"),
                "female": t.get("girl", "Girl"),
                "m": t.get("boy", "Boy"),
                "f": t.get("girl", "Girl"),
                "1": t.get("boy", "Boy"),
                "0": t.get("girl", "Girl")
            }
            # Use case-insensitive mapping
            df_analysis["gender"] = df_analysis["stgender"].astype(str).str.lower().map(gender_map).fillna(t.get("unknown", "Unknown"))
        else:
            logger.warning(f"Unexpected data type for stgender: {df_analysis['stgender'].dtype}")
            # If we can't determine the mapping, use the original values
            df_analysis["gender"] = df_analysis["stgender"].astype(str)
        
        # Log gender distribution
        gender_counts = df_analysis["gender"].value_counts()
        logger.info(f"Gender distribution: {gender_counts.to_dict()}")
        
        return df_analysis
    
    except Exception as e:
        logger.error(f"Error preparing gender data: {str(e)}")
        return None

@error_handler('eduinsight.gender_effect')
def _select_variables(df, t):
    """
    Allow user to select variables for analysis.
    
    Args:
        df (pd.DataFrame): Dataset
        t (dict): Translation dictionary
        
    Returns:
        list: Selected column names
    """
    logger.debug("Selecting variables for analysis")
    st.subheader(t.get("select_variables", "Select Variables"))
    
    # Get available assessment columns
    available_columns = []
    for col in egra_columns + egma_columns:
        if col in df.columns:
            available_columns.append(col)
    
    # If no assessment columns found, allow selection from all numeric columns
    if not available_columns:
        logger.info("No predefined assessment columns found, using all numeric columns")
        numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
        available_columns = [col for col in numeric_cols 
                            if col not in ['school', 'stgender', 'ses', 'home_support']]
    
    logger.debug(f"Available columns: {available_columns}")
    
    # Create selection widget
    selected_columns = st.multiselect(
        t.get("variables_all", "Variables:"),
        options=available_columns,
        default=available_columns[:min(3, len(available_columns))],
        format_func=lambda x: t["columns_of_interest"].get(x, x)
    )
    
    logger.info(f"User selected columns: {selected_columns}")
    return selected_columns

@error_handler('eduinsight.gender_effect')
def _calculate_gender_statistics(df_analysis, selected_columns, t):
    """
    Calculate statistics by gender.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with gender mapping
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
        
    Returns:
        tuple: (mean_scores, mean_scores_melted) DataFrames
    """
    logger.debug("Calculating gender statistics")
    
    try:
        # Calculate mean scores by gender
        mean_scores = df_analysis.groupby("gender")[selected_columns].mean().round(2)
        
        # Create melted version for visualization
        mean_scores_melted = mean_scores.reset_index().melt(
            id_vars=["gender"],
            value_vars=selected_columns,
            var_name=t.get("indicator", "Indicator"),
            value_name=t.get("mean_score", "Mean Score")
        )
        
        # Translate column names for display
        mean_scores_melted[t.get("indicator", "Indicator")] = mean_scores_melted[t.get("indicator", "Indicator")].map(
            lambda x: t["columns_of_interest"].get(x, x)
        )
        
        logger.info("Successfully calculated gender statistics")
        return mean_scores, mean_scores_melted
    
    except Exception as e:
        logger.error(f"Error calculating gender statistics: {str(e)}")
        return None, None

@error_handler('eduinsight.gender_effect')
def _show_gender_overview(mean_scores_melted, t):
    """
    Show overview bar chart of mean scores by gender.
    
    Args:
        mean_scores_melted (pd.DataFrame): Melted mean scores for visualization
        t (dict): Translation dictionary
    """
    logger.debug("Creating gender overview visualization")
    
    st.subheader(t.get("mean_score_gender", "Average scores by gender and by task"))
    
    try:
        # Create grouped bar chart
        fig = px.bar(
            mean_scores_melted,
            x=t.get("indicator", "Indicator"),
            y=t.get("mean_score", "Mean Score"),
            color="gender",
            barmode="group",
            title=t.get("mean_score_by_gender", "Mean Scores by Gender"),
            color_discrete_map={
                t.get("boy", "Boy"): "#1F77B4",
                t.get("girl", "Girl"): "#FF7F0E"
            }
        )
        
        # Improve layout
        fig.update_layout(
            xaxis_title=t.get("indicator", "Indicator"),
            yaxis_title=t.get("mean_score", "Mean Score"),
            legend_title=t.get("gender", "Gender"),
            xaxis_tickangle=-45
        )
        
        st.plotly_chart(fig, use_container_width=True)
        logger.info("Gender overview chart created successfully")
    except Exception as e:
        logger.error(f"Error creating gender overview chart: {str(e)}")
        st.warning(t.get("visualization_error", "Could not create visualization"))

@error_handler('eduinsight.gender_effect')
def _show_gender_distributions(df_analysis, selected_columns, t):
    """
    Show detailed distributions of scores by gender.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with gender mapping
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    logger.debug("Creating gender distribution visualizations")
    
    st.subheader(t.get("score_distribution_gender", "Score distribution by gender"))
    
    # Create boxplots for each selected variable by gender
    for i in range(0, len(selected_columns), 2):
        col1, col2 = st.columns(2)
        
        # Left column visualization
        with col1:
            if i < len(selected_columns):
                safe_plot_creation(
                    lambda: _create_boxplot(df_analysis, "gender", selected_columns[i], t),
                    f"boxplot_{selected_columns[i]}",
                    t.get("visualization_error", "Could not create visualization"),
                    "eduinsight.gender_effect"
                )
        
        # Right column visualization
        with col2:
            if i + 1 < len(selected_columns):
                safe_plot_creation(
                    lambda: _create_boxplot(df_analysis, "gender", selected_columns[i + 1], t),
                    f"boxplot_{selected_columns[i + 1]}",
                    t.get("visualization_error", "Could not create visualization"),
                    "eduinsight.gender_effect"
                )
    
    # Show distribution comparison
    st.subheader(t.get("distribution_comparison", "Score Distribution Comparison"))
    
    # Allow user to select a specific variable for detailed distribution comparison
    if selected_columns:
        distribution_var = st.selectbox(
            t.get("select_distribution_var", "Select variable for distribution comparison:"),
            options=selected_columns,
            format_func=lambda x: t["columns_of_interest"].get(x, x)
        )
        
        if distribution_var:
            safe_plot_creation(
                lambda: _create_distribution_comparison(df_analysis, distribution_var, t),
                f"distribution_comparison_{distribution_var}",
                t.get("visualization_error", "Could not create visualization"),
                "eduinsight.gender_effect"
            )

@error_handler('eduinsight.gender_effect')
def _create_boxplot(df, x_var, y_var, t):
    """
    Create a boxplot for a specific variable by gender.
    
    Args:
        df (pd.DataFrame): Dataset
        x_var (str): X-axis variable (grouping)
        y_var (str): Y-axis variable (measure)
        t (dict): Translation dictionary
        
    Returns:
        plotly.Figure: Boxplot figure or None if creation fails
    """
    logger.debug(f"Creating boxplot for {y_var} by {x_var}")
    
    # Get translated variable name
    y_var_name = t["columns_of_interest"].get(y_var, y_var)
    
    try:
        # Create boxplot
        fig = px.box(
            df,
            x=x_var,
            y=y_var,
            color=x_var,
            title=t.get("histogram_title", "Distribution of {}").format(y_var_name),
            labels={
                x_var: t.get("gender", "Gender"),
                y_var: y_var_name
            }
        )
        
        # Remove legend (redundant with x-axis)
        fig.update_layout(showlegend=False)
        
        # Display the plot
        st.plotly_chart(fig, use_container_width=True)
        logger.debug(f"Boxplot created successfully for {y_var}")
        return fig
    except Exception as e:
        logger.error(f"Error creating boxplot for {y_var}: {str(e)}")
        return None

@error_handler('eduinsight.gender_effect')
def _create_distribution_comparison(df_analysis, distribution_var, t):
    """
    Create a detailed distribution comparison for a selected variable.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with gender mapping
        distribution_var (str): Variable for distribution comparison
        t (dict): Translation dictionary
    """
    logger.debug(f"Creating distribution comparison for {distribution_var}")
    
    try:
        # Get translated variable name
        dist_var_name = t["columns_of_interest"].get(distribution_var, distribution_var)
        
        # Create histogram with density curves
        fig = px.histogram(
            df_analysis,
            x=distribution_var,
            color="gender",
            barmode="overlay",
            opacity=0.7,
            histnorm="probability density",
            marginal="box",
            title=t.get("distribution_title", "Distribution of {}").format(dist_var_name),
            labels={
                distribution_var: dist_var_name,
                "gender": t.get("gender", "Gender")
            },
            color_discrete_map={
                t.get("boy", "Boy"): "#1F77B4",
                t.get("girl", "Girl"): "#FF7F0E"
            }
        )
        
        st.plotly_chart(fig, use_container_width=True)
        
        # Calculate and display basic statistics
        stats_cols = st.columns(4)
        
        # Get groups to analyze
        gender_groups = df_analysis["gender"].unique()
        
        if len(gender_groups) >= 2:
            group1, group2 = gender_groups[:2]
            
            # Calculate statistics safely
            try:
                group1_data = df_analysis[df_analysis['gender'] == group1][distribution_var]
                group2_data = df_analysis[df_analysis['gender'] == group2][distribution_var]
                
                # Mean
                with stats_cols[0]:
                    group1_mean = group1_data.mean() if not group1_data.empty else 0
                    group2_mean = group2_data.mean() if not group2_data.empty else 0
                    st.metric(
                        t.get("mean", "Mean"),
                        f"{group1_mean:.2f} / {group2_mean:.2f}",
                        delta=None
                    )
                
                # Median
                with stats_cols[1]:
                    group1_median = group1_data.median() if not group1_data.empty else 0
                    group2_median = group2_data.median() if not group2_data.empty else 0
                    st.metric(
                        t.get("median", "Median"),
                        f"{group1_median:.2f} / {group2_median:.2f}",
                        delta=None
                    )
                
                # Standard Deviation
                with stats_cols[2]:
                    group1_std = group1_data.std() if not group1_data.empty else 0
                    group2_std = group2_data.std() if not group2_data.empty else 0
                    st.metric(
                        t.get("std_dev", "Std Dev"),
                        f"{group1_std:.2f} / {group2_std:.2f}",
                        delta=None
                    )
                
                # Range
                with stats_cols[3]:
                    range1 = (group1_data.max() - group1_data.min()) if not group1_data.empty else 0
                    range2 = (group2_data.max() - group2_data.min()) if not group2_data.empty else 0
                    st.metric(
                        t.get("range", "Range"),
                        f"{range1:.2f} / {range2:.2f}",
                        delta=None
                    )
            except Exception as e:
                logger.error(f"Error calculating statistics: {str(e)}")
                st.warning(t.get("stats_calculation_error", "Error calculating statistics"))
    
    except Exception as e:
        logger.error(f"Error creating distribution comparison: {str(e)}")
        st.warning(t.get("visualization_error", "Could not create visualization"))

@error_handler('eduinsight.gender_effect')
def _show_gender_statistics(df_analysis, selected_columns, t):
    """
    Show statistical tests for gender differences.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with gender mapping
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    logger.debug("Performing statistical tests for gender differences")
    
    st.subheader(t.get("stat_tests", "Statistical Tests (Mann-Whitney)"))
    
    # Get gender groups
    gender_groups = df_analysis["gender"].unique()
    
    if len(gender_groups) < 2:
        logger.warning("Insufficient gender groups for statistical comparison")
        st.warning(t.get("insufficient_groups", "Insufficient gender groups for statistical comparison."))
        return
    
    # Select the first two groups for analysis (typically boy/girl)
    group1, group2 = gender_groups[:2]
    logger.info(f"Comparing groups: {group1} vs {group2}")
    
    # Calculate statistical tests for each selected variable
    results_data = []
    
    for col in selected_columns:
        try:
            # Get data for each group
            group1_data = df_analysis[df_analysis["gender"] == group1][col].dropna()
            group2_data = df_analysis[df_analysis["gender"] == group2][col].dropna()
            
            # Only perform test if we have enough data
            if len(group1_data) > 0 and len(group2_data) > 0:
                # Perform Mann-Whitney U test
                stat, p = stats.mannwhitneyu(group1_data, group2_data, alternative='two-sided')
                
                # Calculate effect size (r = Z / sqrt(N))
                n1, n2 = len(group1_data), len(group2_data)
                Z = stat / np.sqrt((n1 * n2 * (n1 + n2 + 1)) / 12)
                r = Z / np.sqrt(n1 + n2)
                
                # Determine effect size interpretation
                if abs(r) < 0.1:
                    effect_size = t.get("negligible", "Negligible")
                elif abs(r) < 0.3:
                    effect_size = t.get("small", "Small")
                elif abs(r) < 0.5:
                    effect_size = t.get("medium", "Medium")
                else:
                    effect_size = t.get("large", "Large")
                
                # Calculate difference in means
                mean1 = group1_data.mean()
                mean2 = group2_data.mean()
                diff = mean1 - mean2
                
                # Add result to table
                results_data.append({
                    "Variable": t["columns_of_interest"].get(col, col),
                    "Mann-Whitney U": f"{stat:.0f}",
                    "p-value": f"{p:.4f}",
                    "Effect Size": f"{r:.2f} ({effect_size})",
                    "Difference": f"{diff:.2f}",
                    "Significant": p < 0.05,
                    "Higher Group": group1 if mean1 > mean2 else group2
                })
            else:
                logger.warning(f"Insufficient data for statistical test on {col}")
        except Exception as e:
            logger.error(f"Error performing statistical test for {col}: {str(e)}")
            st.warning(f"{t.get('test_failed', 'Test failed for')} {col}: {str(e)}")
    
    if results_data:
        # Create DataFrame for display
        results_df = pd.DataFrame(results_data)
        
        # Create display table without internal columns
        display_cols = ["Variable", "Mann-Whitney U", "p-value", "Effect Size", "Difference", "Higher Group"]
        display_df = results_df[display_cols].copy()
        
        # Add significance indicator
        def format_p_value(p_value_str):
            try:
                p_val = float(p_value_str)
                return f"{p_value_str}{'*' if p_val < 0.05 else ''}"
            except:
                return p_value_str
        
        display_df["p-value"] = display_df["p-value"].apply(format_p_value)
        
        # Display results table
        st.dataframe(display_df, use_container_width=True)
        
        # Add note about significance
        st.markdown(t.get("significance_note", "*p < 0.05 indicates statistical significance"))
        
        # Highlight significant findings
        significant_results = results_df[results_df["Significant"]]
        
        if not significant_results.empty:
            st.subheader(t.get("significant_findings", "Significant Gender Differences"))
            
            for _, row in significant_results.iterrows():
                # Format as "Girls outperform boys in X by Y points (p=0.001, medium effect)"
                higher_gender = row["Higher Group"]
                lower_gender = group2 if higher_gender == group1 else group1
                
                st.markdown(
                    f"- {higher_gender} {t.get('outperform', 'outperform')} {lower_gender} "
                    f"{t.get('in', 'in')} **{row['Variable']}** "
                    f"{t.get('by', 'by')} {abs(float(row['Difference'])):.2f} "
                    f"{t.get('points', 'points')} "
                    f"(p={row['p-value'].replace('*', '')}, {row['Effect Size'].split('(')[1].replace(')', '')} "
                    f"{t.get('effect', 'effect')})"
                )
        else:
            st.info(t.get("no_significant_diff", "No statistically significant gender differences were found."))
    else:
        logger.warning("No statistical test results to display")
        st.warning(t.get("no_test_results", "Could not generate statistical test results."))

@error_handler('eduinsight.gender_effect')
def _show_subject_patterns(df_analysis, selected_columns, t):
    """
    Show gender patterns by subject area (reading vs. math).
    
    Args:
        df_analysis (pd.DataFrame): Dataset with gender mapping
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    logger.debug("Analyzing gender patterns by subject area")
    
    # Check if we have both EGRA and EGMA variables
    egra_selected = [col for col in selected_columns if col in egra_columns]
    egma_selected = [col for col in selected_columns if col in egma_columns]
    
    if egra_selected and egma_selected:
        st.subheader(t.get("subject_patterns", "Gender Patterns by Subject Area"))
        
        # Get gender groups
        gender_groups = df_analysis["gender"].unique()
        
        if len(gender_groups) >= 2:
            group1, group2 = gender_groups[:2]
            
            try:
                # Calculate means for each subject area
                group1_egra_mean = df_analysis[df_analysis["gender"] == group1][egra_selected].mean().mean()
                group1_egma_mean = df_analysis[df_analysis["gender"] == group1][egma_selected].mean().mean()
                group2_egra_mean = df_analysis[df_analysis["gender"] == group2][egra_selected].mean().mean()
                group2_egma_mean = df_analysis[df_analysis["gender"] == group2][egma_selected].mean().mean()
                
                # Create data for visualization
                subject_data = pd.DataFrame({
                    "Gender": [group1, group1, group2, group2],
                    "Subject": [t.get("reading", "Reading"), t.get("math", "Math"), 
                                t.get("reading", "Reading"), t.get("math", "Math")],
                    "Mean Score": [group1_egra_mean, group1_egma_mean, group2_egra_mean, group2_egma_mean]
                })
                
                # Create bar chart
                fig = px.bar(
                    subject_data,
                    x="Subject",
                    y="Mean Score",
                    color="Gender",
                    barmode="group",
                    title=t.get("subject_comparison", "Gender Comparison by Subject Area"),
                    color_discrete_map={
                        group1: "#1F77B4",
                        group2: "#FF7F0E"
                    }
                )
                
                st.plotly_chart(fig, use_container_width=True)
                
                # Calculate differences
                reading_diff = group1_egra_mean - group2_egra_mean
                math_diff = group1_egma_mean - group2_egma_mean
                
                # Show insights
                st.markdown(f"**{t.get('subject_insights', 'Subject-Specific Insights')}:**")
                
                # Reading comparison
                if abs(reading_diff) < 0.1:
                    st.markdown(f"- {t.get('reading_similar', 'Reading performance is similar across genders.')}")
                else:
                    higher_gender = group1 if reading_diff > 0 else group2
                    st.markdown(f"- {t.get('reading_diff', 'In reading skills,')} {higher_gender} {t.get('higher_by', 'outperform others by')} {abs(reading_diff):.2f} {t.get('points', 'points')}.")
                
                # Math comparison
                if abs(math_diff) < 0.1:
                    st.markdown(f"- {t.get('math_similar', 'Math performance is similar across genders.')}")
                else:
                    higher_gender = group1 if math_diff > 0 else group2
                    st.markdown(f"- {t.get('math_diff', 'In math skills,')} {higher_gender} {t.get('higher_by', 'outperform others by')} {abs(math_diff):.2f} {t.get('points', 'points')}.")
                
                # Compare differences across subjects
                if abs(reading_diff) > abs(math_diff):
                    st.markdown(f"- {t.get('greater_reading_gap', 'The gender gap is larger in reading than in math.')}")
                elif abs(math_diff) > abs(reading_diff):
                    st.markdown(f"- {t.get('greater_math_gap', 'The gender gap is larger in math than in reading.')}")
                else:
                    st.markdown(f"- {t.get('equal_gaps', 'The gender gaps in reading and math are similar in magnitude.')}")
            
            except Exception as e:
                logger.error(f"Error analyzing subject patterns: {str(e)}")
                st.warning(t.get("subject_pattern_error", "Could not analyze subject patterns"))
    else:
        logger.info("Insufficient subject data for gender pattern analysis")

@error_handler('eduinsight.gender_effect')
def _provide_export_options(df_analysis, mean_scores, selected_columns, t):
    """
    Provide options to export analysis results.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with gender mapping
        mean_scores (pd.DataFrame): Mean scores by gender
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    logger.debug("Providing export options")
    
    st.subheader(t.get("export_results", "Export Results"))
    
    col1, col2 = st.columns(2)
    
    # CSV Export
    with col1:
        try:
            csv = mean_scores.reset_index().to_csv(index=False).encode('utf-8-sig')
            st.download_button(
                t.get("export_csv", "Download CSV"),
                csv,
                "gender_analysis.csv",
                "text/csv",
                key='download-gender-csv'
            )
            logger.info("CSV export button created successfully")
        except Exception as e:
            logger.error(f"Error creating CSV export: {str(e)}")
            st.error(t.get("export_csv_error", "Error creating CSV export"))
    
    # Word Export
    with col2:
        if st.button(t.get("export_word", "Export to Word")):
            try:
                logger.info("Starting Word report generation")
                doc = _create_gender_effect_word_report(df_analysis, mean_scores, selected_columns, t)
                
                if doc is not None:
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                        doc.save(tmp.name)
                        logger.debug(f"Word report saved to temporary file: {tmp.name}")
                        
                        with open(tmp.name, 'rb') as f:
                            docx = f.read()
                        st.download_button(
                            t.get("download_word", "Download Word Report"),
                            docx,
                            "gender_effect_report.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    os.unlink(tmp.name)
                    logger.info("Word report generated and download button created successfully")
                else:
                    logger.warning("Word report generation returned None")
                    st.warning(t.get("word_report_failed", "Could not generate Word report"))
            except Exception as e:
                logger.error(f"Error generating Word report: {str(e)}")
                st.error(f"{t.get('error_report_generation', 'Error generating report')}: {str(e)}")

@error_handler('eduinsight.gender_effect')
def _create_gender_effect_word_report(df_analysis, mean_scores, selected_columns, t):
    """
    Create a Word document with gender effect analysis.
    
    Args:
        df_analysis (pd.DataFrame): Dataset with gender mapping
        mean_scores (pd.DataFrame): Mean scores by gender
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
        
    Returns:
        Document: Word document with report or None if creation fails
    """
    logger.debug("Creating Word report for gender effect analysis")
    
    try:
        doc = Document()
        
        # Title
        doc.add_heading(t.get("title_gender_effect", "Gender Effect Analysis"), level=1)
        
        # Introduction
        doc.add_paragraph(t.get("gender_effect_intro", 
            "This report analyzes how gender affects student performance, identifying patterns and areas "
            "where gender differences may impact learning outcomes."
        ))
        
        # Gender distribution
        doc.add_heading(t.get("gender_distribution", "Gender Distribution"), level=2)
        
        try:
            # Calculate gender counts and percentages
            gender_counts = df_analysis["gender"].value_counts()
            total_students = len(df_analysis)
            
            # Create table for gender distribution
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # Add headers
            header_cells = table.rows[0].cells
            header_cells[0].text = t.get("gender", "Gender")
            header_cells[1].text = t.get("student_count", "Number of Students")
            header_cells[2].text = t.get("percentage", "Percentage")
            
            # Add data
            for gender, count in gender_counts.items():
                percentage = (count / total_students) * 100
                row_cells = table.add_row().cells
                row_cells[0].text = str(gender)
                row_cells[1].text = str(count)
                row_cells[2].text = f"{percentage:.1f}%"
            
            # Add space after table
            doc.add_paragraph()
        except Exception as e:
            logger.error(f"Error adding gender distribution table: {str(e)}")
            doc.add_paragraph(t.get("gender_dist_error", "Error creating gender distribution table"))
        
        # Mean scores comparison
        doc.add_heading(t.get("mean_score_gender", "Average scores by gender and by task"), level=2)
        
        try:
            # Add mean scores table
            mean_table = doc.add_table(rows=1, cols=len(selected_columns) + 1)
            mean_table.style = 'Table Grid'
            
            # Add headers
            header_cells = mean_table.rows[0].cells
            header_cells[0].text = t.get("gender", "Gender")
            
            for i, col in enumerate(selected_columns):
                header_cells[i + 1].text = t["columns_of_interest"].get(col, col)
            
            # Add data
            for gender, row in mean_scores.iterrows():
                row_cells = mean_table.add_row().cells
                row_cells[0].text = str(gender)
                
                for i, value in enumerate(row):
                    row_cells[i + 1].text = f"{value:.2f}"
            
            # Add space after table
            doc.add_paragraph()
        except Exception as e:
            logger.error(f"Error adding mean scores table: {str(e)}")
            doc.add_paragraph(t.get("mean_scores_error", "Error creating mean scores table"))
        
        # Add bar chart visualization
        with tempfile.TemporaryDirectory() as tmp_dir:
            try:
                # Prepare data for bar chart
                mean_scores_reset = mean_scores.reset_index()
                mean_scores_melted = pd.melt(
                    mean_scores_reset,
                    id_vars=["gender"],
                    value_vars=selected_columns,
                    var_name="Indicator",
                    value_name="Mean Score"
                )
                
                # Translate column names for display
                mean_scores_melted["Indicator"] = mean_scores_melted["Indicator"].map(
                    lambda x: t["columns_of_interest"].get(x, x)
                )
                
                # Create grouped bar chart
                fig = px.bar(
                    mean_scores_melted,
                    x="Indicator",
                    y="Mean Score",
                    color="gender",
                    barmode="group",
                    title=t.get("mean_score_by_gender", "Mean Scores by Gender"),
                    labels={
                        "Indicator": t.get("indicator", "Indicator"),
                        "Mean Score": t.get("mean_score", "Mean Score"),
                        "gender": t.get("gender", "Gender")
                    }
                )
                
                # Improve layout
                fig.update_layout(xaxis_tickangle=-45)
                
                # Save figure to temporary file
                img_path = os.path.join(tmp_dir, "gender_means.png")
                fig.write_image(img_path, width=800, height=500)
                
                # Add figure to document
                doc.add_picture(img_path, width=Inches(6))
            except Exception as e:
                logger.error(f"Error creating mean scores bar chart: {str(e)}")
                doc.add_paragraph(t.get("bar_chart_error", "Error creating mean scores visualization"))
        
        # Add space after figure
        doc.add_paragraph()
        
        # Continue with the rest of the report...
        
        logger.info("Word report created successfully")
        return doc
    
    except Exception as e:
        logger.error(f"Error creating Word report: {str(e)}")
        return None