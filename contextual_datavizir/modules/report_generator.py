"""
Central report generation module for EduInsight.

This module serves as the entry point for all report generation functionality,
delegating to specialized report generators for different analysis types.
"""

import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import specialized report generators
from modules.report_generators.school_report import generate_school_report
from modules.report_generators.language_report import generate_language_report
from modules.report_generators.contextual_report import generate_contextual_report
from modules.report_generators.gender_report import generate_gender_report
from modules.report_generators.ses_report import generate_ses_report

def generate_report(analysis_type, data, t):
    """
    Generate a report for the specified analysis type.
    
    Args:
        analysis_type (str): Type of analysis ('school', 'language', 'contextual', 'gender', 'ses')
        data (dict): Data needed for report generation
        t (dict): Translation dictionary
        
    Returns:
        Document: Word document with the generated report
    """
    # Create base document with consistent styling
    doc = _create_base_document()
    
    # Route to the appropriate specialized report generator
    if analysis_type == 'school':
        return generate_school_report(data, t, doc)
    elif analysis_type == 'language':
        return generate_language_report(data, t, doc)
    elif analysis_type == 'contextual':
        return generate_contextual_report(data, t, doc)
    elif analysis_type == 'gender':
        return generate_gender_report(data, t, doc)
    elif analysis_type == 'ses':
        return generate_ses_report(data, t, doc)
    else:
        # If analysis type is not recognized, create a generic report
        return _create_generic_report(data, t, doc)

def _create_base_document():
    """
    Create a base document with consistent styling.
    
    Returns:
        Document: Word document with consistent styles
    """
    doc = Document()
    
    # Configure document styles for consistent appearance
    styles = doc.styles
    
    # Configure Title style
    style_title = styles['Title']
    style_title.font.name = 'Calibri'
    style_title.font.size = Pt(18)
    style_title.font.bold = True
    style_title.font.color.rgb = RGBColor(0x0, 0x70, 0xC0)
    
    # Configure Heading 1 style
    style_heading1 = styles['Heading 1']
    style_heading1.font.name = 'Calibri'
    style_heading1.font.size = Pt(16)
    style_heading1.font.bold = True
    style_heading1.font.color.rgb = RGBColor(0x0, 0x70, 0xC0)
    
    # Configure Heading 2 style
    style_heading2 = styles['Heading 2']
    style_heading2.font.name = 'Calibri'
    style_heading2.font.size = Pt(14)
    style_heading2.font.bold = True
    style_heading2.font.color.rgb = RGBColor(0x0, 0x70, 0xC0)
    
    # Configure Heading 3 style
    style_heading3 = styles['Heading 3']
    style_heading3.font.name = 'Calibri'
    style_heading3.font.size = Pt(12)
    style_heading3.font.bold = True
    style_heading3.font.color.rgb = RGBColor(0x0, 0x70, 0xC0)
    
    # Configure Normal style
    style_normal = styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    
    # Configure List Bullet style
    style_list = styles['List Bullet']
    style_list.font.name = 'Calibri'
    style_list.font.size = Pt(11)
    
    return doc

def _create_generic_report(data, t, doc):
    """
    Create a generic report when no specialized generator is available.
    
    Args:
        data (dict): Data needed for report generation
        t (dict): Translation dictionary
        doc (Document): Word document to populate
        
    Returns:
        Document: Word document with generic report
    """
    # Add title
    doc.add_heading(t.get("report_title", "Analysis Report"), level=0)
    
    # Add introduction
    doc.add_paragraph(t.get("generic_report_intro", 
        "This report provides an analysis of the selected educational data."
    ))
    
    # Add generic content
    doc.add_heading(t.get("data_overview", "Data Overview"), level=1)
    doc.add_paragraph(t.get("data_overview_text", 
        "The analysis was conducted on the provided dataset. For detailed insights, "
        "please use one of the specialized analysis modules."
    ))
    
    # Add timestamp
    from datetime import datetime
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f"{t.get('generated_on', 'Generated on')}: {current_time}")
    
    return doc

# Utility functions that can be used by any report generator

def add_image_to_document(doc, fig, width=6, title=None):
    """
    Add a plotly figure to a Word document.
    
    Args:
        doc (Document): Word document to add image to
        fig (plotly.Figure): Plotly figure to add
        width (int): Width in inches
        title (str, optional): Optional title for the image
        
    Returns:
        None
    """
    import tempfile
    
    # Create temporary directory for image
    with tempfile.TemporaryDirectory() as tmp_dir:
        # Save figure as image
        img_path = os.path.join(tmp_dir, "figure.png")
        fig.write_image(img_path, width=800, height=500)
        
        # Add title if provided
        if title:
            doc.add_heading(title, level=3)
        
        # Add image to document
        doc.add_picture(img_path, width=Inches(width))
        
        # Add space after image
        doc.add_paragraph()

def add_table_to_document(doc, df, title=None):
    """
    Add a pandas DataFrame as a table to a Word document.
    
    Args:
        doc (Document): Word document to add table to
        df (pd.DataFrame): DataFrame to add as table
        title (str, optional): Optional title for the table
        
    Returns:
        None
    """
    # Add title if provided
    if title:
        doc.add_heading(title, level=3)
    
    # Create table
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = 'Table Grid'
    
    # Add headers
    for i, column in enumerate(df.columns):
        table.cell(0, i).text = str(column)
    
    # Add data
    for i, row in enumerate(df.itertuples(index=False)):
        for j, value in enumerate(row):
            table.cell(i + 1, j).text = str(value)
    
    # Add space after table
    doc.add_paragraph()

def add_correlation_matrix(doc, df, columns, t, title=None):
    """
    Add a correlation matrix visualization to a Word document.
    
    Args:
        doc (Document): Word document to add correlation matrix to
        df (pd.DataFrame): DataFrame containing data
        columns (list): Columns to include in correlation matrix
        t (dict): Translation dictionary
        title (str, optional): Optional title for the correlation matrix
        
    Returns:
        None
    """
    # Calculate correlation matrix
    corr_matrix = df[columns].corr().round(2)
    
    # Create correlation heatmap
    fig = ff.create_annotated_heatmap(
        z=corr_matrix.values,
        x=list(corr_matrix.columns),
        y=list(corr_matrix.index),
        annotation_text=corr_matrix.values.round(2),
        colorscale='Viridis',
        showscale=True,
        zmin=-1, zmax=1
    )
    
    # Update layout
    fig.update_layout(
        title=title or t.get("correlation_matrix", "Correlation Matrix"),
        xaxis_title=t.get("variables", "Variables"),
        yaxis_title=t.get("variables", "Variables"),
        height=600
    )
    
    # Add figure to document
    add_image_to_document(doc, fig, title=title)
    
    # Add interpretation guide
    doc.add_paragraph(t.get("correlation_interpretation", """
    Interpreting the Correlation Matrix:
    - Values close to 1 indicate a strong positive relationship
    - Values close to -1 indicate a strong negative relationship
    - Values close to 0 indicate a weak or no relationship
    """))
