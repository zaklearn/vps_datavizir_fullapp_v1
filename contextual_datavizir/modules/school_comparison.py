import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import plotly.figure_factory as ff
from docx import Document
from docx.shared import Inches
import tempfile
import os
from config import translations, egra_columns, egma_columns

def show_school_comparison(df, language):
    """
    Analyze and compare performance across schools with tabbed interface.
    
    Args:
        df (pd.DataFrame): Dataset containing school data
        language (str): Selected language for UI text
    """
    t = translations[language]  # Get translations
    
    st.title(t.get("title_school_comparison", "School Comparison"))
    
    # Check for required column
    if "school" not in df.columns:
        st.error(t.get("missing_school_column", "School column not found in dataset"))
        return
    
    # Select variables for analysis
    selected_columns = _select_variables(df, t)
    
    if selected_columns:
        try:
            # Calculate statistics by school
            stats_by_school = _calculate_school_statistics(df, selected_columns)
            
            # Create tabs for different analyses
            tab1, tab2, tab3, tab4, tab5 = st.tabs([
                t.get("tab_overview", "Overview"),
                t.get("tab_school_stats", "School Statistics"),
                t.get("tab_distributions", "Distributions"),
                t.get("tab_comparisons", "Detailed Comparisons"),
                t.get("tab_export", "Export & Recommendations")
            ])
            
            with tab1:
                _show_overview(df, selected_columns, t)
            
            with tab2:
                _show_school_statistics(stats_by_school, t)
            
            with tab3:
                _show_school_distributions(df, selected_columns, t)
            
            with tab4:
                _show_detailed_comparisons(df, selected_columns, t)
            
            with tab5:
                _show_recommendations(df, selected_columns, t)
                st.divider()
                _provide_export_options(df, stats_by_school, selected_columns, t)
        
        except Exception as e:
            st.error(f"{t.get('error_message', 'An error occurred')}: {str(e)}")
    else:
        st.warning(t.get("warning_select_indicator", "Please select at least one indicator to analyze."))

def _select_variables(df, t):
    """
    Allow user to select variables for analysis.
    
    Args:
        df (pd.DataFrame): Dataset
        t (dict): Translation dictionary
        
    Returns:
        list: Selected column names
    """
    st.subheader(t.get("select_variables", "Select Variables"))
    
    # Get available assessment columns
    available_columns = []
    for col in egra_columns + egma_columns:
        if col in df.columns:
            available_columns.append(col)
    
    # If no assessment columns found, allow selection from all numeric columns
    if not available_columns:
        numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
        available_columns = [col for col in numeric_cols if col not in ['school', 'stgender', 'ses']]
    
    # Create selection widget
    selected_columns = st.multiselect(
        t.get("variables_all", "Variables:"),
        options=available_columns,
        default=available_columns[:min(3, len(available_columns))],
        format_func=lambda x: t["columns_of_interest"].get(x, x)
    )
    
    return selected_columns

def _calculate_school_statistics(df, selected_columns):
    """
    Calculate descriptive statistics by school.
    
    Args:
        df (pd.DataFrame): Dataset
        selected_columns (list): Columns to analyze
        
    Returns:
        pd.DataFrame: Statistics by school
    """
    # Calculate statistics by school
    stats_by_school = df.groupby("school")[selected_columns].describe().round(2)
    
    # Create a more usable format with flattened column hierarchy
    stats_by_school.columns = ['_'.join(col).strip() for col in stats_by_school.columns]
    stats_by_school = stats_by_school.reset_index()
    
    # Convert school column to string to handle numeric IDs
    stats_by_school["school"] = stats_by_school["school"].astype(str)
    
    return stats_by_school

def _show_overview(df, selected_columns, t):
    """
    Display overview of school comparison analysis.
    
    Args:
        df (pd.DataFrame): Dataset
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    st.header(t.get("overview_title", "Overview"))
    
    # Display key metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric(
            t.get("total_schools", "Total Schools"),
            len(df["school"].unique())
        )
    
    with col2:
        st.metric(
            t.get("total_students", "Total Students"),
            len(df)
        )
    
    with col3:
        # Average students per school
        students_per_school = df.groupby("school").size().mean()
        st.metric(
            t.get("avg_students_per_school", "Avg. Students/School"),
            f"{students_per_school:.1f}"
        )
    
    with col4:
        # Overall average score
        if selected_columns:
            overall_mean = df[selected_columns].mean().mean()
            st.metric(
                t.get("overall_mean_score", "Overall Mean Score"),
                f"{overall_mean:.1f}"
            )
    
    # School size distribution
    st.subheader(t.get("school_size_distribution", "School Size Distribution"))
    school_sizes = df.groupby("school").size().reset_index(name='student_count')
    
    # Convert school column to string to handle numeric IDs
    school_sizes["school"] = school_sizes["school"].astype(str)
    
    fig = px.bar(
        school_sizes.sort_values('student_count', ascending=True),
        x='school',
        y='student_count',
        title=t.get("students_per_school", "Number of Students per School"),
        color='student_count',
        color_continuous_scale='Viridis'
    )
    fig.update_layout(xaxis_tickangle=-45)
    st.plotly_chart(fig, use_container_width=True)
    
    # Performance heatmap
    if selected_columns:
        st.subheader(t.get("performance_heatmap", "Performance Heatmap"))
        
        # Calculate mean scores by school for selected columns
        mean_scores = df.groupby("school")[selected_columns].mean()
        
        # Create translated column names for display
        translated_columns = [t["columns_of_interest"].get(col, col) for col in mean_scores.columns]
        
        # Create heatmap
        school_names = [str(school) for school in mean_scores.index]
        
        fig = ff.create_annotated_heatmap(
            z=mean_scores.values,
            x=translated_columns,
            y=school_names,
            annotation_text=np.around(mean_scores.values, 1),
            colorscale='RdYlGn',
            showscale=True
        )
        
        fig.update_layout(
            title=t.get("school_performance_matrix", "School Performance Matrix"),
            xaxis_title=t.get("indicators", "Indicators"),
            yaxis_title=t.get("schools", "Schools"),
            height=max(400, len(mean_scores) * 30)  # Dynamic height based on number of schools
        )
        
        st.plotly_chart(fig, use_container_width=True)

def _show_school_statistics(stats_by_school, t):
    """
    Display statistics table and visualizations.
    
    Args:
        stats_by_school (pd.DataFrame): Statistics by school
        t (dict): Translation dictionary
    """
    st.header(t.get("school_statistics_title", "School Statistics"))
    
    # Display summary table
    st.subheader(t.get("table_school_comparison", "Descriptive Statistics by School"))
    st.dataframe(stats_by_school, use_container_width=True)
    
    # Add explanation for statistics
    with st.expander(t.get("statistics_explanation", "Understanding these statistics")):
        st.markdown("""
        - **count**: Number of observations (students)
        - **mean**: Average score
        - **std**: Standard deviation (spread of scores)
        - **min**: Minimum score
        - **25%**: First quartile (25% of scores are below this)
        - **50%**: Median (middle score)
        - **75%**: Third quartile (75% of scores are below this)
        - **max**: Maximum score
        """)
    
    # Create comparison charts
    st.subheader(t.get("statistical_comparisons", "Statistical Comparisons"))
    
    # Mean comparison
    mean_cols = [col for col in stats_by_school.columns if col.endswith('_mean')]
    if mean_cols:
        means_df = stats_by_school[['school'] + mean_cols].copy()
        
        # Rename columns for better display
        column_renames = {}
        for col in mean_cols:
            base_col = col.replace('_mean', '')
            column_renames[col] = t["columns_of_interest"].get(base_col, base_col)
        
        means_df = means_df.rename(columns=column_renames)
        
        # Create bar chart
        fig = px.bar(
            means_df,
            x='school',
            y=list(column_renames.values()),
            barmode='group',
            title=t.get("mean_scores_comparison", "Mean Scores Comparison by School")
        )
        fig.update_layout(xaxis_tickangle=-45)
        st.plotly_chart(fig, use_container_width=True)
    
    # Standard deviation comparison
    std_cols = [col for col in stats_by_school.columns if col.endswith('_std')]
    if std_cols:
        std_df = stats_by_school[['school'] + std_cols].copy()
        
        # Rename columns for better display
        column_renames = {}
        for col in std_cols:
            base_col = col.replace('_std', '')
            column_renames[col] = t["columns_of_interest"].get(base_col, base_col)
        
        std_df = std_df.rename(columns=column_renames)
        
        # Create bar chart
        fig = px.bar(
            std_df,
            x='school',
            y=list(column_renames.values()),
            barmode='group',
            title=t.get("std_comparison", "Standard Deviation Comparison (Score Variability)")
        )
        fig.update_layout(xaxis_tickangle=-45)
        st.plotly_chart(fig, use_container_width=True)

def _show_school_distributions(df, selected_columns, t):
    """
    Display distribution visualizations for each school.
    
    Args:
        df (pd.DataFrame): Dataset
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    st.header(t.get("distribution_analysis", "Distribution Analysis"))
    
    # Let user select visualization type
    viz_type = st.radio(
        t.get("select_visualization", "Select visualization type:"),
        ["histogram", "swarm", "density"],
        format_func=lambda x: {
            "histogram": t.get("histogram", "Histogram"),
            "swarm": t.get("swarm", "Swarm Plot"),
            "density": t.get("density", "Density Plot")
        }[x]
    )
    
    # Show visualizations for each selected column
    for col in selected_columns:
        col_name = t["columns_of_interest"].get(col, col)
        st.subheader(f"{t.get('distribution_for', 'Distribution for')} {col_name}")
        
        if viz_type == "histogram":
            # Histogram with distribution overlay
            fig = px.histogram(
                df,
                x=col,
                color="school",
                marginal="box",
                nbins=30,
                title=f"{col_name} {t.get('by_school', 'by School')}",
                barmode="overlay",
                opacity=0.7
            )
            fig.update_layout(bargap=0.1)
            st.plotly_chart(fig, use_container_width=True)
            
        elif viz_type == "swarm":
            # Swarm plot
            fig = px.strip(
                df,
                x="school",
                y=col,
                color="school",
                title=f"{col_name} {t.get('by_school', 'by School')}",
                hover_data={'school': True, col: ':.2f'}
            )
            fig.update_layout(xaxis_tickangle=-45)
            st.plotly_chart(fig, use_container_width=True)
            
        elif viz_type == "density":
            # Density plot
            school_names = [str(school) for school in df["school"].unique()]
            data_groups = [df[df["school"] == school][col].dropna().values for school in df["school"].unique()]
            
            try:
                fig = ff.create_distplot(
                    data_groups,
                    group_labels=school_names,
                    show_hist=False,
                    show_rug=False
                )
                fig.update_layout(
                    title=f"{col_name} {t.get('density_by_school', 'Density by School')}",
                    xaxis_title=col_name,
                    yaxis_title=t.get("density", "Density")
                )
                st.plotly_chart(fig, use_container_width=True)
            except Exception as e:
                st.warning(f"{t.get('density_plot_error', 'Could not create density plot')}: {str(e)}")
                st.info(t.get('try_different_visualization', 'Try selecting a different visualization type.'))
        
        # Add statistical summary
        with st.expander(t.get("statistical_summary", "Statistical Summary")):
            summary_stats = df.groupby("school")[col].agg(['count', 'mean', 'std', 'min', 'max']).round(2)
            st.dataframe(summary_stats)

def _show_detailed_comparisons(df, selected_columns, t):
    """
    Show detailed comparisons between schools.
    
    Args:
        df (pd.DataFrame): Dataset
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    st.header(t.get("detailed_comparisons", "Detailed Comparisons"))
    
    # School performance ranking
    st.subheader(t.get("school_rankings", "School Rankings"))
    
    # Calculate overall performance score (mean of means)
    school_rankings = df.groupby("school")[selected_columns].mean().mean(axis=1).sort_values(ascending=False)
    
    # Create ranking chart
    fig = px.bar(
        school_rankings.reset_index(),
        x='school',
        y=0,
        title=t.get("overall_performance_ranking", "Overall Performance Ranking"),
        labels={'0': t.get("average_score", "Average Score")},
        color=0,
        color_continuous_scale='Viridis'
    )
    fig.update_layout(xaxis_tickangle=-45)
    st.plotly_chart(fig, use_container_width=True)
    
    # Comparative radar chart
    st.subheader(t.get("comparative_analysis", "Comparative Analysis"))
    
    # Let user select schools to compare
    all_schools = df["school"].unique().tolist()
    
    # Convert to string for display but keep original values for filtering
    school_display_map = {str(school): school for school in all_schools}
    
    selected_school_displays = st.multiselect(
        t.get("select_schools_compare", "Select schools to compare:"),
        options=list(school_display_map.keys()),
        default=list(school_display_map.keys())[:3]  # Default to first 3 schools
    )
    
    # Convert back to original values for data filtering
    schools_to_compare = [school_display_map[display] for display in selected_school_displays]
    
    if schools_to_compare and len(schools_to_compare) >= 2:
        # Calculate mean scores for selected schools
        radar_data = []
        for school in schools_to_compare:
            school_data = df[df["school"] == school][selected_columns].mean()
            radar_data.append(school_data)
        
        # Create radar chart
        fig = go.Figure()
        
        for i, school in enumerate(schools_to_compare):
            fig.add_trace(go.Scatterpolar(
                r=radar_data[i].values,
                theta=[t["columns_of_interest"].get(col, col) for col in selected_columns],
                fill='toself',
                name=str(school)  # Convert to string to handle numeric school IDs
            ))
        
        fig.update_layout(
            polar=dict(
                radialaxis=dict(
                    visible=True,
                    range=[0, max([data.max() for data in radar_data])]
                )
            ),
            showlegend=True,
            title=t.get("school_comparison_radar", "School Comparison Radar Chart")
        )
        st.plotly_chart(fig, use_container_width=True)
    
    # Performance gap analysis
    st.subheader(t.get("performance_gap_analysis", "Performance Gap Analysis"))
    
    # Calculate gaps between schools
    means_by_school = df.groupby("school")[selected_columns].mean()
    gaps = []
    
    for col in selected_columns:
        max_score = means_by_school[col].max()
        min_score = means_by_school[col].min()
        gap = max_score - min_score
        best_school = means_by_school[col].idxmax()
        worst_school = means_by_school[col].idxmin()
        
        gaps.append({
            t.get("indicator", "Indicator"): t["columns_of_interest"].get(col, col),
            t.get("gap", "Gap"): gap,
            t.get("best_school", "Best School"): best_school,
            t.get("worst_school", "Worst School"): worst_school
        })
    
    gaps_df = pd.DataFrame(gaps)
    
    # Display gaps table
    st.dataframe(gaps_df.sort_values(t.get("gap", "Gap"), ascending=False), use_container_width=True)
    
    # Visualize gaps
    fig = px.bar(
        gaps_df.sort_values(t.get("gap", "Gap"), ascending=True),
        y=t.get("indicator", "Indicator"),
        x=t.get("gap", "Gap"),
        orientation='h',
        title=t.get("performance_gaps_by_indicator", "Performance Gaps by Indicator"),
        color=t.get("gap", "Gap"),
        color_continuous_scale='Reds'
    )
    st.plotly_chart(fig, use_container_width=True)

def _show_recommendations(df, selected_columns, t):
    """
    Display recommendations based on school comparison analysis.
    
    Args:
        df (pd.DataFrame): Dataset
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    st.header(t.get("recommendations", "Recommendations"))
    
    # Calculate school performance metrics
    means_by_school = df.groupby("school")[selected_columns].mean()
    overall_performance = means_by_school.mean(axis=1).sort_values(ascending=False)
    
    # Identify top and bottom performing schools
    top_schools = overall_performance.head(3).index.tolist()
    bottom_schools = overall_performance.tail(3).index.tolist()
    
    # General recommendations
    st.subheader(t.get("general_recommendations", "General Recommendations"))
    
    general_recs = [
        t.get("rec1", "Share best practices from high-performing schools with lower-performing schools"),
        t.get("rec2", "Investigate factors behind performance differences (resources, teaching methods, etc.)"),
        t.get("rec3", "Develop targeted interventions for specific areas of weakness identified in each school"),
        t.get("rec4", "Establish mentoring relationships between schools with complementary strengths")
    ]
    
    for rec in general_recs:
        st.markdown(f"- {rec}")
    
    # School-specific recommendations
    st.subheader(t.get("school_specific_recommendations", "School-Specific Recommendations"))
    
    # Top performing schools
    st.markdown(f"### {t.get('top_performing_schools', 'Top Performing Schools')}")
    for school in top_schools:
        st.markdown(f"**{school}**")
        # Find their weakest areas
        school_scores = means_by_school.loc[school]
        weakest_areas = school_scores.nsmallest(2)
        
        for area, score in weakest_areas.items():
            area_name = t["columns_of_interest"].get(area, area)
            st.markdown(f"- {t.get('improve_further', 'Focus on improving')}: {area_name} ({score:.1f})")
        
        st.markdown(f"- {t.get('share_practices', 'Document and share successful practices')}")
    
    # Bottom performing schools
    st.markdown(f"### {t.get('bottom_performing_schools', 'Bottom Performing Schools')}")
    for school in bottom_schools:
        st.markdown(f"**{school}**")
        # Find their strongest and weakest areas
        school_scores = means_by_school.loc[school]
        strongest_area = school_scores.idxmax()
        weakest_area = school_scores.idxmin()
        
        st.markdown(f"- {t.get('leverage_strength', 'Leverage strength in')}: {t['columns_of_interest'].get(strongest_area, strongest_area)}")
        st.markdown(f"- {t.get('urgent_improvement', 'Urgent improvement needed in')}: {t['columns_of_interest'].get(weakest_area, weakest_area)}")
        st.markdown(f"- {t.get('partner_with', 'Consider partnering with higher-performing schools')}")
    
    # Add visualization of improvement priorities
    st.subheader(t.get("improvement_priorities", "Improvement Priorities"))
    
    # Create a matrix of school vs. indicator performance
    improvement_matrix = means_by_school.copy()
    
    # Normalize scores to 0-1 range for visualization
    improvement_matrix = (improvement_matrix - improvement_matrix.min()) / (improvement_matrix.max() - improvement_matrix.min())
    
    # Create heatmap for improvement priorities (inverted colors for urgency)
    school_names = [str(school) for school in improvement_matrix.index]
    
    fig = ff.create_annotated_heatmap(
        z=1 - improvement_matrix.values,  # Invert so red indicates urgency
        x=[t["columns_of_interest"].get(col, col) for col in improvement_matrix.columns],
        y=school_names,
        annotation_text=np.around(means_by_school.values, 1),
        colorscale='RdYlGn_r',  # Reverse colorscale
        showscale=True
    )
    
    fig.update_layout(
        title=t.get("improvement_priority_matrix", "Improvement Priority Matrix (Red = Urgent)"),
        xaxis_title=t.get("indicators", "Indicators"),
        yaxis_title=t.get("schools", "Schools"),
        height=max(400, len(means_by_school) * 30)
    )
    
    st.plotly_chart(fig, use_container_width=True)

def _provide_export_options(df, stats_by_school, selected_columns, t):
    """
    Provide options to export analysis results.
    
    Args:
        df (pd.DataFrame): Original dataset
        stats_by_school (pd.DataFrame): Statistics by school
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
    """
    st.header(t.get("export_options", "Export Options"))
    
    col1, col2 = st.columns(2)
    
    # CSV Export
    with col1:
        # Prepare comprehensive export data
        export_data = _prepare_export_data(df, stats_by_school, selected_columns, t)
        
        csv = export_data.to_csv(index=False).encode('utf-8-sig')
        st.download_button(
            t.get("export_school_comparison_csv", "Download CSV"),
            csv,
            "school_comparison_stats.csv",
            "text/csv",
            key='download-school-stats-csv'
        )
    
    # Word Export
    with col2:
        if st.button(t.get("export_school_comparison_word", "Export to Word")):
            try:
                doc = _create_school_comparison_word_report(df, stats_by_school, selected_columns, t)
                
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    doc.save(tmp.name)
                    with open(tmp.name, 'rb') as f:
                        docx = f.read()
                    st.download_button(
                        t.get("download_school_comparison_word", "Download Word Report"),
                        docx,
                        "school_comparison_report.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                os.unlink(tmp.name)
            except Exception as e:
                st.error(f"{t.get('error_report_generation', 'Error generating report')}: {str(e)}")

def _prepare_export_data(df, stats_by_school, selected_columns, t):
    """
    Prepare comprehensive data for export.
    
    Args:
        df (pd.DataFrame): Original dataset
        stats_by_school (pd.DataFrame): Statistics by school
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
        
    Returns:
        pd.DataFrame: Data ready for export
    """
    export_data = stats_by_school.copy()
    
    # Add rankings
    means_by_school = df.groupby("school")[selected_columns].mean()
    overall_performance = means_by_school.mean(axis=1).sort_values(ascending=False)
    
    export_data = export_data.merge(
        overall_performance.reset_index().rename(columns={0: 'overall_rank'}),
        on='school'
    )
    
    # Add school size
    school_sizes = df.groupby("school").size().reset_index(name='student_count')
    export_data = export_data.merge(school_sizes, on='school')
    
    return export_data

def _create_school_comparison_word_report(df, stats_by_school, selected_columns, t):
    """
    Create a Word document with school comparison analysis.
    
    Args:
        df (pd.DataFrame): Original dataset
        stats_by_school (pd.DataFrame): Statistics by school
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
        
    Returns:
        Document: Word document with report
    """
    doc = Document()
    
    # Title
    doc.add_heading(t.get("title_school_comparison", "School Comparison"), level=1)
    
    # Introduction
    doc.add_paragraph(t.get("school_comparison_intro", 
        "This report analyzes performance differences across schools, identifying patterns and areas for improvement."
    ))
    
    # Overview
    doc.add_heading(t.get("overview", "Overview"), level=2)
    
    total_schools = len(df["school"].unique())
    total_students = len(df)
    students_per_school = df.groupby("school").size().mean()
    
    doc.add_paragraph(f"{t.get('total_schools', 'Total Schools')}: {total_schools}")
    doc.add_paragraph(f"{t.get('total_students', 'Total Students')}: {total_students}")
    doc.add_paragraph(f"{t.get('avg_students_per_school', 'Average Students per School')}: {students_per_school:.1f}")
    
    # Statistics table
    doc.add_heading(t.get("table_school_comparison", "Descriptive Statistics by School"), level=2)
    
    # Create table for overall stats
    table = doc.add_table(rows=1, cols=len(stats_by_school.columns))
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, col in enumerate(stats_by_school.columns):
        header_cells[i].text = str(col)
    
    # Add data
    for _, row in stats_by_school.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    # Add space after table
    doc.add_paragraph()
    
    # School Rankings
    doc.add_heading(t.get("school_rankings", "School Rankings"), level=2)
    
    means_by_school = df.groupby("school")[selected_columns].mean()
    overall_performance = means_by_school.mean(axis=1).sort_values(ascending=False)
    
    ranking_table = doc.add_table(rows=1, cols=3)
    ranking_table.style = 'Table Grid'
    
    # Headers
    header_cells = ranking_table.rows[0].cells
    header_cells[0].text = t.get("rank", "Rank")
    header_cells[1].text = t.get("school", "School")
    header_cells[2].text = t.get("average_score", "Average Score")
    
    # Data
    for i, (school, score) in enumerate(overall_performance.items()):
        row_cells = ranking_table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = school
        row_cells[2].text = f"{score:.2f}"
    
    # Add space after table
    doc.add_paragraph()
    
    # Performance Gaps
    doc.add_heading(t.get("performance_gaps", "Performance Gaps"), level=2)
    
    gaps_table = doc.add_table(rows=1, cols=5)
    gaps_table.style = 'Table Grid'
    
    # Headers
    header_cells = gaps_table.rows[0].cells
    header_cells[0].text = t.get("indicator", "Indicator")
    header_cells[1].text = t.get("gap", "Gap")
    header_cells[2].text = t.get("best_school", "Best School")
    header_cells[3].text = t.get("best_score", "Best Score")
    header_cells[4].text = t.get("worst_school", "Worst School")
    
    # Calculate gaps
    for col in selected_columns:
        max_score = means_by_school[col].max()
        min_score = means_by_school[col].min()
        gap = max_score - min_score
        best_school = means_by_school[col].idxmax()
        worst_school = means_by_school[col].idxmin()
        
        row_cells = gaps_table.add_row().cells
        row_cells[0].text = t["columns_of_interest"].get(col, col)
        row_cells[1].text = f"{gap:.2f}"
        row_cells[2].text = best_school
        row_cells[3].text = f"{max_score:.2f}"
        row_cells[4].text = worst_school
    
    # Add recommendations
    doc.add_heading(t.get("recommendations", "Recommendations"), level=2)
    
    general_recs = [
        t.get("rec1", "Share best practices from high-performing schools with lower-performing schools"),
        t.get("rec2", "Investigate factors behind performance differences (resources, teaching methods, etc.)"),
        t.get("rec3", "Develop targeted interventions for specific areas of weakness identified in each school"),
        t.get("rec4", "Establish mentoring relationships between schools with complementary strengths")
    ]
    
    for rec in general_recs:
        doc.add_paragraph(rec, style='List Bullet')
    
    # School-specific recommendations
    doc.add_heading(t.get("school_specific_recommendations", "School-Specific Recommendations"), level=2)
    
    # Top performing schools
    top_schools = overall_performance.head(3).index.tolist()
    doc.add_heading(t.get("top_performing_schools", "Top Performing Schools"), level=3)
    
    for school in top_schools:
        doc.add_paragraph(f"{school}:", style='Heading 4')
        doc.add_paragraph(t.get("share_practices", "Document and share successful practices"), style='List Bullet')
        
        # Find weakest area
        school_scores = means_by_school.loc[school]
        weakest_area = school_scores.idxmin()
        weakest_score = school_scores.min()
        doc.add_paragraph(f"{t.get('improve_further', 'Focus on improving')}: {t['columns_of_interest'].get(weakest_area, weakest_area)} ({weakest_score:.1f})", style='List Bullet')
    
    # Bottom performing schools
    bottom_schools = overall_performance.tail(3).index.tolist()
    doc.add_heading(t.get("bottom_performing_schools", "Bottom Performing Schools"), level=3)
    
    for school in bottom_schools:
        doc.add_paragraph(f"{school}:", style='Heading 4')
        doc.add_paragraph(t.get("partner_with", "Consider partnering with higher-performing schools"), style='List Bullet')
        
        # Find best area
        school_scores = means_by_school.loc[school]
        best_area = school_scores.idxmax()
        best_score = school_scores.max()
        doc.add_paragraph(f"{t.get('leverage_strength', 'Leverage strength in')}: {t['columns_of_interest'].get(best_area, best_area)} ({best_score:.1f})", style='List Bullet')
    
    return doc