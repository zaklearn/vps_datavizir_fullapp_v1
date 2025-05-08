import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import tempfile
import os
import json
import plotly.io as pio
from pathlib import Path

class AnalysisReportGenerator:
    """
    Utility class for generating comprehensive reports from educational assessment analyses.
    
    This class provides methods to create formatted reports for different types of analyses,
    including statistical analysis, zero scores analysis, school comparison, gender effect,
    and international standards comparison.
    """
    
    def __init__(self):
        """Initialize the report generator with predefined interpretation rules."""
        # Define interpretation rules for different analyses
        self.interpretation_rules = {
            # Rules for statistical analysis (analyse1)
            "statistical": {
                "score_ranges": {
                    "very_low": {"max": 30, "description_en": "very low level", "description_fr": "niveau très faible"},
                    "low": {"min": 30, "max": 50, "description_en": "low level", "description_fr": "niveau faible"},
                    "average": {"min": 50, "max": 70, "description_en": "average level", "description_fr": "niveau moyen"},
                    "good": {"min": 70, "max": 85, "description_en": "good level", "description_fr": "bon niveau"},
                    "excellent": {"min": 85, "description_en": "excellent level", "description_fr": "excellent niveau"}
                },
                "recommendations": {
                    "very_low": {
                        "en": "Urgent reinforcement of basic skills is necessary",
                        "fr": "Un renforcement urgent des compétences de base est nécessaire"
                    },
                    "low": {
                        "en": "Additional support is recommended",
                        "fr": "Un soutien supplémentaire est recommandé"
                    },
                    "average": {
                        "en": "Continue strengthening skills",
                        "fr": "Continuer le renforcement des compétences"
                    },
                    "good": {
                        "en": "Maintain current good practices",
                        "fr": "Maintenir les bonnes pratiques actuelles"
                    },
                    "excellent": {
                        "en": "Consider extending with more advanced content",
                        "fr": "Possibilité d'approfondir avec des contenus plus avancés"
                    }
                }
            },
            
            # Rules for zero scores analysis (analyse2)
            "zero_scores": {
                "competences": {
                    "clpm": {
                        "type_en": "decoding", "type_fr": "décodage",
                        "thresholds": {"critical": 20, "concerning": 10, "watch": 5},
                        "importance_en": "fundamental", "importance_fr": "fondamentale",
                        "description_en": "rapid letter recognition ability", "description_fr": "capacité de reconnaissance rapide des lettres"
                    },
                    "phoneme": {
                        "type_en": "decoding", "type_fr": "décodage",
                        "thresholds": {"critical": 30, "concerning": 20, "watch": 10},
                        "importance_en": "fundamental", "importance_fr": "fondamentale",
                        "description_en": "phonemic awareness", "description_fr": "conscience phonémique"
                    },
                    "sound_word": {
                        "type_en": "reading", "type_fr": "lecture",
                        "thresholds": {"critical": 25, "concerning": 15, "watch": 10},
                        "importance_en": "essential", "importance_fr": "essentielle",
                        "description_en": "word reading accuracy", "description_fr": "précision de lecture des mots"
                    },
                    "cwpm": {
                        "type_en": "fluency", "type_fr": "fluidité",
                        "thresholds": {"critical": 20, "concerning": 15, "watch": 10},
                        "importance_en": "essential", "importance_fr": "essentielle",
                        "description_en": "reading fluency", "description_fr": "fluidité de lecture"
                    },
                    "listening": {
                        "type_en": "comprehension", "type_fr": "compréhension",
                        "thresholds": {"critical": 40, "concerning": 30, "watch": 20},
                        "importance_en": "critical", "importance_fr": "critique",
                        "description_en": "listening comprehension", "description_fr": "compréhension orale"
                    },
                    "orf": {
                        "type_en": "fluency", "type_fr": "fluidité",
                        "thresholds": {"critical": 15, "concerning": 10, "watch": 5},
                        "importance_en": "essential", "importance_fr": "essentielle",
                        "description_en": "oral reading fluency", "description_fr": "fluidité de lecture à voix haute"
                    },
                    "comprehension": {
                        "type_en": "comprehension", "type_fr": "compréhension",
                        "thresholds": {"critical": 40, "concerning": 30, "watch": 20},
                        "importance_en": "critical", "importance_fr": "critique",
                        "description_en": "reading comprehension", "description_fr": "compréhension en lecture"
                    }
                }
            },
            
            # Rules for school comparison (analyse3)
            "school_comparison": {
                "significance_levels": {
                    "high": {"max": 0.01, "description_en": "highly significant", "description_fr": "hautement significatif"},
                    "moderate": {"min": 0.01, "max": 0.05, "description_en": "moderately significant", "description_fr": "modérément significatif"},
                    "not_significant": {"min": 0.05, "description_en": "not significant", "description_fr": "non significatif"}
                },
                "recommendations": {
                    "high": {
                        "en": "Investigate practices in high-performing schools for potential system-wide adoption",
                        "fr": "Étudier les pratiques des écoles performantes pour une adoption potentielle à l'échelle du système"
                    },
                    "moderate": {
                        "en": "Consider targeted interventions to address performance gaps between schools",
                        "fr": "Envisager des interventions ciblées pour combler les écarts de performance entre les écoles"
                    },
                    "not_significant": {
                        "en": "Continue monitoring performance across schools to ensure consistent quality",
                        "fr": "Continuer à surveiller les performances entre les écoles pour garantir une qualité constante"
                    }
                }
            },
            
            # Rules for gender effect analysis (analyse10)
            "gender_effect": {
                "significance_levels": {
                    "high": {"max": 0.01, "description_en": "highly significant", "description_fr": "hautement significatif"},
                    "moderate": {"min": 0.01, "max": 0.05, "description_en": "moderately significant", "description_fr": "modérément significatif"},
                    "not_significant": {"min": 0.05, "description_en": "not significant", "description_fr": "non significatif"}
                },
                "recommendations": {
                    "high": {
                        "en": "Implement gender-responsive teaching strategies to address significant performance gaps",
                        "fr": "Mettre en œuvre des stratégies d'enseignement tenant compte du genre pour combler les écarts de performance significatifs"
                    },
                    "moderate": {
                        "en": "Monitor gender differences and provide targeted support where needed",
                        "fr": "Surveiller les différences entre les genres et fournir un soutien ciblé si nécessaire"
                    },
                    "not_significant": {
                        "en": "Continue gender-inclusive teaching practices",
                        "fr": "Poursuivre les pratiques d'enseignement inclusives en matière de genre"
                    }
                }
            },
            
            # Rules for international standards comparison (analyse12)
            "international_comparison": {
                "achievement_levels": {
                    "critical": {"max": 70, "description_en": "critical gap", "description_fr": "écart critique"},
                    "concerning": {"min": 70, "max": 85, "description_en": "concerning gap", "description_fr": "écart préoccupant"},
                    "approaching": {"min": 85, "max": 100, "description_en": "approaching standard", "description_fr": "proche du standard"},
                    "meeting": {"min": 100, "description_en": "meeting or exceeding standard", "description_fr": "atteint ou dépasse le standard"}
                },
                "recommendations": {
                    "critical": {
                        "en": "Implement intensive intervention programs to address critical performance gaps",
                        "fr": "Mettre en œuvre des programmes d'intervention intensive pour combler les écarts de performance critiques"
                    },
                    "concerning": {
                        "en": "Strengthen instructional approaches and provide targeted support",
                        "fr": "Renforcer les approches pédagogiques et fournir un soutien ciblé"
                    },
                    "approaching": {
                        "en": "Continue current strategies with minor adjustments to reach standards",
                        "fr": "Poursuivre les stratégies actuelles avec des ajustements mineurs pour atteindre les standards"
                    },
                    "meeting": {
                        "en": "Maintain successful practices and consider setting higher goals",
                        "fr": "Maintenir les pratiques réussies et envisager de fixer des objectifs plus élevés"
                    }
                }
            }
        }
    
    def _get_competence_level(self, competence, percentage, analysis_type, language="en"):
        """
        Determines the concern level for a competence based on zero score percentage.
        
        Args:
            competence (str): The competence code (e.g., "clpm", "phoneme")
            percentage (float): Percentage of zero scores
            analysis_type (str): Type of analysis (e.g., "zero_scores")
            language (str): Language for output ("en" or "fr")
            
        Returns:
            str: Concern level (CRITICAL, CONCERNING, WATCH, or SATISFACTORY)
        """
        if analysis_type == "zero_scores":
            if competence in self.interpretation_rules[analysis_type]["competences"]:
                thresholds = self.interpretation_rules[analysis_type]["competences"][competence]["thresholds"]
                
                if percentage >= thresholds["critical"]:
                    return "CRITICAL" if language == "en" else "CRITIQUE"
                elif percentage >= thresholds["concerning"]:
                    return "CONCERNING" if language == "en" else "PRÉOCCUPANT"
                elif percentage >= thresholds["watch"]:
                    return "WATCH" if language == "en" else "À SURVEILLER"
                return "SATISFACTORY" if language == "en" else "SATISFAISANT"
        
        return "UNKNOWN" if language == "en" else "INCONNU"
    
    def _get_progression_chain(self, scores_data, analysis_type="zero_scores", language="en"):
        """
        Analyzes the progression chain of competencies based on scores.
        
        Args:
            scores_data (dict): Dictionary with competence codes as keys and scores as values
            analysis_type (str): Type of analysis
            language (str): Language for output ("en" or "fr")
            
        Returns:
            dict: Progression chain by skill type
        """
        lang_suffix = "_en" if language == "en" else "_fr"
        progression = {
            "decoding" if language == "en" else "décodage": [],
            "reading" if language == "en" else "lecture": [],
            "fluency" if language == "en" else "fluidité": [],
            "comprehension": []  # Same in both languages
        }
        
        if analysis_type == "zero_scores":
            for comp, score in scores_data.items():
                if comp in self.interpretation_rules[analysis_type]["competences"]:
                    comp_type = self.interpretation_rules[analysis_type]["competences"][comp][f"type{lang_suffix}"]
                    progression[comp_type].append((comp, score))
        
        return progression
    
    def _get_impact_description(self, level, percentage, language="en"):
        """
        Generates an impact description based on the concern level.
        
        Args:
            level (str): Concern level (CRITICAL, CONCERNING, WATCH, SATISFACTORY)
            percentage (float): Percentage value
            language (str): Language for output ("en" or "fr")
            
        Returns:
            str: Impact description
        """
        if language == "en":
            if level == "CRITICAL":
                return "Severe impact on learning progression"
            elif level == "CONCERNING":
                return "Significant impact requiring rapid intervention"
            elif level == "WATCH":
                return "Moderate impact requiring regular monitoring"
            return "Minimal impact - maintain vigilance"
        else:  # French
            if level == "CRITIQUE":
                return "Impact sévère sur la progression des apprentissages"
            elif level == "PRÉOCCUPANT":
                return "Impact significatif nécessitant une intervention rapide"
            elif level == "À SURVEILLER":
                return "Impact modéré nécessitant un suivi régulier"
            return "Impact minimal - maintenir la vigilance"
    
    def _get_specific_recommendations(self, level, comp_type, language="en"):
        """
        Generates specific recommendations based on concern level and competence type.
        
        Args:
            level (str): Concern level (CRITICAL, CONCERNING, WATCH, SATISFACTORY)
            comp_type (str): Competence type (decoding, reading, fluency, comprehension)
            language (str): Language for output ("en" or "fr")
            
        Returns:
            list: List of specific recommendations
        """
        if language == "en":
            if level == "CRITICAL":
                if comp_type == "comprehension":
                    return [
                        "Intensive reinforcement of oral comprehension",
                        "Daily guided comprehension activities",
                        "Individualized student monitoring"
                    ]
                elif comp_type == "decoding":
                    return [
                        "Daily phonological awareness exercises",
                        "Systematic decoding activities",
                        "Enhanced visual support"
                    ]
            return ["Maintain regular monitoring", "Adapt activities as needed"]
        else:  # French
            if level == "CRITIQUE":
                if comp_type == "compréhension":
                    return [
                        "Renforcement intensif de la compréhension orale",
                        "Activités quotidiennes de compréhension guidée",
                        "Suivi individualisé des élèves"
                    ]
                elif comp_type == "décodage":
                    return [
                        "Exercices quotidiens de conscience phonologique",
                        "Activités systématiques de décodage",
                        "Support visuel renforcé"
                    ]
            return ["Maintenir le suivi régulier", "Adapter les activités au besoin"]
    
    def _determine_level(self, score):
        """Determines level based on score for statistical analysis."""
        if score < 30:
            return "very_low"
        elif score < 50:
            return "low"
        elif score < 70:
            return "average"
        elif score < 85:
            return "good"
        return "excellent"
    
    def generate_report(self, analysis_type, results, language="en"):
        """
        Generate a report based on the type of analysis and results.
        
        Args:
            analysis_type (str): Type of analysis ("statistical", "zero_scores", etc.)
            results (dict): Dictionary containing analysis results
            language (str): Language for the report ("en" or "fr")
            
        Returns:
            dict: Report content with summary, recommendations, and details
        """
        if analysis_type == "statistical":
            return self._generate_statistical_report(results, language)
        elif analysis_type == "zero_scores":
            return self._generate_zero_scores_report(results, language)
        elif analysis_type == "school_comparison":
            return self._generate_school_comparison_report(results, language)
        elif analysis_type == "gender_effect":
            return self._generate_gender_effect_report(results, language)
        elif analysis_type == "international_comparison":
            return self._generate_international_comparison_report(results, language)
        else:
            return self._generate_default_report(results, language)
    
    def _generate_statistical_report(self, results, language):
        """
        Generates a report for statistical analysis.
        
        Args:
            results (dict): Dictionary containing statistical analysis results
            language (str): Language for the report ("en" or "fr")
            
        Returns:
            dict: Report content with summary, recommendations, and details
        """
        global_stats = results.get("global_stats", {})
        mean_scores = results.get("analysis_context", {}).get("mean_scores", {})
        
        summary = "Statistical Analysis Summary:\n\n" if language == "en" else "Résumé de l'analyse statistique :\n\n"
        recommendations = []
        details = {}
        
        # Analysis for each indicator
        for indicator, mean in mean_scores.items():
            level = self._determine_level(mean)
            rules = self.interpretation_rules["statistical"]
            
            # Get description in appropriate language
            description_key = f"description_{language}"
            default_description = "level" if language == "en" else "niveau"
            description = rules["score_ranges"][level].get(description_key, default_description)
            
            # Add to summary
            summary += f"• {indicator}: {mean:.1f} - {description}\n"
            
            # Add recommendation in appropriate language
            rec_text = rules["recommendations"][level][language]
            recommendations.append(f"For {indicator}: {rec_text}" if language == "en" else f"Pour {indicator} : {rec_text}")
            
            # Add details
            details[indicator] = {
                "level": description,
                "mean_score": mean,
                "interpretation": rec_text
            }
        
        return {
            "summary": summary,
            "recommendations": recommendations,
            "details": details
        }
    
    def _generate_zero_scores_report(self, results, language):
        """
        Generates a detailed interpretation of zero scores.
        
        Args:
            results (dict): Dictionary containing zero scores analysis results
            language (str): Language for the report ("en" or "fr")
            
        Returns:
            dict: Report content with summary, recommendations, and details
        """
        zero_scores_data = results.get("zero_scores_data", {})
        total_students = results.get("total_students", 0)
        
        # Analyze competencies by concern level
        levels = {
            "CRITICAL" if language == "en" else "CRITIQUE": [],
            "CONCERNING" if language == "en" else "PRÉOCCUPANT": [],
            "WATCH" if language == "en" else "À SURVEILLER": [],
            "SATISFACTORY" if language == "en" else "SATISFAISANT": []
        }
        
        for competence, percentage in zero_scores_data.items():
            level = self._get_competence_level(competence, percentage, "zero_scores", language)
            levels[level].append((competence, percentage))
        
        # Build summary
        if language == "en":
            summary = "DETAILED ANALYSIS OF ZERO SCORES\n\n"
            summary += "1. MAJOR ATTENTION POINTS:\n"
            
            if levels["CRITICAL"]:
                summary += "\n🚨 CRITICAL SITUATIONS:\n"
                for comp, score in levels["CRITICAL"]:
                    if comp in self.interpretation_rules["zero_scores"]["competences"]:
                        desc = self.interpretation_rules["zero_scores"]["competences"][comp]["description_en"]
                        summary += f"• {comp} ({score:.1f}% zero scores) - {desc}\n"
                        summary += "  Impact: Major obstacle to learning progression\n"
            
            if levels["CONCERNING"]:
                summary += "\n⚠️ CONCERNING SITUATIONS:\n"
                for comp, score in levels["CONCERNING"]:
                    if comp in self.interpretation_rules["zero_scores"]["competences"]:
                        desc = self.interpretation_rules["zero_scores"]["competences"][comp]["description_en"]
                        summary += f"• {comp} ({score:.1f}% zero scores) - {desc}\n"
        else:  # French
            summary = "ANALYSE DÉTAILLÉE DES SCORES ZÉRO\n\n"
            summary += "1. POINTS D'ATTENTION MAJEURS :\n"
            
            if levels["CRITIQUE"]:
                summary += "\n🚨 SITUATIONS CRITIQUES :\n"
                for comp, score in levels["CRITIQUE"]:
                    if comp in self.interpretation_rules["zero_scores"]["competences"]:
                        desc = self.interpretation_rules["zero_scores"]["competences"][comp]["description_fr"]
                        summary += f"• {comp} ({score:.1f}% de scores zéro) - {desc}\n"
                        summary += "  Impact : Obstacle majeur à la progression des apprentissages\n"
            
            if levels["PRÉOCCUPANT"]:
                summary += "\n⚠️ SITUATIONS PRÉOCCUPANTES :\n"
                for comp, score in levels["PRÉOCCUPANT"]:
                    if comp in self.interpretation_rules["zero_scores"]["competences"]:
                        desc = self.interpretation_rules["zero_scores"]["competences"][comp]["description_fr"]
                        summary += f"• {comp} ({score:.1f}% de scores zéro) - {desc}\n"
        
        # Analyze the progression chain
        progression = self._get_progression_chain(zero_scores_data, "zero_scores", language)
        
        if language == "en":
            summary += "\n2. ANALYSIS BY SKILL CHAIN:\n"
            
            if progression["decoding"]:
                avg_decoding = sum(score for _, score in progression["decoding"]) / len(progression["decoding"])
                summary += f"\n📚 Decoding skills (average: {avg_decoding:.1f}% zero scores)\n"
                
            if progression["comprehension"]:
                avg_comprehension = sum(score for _, score in progression["comprehension"]) / len(progression["comprehension"])
                summary += f"\n🎯 Comprehension skills (average: {avg_comprehension:.1f}% zero scores)\n"
        else:  # French
            summary += "\n2. ANALYSE PAR CHAÎNE DE COMPÉTENCES :\n"
            
            if progression["décodage"]:
                avg_decoding = sum(score for _, score in progression["décodage"]) / len(progression["décodage"])
                summary += f"\n📚 Compétences de décodage (moyenne : {avg_decoding:.1f}% de scores zéro)\n"
                
            if progression["compréhension"]:
                avg_comprehension = sum(score for _, score in progression["compréhension"]) / len(progression["compréhension"])
                summary += f"\n🎯 Compétences de compréhension (moyenne : {avg_comprehension:.1f}% de scores zéro)\n"
        
        # Generate recommendations
        if language == "en":
            recommendations = ["PRIORITY RECOMMENDATIONS:"]
            
            if levels["CRITICAL"]:
                for comp, score in levels["CRITICAL"]:
                    recommendations.append(f"🚨 For {comp}:")
                    recommendations.append(f"  • Immediate implementation of intensive remediation program")
                    recommendations.append(f"  • In-depth diagnostic assessment to identify blockages")
                    recommendations.append(f"  • Specific teacher training on this skill")
            
            if progression["comprehension"]:
                if any(score > 40 for _, score in progression["comprehension"]):
                    recommendations.append("\n📘 For comprehension:")
                    recommendations.append("  • Complete revision of pedagogical approach")
                    recommendations.append("  • Reinforcement of oral and written comprehension activities")
                    recommendations.append("  • Implementation of ability grouping")
        else:  # French
            recommendations = ["RECOMMANDATIONS PRIORITAIRES :"]
            
            if levels["CRITIQUE"]:
                for comp, score in levels["CRITIQUE"]:
                    recommendations.append(f"🚨 Pour {comp} :")
                    recommendations.append(f"  • Mise en place immédiate d'un programme de remédiation intensive")
                    recommendations.append(f"  • Évaluation diagnostique approfondie pour identifier les blocages")
                    recommendations.append(f"  • Formation spécifique des enseignants sur cette compétence")
            
            if progression["compréhension"]:
                if any(score > 40 for _, score in progression["compréhension"]):
                    recommendations.append("\n📘 Pour la compréhension :")
                    recommendations.append("  • Révision complète de l'approche pédagogique")
                    recommendations.append("  • Renforcement des activités de compréhension orale et écrite")
                    recommendations.append("  • Mise en place de groupes de niveaux")
        
        # Generate details
        details = {}
        for competence, percentage in zero_scores_data.items():
            if competence in self.interpretation_rules["zero_scores"]["competences"]:
                level = self._get_competence_level(competence, percentage, "zero_scores", language)
                comp_info = self.interpretation_rules["zero_scores"]["competences"][competence]
                comp_type = comp_info[f"type_{language}"]
                
                details[competence] = {
                    "level": level,
                    "percentage": f"{percentage:.1f}%",
                    "type": comp_type,
                    "importance": comp_info[f"importance_{language}"],
                    "impact": self._get_impact_description(level, percentage, language),
                    "recommendations": self._get_specific_recommendations(level, comp_type, language)
                }
        
        return {
            "summary": summary,
            "recommendations": recommendations,
            "details": details
        }
    
    def _generate_school_comparison_report(self, results, language):
        """
        Generates a report for school comparison analysis.
        
        Args:
            results (dict): Dictionary containing school comparison results
            language (str): Language for the report ("en" or "fr")
            
        Returns:
            dict: Report content with summary, recommendations, and details
        """
        # Implementation for school comparison reporting would go here
        return self._generate_default_report(results, language)
    
    def _generate_gender_effect_report(self, results, language):
        """
        Generates a report for gender effect analysis.
        
        Args:
            results (dict): Dictionary containing gender effect analysis results
            language (str): Language for the report ("en" or "fr")
            
        Returns:
            dict: Report content with summary, recommendations, and details
        """
        # Implementation for gender effect reporting would go here
        return self._generate_default_report(results, language)
    
    def _generate_international_comparison_report(self, results, language):
        """
        Generates a report for international comparison analysis.
        
        Args:
            results (dict): Dictionary containing international comparison results
            language (str): Language for the report ("en" or "fr")
            
        Returns:
            dict: Report content with summary, recommendations, and details
        """
        # Implementation for international comparison reporting would go here
        return self._generate_default_report(results, language)
    
    def _generate_default_report(self, results, language):
        """
        Generates a default report when specific report types are not implemented.
        
        Args:
            results (dict): Analysis results
            language (str): Language for the report ("en" or "fr")
            
        Returns:
            dict: Report content with summary, recommendations, and details
        """
        if language == "en":
            return {
                "summary": "Analysis of results completed.",
                "recommendations": ["Monitor progress regularly"],
                "details": {}
            }
        else:  # French
            return {
                "summary": "Analyse des résultats effectuée.",
                "recommendations": ["Suivre régulièrement les progrès"],
                "details": {}
            }
    
    def show_report_ui(self, report, language):
        """
        Displays the report in the Streamlit interface.
        
        Args:
            report (dict): The report to display
            language (str): Language of the report ("en" or "fr")
        """
        st.markdown("### 📊 " + ("Summary" if language == "en" else "Résumé"))
        st.write(report["summary"])
        
        st.markdown("### 📋 " + ("Recommendations" if language == "en" else "Recommandations"))
        for rec in report["recommendations"]:
            st.markdown(f"• {rec}")
        
        if report["details"]:
            st.markdown("### 📑 " + ("Details by Indicator" if language == "en" else "Détails par indicateur"))
            for indicator, detail in report["details"].items():
                st.markdown(f"**{indicator}**")
                for key, value in detail.items():
                    if isinstance(value, list):
                        st.markdown(f"- {key}:")
                        for item in value:
                            st.markdown(f"  • {item}")
                    else:
                        st.markdown(f"- {key}: {value}")
    
    def create_word_report(self, report, title, language="en"):
        """
        Creates a Word document from a report.
        
        Args:
            report (dict): The report to convert to Word
            title (str): Title for the report
            language (str): Language of the report ("en" or "fr")
            
        Returns:
            docx.Document: Word document with the report
        """
        doc = Document()
        
        # Add title
        doc.add_heading(title, level=1)
        
        # Add summary
        doc.add_heading("Summary" if language == "en" else "Résumé", level=2)
        for paragraph in report["summary"].split('\n'):
            if paragraph.strip():
                if paragraph.startswith('•'):
                    doc.add_paragraph(paragraph, style='List Bullet')
                else:
                    doc.add_paragraph(paragraph)
        
        # Add recommendations
        doc.add_heading("Recommendations" if language == "en" else "Recommandations", level=2)
        for rec in report["recommendations"]:
            if rec.startswith('•') or rec.startswith('  •'):
                # Handle nested bullets
                indent_level = rec.count('  ')
                clean_rec = rec.lstrip('• ')
                para = doc.add_paragraph(clean_rec, style='List Bullet')
                para.paragraph_format.left_indent = Pt(18 * (indent_level//2))
            else:
                doc.add_paragraph(rec)
        
        # Add details if any
        if report["details"]:
            doc.add_heading("Details" if language == "en" else "Détails", level=2)
            
            for indicator, detail in report["details"].items():
                doc.add_heading(indicator, level=3)
                
                for key, value in detail.items():
                    if isinstance(value, list):
                        # Handle lists of values
                        p = doc.add_paragraph(f"{key}:")
                        for item in value:
                            doc.add_paragraph(item, style='List Bullet')
                    else:
                        # Handle single values
                        doc.add_paragraph(f"{key}: {value}")
        
        # Add footer with date
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"{'Report generated on' if language == 'en' else 'Rapport généré le'}: {pd.Timestamp.now().strftime('%Y-%m-%d')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return doc

    def export_report_to_word(self, report, title, output_path, language="en"):
        """
        Exports a report to a Word document file.
        
        Args:
            report (dict): The report to export
            title (str): Title for the report
            output_path (str): Path where to save the Word document
            language (str): Language of the report ("en" or "fr")
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            doc = self.create_word_report(report, title, language)
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Error exporting report to Word: {str(e)}")
            return False
    
    def generate_visualization(self, data, chart_type, title, language="en"):
        """
        Creates a visualization based on report data.
        
        Args:
            data (dict or pd.DataFrame): Data for visualization
            chart_type (str): Type of chart to create ('bar', 'scatter', etc.)
            title (str): Chart title
            language (str): Language for the chart ("en" or "fr")
            
        Returns:
            plotly.graph_objs._figure.Figure: Plotly figure
        """
        # Implementation would vary based on the visualization requirements
        # This is a placeholder for extending the class with visualization capabilities
        pass


# Example usage:
# 
# # Create report generator
# report_gen = AnalysisReportGenerator()
# 
# # Generate report for zero scores analysis
# results = {
#     "zero_scores_data": {
#         "clpm": 25.5,
#         "phoneme": 15.2,
#         "comprehension": 32.1
#     },
#     "total_students": 120
# }
# 
# report = report_gen.generate_report("zero_scores", results, "en")
# 
# # Display in Streamlit
# report_gen.show_report_ui(report, "en")
# 
# # Export to Word
# report_gen.export_report_to_word(report, "Zero Scores Analysis", "zero_scores_report.docx", "en") "