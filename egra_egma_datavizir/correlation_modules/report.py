"""
Module for generating Word reports for correlation analysis.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
import pandas as pd
import plotly.figure_factory as ff
from config import egra_columns, egma_columns

def create_correlation_word_report(corr_matrix, df_strong, t, df):
    """
    Creates a Word document from correlation analysis.
    
    Args:
        corr_matrix (pandas.DataFrame): Correlation matrix
        df_strong (pandas.DataFrame): DataFrame with significant correlations
        t (dict): Translation dictionary for UI elements
        df (pandas.DataFrame): Original data
        
    Returns:
        docx.Document: Word document with the report
    """
    doc = Document()
    
    # Add title
    title = t.get("title_correlation", "Correlation Analysis")
    doc.add_heading(title, level=1)
    
    # Add summary
    add_executive_summary(doc, t)
    
    # Add introduction and interpretation guide
    add_interpretation_guide(doc, t)
    
    # Add correlation matrix section
    add_correlation_matrix_section(doc, corr_matrix, t)
    
    # Add significant correlations section
    if not df_strong.empty:
        add_significant_correlations_section(doc, df_strong, t)
    
    # Add educational interpretation section
    if not df_strong.empty:
        add_educational_interpretation_section(doc, df_strong, t)
    
    # Add no correlations section if applicable
    if df_strong.empty:
        add_no_correlations_section(doc, t)
    
    # Set up headers and footers
    setup_headers_and_footers(doc, title, t)
    
    return doc

def add_executive_summary(doc, t):
    """Add executive summary section."""
    doc.add_heading(t.get("executive_summary", "Executive Summary"), level=1)
    summary_text = t.get("correlation_summary", 
                       "This report analyzes the relationships between different assessment tasks to identify patterns and connections. " + 
                       "Understanding these relationships can help in designing integrated instructional approaches and identifying " + 
                       "skill areas that develop together or may require separate instructional focus.")
    doc.add_paragraph(summary_text)

def add_interpretation_guide(doc, t):
    """Add interpretation guide section."""
    doc.add_heading(t.get("correlation_interpretation", "Interpretation Guide"), level=2)
    doc.add_paragraph(t.get("correlation_point1", "A correlation close to 1 indicates a strong positive relationship"), style='List Bullet')
    doc.add_paragraph(t.get("correlation_point2", "A correlation close to -1 indicates a strong negative relationship"), style='List Bullet')
    doc.add_paragraph(t.get("correlation_point3", "A correlation close to 0 indicates a weak or no relationship"), style='List Bullet')
    doc.add_paragraph(t.get("correlation_point4", "Correlations > 0.5 or < -0.5 are considered significant"), style='List Bullet')

def add_correlation_matrix_section(doc, corr_matrix, t):
    """Add correlation matrix section with table."""
    doc.add_heading(t.get("correlation_matrix", "Correlation Matrix"), level=1)
    
    # Add correlation matrix table
    # Add +1 to rows and columns for headers
    table = doc.add_table(rows=len(corr_matrix.index) + 1, cols=len(corr_matrix.columns) + 1)
    table.style = 'Table Grid'
    
    # Add column headers
    header_cells = table.rows[0].cells
    header_cells[0].text = ""  # Top-left cell
    for i, col in enumerate(corr_matrix.columns):
        header_cells[i + 1].text = t["columns_of_interest"].get(col, col)
    
    # Add row headers and data
    for i, idx in enumerate(corr_matrix.index):
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = t["columns_of_interest"].get(idx, idx)
        for j, col in enumerate(corr_matrix.columns):
            value = corr_matrix.loc[idx, col]
            cell_text = f"{value:.2f}"
            row_cells[j + 1].text = cell_text

def add_significant_correlations_section(doc, df_strong, t):
    """Add significant correlations section."""
    doc.add_heading(t.get("strong_correlations", "Significant Correlations"), level=1)
    
    # Create significant correlations table
    table = doc.add_table(rows=len(df_strong) + 1, cols=3)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = t.get("task_1", "Task 1")
    header_cells[1].text = t.get("task_2", "Task 2")
    header_cells[2].text = t.get("correlation", "Correlation")
    
    # Add data rows
    for i, (_, row) in enumerate(df_strong.iterrows(), 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = t["columns_of_interest"].get(row["task1"], row["task1"])
        row_cells[1].text = t["columns_of_interest"].get(row["task2"], row["task2"])
        row_cells[2].text = f"{row['correlation']:.2f}"

def add_educational_interpretation_section(doc, df_strong, t):
    """Add educational interpretation section."""
    doc.add_heading(t.get("educational_interpretation", "Educational Interpretation"), level=1)
    
    # Group correlations by type (EGRA-EGRA, EGMA-EGMA, or EGRA-EGMA)
    egra_egra = []
    egma_egma = []
    egra_egma = []
    
    for _, row in df_strong.iterrows():
        task1, task2 = row["task1"], row["task2"]
        correlation = row["correlation"]
        
        if task1 in egra_columns and task2 in egra_columns:
            egra_egra.append((task1, task2, correlation))
        elif task1 in egma_columns and task2 in egma_columns:
            egma_egma.append((task1, task2, correlation))
        else:
            egra_egma.append((task1, task2, correlation))
    
    # Add cross-domain interpretations
    if egra_egma:
        doc.add_heading(t.get("cross_domain_title", "Reading-Math Connections"), level=2)
        doc.add_paragraph(t.get("cross_domain_intro", "These correlations indicate connections between reading and math skills:"))
        
        for task1, task2, correlation in egra_egma:
            # Determine which is EGRA and which is EGMA
            egra_task = task1 if task1 in egra_columns else task2
            egma_task = task2 if task2 in egma_columns else task1
            
            # Get translated names
            egra_name = t["columns_of_interest"].get(egra_task, egra_task)
            egma_name = t["columns_of_interest"].get(egma_task, egma_task)
            
            # Format correlation
            corr_str = f"{correlation:.2f}"
            
            # Create paragraph with interpretation
            p = doc.add_paragraph()
            p.add_run(f"{egra_name} & {egma_name} ({corr_str})").bold = True
            
            # Provide task-specific interpretations
            interpretation = get_cross_domain_interpretation(egra_task, egma_task, correlation, t)
            doc.add_paragraph(interpretation, style='List Bullet')
            
            # Instructional implication
            implication = get_instructional_implication(egra_task, egma_task, correlation, t)
            doc.add_paragraph(implication, style='List Bullet')
    
    # Add reading interpretations
    if egra_egra:
        doc.add_heading(t.get("reading_correlations_title", "Within Reading Skills"), level=2)
        doc.add_paragraph(t.get("reading_correlations_intro", "These correlations show relationships between different reading skills:"))
        
        for task1, task2, correlation in egra_egra:
            # Get translated names
            task1_name = t["columns_of_interest"].get(task1, task1)
            task2_name = t["columns_of_interest"].get(task2, task2)
            
            # Format correlation
            corr_str = f"{correlation:.2f}"
            
            # Display correlation with interpretation
            p = doc.add_paragraph()
            p.add_run(f"{task1_name} & {task2_name} ({corr_str})").bold = True
            
            # Provide reading-specific interpretation
            interpretation = get_reading_interpretation(task1, task2, correlation, t)
            doc.add_paragraph(interpretation, style='List Bullet')
    
    # Add math interpretations
    if egma_egma:
        doc.add_heading(t.get("math_correlations_title", "Within Math Skills"), level=2)
        doc.add_paragraph(t.get("math_correlations_intro", "These correlations indicate relationships between different math skills:"))
        
        for task1, task2, correlation in egma_egma:
            # Get translated names
            task1_name = t["columns_of_interest"].get(task1, task1)
            task2_name = t["columns_of_interest"].get(task2, task2)
            
            # Format correlation
            corr_str = f"{correlation:.2f}"
            
            # Display correlation with interpretation
            p = doc.add_paragraph()
            p.add_run(f"{task1_name} & {task2_name} ({corr_str})").bold = True
            
            # Provide math-specific interpretation
            interpretation = get_math_interpretation(task1, task2, correlation, t)
            doc.add_paragraph(interpretation, style='List Bullet')
    
    # Add instructional recommendations
    doc.add_heading(t.get("instructional_recommendations", "Instructional Recommendations"), level=2)
    
    # Add recommendations based on correlation patterns
    if egra_egma:
        p = doc.add_paragraph()
        p.add_run(t.get("integrated_instruction", "Integrated Instruction")).bold = True
        doc.add_paragraph(t.get("integrated_instruction_text", "Consider integrated lessons that combine reading and math skills that show strong correlations"), style='List Bullet')
        doc.add_paragraph(t.get("integrated_instruction_example", "For example, use story problems that incorporate both mathematical thinking and reading comprehension"), style='List Bullet')
    
    if egra_egra:
        p = doc.add_paragraph()
        p.add_run(t.get("reading_instruction", "Reading Instruction")).bold = True
        if any(correlation < 0 for _, _, correlation in egra_egra):
            doc.add_paragraph(t.get("reading_negative_text", "Pay attention to negative correlations that may indicate competing skills or resource allocation"), style='List Bullet')
        doc.add_paragraph(t.get("reading_positive_text", "Build on strong positive correlations by teaching related skills together"), style='List Bullet')
    
    if egma_egma:
        p = doc.add_paragraph()
        p.add_run(t.get("math_instruction", "Math Instruction")).bold = True
        if any(correlation < 0 for _, _, correlation in egma_egma):
            doc.add_paragraph(t.get("math_negative_text", "Address negative correlations by ensuring students understand connections between skills"), style='List Bullet')
        doc.add_paragraph(t.get("math_positive_text", "Leverage strong positive correlations to build mathematical understanding"), style='List Bullet')

def add_no_correlations_section(doc, t):
    """Add section for when no significant correlations are found."""
    doc.add_heading(t.get("no_significant_correlations", "No Significant Correlations"), level=1)
    doc.add_paragraph(t.get("no_strong_correlation", "No significant correlations (>|0.5|) were found."))
    doc.add_paragraph(t.get("weak_correlation_note", "This suggests that the measured skills may be developing independently or that the assessment measures distinct constructs."))

def setup_headers_and_footers(doc, title, t):
    """Set up headers and footers for the document."""
    # Get sections
    sections = doc.sections
    
    # Setup for each section
    for section in sections:
        # Header
        header = section.header
        header_paragraph = header.paragraphs[0]
        header_text = title
        header_paragraph.text = header_text
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Format header text
        for run in header_paragraph.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
        
        # Footer
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        
        # Add page numbers and date
        from datetime import datetime
        date_str = datetime.now().strftime('%Y-%m-%d')
        footer_text = f"{t.get('report_date', 'Report generated on')}: {date_str}"
        footer_paragraph.text = footer_text
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Format footer text
        for run in footer_paragraph.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)  # Gray

def get_cross_domain_interpretation(egra_task, egma_task, correlation, t):
    """
    Provides interpretation for cross-domain correlations.
    
    Args:
        egra_task (str): EGRA task code
        egma_task (str): EGMA task code
        correlation (float): Correlation value
        t (dict): Translation dictionary
        
    Returns:
        str: Interpretation text
    """
    # Define mapping of task pairs to interpretations
    cross_domain_interpretations = {
        # Letter identification and number identification
        ("clpm", "number_id"): {
            "positive": t.get("clpm_number_id_pos", "Symbol recognition is a common cognitive skill underlying both letter and number identification."),
            "negative": t.get("clpm_number_id_neg", "Students may be showing a preference or strength in one symbol system over the other.")
        },
        # Reading fluency and math fluency
        ("cwpm", "addition"): {
            "positive": t.get("cwpm_addition_pos", "Fluency in both domains suggests strong procedural knowledge and automaticity."),
            "negative": t.get("cwpm_addition_neg", "Students might be experiencing trade-offs in developing fluency in different domains.")
        },
        ("cwpm", "subtraction"): {
            "positive": t.get("cwpm_subtraction_pos", "Processing speed may be contributing to both reading and calculation fluency."),
            "negative": t.get("cwpm_subtraction_neg", "Different cognitive processes might be dominant in each domain.")
        },
        # Comprehension and problem solving
        ("comprehension", "problems"): {
            "positive": t.get("comprehension_problems_pos", "Reading comprehension and math problem solving both require similar analytical thinking skills."),
            "negative": t.get("comprehension_problems_neg", "Students may be developing different approaches to comprehension in different domains.")
        }
    }
    
    # Check for specific task pair interpretation
    for (task1, task2), interpretations in cross_domain_interpretations.items():
        if (egra_task == task1 and egma_task == task2) or (egra_task == task2 and egma_task == task1):
            if correlation > 0:
                return interpretations["positive"]
            else:
                return interpretations["negative"]
    
    # Default interpretations based on sign
    if correlation > 0:
        return t.get("cross_domain_positive", "This positive correlation suggests that these reading and math skills share underlying cognitive processes or instructional influences.")
    else:
        return t.get("cross_domain_negative", "This negative correlation might indicate competing cognitive resources or different developmental trajectories for these skills.")

def get_instructional_implication(egra_task, egma_task, correlation, t):
    """
    Provides instructional implications for cross-domain correlations.
    
    Args:
        egra_task (str): EGRA task code
        egma_task (str): EGMA task code
        correlation (float): Correlation value
        t (dict): Translation dictionary
        
    Returns:
        str: Instructional implication text
    """
    # Define mapping of task pairs to instructional implications
    instructional_implications = {
        # Letter identification and number identification
        ("clpm", "number_id"): {
            "positive": t.get("clpm_number_id_implication", "Consider using similar visual recognition strategies for both letters and numbers."),
            "negative": t.get("clpm_number_id_neg_implication", "Analyze different approaches needed for letter vs. number symbol recognition.")
        },
        # Reading fluency and math fluency
        ("cwpm", "addition"): {
            "positive": t.get("cwpm_addition_implication", "Use similar fluency-building techniques across reading and math practice."),
            "negative": t.get("cwpm_addition_neg_implication", "Consider separate instructional approaches for building fluency in reading vs. math.")
        },
        ("cwpm", "subtraction"): {
            "positive": t.get("cwpm_subtraction_implication", "Focus on building automaticity in both domains through regular practice."),
            "negative": t.get("cwpm_subtraction_neg_implication", "Analyze which domain needs more foundational support.")
        },
        # Comprehension and problem solving
        ("comprehension", "problems"): {
            "positive": t.get("comprehension_problems_implication", "Teach similar analytical strategies for both reading comprehension and math problem solving."),
            "negative": t.get("comprehension_problems_neg_implication", "Investigate domain-specific challenges students face in comprehension tasks.")
        }
    }
    
    # Check for specific task pair implication
    for (task1, task2), implications in instructional_implications.items():
        if (egra_task == task1 and egma_task == task2) or (egra_task == task2 and egma_task == task1):
            if correlation > 0:
                return implications["positive"]
            else:
                return implications["negative"]
    
    # Default implications based on sign and general categories
    # Decoding tasks
    if egra_task in ["clpm", "phoneme", "sound_word"]:
        if egma_task in ["number_id", "discrimin"]:
            if correlation > 0:
                return t.get("symbol_processing_implication", "Focus on symbol recognition activities that benefit both letter and number processing.")
            else:
                return t.get("symbol_processing_neg_implication", "Consider different approaches for teaching symbol recognition in each domain.")
    
    # Fluency tasks
    if egra_task in ["cwpm", "orf"]:
        if egma_task in ["addition", "subtraction"]:
            if correlation > 0:
                return t.get("fluency_implication", "Implement timed practice activities for both reading passages and math facts.")
            else:
                return t.get("fluency_neg_implication", "Analyze whether fluency barriers differ between reading and math.")
    
    # Comprehension tasks
    if egra_task in ["comprehension", "listening"]:
        if egma_task in ["problems"]:
            if correlation > 0:
                return t.get("comprehension_implication", "Teach explicit strategies for understanding both text and word problems.")
            else:
                return t.get("comprehension_neg_implication", "Identify domain-specific comprehension challenges.")
    
    # Default generic implication
    if correlation > 0:
        return t.get("general_positive_implication", "Consider instructional activities that develop both skills together.")
    else:
        return t.get("general_negative_implication", "Analyze whether instructional approaches need to be differentiated for each skill.")

def get_reading_interpretation(task1, task2, correlation, t):
    """
    Provides interpretations for reading skill correlations.
    
    Args:
        task1 (str): First EGRA task code
        task2 (str): Second EGRA task code
        correlation (float): Correlation value
        t (dict): Translation dictionary
        
    Returns:
        str: Interpretation text
    """
    # Define reading skill interpretations by task pairs
    reading_interpretations = {
        # Letter recognition and phoneme awareness
        ("clpm", "phoneme"): t.get("clpm_phoneme", "Letter recognition and phonemic awareness are foundational alphabetic skills that typically develop together."),
        
        # Letter recognition and word reading
        ("clpm", "sound_word"): t.get("clpm_sound_word", "Letter recognition directly supports word reading through grapheme-phoneme mapping."),
        
        # Phoneme awareness and word reading
        ("phoneme", "sound_word"): t.get("phoneme_sound_word", "Phonemic awareness is essential for decoding and sounding out words."),
        
        # Word reading and fluency
        ("sound_word", "cwpm"): t.get("sound_word_cwpm", "Word reading accuracy contributes to reading fluency development."),
        
        # Fluency and comprehension
        ("cwpm", "comprehension"): t.get("cwpm_comprehension", "Reading fluency supports comprehension by freeing cognitive resources for understanding."),
        ("orf", "comprehension"): t.get("orf_comprehension", "Oral reading fluency and comprehension are closely linked reading skills."),
        
        # Listening and reading comprehension
        ("listening", "comprehension"): t.get("listening_comprehension", "Listening comprehension and reading comprehension both rely on language understanding skills.")
    }
    
    # Sort tasks alphabetically to match dictionary keys
    sorted_pair = tuple(sorted([task1, task2]))
    
    # Look for specific interpretation
    if sorted_pair in reading_interpretations:
        if correlation > 0:
            return reading_interpretations[sorted_pair]
        else:
            return t.get("reading_unexpected_negative", "This negative correlation is unexpected and may indicate measurement issues or competing instructional focus.")
    
    # Default interpretation based on reading development theory
    # Check if tasks are part of the same reading component
    decoding_tasks = ["clpm", "phoneme", "sound_word"]
    fluency_tasks = ["cwpm", "orf"]
    comprehension_tasks = ["listening", "comprehension"]
    
    if (task1 in decoding_tasks and task2 in decoding_tasks):
        return t.get("within_decoding", "These skills both contribute to early decoding development and typically develop together.")
    
    elif (task1 in fluency_tasks and task2 in fluency_tasks):
        return t.get("within_fluency", "These fluency measures reflect related aspects of reading rate, accuracy, and prosody.")
    
    elif (task1 in comprehension_tasks and task2 in comprehension_tasks):
        return t.get("within_comprehension", "These skills both measure aspects of language and text understanding.")
    
    # Check for connections across reading components
    elif (task1 in decoding_tasks and task2 in fluency_tasks) or (task1 in fluency_tasks and task2 in decoding_tasks):
        return t.get("decoding_fluency", "Decoding skills are foundational to fluency development.")
    
    elif (task1 in fluency_tasks and task2 in comprehension_tasks) or (task1 in comprehension_tasks and task2 in fluency_tasks):
        return t.get("fluency_comprehension", "Fluency supports comprehension by reducing cognitive load during reading.")
    
    elif (task1 in decoding_tasks and task2 in comprehension_tasks) or (task1 in comprehension_tasks and task2 in decoding_tasks):
        return t.get("decoding_comprehension", "Basic decoding skills are necessary but not sufficient for comprehension.")
    
    # Default
    return t.get("general_reading", "These reading skills show a relationship that suggests shared cognitive or instructional factors.")

def get_math_interpretation(task1, task2, correlation, t):
    """
    Provides interpretations for math skill correlations.
    
    Args:
        task1 (str): First EGMA task code
        task2 (str): Second EGMA task code
        correlation (float): Correlation value
        t (dict): Translation dictionary
        
    Returns:
        str: Interpretation text
    """
    # Define math skill interpretations by task pairs
    math_interpretations = {
        # Number identification and discrimination
        ("number_id", "discrimin"): t.get("number_id_discrimin", "Number identification and magnitude comparison both involve number sense."),
        
        # Number identification and missing number
        ("number_id", "missing_number"): t.get("number_id_missing", "Number identification supports sequential number pattern recognition."),
        
        # Number discrimination and missing number
        ("discrimin", "missing_number"): t.get("discrimin_missing", "Number magnitude comparison and sequence skills both involve number relationships."),
        
        # Basic operations
        ("addition", "subtraction"): t.get("addition_subtraction", "Addition and subtraction are related operations that often develop together."),
        
        # Operations and problem solving
        ("addition", "problems"): t.get("addition_problems", "Addition skills contribute to problem-solving ability."),
        ("subtraction", "problems"): t.get("subtraction_problems", "Subtraction skills are important for problem-solving applications.")
    }
    
    # Sort tasks alphabetically to match dictionary keys
    sorted_pair = tuple(sorted([task1, task2]))
    
    # Look for specific interpretation
    if sorted_pair in math_interpretations:
        if correlation > 0:
            return math_interpretations[sorted_pair]
        else:
            return t.get("math_unexpected_negative", "This negative correlation is unexpected and may indicate measurement issues or competing instructional focus.")
    
    # Default interpretation based on math domain relationships
    # Group math tasks by conceptual domain
    number_sense_tasks = ["number_id", "discrimin", "missing_number"]
    operations_tasks = ["addition", "subtraction"]
    application_tasks = ["problems"]
    
    if (task1 in number_sense_tasks and task2 in number_sense_tasks):
        return t.get("within_number_sense", "These skills both contribute to foundational number sense development.")
    
    elif (task1 in operations_tasks and task2 in operations_tasks):
        return t.get("within_operations", "These operation skills are conceptually related and rely on similar numerical understanding.")
    
    elif (task1 in number_sense_tasks and task2 in operations_tasks) or (task1 in operations_tasks and task2 in number_sense_tasks):
        return t.get("number_sense_operations", "Number sense provides an important foundation for operational understanding.")
    
    elif (task1 in operations_tasks and task2 in application_tasks) or (task1 in application_tasks and task2 in operations_tasks):
        return t.get("operations_applications", "Operational fluency supports problem-solving applications.")
    
    elif (task1 in number_sense_tasks and task2 in application_tasks) or (task1 in application_tasks and task2 in number_sense_tasks):
        return t.get("number_sense_applications", "Number sense contributes to problem-solving capacity.")
    
    # Default
    return t.get("general_math", "These math skills show a relationship that suggests shared cognitive or instructional factors.")
