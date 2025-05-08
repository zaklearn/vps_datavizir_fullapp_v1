import streamlit as st
import pandas as pd
import plotly.express as px
from docx import Document
from docx.shared import Inches
import tempfile
import os
from config import translations  # Import translation dictionary

def show_statistical_overview(df, language):
    """
    Displays descriptive statistics and visualizations for selected variables.
    
    Args:
        df (pandas.DataFrame): The data to analyze
        language (str): Selected language for UI elements (en/fr)
    """
    t = translations[language]  # Get translations for selected language
    
    # Define EGRA and EGMA column lists for organization
    egra_columns = ["clpm", "phoneme", "sound_word", "cwpm", "listening", "orf", "comprehension"]
    egma_columns = ["number_id", "discrimin", "missing_number", "addition", "subtraction", "problems"]
    
    # Ensure the DataFrame contains at least some of the required columns
    available_columns = [col for col in egra_columns + egma_columns if col in df.columns]
    
    if not available_columns:
        st.error(t.get("no_assessment_columns", "No assessment columns found in the data."))
        return
    
    st.subheader(t.get("select_variables", "📊 Select Variables"))
    
    # Create two columns to organize variable selection
    col1, col2 = st.columns(2)
    
    with col1:
        # Filter available EGRA columns
        available_egra = [col for col in egra_columns if col in df.columns]
        selected_egra = st.multiselect(
            t.get("egra_variables", "EGRA Variables:"),
            options=available_egra,
            default=available_egra[:3] if len(available_egra) > 3 else available_egra,
            format_func=lambda x: t["columns_of_interest"].get(x, x)
        )
    
    with col2:
        # Filter available EGMA columns
        available_egma = [col for col in egma_columns if col in df.columns]
        selected_egma = st.multiselect(
            t.get("egma_variables", "EGMA Variables:"),
            options=available_egma,
            default=available_egma[:3] if len(available_egma) > 3 else available_egma,
            format_func=lambda x: t["columns_of_interest"].get(x, x)
        )
    
    # Combine selected variables
    selected_columns = selected_egra + selected_egma
    
    if selected_columns:
        try:
            # Extract the selected columns
            df_filtered = df[selected_columns].copy()
            
            # Calculate descriptive statistics for selected columns
            st.subheader(t.get("table_statistics", "📋 Descriptive Statistics"))
            stats_summary = df_filtered.describe(percentiles=[.25, .5, .75, .9]).round(2)
            st.dataframe(stats_summary)
            
            # Export options
            col1, col2 = st.columns(2)
            
            # CSV Export
            with col1:
                csv = stats_summary.to_csv(index=True).encode('utf-8-sig')
                st.download_button(
                    t.get("export_csv", "📥 Download CSV"),
                    csv,
                    "descriptive_statistics.csv",
                    "text/csv",
                    key='download-csv'
                )
            
            # Word Export
            with col2:
                if st.button(t.get("export_word", "📄 Export to Word")):
                    try:
                        doc = create_word_report(df_filtered, stats_summary, selected_columns, t)
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                            doc.save(tmp.name)
                            with open(tmp.name, 'rb') as f:
                                docx = f.read()
                            st.download_button(
                                t.get("download_word", "📥 Download Word Report"),
                                docx,
                                "descriptive_statistics.docx",
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        os.unlink(tmp.name)
                    except Exception as e:
                        st.error(f"Error creating Word report: {str(e)}")
            
            # Detailed display by indicator
            st.subheader(t.get("distribution_scores", "📊 Distribution of Scores"))
            
            for column in selected_columns:
                # Create columns for each indicator
                col1, col2 = st.columns(2)
                
                with col1:
                    # Indicator statistics table
                    st.write(f"**{t['columns_of_interest'].get(column, column)}**")
                    indicator_stats = df_filtered[column].describe(percentiles=[.25, .5, .75, .9]).round(2)
                    st.dataframe(indicator_stats)
                
                with col2:
                    # Generate histogram with box plot
                    try:
                        fig = px.histogram(
                            df_filtered,
                            x=column,
                            nbins=20,
                            marginal="box",
                            opacity=0.7,
                            title=t.get("histogram_title", "Distribution of {}").format(
                                t["columns_of_interest"].get(column, column)
                            ),
                            color_discrete_sequence=px.colors.sequential.Viridis
                        )
                        fig.update_layout(showlegend=False)
                        st.plotly_chart(fig, use_container_width=True)
                    except Exception as e:
                        st.error(f"Error creating visualization for {column}: {str(e)}")
                
                st.markdown("---")  # Separator between indicators
        
        except Exception as e:
            st.error(f"Error in statistical analysis: {str(e)}")
    
    else:
        st.warning(t.get("warning_select_variable", "Please select at least one variable to analyze."))

def create_word_report(df_filtered, stats_summary, selected_columns, t):
    """
    Creates a Word report with statistics and graphs.
    
    Args:
        df_filtered (pandas.DataFrame): Filtered data for selected columns
        stats_summary (pandas.DataFrame): Summary statistics
        selected_columns (list): List of selected column names
        t (dict): Translation dictionary
        
    Returns:
        docx.Document: Word document with the report
    """
    doc = Document()
    
    # Main title
    doc.add_heading(t.get("title_statistics", "Descriptive Statistics Analysis"), level=1)
    
    # Global statistics table
    doc.add_heading(t.get("table_statistics", "Descriptive Statistics Table"), level=2)
    
    # Create table with statistics
    # Add +1 to columns for row headers
    table = doc.add_table(rows=len(stats_summary.index) + 1, cols=len(stats_summary.columns) + 1)
    table.style = 'Table Grid'
    
    # Add column headers
    header_cells = table.rows[0].cells
    header_cells[0].text = "Statistic"
    for idx, col in enumerate(stats_summary.columns):
        header_cells[idx + 1].text = t["columns_of_interest"].get(col, col)
    
    # Add row data
    for i, stat_name in enumerate(stats_summary.index):
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = str(stat_name)
        for j, col in enumerate(stats_summary.columns):
            row_cells[j + 1].text = str(stats_summary.loc[stat_name, col])
    
    doc.add_paragraph("")  # Space
    
    # Section for each indicator
    doc.add_heading(t.get("distribution_scores", "Distribution of Scores"), level=2)
    
    # Create temporary directory for images
    with tempfile.TemporaryDirectory() as tmp_dir:
        for col in selected_columns:
            # Indicator title
            doc.add_heading(t["columns_of_interest"].get(col, col), level=3)
            
            # Specific indicator statistics
            indicator_stats = df_filtered[col].describe(percentiles=[.25, .5, .75, .9]).round(2)
            
            # Create table for this indicator
            table = doc.add_table(rows=len(indicator_stats.index) + 1, cols=2)
            table.style = 'Table Grid'
            
            # Headers
            header_cells = table.rows[0].cells
            header_cells[0].text = "Statistic"
            header_cells[1].text = "Value"
            
            # Fill statistics
            for i, stat_name in enumerate(indicator_stats.index):
                row_cells = table.rows[i + 1].cells
                row_cells[0].text = str(stat_name)
                row_cells[1].text = str(indicator_stats.loc[stat_name])
            
            doc.add_paragraph("")  # Space
            
            # Create and add histogram
            try:
                fig = px.histogram(
                    df_filtered,
                    x=col,
                    nbins=20,
                    marginal="box",
                    opacity=0.7,
                    title=t.get("histogram_title", "Distribution of {}").format(
                        t["columns_of_interest"].get(col, col)
                    ),
                    color_discrete_sequence=px.colors.sequential.Viridis
                )
                
                img_path = os.path.join(tmp_dir, f"{col}_histogram.png")
                fig.write_image(img_path, width=1000, height=600)
                
                # Add image to document
                doc.add_picture(img_path, width=Inches(6))
                doc.add_paragraph("")  # Space after image
            except Exception as e:
                # Add error message if visualization fails
                doc.add_paragraph(f"Error generating visualization: {str(e)}")
    
    return doc