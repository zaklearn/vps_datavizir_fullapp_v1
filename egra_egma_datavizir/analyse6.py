import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
from config import translations, egra_columns, egma_columns

def cronbach_alpha(data):
    """
    Calculates Cronbach's Alpha reliability coefficient.
    
    Args:
        data (pandas.DataFrame): DataFrame with items/variables as columns
        
    Returns:
        float or None: Cronbach's Alpha coefficient or None if calculation is not possible
    """
    # Drop rows with any NaN values
    items = data.dropna(axis=0)
    
    # Need at least 2 items and 3 respondents for meaningful calculation
    if items.shape[1] < 2 or items.shape[0] < 3:
        return None
    
    # Calculate item variances and total variance
    item_variances = items.var(axis=0, ddof=1)
    total_variance = items.sum(axis=1).var(ddof=1)
    
    # Avoid division by zero
    if total_variance == 0:
        return None
    
    # Calculate Cronbach's Alpha
    n_items = items.shape[1]
    alpha = (n_items / (n_items - 1)) * (1 - item_variances.sum() / total_variance)
    
    return alpha

def interpret_alpha(alpha):
    """
    Interprets the value of Cronbach's Alpha coefficient.
    
    Args:
        alpha (float or None): Cronbach's Alpha coefficient
        
    Returns:
        tuple: (interpretation category, description, color code)
    """
    if alpha is None:
        return ("insufficient_data", "Insufficient data", "#CCCCCC")
    elif alpha >= 0.9:
        return ("excellent", "Excellent reliability", "#1E8449")  # Dark green
    elif alpha >= 0.8:
        return ("good", "Good reliability", "#58D68D")  # Green
    elif alpha >= 0.7:
        return ("acceptable", "Acceptable reliability", "#F4D03F")  # Yellow
    elif alpha >= 0.6:
        return ("questionable", "Questionable reliability", "#F39C12")  # Orange
    elif alpha >= 0.5:
        return ("poor", "Poor reliability", "#E74C3C")  # Red
    else:
        return ("unacceptable", "Unacceptable reliability", "#922B21")  # Dark red

def show_reliability(df, language):
    """
    Analyzes and displays test reliability using Cronbach's Alpha.
    
    Args:
        df (pandas.DataFrame): The data to analyze
        language (str): Selected language for UI elements
    """
    t = translations[language]  # Get translations for selected language
    
    st.markdown(f"""
    ### {t.get("title_cronbach", "Test Reliability Analysis (Cronbach's Alpha)")}
    
    🔍 **{t.get("cronbach_intro", "Objective: Evaluate the internal consistency of EGRA and EGMA assessments.")}**
    
    📌 **{t.get("cronbach_interpretation", "Interpretation Guide:")}**
    - {t.get("alpha_excellent", "α ≥ 0.9 → Excellent reliability")}
    - {t.get("alpha_good", "0.8 ≤ α < 0.9 → Good reliability")}
    - {t.get("alpha_acceptable", "0.7 ≤ α < 0.8 → Acceptable reliability")}
    - {t.get("alpha_questionable", "0.6 ≤ α < 0.7 → Questionable reliability")}
    - {t.get("alpha_poor", "0.5 ≤ α < 0.6 → Poor reliability")}
    - {t.get("alpha_unacceptable", "α < 0.5 → Unacceptable reliability")}
    """)
    
    # Check if the required columns are present
    available_egra = [col for col in egra_columns if col in df.columns]
    available_egma = [col for col in egma_columns if col in df.columns]
    
    if not available_egra and not available_egma:
        st.error(t.get("no_assessment_columns", "No EGRA or EGMA assessment columns found in the data."))
        return
    
    try:
        # Check if language_teaching column exists for separating English and Dutch
        has_language_column = "language_teaching" in df.columns
        
        # Calculate Cronbach's Alpha for different assessment groups
        alpha_results = []
        
        # EGRA by language (if language column exists)
        if has_language_column and available_egra:
            # Filter by language
            for language_code, language_name in [("English", t.get("english", "English")), 
                                                 ("Dutch", t.get("dutch", "Dutch"))]:
                df_language = df[df["language_teaching"] == language_code]
                
                if not df_language.empty:
                    # Calculate alpha for this language group
                    alpha_value = cronbach_alpha(df_language[available_egra])
                    interpretation, description, color = interpret_alpha(alpha_value)
                    
                    alpha_results.append({
                        "test_group": f"EGRA ({language_name})",
                        "n_items": len(available_egra),
                        "n_students": len(df_language),
                        "alpha": alpha_value,
                        "interpretation": t.get(interpretation, description),
                        "color": color
                    })
        
        # EGRA (all students) if no language column or as additional info
        if available_egra:
            alpha_value = cronbach_alpha(df[available_egra])
            interpretation, description, color = interpret_alpha(alpha_value)
            
            alpha_results.append({
                "test_group": t.get("egra_all", "EGRA (All Students)"),
                "n_items": len(available_egra),
                "n_students": len(df),
                "alpha": alpha_value,
                "interpretation": t.get(interpretation, description),
                "color": color
            })
        
        # EGMA (all students)
        if available_egma:
            alpha_value = cronbach_alpha(df[available_egma])
            interpretation, description, color = interpret_alpha(alpha_value)
            
            alpha_results.append({
                "test_group": t.get("egma_all", "EGMA (All Students)"),
                "n_items": len(available_egma),
                "n_students": len(df),
                "alpha": alpha_value,
                "interpretation": t.get(interpretation, description),
                "color": color
            })
        
        # Convert results to DataFrame for display
        df_results = pd.DataFrame(alpha_results)
        
        # Display results table
        st.subheader(t.get("cronbach_results", "📊 Cronbach's Alpha Results"))
        
        # Format the table for display
        formatted_results = df_results.copy()
        formatted_results.columns = [
            t.get("test_group", "Test Group"),
            t.get("n_items", "Number of Items"),
            t.get("n_students", "Number of Students"),
            t.get("cronbach_alpha", "Cronbach's Alpha"),
            t.get("reliability", "Reliability"),
            "color"  # Hidden column for visualization
        ]
        
        # Format the Alpha value to 3 decimal places or "N/A"
        formatted_results[t.get("cronbach_alpha", "Cronbach's Alpha")] = formatted_results[
            t.get("cronbach_alpha", "Cronbach's Alpha")
        ].apply(lambda x: f"{x:.3f}" if x is not None else t.get("not_available", "N/A"))
        
        # Display the table without the color column
        st.dataframe(formatted_results.drop(columns=["color"]))
        
        # Visualization of results
        st.subheader(t.get("reliability_visualization", "📈 Reliability Visualization"))
        
        # Create a bar chart of alpha values
        chart_data = formatted_results.copy()
        
        # Create a numeric version of alpha for the chart
        chart_data["alpha_numeric"] = df_results["alpha"].apply(
            lambda x: x if x is not None else 0
        )
        
        # Add a column for the color scale
        chart_data["reliability_level"] = df_results["alpha"].apply(
            lambda x: interpret_alpha(x)[0] if x is not None else "insufficient_data"
        )
        
        fig = px.bar(
            chart_data,
            x=t.get("test_group", "Test Group"),
            y="alpha_numeric",
            color="reliability_level",
            labels={
                "alpha_numeric": t.get("cronbach_alpha", "Cronbach's Alpha"),
                "reliability_level": t.get("reliability_level", "Reliability Level")
            },
            color_discrete_map={
                "excellent": "#1E8449",
                "good": "#58D68D",
                "acceptable": "#F4D03F",
                "questionable": "#F39C12",
                "poor": "#E74C3C",
                "unacceptable": "#922B21",
                "insufficient_data": "#CCCCCC"
            },
            title=t.get("reliability_comparison", "Reliability Comparison")
        )
        
        # Add threshold lines
        fig.add_hline(y=0.9, line_dash="dash", line_color="#1E8449", annotation_text=t.get("excellent_threshold", "Excellent (0.9)"))
        fig.add_hline(y=0.8, line_dash="dash", line_color="#58D68D", annotation_text=t.get("good_threshold", "Good (0.8)"))
        fig.add_hline(y=0.7, line_dash="dash", line_color="#F4D03F", annotation_text=t.get("acceptable_threshold", "Acceptable (0.7)"))
        fig.add_hline(y=0.6, line_dash="dash", line_color="#F39C12", annotation_text=t.get("questionable_threshold", "Questionable (0.6)"))
        fig.add_hline(y=0.5, line_dash="dash", line_color="#E74C3C", annotation_text=t.get("poor_threshold", "Poor (0.5)"))
        
        # Update layout
        fig.update_layout(
            yaxis_range=[0, 1],
            xaxis_title=t.get("test_group", "Test Group"),
            yaxis_title=t.get("cronbach_alpha", "Cronbach's Alpha"),
            height=500
        )
        
        st.plotly_chart(fig, use_container_width=True)
        
        # Export options
        col1, col2 = st.columns(2)
        
        # CSV Export
        with col1:
            csv = formatted_results.drop(columns=["color"]).to_csv(index=False).encode('utf-8-sig')
            st.download_button(
                t.get("export_cronbach_csv", "📥 Download CSV"),
                csv,
                "cronbach_alpha_results.csv",
                "text/csv",
                key='download-cronbach-csv'
            )
        
        # Word Export
        with col2:
            if st.button(t.get("export_cronbach_word", "📄 Export to Word")):
                try:
                    doc = create_reliability_word_report(df_results, fig, t)
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                        doc.save(tmp.name)
                        with open(tmp.name, 'rb') as f:
                            docx = f.read()
                        st.download_button(
                            t.get("download_cronbach_word", "📥 Download Word Report"),
                            docx,
                            "reliability_analysis.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    os.unlink(tmp.name)
                except Exception as e:
                    st.error(f"Error creating Word report: {str(e)}")
        
        # Educational Interpretation Section
        st.subheader(t.get("educational_interpretation", "🔍 Educational Interpretation"))
        
        # Determine overall assessment quality based on alpha results
        egra_alphas = [result["alpha"] for result in alpha_results if "EGRA" in result["test_group"]]
        egma_alphas = [result["alpha"] for result in alpha_results if "EGMA" in result["test_group"]]
        
        # Filter out None values
        valid_egra_alphas = [alpha for alpha in egra_alphas if alpha is not None]
        valid_egma_alphas = [alpha for alpha in egma_alphas if alpha is not None]
        
        # Calculate average alphas if possible
        avg_egra_alpha = sum(valid_egra_alphas) / len(valid_egra_alphas) if valid_egra_alphas else None
        avg_egma_alpha = sum(valid_egma_alphas) / len(valid_egma_alphas) if valid_egma_alphas else None
        
        # Provide interpretation for EGRA
        if valid_egra_alphas:
            st.markdown(f"**{t.get('egra_reliability', 'EGRA Reliability Interpretation')}:**")
            
            if avg_egra_alpha >= 0.7:
                st.markdown(t.get('egra_reliability_good', """
                The EGRA assessment demonstrates acceptable to good reliability. This indicates:
                - The test items consistently measure the same construct
                - Results can be considered dependable for instructional decisions
                - Different items within the test work together coherently
                """))
            elif avg_egra_alpha >= 0.6:
                st.markdown(t.get('egra_reliability_questionable', """
                The EGRA assessment shows questionable reliability. This suggests:
                - There may be some inconsistency in what the test items measure
                - Results should be interpreted with caution
                - Consider using additional assessments to confirm findings
                - Test revision may be beneficial in future implementations
                """))
            else:
                st.markdown(t.get('egra_reliability_poor', """
                The EGRA assessment demonstrates poor reliability. This indicates:
                - Significant inconsistency in what the test items measure
                - Results should not be the sole basis for important decisions
                - The test likely needs revision or replacement
                - Individual items may be measuring different constructs
                """))
        
        # Provide interpretation for EGMA
        if valid_egma_alphas:
            st.markdown(f"**{t.get('egma_reliability', 'EGMA Reliability Interpretation')}:**")
            
            if avg_egma_alpha >= 0.7:
                st.markdown(t.get('egma_reliability_good', """
                The EGMA assessment demonstrates acceptable to good reliability. This indicates:
                - The test items consistently measure the same mathematical constructs
                - Results can be considered dependable for instructional decisions
                - Different items within the test work together coherently
                """))
            elif avg_egma_alpha >= 0.6:
                st.markdown(t.get('egma_reliability_questionable', """
                The EGMA assessment shows questionable reliability. This suggests:
                - There may be some inconsistency in what the test items measure
                - Results should be interpreted with caution
                - Consider using additional assessments to confirm findings
                - Test revision may be beneficial in future implementations
                """))
            else:
                st.markdown(t.get('egma_reliability_poor', """
                The EGMA assessment demonstrates poor reliability. This indicates:
                - Significant inconsistency in what the test items measure
                - Results should not be the sole basis for important decisions
                - The test likely needs revision or replacement
                - Individual items may be measuring different constructs
                """))
        
        # Provide recommendations based on results
        st.markdown(f"**{t.get('reliability_recommendations', 'Recommendations Based on Reliability Results')}:**")
        
        all_alphas = valid_egra_alphas + valid_egma_alphas
        
        if not all_alphas:
            st.markdown(t.get('reliability_no_data', """
            - Insufficient data to calculate reliability
            - Consider collecting more complete data in future assessments
            - Ensure proper test administration procedures are followed
            """))
        elif all(alpha >= 0.7 for alpha in all_alphas):
            st.markdown(t.get('reliability_all_good', """
            - Continue using the current assessment tools with confidence
            - Results can be used for instructional planning and student evaluation
            - Consider conducting validity studies to further strengthen the assessment system
            - Monitor reliability in future administrations to ensure consistency
            """))
        elif any(alpha < 0.6 for alpha in all_alphas):
            st.markdown(t.get('reliability_some_poor', """
            - Review and potentially revise assessment items with low reliability
            - Use multiple measures when making important educational decisions
            - Provide additional training to test administrators to ensure consistent procedures
            - Consider item analysis to identify specific problematic questions
            """))
        else:
            st.markdown(t.get('reliability_mixed', """
            - Use results cautiously, especially for tests with lower reliability
            - Consider strengthening assessments with questionable reliability
            - Continue using assessments with good reliability
            - Supplement assessment data with teacher observations and other measures
            """))
        
    except Exception as e:
        st.error(f"Error calculating reliability: {str(e)}")

def create_reliability_word_report(df_results, fig, t):
    """
    Creates a Word report with reliability analysis results.
    
    Args:
        df_results (pandas.DataFrame): DataFrame with reliability results
        fig (plotly.graph_objs._figure.Figure): Reliability visualization figure
        t (dict): Translation dictionary
        
    Returns:
        docx.Document: Word document with the report
    """
    doc = Document()
    
    # Title
    doc.add_heading(t.get("title_cronbach", "Test Reliability Analysis (Cronbach's Alpha)"), level=1)
    
    # Introduction and interpretation guide
    doc.add_paragraph(t.get("cronbach_intro", "Objective: Evaluate the internal consistency of EGRA and EGMA assessments."))
    
    doc.add_heading(t.get("cronbach_interpretation", "Interpretation Guide"), level=2)
    doc.add_paragraph(t.get("alpha_excellent", "α ≥ 0.9 → Excellent reliability"), style='List Bullet')
    doc.add_paragraph(t.get("alpha_good", "0.8 ≤ α < 0.9 → Good reliability"), style='List Bullet')
    doc.add_paragraph(t.get("alpha_acceptable", "0.7 ≤ α < 0.8 → Acceptable reliability"), style='List Bullet')
    doc.add_paragraph(t.get("alpha_questionable", "0.6 ≤ α < 0.7 → Questionable reliability"), style='List Bullet')
    doc.add_paragraph(t.get("alpha_poor", "0.5 ≤ α < 0.6 → Poor reliability"), style='List Bullet')
    doc.add_paragraph(t.get("alpha_unacceptable", "α < 0.5 → Unacceptable reliability"), style='List Bullet')
    
    # Results table
    doc.add_heading(t.get("cronbach_results", "Cronbach's Alpha Results"), level=2)
    
    # Create table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = t.get("test_group", "Test Group")
    header_cells[1].text = t.get("n_items", "Number of Items")
    header_cells[2].text = t.get("n_students", "Number of Students")
    header_cells[3].text = t.get("cronbach_alpha", "Cronbach's Alpha")
    header_cells[4].text = t.get("reliability", "Reliability")
    
    # Apply header formatting
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
    
    # Add data rows
    for _, row in df_results.iterrows():
        # Add row to table
        row_cells = table.add_row().cells
        
        # Fill in values
        row_cells[0].text = row["test_group"]
        row_cells[1].text = str(row["n_items"])
        row_cells[2].text = str(row["n_students"])
        
        # Format alpha value
        if row["alpha"] is not None:
            row_cells[3].text = f"{row['alpha']:.3f}"
        else:
            row_cells[3].text = t.get("not_available", "N/A")
        
        row_cells[4].text = row["interpretation"]
    
    # Add visualization
    doc.add_heading(t.get("reliability_visualization", "Reliability Visualization"), level=2)
    
    # Save and add chart
    with tempfile.TemporaryDirectory() as tmp_dir:
        img_path = os.path.join(tmp_dir, "reliability_chart.png")
        fig.write_image(img_path, width=800, height=500)
        doc.add_picture(img_path, width=Inches(6))
    
    # Educational interpretation
    doc.add_heading(t.get("educational_interpretation", "Educational Interpretation"), level=2)
    
    # Calculate average alphas for interpretation
    egra_alphas = [row["alpha"] for _, row in df_results.iterrows() if "EGRA" in row["test_group"]]
    egma_alphas = [row["alpha"] for _, row in df_results.iterrows() if "EGMA" in row["test_group"]]
    
    # Filter out None values
    valid_egra_alphas = [alpha for alpha in egra_alphas if alpha is not None]
    valid_egma_alphas = [alpha for alpha in egma_alphas if alpha is not None]
    
    # Calculate average alphas if possible
    avg_egra_alpha = sum(valid_egra_alphas) / len(valid_egra_alphas) if valid_egra_alphas else None
    avg_egma_alpha = sum(valid_egma_alphas) / len(valid_egma_alphas) if valid_egma_alphas else None
    
    # EGRA interpretation
    if valid_egra_alphas:
        p = doc.add_paragraph()
        p.add_run(t.get('egra_reliability', 'EGRA Reliability Interpretation')).bold = True
        
        if avg_egra_alpha >= 0.7:
            doc.add_paragraph(t.get('egra_reliability_good', """
            The EGRA assessment demonstrates acceptable to good reliability. This indicates:
            - The test items consistently measure the same construct
            - Results can be considered dependable for instructional decisions
            - Different items within the test work together coherently
            """))
        elif avg_egra_alpha >= 0.6:
            doc.add_paragraph(t.get('egra_reliability_questionable', """
            The EGRA assessment shows questionable reliability. This suggests:
            - There may be some inconsistency in what the test items measure
            - Results should be interpreted with caution
            - Consider using additional assessments to confirm findings
            - Test revision may be beneficial in future implementations
            """))
        else:
            doc.add_paragraph(t.get('egra_reliability_poor', """
            The EGRA assessment demonstrates poor reliability. This indicates:
            - Significant inconsistency in what the test items measure
            - Results should not be the sole basis for important decisions
            - The test likely needs revision or replacement
            - Individual items may be measuring different constructs
            """))
    
    # EGMA interpretation
    if valid_egma_alphas:
        p = doc.add_paragraph()
        p.add_run(t.get('egma_reliability', 'EGMA Reliability Interpretation')).bold = True
        
        if avg_egma_alpha >= 0.7:
            doc.add_paragraph(t.get('egma_reliability_good', """
            The EGMA assessment demonstrates acceptable to good reliability. This indicates:
            - The test items consistently measure the same mathematical constructs
            - Results can be considered dependable for instructional decisions
            - Different items within the test work together coherently
            """))
        elif avg_egma_alpha >= 0.6:
            doc.add_paragraph(t.get('egma_reliability_questionable', """
            The EGMA assessment shows questionable reliability. This suggests:
            - There may be some inconsistency in what the test items measure
            - Results should be interpreted with caution
            - Consider using additional assessments to confirm findings
            - Test revision may be beneficial in future implementations
            """))
        else:
            doc.add_paragraph(t.get('egma_reliability_poor', """
            The EGMA assessment demonstrates poor reliability. This indicates:
            - Significant inconsistency in what the test items measure
            - Results should not be the sole basis for important decisions
            - The test likely needs revision or replacement
            - Individual items may be measuring different constructs
            """))
    
    # Recommendations
    p = doc.add_paragraph()
    p.add_run(t.get('reliability_recommendations', 'Recommendations Based on Reliability Results')).bold = True
    
    all_alphas = valid_egra_alphas + valid_egma_alphas
    
    if not all_alphas:
        doc.add_paragraph(t.get('reliability_no_data', """
        - Insufficient data to calculate reliability
        - Consider collecting more complete data in future assessments
        - Ensure proper test administration procedures are followed
        """))
    elif all(alpha >= 0.7 for alpha in all_alphas):
        doc.add_paragraph(t.get('reliability_all_good', """
        - Continue using the current assessment tools with confidence
        - Results can be used for instructional planning and student evaluation
        - Consider conducting validity studies to further strengthen the assessment system
        - Monitor reliability in future administrations to ensure consistency
        """))
    elif any(alpha < 0.6 for alpha in all_alphas):
        doc.add_paragraph(t.get('reliability_some_poor', """
        - Review and potentially revise assessment items with low reliability
        - Use multiple measures when making important educational decisions
        - Provide additional training to test administrators to ensure consistent procedures
        - Consider item analysis to identify specific problematic questions
        """))
    else:
        doc.add_paragraph(t.get('reliability_mixed', """
        - Use results cautiously, especially for tests with lower reliability
        - Consider strengthening assessments with questionable reliability
        - Continue using assessments with good reliability
        - Supplement assessment data with teacher observations and other measures
        """))
    
    # Footer with date
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = t.get('report_date', 'Report generated on: ') + pd.Timestamp.now().strftime('%Y-%m-%d')
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc