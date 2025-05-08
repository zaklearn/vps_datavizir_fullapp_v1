import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import scipy.stats as stats
from docx import Document
from docx.shared import Inches
import tempfile
import os
from config import translations, egra_columns, egma_columns

def show_language_comparison(df, language):
    """
    Compares performance between students taught in English versus Dutch.
    
    Args:
        df (pandas.DataFrame): The data to analyze
        language (str): Selected language for UI elements
    """
    t = translations[language]  # Get translations for selected language
    
    st.markdown(f"""
    ### {t.get("title_language_comparison", "Language of Instruction Comparison")}
    
    🔍 **{t.get("language_comparison_intro", "Objective: Compare performance between students taught in English versus Dutch.")}**
    """)
    
    # Check if language_teaching column exists
    if "language_teaching" not in df.columns:
        st.error(t.get("no_language_column", "Error: No 'language_teaching' column found in the data."))
        return
    
    # Get available assessment columns
    available_egra = [col for col in egra_columns if col in df.columns]
    available_egma = [col for col in egma_columns if col in df.columns]
    available_columns = available_egra + available_egma
    
    if not available_columns:
        st.error(t.get("no_assessment_columns", "No assessment columns found in the data."))
        return
    
    # Prepare data - filter by language and handle missing values
    df_analysis = df.copy()
    
    # Check for English and Dutch in the data
    language_values = df_analysis['language_teaching'].dropna().unique()
    
    # Map similar values and handle different cases
    language_map = {}
    for val in language_values:
        if isinstance(val, str):
            if val.lower() in ['english', 'eng', 'en', 'anglais']:
                language_map[val] = "English"
            elif val.lower() in ['dutch', 'nederlands', 'nl', 'néerlandais']:
                language_map[val] = "Dutch"
            else:
                language_map[val] = val  # Keep as is
    
    # Apply mapping if needed
    if language_map:
        df_analysis["language_teaching"] = df_analysis["language_teaching"].map(
            lambda x: language_map.get(x, x) if x in language_map else x
        )
    
    # Get unique languages after mapping
    languages = df_analysis["language_teaching"].dropna().unique()
    
    # Check if we have both English and Dutch
    has_english = "English" in languages
    has_dutch = "Dutch" in languages
    
    if not (has_english and has_dutch):
        st.warning(t.get("missing_languages", "Warning: Data does not contain both English and Dutch instruction. Available languages: {}").format(
            ", ".join(languages)
        ))
        if len(languages) < 1:
            return  # No language data available
    
    # Allow users to select columns for analysis
    st.subheader(t.get("select_variables", "📋 Select Variables for Analysis"))
    
    # Create two columns for selection
    col1, col2 = st.columns(2)
    
    with col1:
        selected_egra = st.multiselect(
            t.get("egra_variables", "EGRA Variables:"),
            options=available_egra,
            default=available_egra[:min(3, len(available_egra))],  # Default select up to 3 EGRA variables
            format_func=lambda x: t["columns_of_interest"].get(x, x)
        )
    
    with col2:
        selected_egma = st.multiselect(
            t.get("egma_variables", "EGMA Variables:"),
            options=available_egma,
            default=available_egma[:min(3, len(available_egma))],  # Default select up to 3 EGMA variables
            format_func=lambda x: t["columns_of_interest"].get(x, x)
        )
    
    selected_columns = selected_egra + selected_egma
    
    if selected_columns:
        try:
            # Filter for English and Dutch (if both exist)
            df_english = df_analysis[df_analysis["language_teaching"] == "English"] if has_english else pd.DataFrame()
            df_dutch = df_analysis[df_analysis["language_teaching"] == "Dutch"] if has_dutch else pd.DataFrame()
            
            # Check if we have data for at least one language
            if df_english.empty and df_dutch.empty:
                st.error(t.get("no_language_data", "No data available for English or Dutch language of instruction."))
                return
            
            # Calculate mean scores by language
            mean_scores = {}
            sample_sizes = {}
            
            if not df_english.empty:
                mean_scores["English"] = df_english[selected_columns].mean().round(2)
                sample_sizes["English"] = len(df_english)
            
            if not df_dutch.empty:
                mean_scores["Dutch"] = df_dutch[selected_columns].mean().round(2)
                sample_sizes["Dutch"] = len(df_dutch)
            
            # Convert to DataFrame for display
            mean_scores_df = pd.DataFrame(mean_scores).T
            mean_scores_df["Sample Size"] = pd.Series(sample_sizes)
            
            # Display the language performance table
            st.subheader(t.get("language_performance_results", "📊 Mean Scores by Language of Instruction"))
            st.dataframe(mean_scores_df)
            
            # Create bar chart for mean scores by language
            st.subheader(t.get("mean_scores_chart", "📊 Mean Scores Comparison by Language"))
            
            # Reshape data for plotting
            plot_data = []
            for lang, scores in mean_scores.items():
                for col in selected_columns:
                    plot_data.append({
                        "language": lang,
                        "variable": t["columns_of_interest"].get(col, col),
                        "score": scores[col],
                        "variable_code": col  # Keep original code for ordering
                    })
            
            plot_df = pd.DataFrame(plot_data)
            
            # Create bar chart
            fig = px.bar(
                plot_df,
                x="variable",
                y="score",
                color="language",
                barmode="group",
                title=t.get("language_comparison_title", "Performance Comparison by Language of Instruction"),
                labels={
                    "variable": t.get("assessment_task", "Assessment Task"),
                    "score": t.get("mean_score", "Mean Score"),
                    "language": t.get("language_of_instruction", "Language of Instruction")
                },
                color_discrete_map={
                    "English": "#3498DB",  # Blue
                    "Dutch": "#F39C12"    # Orange
                }
            )
            
            # Update layout
            fig.update_layout(
                xaxis_tickangle=-45,
                legend_title=t.get("language_of_instruction", "Language of Instruction"),
                height=600
            )
            
            st.plotly_chart(fig, use_container_width=True)
            
            # Distribution plots (Box plots) by language for each variable
            st.subheader(t.get("distribution_by_language", "📈 Score Distributions by Language of Instruction"))
            
            # Create box plots for each selected variable
            for i in range(0, len(selected_columns), 2):
                # Create two columns for displaying plots side by side
                col1, col2 = st.columns(2)
                
                with col1:
                    if i < len(selected_columns):
                        column = selected_columns[i]
                        col_name = t["columns_of_interest"].get(column, column)
                        
                        # Create box plot
                        box_fig = px.box(
                            df_analysis[df_analysis["language_teaching"].isin(["English", "Dutch"])],
                            x="language_teaching",
                            y=column,
                            color="language_teaching",
                            title=f"{col_name}",
                            labels={
                                "language_teaching": t.get("language_of_instruction", "Language of Instruction"),
                                column: col_name
                            },
                            color_discrete_map={
                                "English": "#3498DB",
                                "Dutch": "#F39C12"
                            }
                        )
                        
                        # Update layout
                        box_fig.update_layout(
                            showlegend=False,
                            height=400
                        )
                        
                        st.plotly_chart(box_fig, use_container_width=True)
                
                with col2:
                    if i + 1 < len(selected_columns):
                        column = selected_columns[i + 1]
                        col_name = t["columns_of_interest"].get(column, column)
                        
                        # Create box plot
                        box_fig = px.box(
                            df_analysis[df_analysis["language_teaching"].isin(["English", "Dutch"])],
                            x="language_teaching",
                            y=column,
                            color="language_teaching",
                            title=f"{col_name}",
                            labels={
                                "language_teaching": t.get("language_of_instruction", "Language of Instruction"),
                                column: col_name
                            },
                            color_discrete_map={
                                "English": "#3498DB",
                                "Dutch": "#F39C12"
                            }
                        )
                        
                        # Update layout
                        box_fig.update_layout(
                            showlegend=False,
                            height=400
                        )
                        
                        st.plotly_chart(box_fig, use_container_width=True)
            
            # Statistical significance testing (Mann-Whitney U test)
            st.subheader(t.get("statistical_testing", "📊 Statistical Significance Testing"))
            
            # Only perform tests if we have both English and Dutch data
            if has_english and has_dutch and not df_english.empty and not df_dutch.empty:
                st.markdown(t.get("mann_whitney_explanation", """
                The analysis below uses the Mann-Whitney U test, a non-parametric method for comparing two independent groups.
                A p-value < 0.05 indicates statistically significant differences between English and Dutch instruction.
                """))
                
                # Run Mann-Whitney test for each variable
                test_results = []
                
                for col in selected_columns:
                    col_name = t["columns_of_interest"].get(col, col)
                    
                    # Get data for English and Dutch
                    english_data = df_english[col].dropna()
                    dutch_data = df_dutch[col].dropna()
                    
                    # Calculate means for effect direction
                    english_mean = english_data.mean()
                    dutch_mean = dutch_data.mean()
                    
                    # Perform Mann-Whitney test if we have data for both groups
                    if len(english_data) > 0 and len(dutch_data) > 0:
                        try:
                            u_stat, p_value = stats.mannwhitneyu(english_data, dutch_data, alternative='two-sided')
                            
                            # Determine which language performed better
                            better_language = "English" if english_mean > dutch_mean else "Dutch"
                            
                            test_results.append({
                                "variable": col_name,
                                "english_mean": english_mean,
                                "dutch_mean": dutch_mean,
                                "difference": abs(english_mean - dutch_mean),
                                "percent_diff": abs(english_mean - dutch_mean) / ((english_mean + dutch_mean) / 2) * 100 if english_mean + dutch_mean > 0 else 0,
                                "better_language": better_language,
                                "u_statistic": u_stat,
                                "p_value": p_value,
                                "significant": p_value < 0.05
                            })
                        except Exception as e:
                            # Handle errors in statistical testing
                            test_results.append({
                                "variable": col_name,
                                "english_mean": english_mean,
                                "dutch_mean": dutch_mean,
                                "difference": abs(english_mean - dutch_mean),
                                "percent_diff": abs(english_mean - dutch_mean) / ((english_mean + dutch_mean) / 2) * 100 if english_mean + dutch_mean > 0 else 0,
                                "better_language": None,
                                "u_statistic": None,
                                "p_value": None,
                                "significant": None,
                                "error": str(e)
                            })
                
                # Display test results if any tests were performed
                if test_results:
                    test_df = pd.DataFrame(test_results)
                    
                    # Format the display DataFrame
                    display_df = test_df.copy()
                    display_df.columns = [
                        t.get("variable", "Variable"),
                        t.get("english_mean", "English Mean"),
                        t.get("dutch_mean", "Dutch Mean"),
                        t.get("difference", "Difference"),
                        t.get("percent_diff", "% Difference"),
                        t.get("better_language", "Better Performance"),
                        t.get("u_statistic", "U Statistic"),
                        t.get("p_value", "p-value"),
                        t.get("significant", "Significant Difference"),
                        *([t.get("error", "Error")] if "error" in test_df.columns else [])
                    ]
                    
                    # Format values for display
                    display_df[t.get("english_mean", "English Mean")] = display_df[t.get("english_mean", "English Mean")].round(2)
                    display_df[t.get("dutch_mean", "Dutch Mean")] = display_df[t.get("dutch_mean", "Dutch Mean")].round(2)
                    display_df[t.get("difference", "Difference")] = display_df[t.get("difference", "Difference")].round(2)
                    display_df[t.get("percent_diff", "% Difference")] = display_df[t.get("percent_diff", "% Difference")].round(1)
                    
                    # Format p-values for display
                    if "p_value" in test_df.columns:
                        display_df[t.get("p_value", "p-value")] = display_df[t.get("p_value", "p-value")].apply(
                            lambda x: f"{x:.4f}" if x is not None else "N/A"
                        )
                    
                    # Format significant column
                    if "significant" in test_df.columns:
                        display_df[t.get("significant", "Significant Difference")] = display_df[t.get("significant", "Significant Difference")].apply(
                            lambda x: t.get("significant_yes", "Yes") if x else t.get("significant_no", "No") if x is not None else "N/A"
                        )
                    
                    st.dataframe(display_df)
                    
                    # Summary of language differences
                    st.subheader(t.get("language_summary", "Summary of Language of Instruction Differences"))
                    
                    # Count significant differences
                    sig_differences = [r for r in test_results if r.get("significant")]
                    english_advantage = [r for r in sig_differences if r.get("better_language") == "English"]
                    dutch_advantage = [r for r in sig_differences if r.get("better_language") == "Dutch"]
                    
                    if sig_differences:
                        st.markdown(t.get("significant_diff_found", "**Significant differences by language of instruction were found in {} out of {} variables analyzed.**").format(
                            len(sig_differences), len(test_results)
                        ))
                        
                        if english_advantage:
                            st.markdown(t.get("english_advantage", "**English instruction showed significantly better performance in:**"))
                            for r in english_advantage:
                                st.markdown(f"- {r['variable']} ({r['percent_diff']:.1f}% difference)")
                        
                        if dutch_advantage:
                            st.markdown(t.get("dutch_advantage", "**Dutch instruction showed significantly better performance in:**"))
                            for r in dutch_advantage:
                                st.markdown(f"- {r['variable']} ({r['percent_diff']:.1f}% difference)")
                        
                        # Educational implications
                        st.markdown(t.get("educational_implications", "**Educational Implications:**"))
                        
                        # Check which types of variables show language differences
                        reading_diffs = [r for r in sig_differences if any(col in r["variable"] for col in ["Letter", "Phon", "Word", "Reading", "Comprehension"])]
                        math_diffs = [r for r in sig_differences if any(col in r["variable"] for col in ["Number", "Addition", "Subtraction", "Problem"])]
                        
                        if reading_diffs:
                            st.markdown(t.get("reading_implications", "- **Reading skills:** Consider language-specific approaches to reading instruction"))
                        
                        if math_diffs:
                            st.markdown(t.get("math_implications", "- **Math skills:** Review language effects on mathematical performance and instruction"))
                        
                        st.markdown(t.get("general_implications", """
                        - Examine curriculum and materials in each language for possible differences
                        - Consider teacher training factors that might differ by language of instruction
                        - Review assessment approaches for potential language bias
                        - Develop language-specific intervention strategies where needed
                        """))
                    else:
                        st.markdown(t.get("no_significant_diff", """
                        **No statistically significant differences were found between English and Dutch instruction.**
                        
                        This suggests that both languages of instruction are equally effective across the assessed skills,
                        indicating consistent educational quality regardless of instructional language.
                        """))
                
                # Export to Word
                if st.button(t.get("export_language_word", "📄 Export to Word")):
                    try:
                        doc = create_language_comparison_word_report(
                            df_analysis, test_results if has_english and has_dutch else [], 
                            mean_scores, sample_sizes, selected_columns, t
                        )
                        
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                            doc.save(tmp.name)
                            with open(tmp.name, 'rb') as f:
                                docx = f.read()
                            st.download_button(
                                t.get("download_language_word", "📥 Download Word Report"),
                                docx,
                                "language_comparison_analysis.docx",
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        os.unlink(tmp.name)
                    except Exception as e:
                        st.error(f"Error creating Word report: {str(e)}")
            else:
                st.info(t.get("cannot_perform_tests", "Statistical tests cannot be performed because data for both English and Dutch instruction is not available."))
        
        except Exception as e:
            st.error(f"Error in language comparison analysis: {str(e)}")
    
    else:
        st.warning(t.get("warning_select_variable", "Please select at least one variable to analyze."))

def create_language_comparison_word_report(df, test_results, mean_scores, sample_sizes, selected_columns, t):
    """
    Creates a Word report with language comparison analysis.
    
    Args:
        df (pandas.DataFrame): The data to analyze
        test_results (list): Statistical test results
        mean_scores (dict): Mean scores by language
        sample_sizes (dict): Sample sizes by language
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
        
    Returns:
        docx.Document: Word document with the report
    """
    doc = Document()
    
    # Title
    doc.add_heading(t.get("title_language_comparison", "Language of Instruction Comparison"), level=1)
    
    # Introduction
    doc.add_paragraph(t.get("language_comparison_intro", "Objective: Compare performance between students taught in English versus Dutch."))
    
    # Sample information
    doc.add_heading(t.get("sample_info", "Sample Information"), level=2)
    
    # Create language distribution table
    language_table = doc.add_table(rows=len(sample_sizes) + 1, cols=2)
    language_table.style = 'Table Grid'
    
    # Add headers
    header_cells = language_table.rows[0].cells
    header_cells[0].text = t.get("language_of_instruction", "Language of Instruction")
    header_cells[1].text = t.get("sample_size", "Sample Size")
    
    # Add data rows
    for i, (lang, size) in enumerate(sample_sizes.items(), 1):
        row_cells = language_table.rows[i].cells
        row_cells[0].text = lang
        row_cells[1].text = str(size)
    
    # Mean scores by language
    doc.add_heading(t.get("mean_scores_language", "Mean Scores by Language of Instruction"), level=2)
    
    # Create mean scores table
    # First, determine how many columns we need (languages + 1 for variable names)
    table_cols = len(mean_scores) + 1
    table_rows = len(selected_columns) + 1  # +1 for header
    
    mean_scores_table = doc.add_table(rows=table_rows, cols=table_cols)
    mean_scores_table.style = 'Table Grid'
    
    # Add headers
    header_cells = mean_scores_table.rows[0].cells
    header_cells[0].text = t.get("variable", "Variable")
    
    for i, lang in enumerate(mean_scores.keys(), 1):
        header_cells[i].text = lang
    
    # Add data rows
    for i, col in enumerate(selected_columns, 1):
        row_cells = mean_scores_table.rows[i].cells
        row_cells[0].text = t["columns_of_interest"].get(col, col)
        
        for j, lang in enumerate(mean_scores.keys(), 1):
            row_cells[j].text = f"{mean_scores[lang][col]:.2f}"
    
    # Statistical test results
    if test_results:
        doc.add_heading(t.get("statistical_testing", "Statistical Significance Testing"), level=2)
        
        doc.add_paragraph(t.get("mann_whitney_explanation", """
        The analysis uses the Mann-Whitney U test, a non-parametric method for comparing two independent groups.
        A p-value < 0.05 indicates statistically significant differences between English and Dutch instruction.
        """))
        
        # Create test results table
        test_table = doc.add_table(rows=len(test_results) + 1, cols=6)  # Variable, English Mean, Dutch Mean, Difference, Better, Significant
        test_table.style = 'Table Grid'
        
        # Add headers
        header_cells = test_table.rows[0].cells
        header_cells[0].text = t.get("variable", "Variable")
        header_cells[1].text = t.get("english_mean", "English Mean")
        header_cells[2].text = t.get("dutch_mean", "Dutch Mean")
        header_cells[3].text = t.get("difference", "Difference")
        header_cells[4].text = t.get("better_language", "Better Performance")
        header_cells[5].text = t.get("significant", "Significant")
        
        # Add data rows
        for i, result in enumerate(test_results, 1):
            row_cells = test_table.rows[i].cells
            
            row_cells[0].text = result["variable"]
            row_cells[1].text = f"{result['english_mean']:.2f}" if result['english_mean'] is not None else "N/A"
            row_cells[2].text = f"{result['dutch_mean']:.2f}" if result['dutch_mean'] is not None else "N/A"
            row_cells[3].text = f"{result['difference']:.2f}" if result['difference'] is not None else "N/A"
            row_cells[4].text = result["better_language"] if result["better_language"] is not None else "N/A"
            
            if result.get('significant') is not None:
                row_cells[5].text = t.get("significant_yes", "Yes") if result['significant'] else t.get("significant_no", "No")
            else:
                row_cells[5].text = "N/A"
    
    # Distribution plots for each variable
    doc.add_heading(t.get("distribution_language", "Score Distributions by Language of Instruction"), level=2)
    
    with tempfile.TemporaryDirectory() as tmp_dir:
        for column in selected_columns:
            col_name = t["columns_of_interest"].get(column, column)
            doc.add_heading(col_name, level=3)
            
            # Create box plot
            fig = px.box(
                df[df["language_teaching"].isin(["English", "Dutch"])],
                x="language_teaching",
                y=column,
                color="language_teaching",
                labels={
                    "language_teaching": t.get("language_of_instruction", "Language of Instruction"),
                    column: col_name
                },
                color_discrete_map={
                    "English": "#3498DB",
                    "Dutch": "#F39C12"
                }
            )
            
            # Save plot
            img_path = os.path.join(tmp_dir, f"{column}_language_boxplot.png")
            fig.write_image(img_path, width=800, height=400)
            
            # Add plot to document
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph()
    
    # Summary of language differences
    if test_results:
        doc.add_heading(t.get("language_summary", "Summary of Language of Instruction Differences"), level=2)
        
        # Count significant differences
        sig_differences = [r for r in test_results if r.get("significant")]
        english_advantage = [r for r in sig_differences if r.get("better_language") == "English"]
        dutch_advantage = [r for r in sig_differences if r.get("better_language") == "Dutch"]
        
        if sig_differences:
            doc.add_paragraph(t.get("significant_diff_found", "Significant differences by language of instruction were found in {} out of {} variables analyzed.").format(
                len(sig_differences), len(test_results)
            ))
            
            if english_advantage:
                p = doc.add_paragraph()
                p.add_run(t.get("english_advantage", "English instruction showed significantly better performance in:")).bold = True
                
                for r in english_advantage:
                    doc.add_paragraph(f"{r['variable']} ({r['percent_diff']:.1f}% difference)", style='List Bullet')
            
            if dutch_advantage:
                p = doc.add_paragraph()
                p.add_run(t.get("dutch_advantage", "Dutch instruction showed significantly better performance in:")).bold = True
                
                for r in dutch_advantage:
                    doc.add_paragraph(f"{r['variable']} ({r['percent_diff']:.1f}% difference)", style='List Bullet')
            
            # Educational implications
            p = doc.add_paragraph()
            p.add_run(t.get("educational_implications", "Educational Implications:")).bold = True
            
            # Check which types of variables show language differences
            reading_diffs = [r for r in sig_differences if any(col in r["variable"] for col in ["Letter", "Phon", "Word", "Reading", "Comprehension"])]
            math_diffs = [r for r in sig_differences if any(col in r["variable"] for col in ["Number", "Addition", "Subtraction", "Problem"])]
            
            if reading_diffs:
                doc.add_paragraph(t.get("reading_implications", "Reading skills: Consider language-specific approaches to reading instruction"), style='List Bullet')
            
            if math_diffs:
                doc.add_paragraph(t.get("math_implications", "Math skills: Review language effects on mathematical performance and instruction"), style='List Bullet')
            
            doc.add_paragraph(t.get("examine_curriculum", "Examine curriculum and materials in each language for possible differences"), style='List Bullet')
            doc.add_paragraph(t.get("teacher_training", "Consider teacher training factors that might differ by language of instruction"), style='List Bullet')
            doc.add_paragraph(t.get("assessment_bias", "Review assessment approaches for potential language bias"), style='List Bullet')
            doc.add_paragraph(t.get("intervention_strategies", "Develop language-specific intervention strategies where needed"), style='List Bullet')
        else:
            doc.add_paragraph(t.get("no_significant_diff", """
            No statistically significant differences were found between English and Dutch instruction.
            
            This suggests that both languages of instruction are equally effective across the assessed skills,
            indicating consistent educational quality regardless of instructional language.
            """))
    
    # Methodology note
    doc.add_paragraph(t.get("methodology_note", """
    Methodology Note: This analysis compares performance between students taught in English versus Dutch. The Mann-Whitney U test is used because it does not assume normal distribution of the data and is appropriate for comparing two independent groups.
    """), style='Normal')
    
    return doc