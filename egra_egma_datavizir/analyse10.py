import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import scipy.stats as stats
from docx import Document
from docx.shared import Inches
import tempfile
import os
from config import translations, egra_columns, egma_columns

def show_gender_effect(df, language):
    """
    Analyzes and displays gender differences in EGRA and EGMA assessment results.
    
    Args:
        df (pandas.DataFrame): The data to analyze
        language (str): Selected language for UI elements
    """
    t = translations[language]  # Get translations for selected language
    
    st.markdown(f"""
    ### {t.get("title_gender_effect", "Gender Effect Analysis")}
    
    🔍 **{t.get("gender_intro", "Objective: Examine differences in performance between boys and girls across assessment variables.")}**
    """)
    
    # Check if gender column exists
    if "stgender" not in df.columns:
        st.error(t.get("no_gender_column", "Error: No 'stgender' column found in the data."))
        return
    
    # Get available assessment columns
    available_egra = [col for col in egra_columns if col in df.columns]
    available_egma = [col for col in egma_columns if col in df.columns]
    available_columns = available_egra + available_egma
    
    if not available_columns:
        st.error(t.get("no_assessment_columns", "No assessment columns found in the data."))
        return
    """
    # Prepare data - map gender codes to labels and handle missing values
    # Assuming: 1 = Boy, 0 = Girl (common coding in educational datasets)
    df_analysis = df.copy()
    
    # Check the unique values in stgender to determine mapping approach
    gender_values = df_analysis['stgender'].dropna().unique()
    
    # Determine if gender is coded as 0/1 or has string values
    if all(isinstance(x, (int, float)) for x in gender_values):
        # Numeric coding
        df_analysis["gender"] = df_analysis["stgender"].map({
            1: t.get("boy", "Garçon"), 
            0: t.get("girl", "Fille")
        }).fillna(t.get("unknown", "Unknown"))
    else:
        # Try to map using common string values (case-insensitive)
        gender_map = {}
        for val in gender_values:
            if isinstance(val, str):
                if val.lower() in ['boy', 'boys', 'male', 'm', 'homme', 'garçon']:
                    gender_map[val] = t.get("boy", "Boy")
                elif val.lower() in ['girl', 'girls', 'female', 'f', 'femme', 'fille']:
                    gender_map[val] = t.get("girl", "Girl")
                else:
                    gender_map[val] = t.get("unknown", "Unknown")
            else:
                gender_map[val] = t.get("unknown", "Unknown")
        
        df_analysis["gender"] = df_analysis["stgender"].map(gender_map).fillna(t.get("unknown", "Unknown"))
    
    # Check if we have enough data for each gender
    gender_counts = df_analysis["gender"].value_counts()
    if len(gender_counts) < 2 or (t.get("boy", "Boy") not in gender_counts or t.get("girl", "Girl") not in gender_counts):
        st.warning(t.get("insufficient_gender_data", "Warning: Insufficient data for gender comparison. Please check gender coding."))
        st.write(t.get("gender_distribution", "Current gender distribution:"))
        st.write(gender_counts)
        return
    
    # Remove unknown gender for analysis
    df_analysis = df_analysis[df_analysis["gender"] != t.get("unknown", "Unknown")]
    """
    # Prepare data - map gender codes to labels and handle missing values
    df_analysis = df.copy()

    # Map gender values (handles both numeric and string cases)
    gender_map = {
        1: t.get("boy", "Garçon"),
        0: t.get("girl", "Fille"),
        "boy": t.get("boy", "Garçon"), "boys": t.get("boy", "Garçon"), "male": t.get("boy", "Garçon"),
        "m": t.get("boy", "Garçon"), "homme": t.get("boy", "Garçon"), "garçon": t.get("boy", "Garçon"),
        "girl": t.get("girl", "Fille"), "girls": t.get("girl", "Fille"), "female": t.get("girl", "Fille"),
        "f": t.get("girl", "Fille"), "femme": t.get("girl", "Fille"), "fille": t.get("girl", "Fille")
    }

    # Normalize and map gender
    df_analysis["gender"] = df_analysis["stgender"].apply(
        lambda x: gender_map.get(str(x).strip().lower(), t.get("unknown", "Inconnu")) if isinstance(x, str)
        else gender_map.get(x, t.get("unknown", "Inconnu"))
    )

    # Check if we have enough valid gender data
    gender_counts = df_analysis["gender"].value_counts()
    if len(gender_counts) < 2 or not all(label in gender_counts for label in [t.get("boy", "Garçon"), t.get("girl", "Fille")]):
        st.warning(t.get("insufficient_gender_data", "Warning: Insufficient data for gender comparison. Please check gender coding."))
        st.write(t.get("gender_distribution", "Current gender distribution:"))
        st.write(gender_counts)
        return

    # Remove unknowns before analysis
    df_analysis = df_analysis[df_analysis["gender"] != t.get("unknown", "Inconnu")]

    # Allow users to select columns for analysis
    st.subheader(t.get("select_variables", "📋 Select Variables for Analysis"))
    
    # Create two columns for selection
    col1, col2 = st.columns(2)
    
    with col1:
        selected_egra = st.multiselect(
            t.get("egra_variables", "EGRA Variables:"),
            options=available_egra,
            default=available_egra[:min(3, len(available_egra))],  # Default select up to 3 EGRA variables
            format_func=lambda x: t["columns_of_interest"].get(x, x)
        )
    
    with col2:
        selected_egma = st.multiselect(
            t.get("egma_variables", "EGMA Variables:"),
            options=available_egma,
            default=available_egma[:min(3, len(available_egma))],  # Default select up to 3 EGMA variables
            format_func=lambda x: t["columns_of_interest"].get(x, x)
        )
    
    selected_columns = selected_egra + selected_egma
    
    if selected_columns:
        try:
            # Calculate mean scores by gender
            mean_scores_by_gender = df_analysis.groupby("gender")[selected_columns].mean().round(2)
            
            # Calculate sample sizes by gender for reference
            sample_sizes = df_analysis.groupby("gender").size().rename(t.get("sample_size", "Sample Size"))
            
            # Combine with mean scores for display
            performance_table = pd.concat([mean_scores_by_gender, sample_sizes], axis=1)
            
            # Display the gender performance table
            st.subheader(t.get("gender_performance_results", "📊 Mean Scores by Gender"))
            st.dataframe(performance_table, use_container_width=True)
            
            # Create bar chart for mean scores by gender
            st.subheader(t.get("mean_scores_chart", "📊 Mean Scores Comparison by Gender"))
            
            # Reshape data for plotting
            plot_data = []
            for gender, row in mean_scores_by_gender.iterrows():
                for col in selected_columns:
                    plot_data.append({
                        "gender": gender,
                        "variable": t["columns_of_interest"].get(col, col),
                        "score": row[col],
                        "variable_code": col  # Keep original code for ordering
                    })
            
            plot_df = pd.DataFrame(plot_data)
            
            # Create bar chart
            fig = px.bar(
                plot_df,
                x="variable",
                y="score",
                color="gender",
                barmode="group",
                title=t.get("gender_comparison_title", "Performance Comparison by Gender"),
                labels={
                    "variable": t.get("assessment_task", "Assessment Task"),
                    "score": t.get("mean_score", "Mean Score"),
                    "gender": t.get("gender", "Gender")
                },
                color_discrete_map={
                    t.get("boy", "Boy"): "#3498DB",  # Blue
                    t.get("girl", "Girl"): "#E83E8C"  # Pink
                }
            )
            
            # Update layout
            fig.update_layout(
                xaxis_tickangle=-45,
                legend_title=t.get("gender", "Gender"),
                height=600
            )
            
            st.plotly_chart(fig, use_container_width=True)
            
            # Distribution plots (Box plots) by gender for each variable
            st.subheader(t.get("distribution_by_gender", "📈 Score Distributions by Gender"))
            
            # Create box plots for each selected variable
            for i in range(0, len(selected_columns), 2):
                # Create two columns for displaying plots side by side
                col1, col2 = st.columns(2)
                
                with col1:
                    if i < len(selected_columns):
                        column = selected_columns[i]
                        col_name = t["columns_of_interest"].get(column, column)
                        
                        # Create box plot
                        box_fig = px.box(
                            df_analysis,
                            x="gender",
                            y=column,
                            color="gender",
                            title=f"{col_name}",
                            labels={
                                "gender": t.get("gender", "Gender"),
                                column: col_name
                            },
                            color_discrete_map={
                                t.get("boy", "Boy"): "#3498DB",
                                t.get("girl", "Girl"): "#E83E8C"
                            }
                        )
                        
                        # Update layout
                        box_fig.update_layout(
                            showlegend=False,
                            height=400
                        )
                        
                        st.plotly_chart(box_fig, use_container_width=True)
                
                with col2:
                    if i + 1 < len(selected_columns):
                        column = selected_columns[i + 1]
                        col_name = t["columns_of_interest"].get(column, column)
                        
                        # Create box plot
                        box_fig = px.box(
                            df_analysis,
                            x="gender",
                            y=column,
                            color="gender",
                            title=f"{col_name}",
                            labels={
                                "gender": t.get("gender", "Gender"),
                                column: col_name
                            },
                            color_discrete_map={
                                t.get("boy", "Boy"): "#3498DB",
                                t.get("girl", "Girl"): "#E83E8C"
                            }
                        )
                        
                        # Update layout
                        box_fig.update_layout(
                            showlegend=False,
                            height=400
                        )
                        
                        st.plotly_chart(box_fig, use_container_width=True)
            
            # Statistical significance testing (Mann-Whitney U test)
            st.subheader(t.get("statistical_testing", "📊 Statistical Significance Testing"))
            
            st.markdown(t.get("mann_whitney_explanation", """
            The analysis below uses the Mann-Whitney U test, a non-parametric method for comparing two independent groups.
            A p-value < 0.05 indicates statistically significant differences between boys and girls.
            """))
            
            # Run Mann-Whitney test for each variable
            test_results = []
            
            for col in selected_columns:
                col_name = t["columns_of_interest"].get(col, col)
                
                # Get data for boys and girls
                boys_data = df_analysis[df_analysis["gender"] == t.get("boy", "Boy")][col].dropna()
                girls_data = df_analysis[df_analysis["gender"] == t.get("girl", "Girl")][col].dropna()
                
                # Calculate means for effect direction
                boys_mean = boys_data.mean()
                girls_mean = girls_data.mean()
                
                # Perform Mann-Whitney test if we have data for both groups
                if len(boys_data) > 0 and len(girls_data) > 0:
                    try:
                        u_stat, p_value = stats.mannwhitneyu(boys_data, girls_data, alternative='two-sided')
                        
                        # Determine which gender performed better
                        better_gender = t.get("boy", "Boy") if boys_mean > girls_mean else t.get("girl", "Girl")
                        
                        test_results.append({
                            "variable": col_name,
                            "boys_mean": boys_mean,
                            "girls_mean": girls_mean,
                            "difference": abs(boys_mean - girls_mean),
                            "percent_diff": abs(boys_mean - girls_mean) / ((boys_mean + girls_mean) / 2) * 100 if boys_mean + girls_mean > 0 else 0,
                            "better_gender": better_gender,
                            "u_statistic": u_stat,
                            "p_value": p_value,
                            "significant": p_value < 0.05
                        })
                    except Exception as e:
                        # Handle errors in statistical testing
                        test_results.append({
                            "variable": col_name,
                            "boys_mean": boys_mean,
                            "girls_mean": girls_mean,
                            "difference": abs(boys_mean - girls_mean),
                            "percent_diff": abs(boys_mean - girls_mean) / ((boys_mean + girls_mean) / 2) * 100 if boys_mean + girls_mean > 0 else 0,
                            "better_gender": None,
                            "u_statistic": None,
                            "p_value": None,
                            "significant": None,
                            "error": str(e)
                        })
            
            # Display test results if any tests were performed
            if test_results:
                test_df = pd.DataFrame(test_results)
                
                # Format the display DataFrame
                display_df = test_df.copy()
                display_df.columns = [
                    t.get("variable", "Variable"),
                    t.get("boys_mean", "Boys Mean"),
                    t.get("girls_mean", "Girls Mean"),
                    t.get("difference", "Difference"),
                    t.get("percent_diff", "% Difference"),
                    t.get("better_gender", "Better Performance"),
                    t.get("u_statistic", "U Statistic"),
                    t.get("p_value", "p-value"),
                    t.get("significant", "Significant Difference"),
                    *([t.get("error", "Error")] if "error" in test_df.columns else [])
                ]
                
                # Format values for display
                display_df[t.get("boys_mean", "Boys Mean")] = display_df[t.get("boys_mean", "Boys Mean")].round(2)
                display_df[t.get("girls_mean", "Girls Mean")] = display_df[t.get("girls_mean", "Girls Mean")].round(2)
                display_df[t.get("difference", "Difference")] = display_df[t.get("difference", "Difference")].round(2)
                display_df[t.get("percent_diff", "% Difference")] = display_df[t.get("percent_diff", "% Difference")].round(1)
                
                # Format p-values for display
                if "p_value" in test_df.columns:
                    display_df[t.get("p_value", "p-value")] = display_df[t.get("p_value", "p-value")].apply(
                        lambda x: f"{x:.4f}" if x is not None else "N/A"
                    )
                
                # Format significant column
                if "significant" in test_df.columns:
                    display_df[t.get("significant", "Significant Difference")] = display_df[t.get("significant", "Significant Difference")].apply(
                        lambda x: t.get("significant_yes", "Yes") if x else t.get("significant_no", "No") if x is not None else "N/A"
                    )
                
                st.dataframe(display_df)
                
                # Summary of gender differences
                st.subheader(t.get("gender_summary", "Summary of Gender Differences"))
                
                # Count significant differences
                sig_differences = [r for r in test_results if r.get("significant")]
                boy_advantage = [r for r in sig_differences if r.get("better_gender") == t.get("boy", "Boy")]
                girl_advantage = [r for r in sig_differences if r.get("better_gender") == t.get("girl", "Girl")]
                
                if sig_differences:
                    st.markdown(t.get("significant_diff_found", "**Significant gender differences were found in {} out of {} variables analyzed.**").format(
                        len(sig_differences), len(test_results)
                    ))
                    
                    if boy_advantage:
                        st.markdown(t.get("boy_advantage", "**Boys performed significantly better in:**"))
                        for r in boy_advantage:
                            st.markdown(f"- {r['variable']} ({r['percent_diff']:.1f}% difference)")
                    
                    if girl_advantage:
                        st.markdown(t.get("girl_advantage", "**Girls performed significantly better in:**"))
                        for r in girl_advantage:
                            st.markdown(f"- {r['variable']} ({r['percent_diff']:.1f}% difference)")
                    
                    # Educational implications
                    st.markdown(t.get("educational_implications", "**Educational Implications:**"))
                    
                    # Check which types of variables show gender differences
                    reading_diffs = [r for r in sig_differences if any(col in r["variable"] for col in ["Letter", "Phon", "Word", "Reading", "Comprehension"])]
                    math_diffs = [r for r in sig_differences if any(col in r["variable"] for col in ["Number", "Addition", "Subtraction", "Problem"])]
                    
                    if reading_diffs:
                        st.markdown(t.get("reading_implications", "- **Reading skills:** Consider gender-responsive teaching strategies to address observed differences"))
                    
                    if math_diffs:
                        st.markdown(t.get("math_implications", "- **Math skills:** Implement targeted interventions to close gender gaps in mathematical performance"))
                    
                    st.markdown(t.get("general_implications", """
                    - Review teaching materials and methods for potential gender bias
                    - Consider mixed-gender collaborative learning activities
                    - Provide targeted support for the lower-performing gender in specific skill areas
                    - Monitor progress to ensure equitable outcomes
                    """))
                else:
                    st.markdown(t.get("no_significant_diff", """
                    **No statistically significant gender differences were found.**
                    
                    This suggests that both boys and girls are performing at similar levels across the assessed skills,
                    indicating equitable educational outcomes between genders.
                    """))
            
            # Export to Word
            if st.button(t.get("export_gender_word", "📄 Export to Word")):
                try:
                    doc = create_gender_effect_word_report(
                        df_analysis, test_results, mean_scores_by_gender, selected_columns, t
                    )
                    
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                        doc.save(tmp.name)
                        with open(tmp.name, 'rb') as f:
                            docx = f.read()
                        st.download_button(
                            t.get("download_gender_word", "📥 Download Word Report"),
                            docx,
                            "gender_effect_analysis.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    os.unlink(tmp.name)
                except Exception as e:
                    st.error(f"Error creating Word report: {str(e)}")
        
        except Exception as e:
            st.error(f"Error in gender effect analysis: {str(e)}")
    
    else:
        st.warning(t.get("warning_select_variable", "Please select at least one variable to analyze."))

def create_gender_effect_word_report(df, test_results, mean_scores, selected_columns, t):
    """
    Creates a Word report with gender effect analysis.
    
    Args:
        df (pandas.DataFrame): Prepared data with gender mapping
        test_results (list): Statistical test results
        mean_scores (pandas.DataFrame): Mean scores by gender
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
        
    Returns:
        docx.Document: Word document with the report
    """
    doc = Document()
    
    # Title
    doc.add_heading(t.get("title_gender_effect", "Gender Effect Analysis"), level=1)
    
    # Introduction
    doc.add_paragraph(t.get("gender_intro", "Objective: Examine differences in performance between boys and girls across assessment variables."))
    
    # Sample information
    doc.add_heading(t.get("sample_info", "Sample Information"), level=2)
    
    # Create gender distribution table
    gender_counts = df["gender"].value_counts().reset_index()
    gender_counts.columns = [t.get("gender", "Gender"), t.get("count", "Count")]
    
    gender_table = doc.add_table(rows=len(gender_counts) + 1, cols=2)
    gender_table.style = 'Table Grid'
    
    # Add headers
    header_cells = gender_table.rows[0].cells
    header_cells[0].text = t.get("gender", "Gender")
    header_cells[1].text = t.get("count", "Count")
    
    # Add data rows
    for i, (_, row) in enumerate(gender_counts.iterrows(), 1):
        row_cells = gender_table.rows[i].cells
        row_cells[0].text = row[t.get("gender", "Gender")]
        row_cells[1].text = str(row[t.get("count", "Count")])
    
    # Mean scores by gender
    doc.add_heading(t.get("mean_scores_gender", "Mean Scores by Gender"), level=2)
    
    # Reset index to make gender a column
    mean_scores_table = mean_scores.reset_index()
    
    # Create table
    table = doc.add_table(rows=len(mean_scores_table) + 1, cols=len(mean_scores_table.columns))
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = t.get("gender", "Gender")
    
    for i, col in enumerate(mean_scores_table.columns[1:], 1):
        header_cells[i].text = t["columns_of_interest"].get(col, col)
    
    # Add data rows
    for i, (_, row) in enumerate(mean_scores_table.iterrows(), 1):
        row_cells = table.rows[i].cells
        
        for j, value in enumerate(row):
            row_cells[j].text = str(value)
    
    # Statistical test results
    if test_results:
        doc.add_heading(t.get("statistical_testing", "Statistical Significance Testing"), level=2)
        
        doc.add_paragraph(t.get("mann_whitney_explanation", """
        The analysis uses the Mann-Whitney U test, a non-parametric method for comparing two independent groups.
        A p-value < 0.05 indicates statistically significant differences between boys and girls.
        """))
        
        # Create test results table
        test_table = doc.add_table(rows=len(test_results) + 1, cols=6)  # Variable, Boys Mean, Girls Mean, Difference, Better, p-value, Significant
        test_table.style = 'Table Grid'
        
        # Add headers
        header_cells = test_table.rows[0].cells
        header_cells[0].text = t.get("variable", "Variable")
        header_cells[1].text = t.get("boys_mean", "Boys Mean")
        header_cells[2].text = t.get("girls_mean", "Girls Mean")
        header_cells[3].text = t.get("difference", "Difference")
        header_cells[4].text = t.get("better_gender", "Better Performance")
        header_cells[5].text = t.get("significant", "Significant")
        
        # Add data rows
        for i, result in enumerate(test_results, 1):
            row_cells = test_table.rows[i].cells
            
            row_cells[0].text = result["variable"]
            row_cells[1].text = f"{result['boys_mean']:.2f}" if result['boys_mean'] is not None else "N/A"
            row_cells[2].text = f"{result['girls_mean']:.2f}" if result['girls_mean'] is not None else "N/A"
            row_cells[3].text = f"{result['difference']:.2f}" if result['difference'] is not None else "N/A"
            row_cells[4].text = result["better_gender"] if result["better_gender"] is not None else "N/A"
            
            if result.get('significant') is not None:
                row_cells[5].text = t.get("significant_yes", "Yes") if result['significant'] else t.get("significant_no", "No")
            else:
                row_cells[5].text = "N/A"
    
    # Distribution plots for each variable
    doc.add_heading(t.get("distribution_gender", "Score Distributions by Gender"), level=2)
    
    with tempfile.TemporaryDirectory() as tmp_dir:
        for column in selected_columns:
            col_name = t["columns_of_interest"].get(column, column)
            doc.add_heading(col_name, level=3)
            
            # Create box plot
            fig = px.box(
                df,
                x="gender",
                y=column,
                color="gender",
                labels={
                    "gender": t.get("gender", "Gender"),
                    column: col_name
                },
                color_discrete_map={
                    t.get("boy", "Boy"): "#3498DB",
                    t.get("girl", "Girl"): "#E83E8C"
                }
            )
            
            # Save plot
            img_path = os.path.join(tmp_dir, f"{column}_gender_boxplot.png")
            fig.write_image(img_path, width=800, height=400)
            
            # Add plot to document
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph()
    
    # Summary of gender differences
    doc.add_heading(t.get("gender_summary", "Summary of Gender Differences"), level=2)
    
    if test_results:
        # Count significant differences
        sig_differences = [r for r in test_results if r.get("significant")]
        boy_advantage = [r for r in sig_differences if r.get("better_gender") == t.get("boy", "Boy")]
        girl_advantage = [r for r in sig_differences if r.get("better_gender") == t.get("girl", "Girl")]
        
        if sig_differences:
            doc.add_paragraph(t.get("significant_diff_found", "Significant gender differences were found in {} out of {} variables analyzed.").format(
                len(sig_differences), len(test_results)
            ))
            
            if boy_advantage:
                p = doc.add_paragraph()
                p.add_run(t.get("boy_advantage", "Boys performed significantly better in:")).bold = True
                
                for r in boy_advantage:
                    doc.add_paragraph(f"{r['variable']} ({r['percent_diff']:.1f}% difference)", style='List Bullet')
            
            if girl_advantage:
                p = doc.add_paragraph()
                p.add_run(t.get("girl_advantage", "Girls performed significantly better in:")).bold = True
                
                for r in girl_advantage:
                    doc.add_paragraph(f"{r['variable']} ({r['percent_diff']:.1f}% difference)", style='List Bullet')
            
            # Educational implications
            p = doc.add_paragraph()
            p.add_run(t.get("educational_implications", "Educational Implications:")).bold = True
            
            # Check which types of variables show gender differences
            reading_diffs = [r for r in sig_differences if any(col in r["variable"] for col in ["Letter", "Phon", "Word", "Reading", "Comprehension"])]
            math_diffs = [r for r in sig_differences if any(col in r["variable"] for col in ["Number", "Addition", "Subtraction", "Problem"])]
            
            if reading_diffs:
                doc.add_paragraph(t.get("reading_implications", "Reading skills: Consider gender-responsive teaching strategies to address observed differences"), style='List Bullet')
            
            if math_diffs:
                doc.add_paragraph(t.get("math_implications", "Math skills: Implement targeted interventions to close gender gaps in mathematical performance"), style='List Bullet')
            
            doc.add_paragraph(t.get("review_materials", "Review teaching materials and methods for potential gender bias"), style='List Bullet')
            doc.add_paragraph(t.get("collaborative_learning", "Consider mixed-gender collaborative learning activities"), style='List Bullet')
            doc.add_paragraph(t.get("targeted_support", "Provide targeted support for the lower-performing gender in specific skill areas"), style='List Bullet')
            doc.add_paragraph(t.get("monitor_progress", "Monitor progress to ensure equitable outcomes"), style='List Bullet')
        else:
            doc.add_paragraph(t.get("no_significant_diff", """
            No statistically significant gender differences were found.
            
            This suggests that both boys and girls are performing at similar levels across the assessed skills,
            indicating equitable educational outcomes between genders.
            """))
    
    return doc