import streamlit as st
import pandas as pd
import plotly.express as px
from docx import Document
from docx.shared import Inches
import tempfile
import os
from io import BytesIO
from config import translations

def show_zero_scores(df, language):
    """
    Analyzes and displays the proportion of zero scores for EGRA assessment tasks.
    
    Args:
        df (pandas.DataFrame): The data to analyze
        language (str): Selected language for UI elements (en/fr)
    """
    t = translations[language]  # Get translations for selected language
    
    # Define EGRA and EGMA column lists for organization
    egra_columns = ["clpm", "phoneme", "sound_word", "cwpm", "listening", "orf", "comprehension"]
    egma_columns = ["number_id", "discrimin", "missing_number", "addition", "subtraction", "problems"]
    
    # Ensure the DataFrame contains at least some of the required columns
    available_columns = [col for col in egra_columns + egma_columns if col in df.columns]
    
    if not available_columns:
        st.error(t.get("no_assessment_columns", "No assessment columns found in the data."))
        return
    
    st.subheader(t.get("select_variables", "📊 Select Variables"))
    
    # Create two columns to organize variable selection
    col1, col2 = st.columns(2)
    
    with col1:
        # Filter available EGRA columns
        available_egra = [col for col in egra_columns if col in df.columns]
        selected_egra = st.multiselect(
            t.get("egra_variables", "EGRA Variables:"),
            options=available_egra,
            default=available_egra[:3] if len(available_egra) > 3 else available_egra,
            format_func=lambda x: t["columns_of_interest"].get(x, x)
        )
    
    with col2:
        # Filter available EGMA columns
        available_egma = [col for col in egma_columns if col in df.columns]
        selected_egma = st.multiselect(
            t.get("egma_variables", "EGMA Variables:"),
            options=available_egma,
            default=available_egma[:3] if len(available_egma) > 3 else available_egma,
            format_func=lambda x: t["columns_of_interest"].get(x, x)
        )
    
    # Combine selected variables
    selected_columns = selected_egra + selected_egma
    
    if selected_columns:
        try:
            # Calculate zero scores
            zero_scores = (df[selected_columns] == 0).sum()
            total_students = len(df)
            percentage_zero = ((zero_scores / total_students) * 100).round(2)
            
            # Create DataFrame for display and visualization
            df_zero_scores = pd.DataFrame({
                "Task": [t.get("columns_of_interest", {}).get(col, col) for col in selected_columns],
                "Zero_Count": zero_scores.values,
                "Percentage": percentage_zero.values
            })
            
            # Display results table
            st.subheader(t.get("table_zero_scores", "📋 Proportion of Students with Zero Scores"))
            styled_df = df_zero_scores.copy()
            styled_df["Task_Code"] = selected_columns  # Store original column codes for interpretation
            styled_df.columns = [
                t.get("task_column", "Task"), 
                t.get("count_column", "Count of Zeros"),
                t.get("percentage_column", "Percentage of Zero Scores"),
                "Task_Code"  # Hidden column for interpretation
            ]
            st.dataframe(styled_df[styled_df.columns[:-1]])  # Display without the hidden column
            
            # Export options
            col1, col2, col3 = st.columns(3)
            
            # CSV Export
            with col1:
                csv = styled_df[styled_df.columns[:-1]].to_csv(index=False).encode('utf-8-sig')
                st.download_button(
                    t.get("export_zero_scores_csv", "📥 Download CSV"),
                    csv,
                    "zero_scores.csv",
                    "text/csv",
                    key='download-zero-scores-csv'
                )
            
            # Word Export
            with col2:
                if st.button(t.get("export_zero_scores_word", "📄 Export to Word")):
                    try:
                        doc = create_zero_scores_word_report(styled_df, t)
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                            doc.save(tmp.name)
                            with open(tmp.name, 'rb') as f:
                                docx = f.read()
                            st.download_button(
                                t.get("download_zero_scores_word", "📥 Download Word Report"),
                                docx,
                                "zero_scores_report.docx",
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        os.unlink(tmp.name)
                    except Exception as e:
                        st.error(f"Error creating Word report: {str(e)}")
            
            # Interpretation Generation
            with col3:
                if st.button(t.get("generate_interpretation", "🔍 Generate Interpretation")):
                    # Call interpretation function and display results
                    show_interpretation(styled_df, t)
            
            # Visualization
            st.subheader(t.get("zero_scores_chart_title", "📊 Percentage of Students with Zero Scores by Task"))
            try:
                # Sort by percentage for better visualization
                df_zero_scores_sorted = df_zero_scores.sort_values("Percentage", ascending=True)
                
                fig = px.bar(
                    df_zero_scores_sorted,
                    x="Percentage",
                    y="Task",
                    orientation="h",
                    text="Percentage",
                    color="Percentage",
                    color_continuous_scale="Viridis",
                    title=t.get("zero_scores_chart_title", "Percentage of Students with Zero Scores by Task")
                )
                
                # Define critical thresholds with color indication
                fig.add_vline(x=10, line_width=1, line_dash="dash", line_color="yellow")
                fig.add_vline(x=20, line_width=1, line_dash="dash", line_color="orange")
                fig.add_vline(x=30, line_width=1, line_dash="dash", line_color="red")
                
                # Add annotations for thresholds
                fig.add_annotation(
                    x=10, y=0,
                    text=t.get("acceptable_threshold", "Acceptable"),
                    showarrow=False,
                    yshift=-20
                )
                fig.add_annotation(
                    x=20, y=0,
                    text=t.get("concerning_threshold", "Concerning"),
                    showarrow=False,
                    yshift=-20
                )
                fig.add_annotation(
                    x=30, y=0,
                    text=t.get("critical_threshold", "Critical"),
                    showarrow=False,
                    yshift=-20
                )
                
                fig.update_layout(height=400)
                st.plotly_chart(fig, use_container_width=True)
            except Exception as e:
                st.error(f"Error creating visualization: {str(e)}")
            
        except Exception as e:
            st.error(f"Error in zero scores analysis: {str(e)}")
    
    else:
        st.warning(t.get("warning_select_task", "Please select at least one task to analyze."))

def show_interpretation(df_zero_scores, t):
    """
    Displays educational interpretation of zero scores with recommendations.
    
    Args:
        df_zero_scores (pandas.DataFrame): DataFrame with zero scores analysis
        t (dict): Translation dictionary
    """
    # Create categories based on percentage of zero scores
    categories = {
        "critical": {"tasks": [], "threshold": 30, "color": "🔴"},
        "concerning": {"tasks": [], "threshold": 20, "color": "🟠"},
        "monitor": {"tasks": [], "threshold": 10, "color": "🟡"},
        "acceptable": {"tasks": [], "threshold": 0, "color": "🟢"}
    }
    
    # Categorize tasks based on zero score percentages
    for _, row in df_zero_scores.iterrows():
        percentage = row[t.get("percentage_column", "Percentage of Zero Scores")]
        task_name = row[t.get("task_column", "Task")]
        task_code = row["Task_Code"]
        task_info = {"name": task_name, "code": task_code, "percentage": percentage}
        
        if percentage >= categories["critical"]["threshold"]:
            categories["critical"]["tasks"].append(task_info)
        elif percentage >= categories["concerning"]["threshold"]:
            categories["concerning"]["tasks"].append(task_info)
        elif percentage >= categories["monitor"]["threshold"]:
            categories["monitor"]["tasks"].append(task_info)
        else:
            categories["acceptable"]["tasks"].append(task_info)
    
    # Display interpretation
    st.subheader(t.get("interpretation_title", "📝 Educational Interpretation"))
    
    # Summary based on findings
    st.markdown("### " + t.get("summary_title", "Summary"))
    
    # Generate overall summary
    if categories["critical"]["tasks"]:
        summary_color = "🔴"
        summary_status = t.get("critical_status", "Critical areas requiring immediate intervention")
    elif categories["concerning"]["tasks"]:
        summary_color = "🟠"
        summary_status = t.get("concerning_status", "Areas of concern requiring attention")
    elif categories["monitor"]["tasks"]:
        summary_color = "🟡"
        summary_status = t.get("monitor_status", "Some skills need monitoring")
    else:
        summary_color = "🟢"
        summary_status = t.get("acceptable_status", "All skills are at acceptable levels")
    
    st.markdown(f"{summary_color} **{summary_status}**")
    
    # Generate detailed findings
    if categories["critical"]["tasks"] or categories["concerning"]["tasks"]:
        st.markdown(t.get("findings_text", "Analysis shows significant learning gaps in key reading skills:"))
        
        # Detail critical areas
        if categories["critical"]["tasks"]:
            st.markdown(f"#### {categories['critical']['color']} " + 
                       t.get("critical_areas", "Critical Areas (>{}% zero scores)").format(
                           categories["critical"]["threshold"]))
            for task in categories["critical"]["tasks"]:
                st.markdown(f"- **{task['name']}**: {task['percentage']}% " +
                           t.get("zero_score_text", "of students scored zero"))
        
        # Detail concerning areas
        if categories["concerning"]["tasks"]:
            st.markdown(f"#### {categories['concerning']['color']} " + 
                       t.get("concerning_areas", "Concerning Areas (>{}% zero scores)").format(
                           categories["concerning"]["threshold"]))
            for task in categories["concerning"]["tasks"]:
                st.markdown(f"- **{task['name']}**: {task['percentage']}% " +
                           t.get("zero_score_text", "of students scored zero"))
    
    # Recommendations section
    st.markdown("### " + t.get("recommendations_title", "Recommendations"))
    
    # Define skill-specific recommendations
    recommendations = {
        "clpm": {
            "critical": t.get("clpm_critical", "Implement daily letter recognition activities; Use flashcards and letter games; Provide intensive small-group interventions focusing on alphabet knowledge."),
            "concerning": t.get("clpm_concerning", "Increase letter recognition practice; Include more alphabet activities in regular instruction."),
            "monitor": t.get("clpm_monitor", "Continue regular letter recognition activities while monitoring progress.")
        },
        "phoneme": {
            "critical": t.get("phoneme_critical", "Implement intensive phonemic awareness training; Use sound isolation, blending, and segmentation exercises daily; Provide structured small-group interventions."),
            "concerning": t.get("phoneme_concerning", "Strengthen phonemic awareness instruction; Increase sound manipulation activities in regular classroom work."),
            "monitor": t.get("phoneme_monitor", "Maintain regular phonemic awareness activities and monitor student progress.")
        },
        "sound_word": {
            "critical": t.get("sound_word_critical", "Implement systematic decoding instruction; Use decodable texts that match student skill level; Provide additional small-group intervention focusing on word reading strategies."),
            "concerning": t.get("sound_word_concerning", "Strengthen word reading instruction; Increase practice with decodable texts."),
            "monitor": t.get("sound_word_monitor", "Continue word reading instruction while monitoring student progress.")
        },
        "cwpm": {
            "critical": t.get("cwpm_critical", "Implement intensive fluency training with repeated reading procedures; Use partner reading strategies; Provide small-group guided oral reading interventions."),
            "concerning": t.get("cwpm_concerning", "Increase fluency practice opportunities; Incorporate more guided reading activities."),
            "monitor": t.get("cwpm_monitor", "Maintain regular fluency practice and monitor student progress.")
        },
        "listening": {
            "critical": t.get("listening_critical", "Implement daily listening comprehension activities; Use structured listening exercises with explicit instruction; Provide regular read-alouds with comprehension questions."),
            "concerning": t.get("listening_concerning", "Strengthen listening comprehension instruction; Increase read-aloud activities with discussion."),
            "monitor": t.get("listening_monitor", "Continue regular listening comprehension activities and monitor progress.")
        },
        "orf": {
            "critical": t.get("orf_critical", "Implement intensive oral reading fluency intervention; Use repeated reading techniques; Provide teacher-guided fluency practice in small groups."),
            "concerning": t.get("orf_concerning", "Strengthen oral reading fluency instruction; Increase opportunities for guided oral reading practice."),
            "monitor": t.get("orf_monitor", "Maintain regular oral reading fluency practice and monitor student progress.")
        },
        "comprehension": {
            "critical": t.get("comprehension_critical", "Implement explicit reading comprehension strategy instruction; Use scaffolded reading experiences; Provide small-group interventions focusing on comprehension strategies."),
            "concerning": t.get("comprehension_concerning", "Strengthen comprehension strategy instruction; Increase guided reading with comprehension focus."),
            "monitor": t.get("comprehension_monitor", "Continue comprehension strategy instruction while monitoring progress.")
        }
    }
    
    # Generate recommendations for critical and concerning areas
    if categories["critical"]["tasks"]:
        st.markdown(f"#### {t.get('priority_recommendations', 'Priority Recommendations')}")
        for task in categories["critical"]["tasks"]:
            if task["code"] in recommendations:
                st.markdown(f"**{task['name']}**:")
                st.markdown(f"{recommendations[task['code']]['critical']}")
                st.markdown("")
    
    if categories["concerning"]["tasks"]:
        st.markdown(f"#### {t.get('secondary_recommendations', 'Secondary Recommendations')}")
        for task in categories["concerning"]["tasks"]:
            if task["code"] in recommendations:
                st.markdown(f"**{task['name']}**:")
                st.markdown(f"{recommendations[task['code']]['concerning']}")
                st.markdown("")
    
    # Additional recommendations based on reading development stages
    st.markdown("#### " + t.get("implementation_strategies", "Implementation Strategies"))
    
    # Determine the focus area (decode, fluency, or comprehension) based on the most critical skills
    decode_skills = ["clpm", "phoneme", "sound_word"]
    fluency_skills = ["cwpm", "orf"]
    comprehension_skills = ["listening", "comprehension"]
    
    critical_codes = [task["code"] for task in categories["critical"]["tasks"]]
    concerning_codes = [task["code"] for task in categories["concerning"]["tasks"]]
    
    # Identify which reading development stage needs most attention
    if any(skill in critical_codes for skill in decode_skills):
        st.markdown(t.get("decode_focus", "**Focus on Foundational Skills Development:**"))
        st.markdown("- " + t.get("decode_strat1", "Implement a systematic and explicit phonics program"))
        st.markdown("- " + t.get("decode_strat2", "Use multisensory teaching approaches for letter-sound relationships"))
        st.markdown("- " + t.get("decode_strat3", "Provide daily opportunities for guided practice with immediate feedback"))
    
    if any(skill in critical_codes for skill in fluency_skills):
        st.markdown(t.get("fluency_focus", "**Focus on Reading Fluency Development:**"))
        st.markdown("- " + t.get("fluency_strat1", "Implement repeated reading procedures with appropriate texts"))
        st.markdown("- " + t.get("fluency_strat2", "Use choral reading and partner reading strategies"))
        st.markdown("- " + t.get("fluency_strat3", "Provide regular opportunities for performance reading"))
    
    if any(skill in critical_codes for skill in comprehension_skills):
        st.markdown(t.get("comprehension_focus", "**Focus on Comprehension Strategy Development:**"))
        st.markdown("- " + t.get("comprehension_strat1", "Explicitly teach comprehension strategies (predicting, questioning, summarizing)"))
        st.markdown("- " + t.get("comprehension_strat2", "Use graphic organizers to support understanding of text structure"))
        st.markdown("- " + t.get("comprehension_strat3", "Implement regular discussion groups around text meaning"))
    
    # Assessment recommendations
    st.markdown("#### " + t.get("assessment_recommendations", "Assessment Recommendations"))
    st.markdown("- " + t.get("assessment_rec1", "Conduct regular progress monitoring assessments (every 2-3 weeks for critical skills)"))
    st.markdown("- " + t.get("assessment_rec2", "Use formative assessments to adjust instruction as needed"))
    st.markdown("- " + t.get("assessment_rec3", "Re-assess all skills after 8-10 weeks of intervention"))

def create_zero_scores_word_report(df_zero_scores, t):
    """
    Creates a Word report with zero scores analysis and recommendations.
    
    Args:
        df_zero_scores (pandas.DataFrame): DataFrame with zero scores analysis
        t (dict): Translation dictionary
        
    Returns:
        docx.Document: Word document with the report
    """
    doc = Document()
    
    # Main title
    doc.add_heading(t.get("title_zero_scores", "Zero Scores Analysis"), level=1)
    
    # Data summary section
    doc.add_heading(t.get("table_zero_scores", "Proportion of Students with Zero Scores"), level=2)
    
    # Create table for zero scores data
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = t.get("task_column", "Task")
    header_cells[1].text = t.get("count_column", "Count of Zeros")
    header_cells[2].text = t.get("percentage_column", "Percentage of Zero Scores")
    
    # Add rows
    for _, row in df_zero_scores.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row[t.get("task_column", "Task")]
        row_cells[1].text = str(row[t.get("count_column", "Count of Zeros")])
        row_cells[2].text = f"{row[t.get('percentage_column', 'Percentage of Zero Scores')]}%"
    
    # Create categories for interpretation
    categories = {
        "critical": {"tasks": [], "threshold": 30},
        "concerning": {"tasks": [], "threshold": 20},
        "monitor": {"tasks": [], "threshold": 10},
        "acceptable": {"tasks": [], "threshold": 0}
    }
    
    # Categorize tasks based on zero score percentages
    for _, row in df_zero_scores.iterrows():
        percentage = row[t.get("percentage_column", "Percentage of Zero Scores")]
        task_name = row[t.get("task_column", "Task")]
        task_code = row["Task_Code"]
        task_info = {"name": task_name, "code": task_code, "percentage": percentage}
        
        if percentage >= categories["critical"]["threshold"]:
            categories["critical"]["tasks"].append(task_info)
        elif percentage >= categories["concerning"]["threshold"]:
            categories["concerning"]["tasks"].append(task_info)
        elif percentage >= categories["monitor"]["threshold"]:
            categories["monitor"]["tasks"].append(task_info)
        else:
            categories["acceptable"]["tasks"].append(task_info)
    
    # Add interpretation section
    doc.add_heading(t.get("interpretation_title", "Educational Interpretation"), level=2)
    
    # Add summary paragraph
    if categories["critical"]["tasks"]:
        summary_status = t.get("critical_status", "Critical areas requiring immediate intervention")
    elif categories["concerning"]["tasks"]:
        summary_status = t.get("concerning_status", "Areas of concern requiring attention")
    elif categories["monitor"]["tasks"]:
        summary_status = t.get("monitor_status", "Some skills need monitoring")
    else:
        summary_status = t.get("acceptable_status", "All skills are at acceptable levels")
    
    doc.add_paragraph(summary_status, style='Intense Quote')
    
    # Add detailed findings
    if categories["critical"]["tasks"] or categories["concerning"]["tasks"]:
        doc.add_paragraph(t.get("findings_text", "Analysis shows significant learning gaps in key reading skills:"))
        
        # Detail critical areas
        if categories["critical"]["tasks"]:
            doc.add_heading(t.get("critical_areas", "Critical Areas (>{}% zero scores)").format(
                           categories["critical"]["threshold"]), level=3)
            for task in categories["critical"]["tasks"]:
                doc.add_paragraph(f"{task['name']}: {task['percentage']}% " +
                                 t.get("zero_score_text", "of students scored zero"), style='List Bullet')
        
        # Detail concerning areas
        if categories["concerning"]["tasks"]:
            doc.add_heading(t.get("concerning_areas", "Concerning Areas (>{}% zero scores)").format(
                           categories["concerning"]["threshold"]), level=3)
            for task in categories["concerning"]["tasks"]:
                doc.add_paragraph(f"{task['name']}: {task['percentage']}% " +
                                 t.get("zero_score_text", "of students scored zero"), style='List Bullet')
    
    # Recommendations section
    doc.add_heading(t.get("recommendations_title", "Recommendations"), level=2)
    
    # Define skill-specific recommendations
    recommendations = {
        "clpm": {
            "critical": t.get("clpm_critical", "Implement daily letter recognition activities; Use flashcards and letter games; Provide intensive small-group interventions focusing on alphabet knowledge."),
            "concerning": t.get("clpm_concerning", "Increase letter recognition practice; Include more alphabet activities in regular instruction."),
            "monitor": t.get("clpm_monitor", "Continue regular letter recognition activities while monitoring progress.")
        },
        "phoneme": {
            "critical": t.get("phoneme_critical", "Implement intensive phonemic awareness training; Use sound isolation, blending, and segmentation exercises daily; Provide structured small-group interventions."),
            "concerning": t.get("phoneme_concerning", "Strengthen phonemic awareness instruction; Increase sound manipulation activities in regular classroom work."),
            "monitor": t.get("phoneme_monitor", "Maintain regular phonemic awareness activities and monitor student progress.")
        },
        "sound_word": {
            "critical": t.get("sound_word_critical", "Implement systematic decoding instruction; Use decodable texts that match student skill level; Provide additional small-group intervention focusing on word reading strategies."),
            "concerning": t.get("sound_word_concerning", "Strengthen word reading instruction; Increase practice with decodable texts."),
            "monitor": t.get("sound_word_monitor", "Continue word reading instruction while monitoring student progress.")
        },
        "cwpm": {
            "critical": t.get("cwpm_critical", "Implement intensive fluency training with repeated reading procedures; Use partner reading strategies; Provide small-group guided oral reading interventions."),
            "concerning": t.get("cwpm_concerning", "Increase fluency practice opportunities; Incorporate more guided reading activities."),
            "monitor": t.get("cwpm_monitor", "Maintain regular fluency practice and monitor student progress.")
        },
        "listening": {
            "critical": t.get("listening_critical", "Implement daily listening comprehension activities; Use structured listening exercises with explicit instruction; Provide regular read-alouds with comprehension questions."),
            "concerning": t.get("listening_concerning", "Strengthen listening comprehension instruction; Increase read-aloud activities with discussion."),
            "monitor": t.get("listening_monitor", "Continue regular listening comprehension activities and monitor progress.")
        },
        "orf": {
            "critical": t.get("orf_critical", "Implement intensive oral reading fluency intervention; Use repeated reading techniques; Provide teacher-guided fluency practice in small groups."),
            "concerning": t.get("orf_concerning", "Strengthen oral reading fluency instruction; Increase opportunities for guided oral reading practice."),
            "monitor": t.get("orf_monitor", "Maintain regular oral reading fluency practice and monitor student progress.")
        },
        "comprehension": {
            "critical": t.get("comprehension_critical", "Implement explicit reading comprehension strategy instruction; Use scaffolded reading experiences; Provide small-group interventions focusing on comprehension strategies."),
            "concerning": t.get("comprehension_concerning", "Strengthen comprehension strategy instruction; Increase guided reading with comprehension focus."),
            "monitor": t.get("comprehension_monitor", "Continue comprehension strategy instruction while monitoring progress.")
        }
    }
    
    # Generate recommendations for critical and concerning areas
    if categories["critical"]["tasks"]:
        doc.add_heading(t.get('priority_recommendations', 'Priority Recommendations'), level=3)
        for task in categories["critical"]["tasks"]:
            if task["code"] in recommendations:
                para = doc.add_paragraph()
                para.add_run(f"{task['name']}:").bold = True
                doc.add_paragraph(f"{recommendations[task['code']]['critical']}")
    
    if categories["concerning"]["tasks"]:
        doc.add_heading(t.get('secondary_recommendations', 'Secondary Recommendations'), level=3)
        for task in categories["concerning"]["tasks"]:
            if task["code"] in recommendations:
                para = doc.add_paragraph()
                para.add_run(f"{task['name']}:").bold = True
                doc.add_paragraph(f"{recommendations[task['code']]['concerning']}")
    
    # Create visualization for Word document
    with tempfile.TemporaryDirectory() as tmp_dir:
        # Sort by percentage for better visualization
        df_visualization = df_zero_scores.copy()
        df_visualization = df_visualization.sort_values(t.get("percentage_column", "Percentage of Zero Scores"), ascending=True)
        
        try:
            fig = px.bar(
                df_visualization,
                x=t.get("percentage_column", "Percentage of Zero Scores"),
                y=t.get("task_column", "Task"),
                orientation="h",
                text=t.get("percentage_column", "Percentage of Zero Scores"),
                color=t.get("percentage_column", "Percentage of Zero Scores"),
                color_continuous_scale="Viridis",
                title=t.get("zero_scores_chart_title", "Percentage of Students with Zero Scores by Task")
            )
            
            # Define critical thresholds with color indication
            fig.add_vline(x=10, line_width=1, line_dash="dash", line_color="yellow")
            fig.add_vline(x=20, line_width=1, line_dash="dash", line_color="orange")
            fig.add_vline(x=30, line_width=1, line_dash="dash", line_color="red")
            
            # Add annotations for thresholds
            fig.add_annotation(
                x=10, y=0,
                text=t.get("acceptable_threshold", "Acceptable"),
                showarrow=False,
                yshift=-20
            )
            fig.add_annotation(
                x=20, y=0,
                text=t.get("concerning_threshold", "Concerning"),
                showarrow=False,
                yshift=-20
            )
            fig.add_annotation(
                x=30, y=0,
                text=t.get("critical_threshold", "Critical"),
                showarrow=False,
                yshift=-20
            )
            
            # Save visualization
            img_path = os.path.join(tmp_dir, "zero_scores_chart.png")
            fig.write_image(img_path, width=1000, height=600)
            
            # Add visualization to document
            doc.add_heading(t.get("visualization_title", "Visualization"), level=2)
            doc.add_picture(img_path, width=Inches(6))
            
        except Exception as e:
            doc.add_paragraph(f"Error creating visualization: {str(e)}")
    
    return doc