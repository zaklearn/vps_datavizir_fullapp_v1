import os
import tempfile
from datetime import datetime
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class WordReportGenerator:
    """
    Comprehensive system for generating professional Word reports from analysis results.
    
    This class provides methods to create Word documents with formatted tables,
    embedded visualizations, headers, footers, and appropriate text styling.
    It supports both English and French content and includes customizable templates
    for different report types.
    """
    
    def __init__(self, language="en"):
        """
        Initialize the Word report generator.
        
        Args:
            language (str): The language for the report ("en" or "fr")
        """
        self.language = language
        self.doc = None
        self.temp_dir = None
        self.image_count = 0
        self.table_count = 0
        
        # Translation dictionary for report elements
        self.translations = {
            "en": {
                "report_title": "Assessment Analysis Report",
                "table_of_contents": "Table of Contents",
                "executive_summary": "Executive Summary",
                "methodology": "Methodology",
                "results": "Results",
                "interpretation": "Interpretation",
                "recommendations": "Recommendations",
                "conclusion": "Conclusion",
                "appendix": "Appendix",
                "figure": "Figure",
                "table": "Table",
                "page": "Page",
                "of": "of",
                "date_generated": "Report generated on",
                "confidential": "CONFIDENTIAL",
                "statistical_overview": "Statistical Overview",
                "zero_scores_analysis": "Zero Scores Analysis",
                "school_comparison": "School Comparison",
                "gender_analysis": "Gender Analysis",
                "international_comparison": "International Standards Comparison",
                "mean": "Mean",
                "median": "Median",
                "std_dev": "Standard Deviation",
                "min": "Minimum",
                "max": "Maximum",
                "interpretation_high": "High performance area",
                "interpretation_average": "Average performance area",
                "interpretation_low": "Low performance area, improvement needed",
                "interpretation_critical": "Critical area requiring immediate attention",
                "recommendation_high": "Maintain current instructional practices",
                "recommendation_average": "Continue regular instruction with some enhancements",
                "recommendation_low": "Implement targeted interventions",
                "recommendation_critical": "Comprehensive intervention program required"
            },
            "fr": {
                "report_title": "Rapport d'Analyse d'Évaluation",
                "table_of_contents": "Table des Matières",
                "executive_summary": "Résumé Exécutif",
                "methodology": "Méthodologie",
                "results": "Résultats",
                "interpretation": "Interprétation",
                "recommendations": "Recommandations",
                "conclusion": "Conclusion",
                "appendix": "Annexe",
                "figure": "Figure",
                "table": "Tableau",
                "page": "Page",
                "of": "sur",
                "date_generated": "Rapport généré le",
                "confidential": "CONFIDENTIEL",
                "statistical_overview": "Aperçu Statistique",
                "zero_scores_analysis": "Analyse des Scores Zéro",
                "school_comparison": "Comparaison entre Écoles",
                "gender_analysis": "Analyse selon le Genre",
                "international_comparison": "Comparaison aux Standards Internationaux",
                "mean": "Moyenne",
                "median": "Médiane",
                "std_dev": "Écart-Type",
                "min": "Minimum",
                "max": "Maximum",
                "interpretation_high": "Zone de performance élevée",
                "interpretation_average": "Zone de performance moyenne",
                "interpretation_low": "Zone de performance faible, amélioration nécessaire",
                "interpretation_critical": "Zone critique nécessitant une attention immédiate",
                "recommendation_high": "Maintenir les pratiques d'enseignement actuelles",
                "recommendation_average": "Poursuivre l'enseignement régulier avec quelques améliorations",
                "recommendation_low": "Mettre en œuvre des interventions ciblées",
                "recommendation_critical": "Programme d'intervention complet requis"
            }
        }
    
    def get_text(self, key, default=""):
        """Get translated text for the current language."""
        return self.translations.get(self.language, {}).get(key, default)
    
    def format_date(self, date=None):
        """Format date according to language conventions."""
        if date is None:
            date = datetime.now()
            
        if self.language == "fr":
            # French date format
            return date.strftime("%d %B %Y").replace("January", "Janvier").replace("February", "Février") \
                .replace("March", "Mars").replace("April", "Avril").replace("May", "Mai") \
                .replace("June", "Juin").replace("July", "Juillet").replace("August", "Août") \
                .replace("September", "Septembre").replace("October", "Octobre") \
                .replace("November", "Novembre").replace("December", "Décembre")
        else:
            # English date format
            return date.strftime("%B %d, %Y")
    
    def format_number(self, number, decimal_places=2):
        """Format number according to language conventions."""
        if self.language == "fr":
            # French number format (comma for decimal separator, space for thousands)
            formatted = f"{number:.{decimal_places}f}".replace(".", ",")
            parts = formatted.split(",")
            if len(parts[0]) > 3:
                # Add thousand separators
                parts[0] = " ".join([parts[0][max(0, i-3):i] for i in range(len(parts[0]), 0, -3)][::-1])
            return ",".join(parts)
        else:
            # English number format (period for decimal separator, comma for thousands)
            return f"{number:,.{decimal_places}f}"
    
    def create_new_report(self, title=None, include_toc=True, confidential=False):
        """
        Create a new Word document with initial setup.
        
        Args:
            title (str, optional): Report title
            include_toc (bool): Whether to include a table of contents
            confidential (bool): Whether to mark the document as confidential
            
        Returns:
            Document: The created document
        """
        # Create new document
        self.doc = Document()
        
        # Create temporary directory for images
        if self.temp_dir is None:
            self.temp_dir = tempfile.TemporaryDirectory()
        
        # Reset counters
        self.image_count = 0
        self.table_count = 0
        
        # Set up default styles
        self._setup_styles()
        
        # Add title page
        if title is None:
            title = self.get_text("report_title")
        
        self._add_title_page(title, confidential)
        
        # Add table of contents if requested
        if include_toc:
            self._add_table_of_contents()
        
        # Add page break after TOC
        self.doc.add_page_break()
        
        return self.doc
    
    def _setup_styles(self):
        """Set up document styles for consistent formatting."""
        # Modify built-in styles
        styles = self.doc.styles
        
        # Title style
        title_style = styles['Title']
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(28)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 62, 117)  # Dark blue
        
        # Heading styles
        for i in range(1, 4):
            heading_style = styles[f'Heading {i}']
            heading_style.font.name = 'Calibri'
            heading_style.font.size = Pt(20 - (i * 2))  # 18, 16, 14
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(0, 82, 147)  # Blue
            
        # Normal text style
        normal_style = styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(8)
        normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # Caption style for figures and tables
        caption_style = styles['Caption']
        caption_style.font.name = 'Calibri'
        caption_style.font.size = Pt(10)
        caption_style.font.italic = True
        
        # Create custom styles for interpretation sections
        self._create_custom_style('Interpretation', 'Calibri', 11, True, RGBColor(0, 128, 0))  # Green
        self._create_custom_style('Recommendation', 'Calibri', 11, True, RGBColor(192, 0, 0))  # Red
        self._create_custom_style('Highlight', 'Calibri', 11, True, RGBColor(128, 0, 128))  # Purple
    
    def _create_custom_style(self, name, font_name, font_size, is_bold, color):
        """
        Create a custom paragraph style.
        
        Args:
            name (str): Style name
            font_name (str): Font name
            font_size (int): Font size in points
            is_bold (bool): Whether the font is bold
            color (RGBColor): Font color
        """
        styles = self.doc.styles
        
        # Check if style already exists
        try:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except:
            # Style already exists, get it
            style = styles[name]
        
        # Set style properties
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.font.bold = is_bold
        style.font.color.rgb = color
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    def _add_title_page(self, title, confidential=False):
        """
        Add a title page to the document.
        
        Args:
            title (str): Report title
            confidential (bool): Whether to mark the document as confidential
        """
        # Add title
        title_paragraph = self.doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 62, 117)  # Dark blue
        
        # Add space
        for _ in range(10):
            self.doc.add_paragraph()
        
        # Add date
        date_paragraph = self.doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_text = f"{self.get_text('date_generated')}: {self.format_date()}"
        date_run = date_paragraph.add_run(date_text)
        date_run.font.name = 'Calibri'
        date_run.font.size = Pt(12)
        date_run.font.bold = True
        
        # Add confidential marking if requested
        if confidential:
            conf_paragraph = self.doc.add_paragraph()
            conf_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            conf_run = conf_paragraph.add_run(self.get_text('confidential'))
            conf_run.font.name = 'Calibri'
            conf_run.font.size = Pt(14)
            conf_run.font.bold = True
            conf_run.font.color.rgb = RGBColor(192, 0, 0)  # Red
        
        # Add page break after title page
        self.doc.add_page_break()
    
    def _add_table_of_contents(self):
        """Add a table of contents to the document."""
        # Add heading for TOC
        self.doc.add_heading(self.get_text('table_of_contents'), level=1)
        
        # Add paragraph that will be used for TOC (will be populated when document is opened in Word)
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        
        # Add TOC field
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar3)
    
    def add_section(self, title, level=1):
        """
        Add a new section to the report.
        
        Args:
            title (str): Section title
            level (int): Heading level (1-3)
        """
        # Ensure valid heading level
        if level < 1:
            level = 1
        elif level > 3:
            level = 3
            
        # Add heading
        self.doc.add_heading(title, level=level)
    
    def add_executive_summary(self, text):
        """
        Add an executive summary section to the report.
        
        Args:
            text (str): Executive summary text
        """
        self.add_section(self.get_text('executive_summary'), level=1)
        self.doc.add_paragraph(text)
    
    def add_methodology(self, text):
        """
        Add a methodology section to the report.
        
        Args:
            text (str): Methodology text
        """
        self.add_section(self.get_text('methodology'), level=1)
        self.doc.add_paragraph(text)
    
    def add_paragraph(self, text, style=None):
        """
        Add a paragraph to the report.
        
        Args:
            text (str): Paragraph text
            style (str, optional): Paragraph style name
        """
        paragraph = self.doc.add_paragraph(text)
        if style is not None:
            paragraph.style = style
        return paragraph
    
    def add_bullet_point(self, text, level=0, style=None):
        """
        Add a bullet point to the report.
        
        Args:
            text (str): Bullet point text
            level (int): Indentation level (0-2)
            style (str, optional): Paragraph style name
        """
        paragraph = self.doc.add_paragraph(text, style='List Bullet')
        if style is not None:
            paragraph.style = style
            
        # Adjust indentation for nested bullets
        if level > 0:
            paragraph.paragraph_format.left_indent = Inches(0.25 * level)
            
        return paragraph
    
    def add_numbered_point(self, text, level=0, style=None):
        """
        Add a numbered point to the report.
        
        Args:
            text (str): Numbered point text
            level (int): Indentation level (0-2)
            style (str, optional): Paragraph style name
        """
        paragraph = self.doc.add_paragraph(text, style='List Number')
        if style is not None:
            paragraph.style = style
            
        # Adjust indentation for nested numbering
        if level > 0:
            paragraph.paragraph_format.left_indent = Inches(0.25 * level)
            
        return paragraph
    
    def add_table(self, data, headers=None, title=None, autofit=True, style='Table Grid'):
        """
        Add a table to the report.
        
        Args:
            data (list or pandas.DataFrame): Table data
            headers (list, optional): Column headers
            title (str, optional): Table title/caption
            autofit (bool): Whether to autofit columns to content
            style (str): Table style name
            
        Returns:
            Table: The created table
        """
        # Convert pandas DataFrame to list format if needed
        if hasattr(data, 'columns') and hasattr(data, 'values'):
            # It's a pandas DataFrame
            if headers is None:
                headers = [str(col) for col in data.columns]
            data_values = data.values.tolist()
            
            # Handle DataFrames with index
            if data.index.name is not None or not all(isinstance(idx, int) and idx == pos for pos, idx in enumerate(data.index)):
                # Add index as first column
                if headers:
                    headers.insert(0, str(data.index.name) if data.index.name is not None else "Index")
                data_values = [[idx] + row for idx, row in zip(data.index, data_values)]
        else:
            data_values = data
        
        # Determine table dimensions
        rows = len(data_values) + (1 if headers else 0)
        cols = len(headers) if headers else (len(data_values[0]) if data_values else 0)
        
        if rows == 0 or cols == 0:
            # No data to display
            self.doc.add_paragraph("No data available for table.")
            return None
        
        # Create table
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = style
        
        # Add headers if provided
        if headers:
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = str(header)
                # Format header cell
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Add data rows
        start_row = 1 if headers else 0
        for i, row_data in enumerate(data_values):
            row = table.rows[i + start_row]
            for j, value in enumerate(row_data):
                cell = row.cells[j]
                
                # Format numbers according to language conventions
                if isinstance(value, (int, float)) and not isinstance(value, bool):
                    cell.text = self.format_number(value)
                else:
                    cell.text = str(value)
        
        # Autofit columns to content if requested
        if autofit:
            table.autofit = True
        
        # Add caption if provided
        if title:
            self.table_count += 1
            caption_text = f"{self.get_text('table')} {self.table_count}: {title}"
            caption = self.doc.add_paragraph(caption_text, style='Caption')
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return table
    
    def add_figure(self, figure, title=None, width=6, height=None, dpi=300):
        """
        Add a matplotlib figure to the report.
        
        Args:
            figure (matplotlib.figure.Figure): Matplotlib figure
            title (str, optional): Figure title/caption
            width (float): Width in inches
            height (float, optional): Height in inches
            dpi (int): Resolution in dots per inch
            
        Returns:
            str: Path to the saved image
        """
        if self.temp_dir is None:
            self.temp_dir = tempfile.TemporaryDirectory()
            
        # Create unique filename
        self.image_count += 1
        img_path = os.path.join(self.temp_dir.name, f"figure_{self.image_count}.png")
        
        # Save figure to file
        if height:
            figure.set_size_inches(width, height)
        else:
            figure.set_size_inches(width, width * 0.75)  # 4:3 aspect ratio if height not specified
            
        figure.savefig(img_path, dpi=dpi, bbox_inches='tight')
        plt.close(figure)
        
        # Add figure to document
        self.doc.add_picture(img_path, width=Inches(width))
        
        # Add caption if provided
        if title:
            caption_text = f"{self.get_text('figure')} {self.image_count}: {title}"
            caption = self.doc.add_paragraph(caption_text, style='Caption')
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return img_path
    
    def add_plotly_figure(self, fig, title=None, width=6, height=4, scale=1):
        """
        Add a plotly figure to the report.
        
        Args:
            fig (plotly.graph_objects.Figure): Plotly figure
            title (str, optional): Figure title/caption
            width (float): Width in inches
            height (float): Height in inches
            scale (float): Scale factor for resolution
            
        Returns:
            str: Path to the saved image
        """
        if self.temp_dir is None:
            self.temp_dir = tempfile.TemporaryDirectory()
            
        # Create unique filename
        self.image_count += 1
        img_path = os.path.join(self.temp_dir.name, f"figure_{self.image_count}.png")
        
        # Export plotly figure as image
        fig.write_image(img_path, width=width*100, height=height*100, scale=scale)
        
        # Add figure to document
        self.doc.add_picture(img_path, width=Inches(width))
        
        # Add caption if provided
        if title:
            caption_text = f"{self.get_text('figure')} {self.image_count}: {title}"
            caption = self.doc.add_paragraph(caption_text, style='Caption')
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return img_path
    
    def add_image(self, image_path, title=None, width=6):
        """
        Add an image from file to the report.
        
        Args:
            image_path (str): Path to the image file
            title (str, optional): Image title/caption
            width (float): Width in inches
            
        Returns:
            None
        """
        # Add image to document
        self.image_count += 1
        self.doc.add_picture(image_path, width=Inches(width))
        
        # Add caption if provided
        if title:
            caption_text = f"{self.get_text('figure')} {self.image_count}: {title}"
            caption = self.doc.add_paragraph(caption_text, style='Caption')
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_page_break(self):
        """Add a page break to the report."""
        self.doc.add_page_break()
    
    def add_interpretation(self, text, indicator=None, score=None):
        """
        Add an interpretation section with optional score-based formatting.
        
        Args:
            text (str): Interpretation text
            indicator (str, optional): Name of the indicator being interpreted
            score (float, optional): Score value for color-coded interpretation
            
        Returns:
            Paragraph: The created paragraph
        """
        # Add heading if indicator name is provided
        if indicator:
            self.doc.add_heading(f"{self.get_text('interpretation')}: {indicator}", level=3)
        
        # Determine interpretation style based on score
        style = 'Normal'
        if score is not None:
            if score >= 80:
                text = f"{self.get_text('interpretation_high')}: {text}"
                style = 'Interpretation'
            elif score >= 60:
                text = f"{self.get_text('interpretation_average')}: {text}"
                style = 'Interpretation'
            elif score >= 40:
                text = f"{self.get_text('interpretation_low')}: {text}"
                style = 'Highlight'
            else:
                text = f"{self.get_text('interpretation_critical')}: {text}"
                style = 'Recommendation'
        
        # Add interpretation paragraph
        return self.add_paragraph(text, style=style)
    
    def add_recommendation(self, text, indicator=None, score=None):
        """
        Add a recommendation section with optional score-based formatting.
        
        Args:
            text (str): Recommendation text
            indicator (str, optional): Name of the indicator
            score (float, optional): Score value for different recommendations
            
        Returns:
            Paragraph: The created paragraph
        """
        # Add heading if indicator name is provided
        if indicator:
            self.doc.add_heading(f"{self.get_text('recommendations')}: {indicator}", level=3)
        
        # Determine recommendation based on score
        if score is not None:
            if score >= 80:
                recommendation = self.get_text('recommendation_high')
            elif score >= 60:
                recommendation = self.get_text('recommendation_average')
            elif score >= 40:
                recommendation = self.get_text('recommendation_low')
            else:
                recommendation = self.get_text('recommendation_critical')
                
            text = f"{recommendation}: {text}"
        
        # Add recommendation paragraph
        return self.add_paragraph(text, style='Recommendation')
    
    def add_appendix(self, title, content):
        """
        Add an appendix section to the report.
        
        Args:
            title (str): Appendix title
            content (str): Appendix content
            
        Returns:
            None
        """
        # Add page break before appendix
        self.doc.add_page_break()
        
        # Add appendix title
        self.doc.add_heading(f"{self.get_text('appendix')}: {title}", level=1)
        
        # Add content
        self.doc.add_paragraph(content)
    
    def setup_headers_and_footers(self, include_page_numbers=True, include_title=True, title=None):
        """
        Set up headers and footers for the document.
        
        Args:
            include_page_numbers (bool): Whether to include page numbers in footer
            include_title (bool): Whether to include title in header
            title (str, optional): Document title for header
            
        Returns:
            None
        """
        # Get sections
        sections = self.doc.sections
        
        # Setup for each section
        for section in sections:
            # Header
            if include_title:
                header = section.header
                header_paragraph = header.paragraphs[0]
                header_text = title if title else self.get_text('report_title')
                header_paragraph.text = header_text
                header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Format header text
                for run in header_paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
            
            # Footer
            if include_page_numbers:
                footer = section.footer
                footer_paragraph = footer.paragraphs[0]
                
                # Add page numbers field code
                page_nums_fmt = f"{self.get_text('page')} {{0}} {self.get_text('of')} {{1}}"
                add_page_number(footer_paragraph, page_nums_fmt)
                footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Format footer text
                for run in footer_paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
    
    def save(self, filename):
        """
        Save the report to a file.
        
        Args:
            filename (str): Output filename
            
        Returns:
            None
        """
        # Make sure directory exists
        os.makedirs(os.path.dirname(os.path.abspath(filename)), exist_ok=True)
        
        # Save document
        self.doc.save(filename)
        
        # Clean up temporary files
        self.cleanup()
    
    def get_document(self):
        """
        Get the document object.
        
        Returns:
            Document: The document object
        """
        return self.doc
    
    def cleanup(self):
        """Clean up temporary files."""
        if self.temp_dir:
            try:
                self.temp_dir.cleanup()
                self.temp_dir = None
            except:
                pass


# Helper function to add page numbers
def add_page_number(paragraph, fmt="{0} of {1}"):
    """
    Add page numbers field code to a paragraph.
    
    Args:
        paragraph: The paragraph to add page numbers to
        fmt (str): Format string for page numbers
        
    Returns:
        None
    """
    run = paragraph.add_run()
    
    # Add field codes for page numbers
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText1 = OxmlElement('w:instrText')
    instrText1.set(qn('xml:space'), 'preserve')
    instrText1.text = f' PAGE '
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    # Add "of" text
    text_run = paragraph.add_run(f" {fmt.split('{0}')[1].split('{1}')[0]} ")
    
    # Add total pages field
    run2 = paragraph.add_run()
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    
    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = f' NUMPAGES '
    
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    # Assemble all elements
    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText1)
    r_element.append(fldChar2)
    
    r_element2 = run2._r
    r_element2.append(fldChar3)
    r_element2.append(instrText2)
    r_element2.append(fldChar4)


# Report template classes for specific analysis types
class StatisticalReportTemplate:
    """Template for statistical analysis reports."""
    
    @staticmethod
    def create_report(report_gen, results, title=None):
        """
        Create a statistical analysis report.
        
        Args:
            report_gen (WordReportGenerator): Report generator instance
            results (dict): Statistical analysis results
            title (str, optional): Report title
            
        Returns:
            Document: The created document
        """
        # Get language
        language = report_gen.language
        
        # Create new document
        if title is None:
            title = report_gen.get_text('statistical_overview')
        doc = report_gen.create_new_report(title)
        
        # Add executive summary
        summary_text = results.get('summary', 'This report provides a statistical overview of the assessment results.')
        report_gen.add_executive_summary(summary_text)
        
        # Add methodology section
        methodology_text = (
            "This analysis calculates descriptive statistics for the selected assessment variables. "
            "It includes measures of central tendency (mean, median), dispersion (standard deviation, range), "
            "and distribution characteristics for each variable."
        )
        report_gen.add_methodology(methodology_text)
        
        # Add results section
        report_gen.add_section(report_gen.get_text('results'), level=1)
        
        # Add statistics table
        stats_data = results.get('statistics')
        if stats_data is not None:
            report_gen.add_table(
                stats_data, 
                title=report_gen.get_text('descriptive_statistics', 'Descriptive Statistics')
            )
        
        # Add visualizations
        visualizations = results.get('visualizations', [])
        for viz in visualizations:
            if 'figure' in viz and 'title' in viz:
                report_gen.add_plotly_figure(viz['figure'], viz['title'])
        
        # Add interpretations
        report_gen.add_section(report_gen.get_text('interpretation'), level=1)
        interpretations = results.get('interpretations', [])
        for interp in interpretations:
            indicator = interp.get('indicator')
            text = interp.get('text', '')
            score = interp.get('score')
            report_gen.add_interpretation(text, indicator, score)
        
        # Add recommendations
        report_gen.add_section(report_gen.get_text('recommendations'), level=1)
        recommendations = results.get('recommendations', [])
        for rec in recommendations:
            text = rec.get('text', '')
            indicator = rec.get('indicator')
            score = rec.get('score')
            report_gen.add_recommendation(text, indicator, score)
        
        # Set up headers and footers
        report_gen.setup_headers_and_footers(title=title)
        
        return doc


class ZeroScoresReportTemplate:
    """Template for zero scores analysis reports."""
    
    @staticmethod
    def create_report(report_gen, results, title=None):
        """
        Create a zero scores analysis report.
        
        Args:
            report_gen (WordReportGenerator): Report generator instance
            results (dict): Zero scores analysis results
            title (str, optional): Report title
            
        Returns:
            Document: The created document
        """
        # Get language
        language = report_gen.language
        
        # Create new document
        if title is None:
            title = report_gen.get_text('zero_scores_analysis')
        doc = report_gen.create_new_report(title)
        
        # Add executive summary
        summary_text = results.get('summary', 'This report analyzes the percentage of students scoring zero on each assessment task.')
        report_gen.add_executive_summary(summary_text)
        
        # Add methodology section
        methodology_text = (
            "This analysis focuses on the percentage of students who scored zero on each assessment task. "
            "Zero scores often indicate a lack of basic skills or understanding in a particular area, "
            "and can help identify critical intervention needs."
        )
        report_gen.add_methodology(methodology_text)
        
        # Add results section
        report_gen.add_section(report_gen.get_text('results'), level=1)
        
        # Add zero scores table
        zero_scores_data = results.get('zero_scores_data')
        if zero_scores_data is not None:
            report_gen.add_table(
                zero_scores_data, 
                title=report_gen.get_text('zero_scores_table', 'Percentage of Zero Scores by Task')
            )
        
        # Add visualizations
        visualizations = results.get('visualizations', [])
        for viz in visualizations:
            if 'figure' in viz and 'title' in viz:
                report_gen.add_plotly_figure(viz['figure'], viz['title'])
        
        # Add critical areas section
        report_gen.add_section(report_gen.get_text('critical_areas', 'Critical Areas'), level=2)
        critical_areas = results.get('critical_areas', [])
        for area in critical_areas:
            task = area.get('task', '')
            percentage = area.get('percentage', 0)
            report_gen.add_bullet_point(
                f"{task}: {report_gen.format_number(percentage)}% " + 
                report_gen.get_text('zero_score_text', 'of students scored zero')
            )
        
        # Add interpretations
        report_gen.add_section(report_gen.get_text('interpretation'), level=1)
        report_gen.add_paragraph(results.get('interpretation_summary', ''))
        
        # Add task-specific interpretations
        task_interpretations = results.get('task_interpretations', [])
        for interp in task_interpretations:
            task = interp.get('task', '')
            text = interp.get('text', '')
            level = interp.get('level', '')
            
            # Determine style based on level
            style = 'Normal'
            if level == 'CRITICAL' or level == 'CRITIQUE':
                style = 'Recommendation'
            elif level == 'CONCERNING' or level == 'PRÉOCCUPANT':
                style = 'Highlight'
            elif level == 'WATCH' or level == 'À SURVEILLER':
                style = 'Interpretation'
                
            report_gen.add_section(task, level=3)
            report_gen.add_paragraph(text, style=style)
        
        # Add recommendations
        report_gen.add_section(report_gen.get_text('recommendations'), level=1)
        recommendations = results.get('recommendations', [])
        for rec in recommendations:
            report_gen.add_bullet_point(rec)
        
        # Set up headers and footers
        report_gen.setup_headers_and_footers(title=title)
        
        return doc


class SchoolComparisonReportTemplate:
    """Template for school comparison reports."""
    
    @staticmethod
    def create_report(report_gen, results, title=None):
        """
        Create a school comparison report.
        
        Args:
            report_gen (WordReportGenerator): Report generator instance
            results (dict): School comparison results
            title (str, optional): Report title
            
        Returns:
            Document: The created document
        """
        # Get language
        language = report_gen.language
        
        # Create new document
        if title is None:
            title = report_gen.get_text('school_comparison')
        doc = report_gen.create_new_report(title)
        
        # Add executive summary
        summary_text = results.get('summary', 'This report compares performance across different schools.')
        report_gen.add_executive_summary(summary_text)
        
        # Add methodology section
        methodology_text = (
            "This analysis compares assessment results across different schools. "
            "It includes statistical tests to determine if observed differences are significant, "
            "and identifies areas where certain schools may need additional support."
        )
        report_gen.add_methodology(methodology_text)
        
        # Add results section
        report_gen.add_section(report_gen.get_text('results'), level=1)
        
        # Add school comparison table
        comparison_data = results.get('comparison_data')
        if comparison_data is not None:
            report_gen.add_table(
                comparison_data, 
                title=report_gen.get_text('school_comparison_table', 'Performance Comparison by School')
            )
        
        # Add visualizations
        visualizations = results.get('visualizations', [])
        for viz in visualizations:
            if 'figure' in viz and 'title' in viz:
                report_gen.add_plotly_figure(viz['figure'], viz['title'])
        
        # Add statistical test results
        test_results = results.get('statistical_tests')
        if test_results is not None:
            report_gen.add_section(report_gen.get_text('statistical_tests', 'Statistical Tests'), level=2)
            report_gen.add_table(
                test_results,
                title=report_gen.get_text('significance_tests', 'Significance of Differences Between Schools')
            )
        
        # Add interpretations
        report_gen.add_section(report_gen.get_text('interpretation'), level=1)
        interpretations = results.get('interpretations', [])
        for interp in interpretations:
            indicator = interp.get('indicator')
            text = interp.get('text', '')
            report_gen.add_section(indicator, level=3)
            report_gen.add_paragraph(text)
        
        # Add recommendations
        report_gen.add_section(report_gen.get_text('recommendations'), level=1)
        
        # Add school-specific recommendations
        school_recommendations = results.get('school_recommendations', {})
        for school, recs in school_recommendations.items():
            report_gen.add_section(school, level=2)
            for rec in recs:
                report_gen.add_bullet_point(rec)
        
        # Add general recommendations
        general_recommendations = results.get('general_recommendations', [])
        if general_recommendations:
            report_gen.add_section(report_gen.get_text('general_recommendations', 'General Recommendations'), level=2)
            for rec in general_recommendations:
                report_gen.add_bullet_point(rec)
        
        # Set up headers and footers
        report_gen.setup_headers_and_footers(title=title)
        
        return doc


class GenderAnalysisReportTemplate:
    """Template for gender analysis reports."""
    
    @staticmethod
    def create_report(report_gen, results, title=None):
        """
        Create a gender analysis report.
        
        Args:
            report_gen (WordReportGenerator): Report generator instance
            results (dict): Gender analysis results
            title (str, optional): Report title
            
        Returns:
            Document: The created document
        """
        # Get language
        language = report_gen.language
        
        # Create new document
        if title is None:
            title = report_gen.get_text('gender_analysis')
        doc = report_gen.create_new_report(title)
        
        # Add executive summary
        summary_text = results.get('summary', 'This report analyzes differences in performance between boys and girls.')
        report_gen.add_executive_summary(summary_text)
        
        # Add methodology section
        methodology_text = (
            "This analysis compares assessment results between boys and girls. "
            "It includes statistical tests to determine if observed differences are significant, "
            "and provides insights into potential gender disparities in educational outcomes."
        )
        report_gen.add_methodology(methodology_text)
        
        # Add results section
        report_gen.add_section(report_gen.get_text('results'), level=1)
        
        # Add gender comparison table
        comparison_data = results.get('comparison_data')
        if comparison_data is not None:
            report_gen.add_table(
                comparison_data, 
                title=report_gen.get_text('gender_comparison_table', 'Performance Comparison by Gender')
            )
        
        # Add visualizations
        visualizations = results.get('visualizations', [])
        for viz in visualizations:
            if 'figure' in viz and 'title' in viz:
                report_gen.add_plotly_figure(viz['figure'], viz['title'])
        
        # Add statistical test results
        test_results = results.get('statistical_tests')
        if test_results is not None:
            report_gen.add_section(report_gen.get_text('statistical_tests', 'Statistical Tests'), level=2)
            report_gen.add_table(
                test_results,
                title=report_gen.get_text('significance_tests', 'Significance of Gender Differences')
            )
        
        # Add significant differences section
        report_gen.add_section(report_gen.get_text('significant_differences', 'Significant Gender Differences'), level=2)
        
        # Add areas where boys perform better
        boys_better = results.get('boys_advantage', [])
        if boys_better:
            report_gen.add_section(report_gen.get_text('boys_advantage', 'Boys Perform Better'), level=3)
            for area in boys_better:
                task = area.get('task', '')
                diff = area.get('difference', 0)
                report_gen.add_bullet_point(f"{task}: {report_gen.format_number(diff, 2)} " + 
                                          report_gen.get_text('points_higher', 'points higher'))
        
        # Add areas where girls perform better
        girls_better = results.get('girls_advantage', [])
        if girls_better:
            report_gen.add_section(report_gen.get_text('girls_advantage', 'Girls Perform Better'), level=3)
            for area in girls_better:
                task = area.get('task', '')
                diff = area.get('difference', 0)
                report_gen.add_bullet_point(f"{task}: {report_gen.format_number(diff, 2)} " + 
                                          report_gen.get_text('points_higher', 'points higher'))
        
        # Add interpretations
        report_gen.add_section(report_gen.get_text('interpretation'), level=1)
        interpretations = results.get('interpretations', [])
        for interp in interpretations:
            text = interp.get('text', '')
            report_gen.add_paragraph(text)
        
        # Add recommendations
        report_gen.add_section(report_gen.get_text('recommendations'), level=1)
        recommendations = results.get('recommendations', [])
        for rec in recommendations:
            report_gen.add_bullet_point(rec)
        
        # Set up headers and footers
        report_gen.setup_headers_and_footers(title=title)
        
        return doc


class InternationalComparisonReportTemplate:
    """Template for international comparison reports."""
    
    @staticmethod
    def create_report(report_gen, results, title=None):
        """
        Create an international comparison report.
        
        Args:
            report_gen (WordReportGenerator): Report generator instance
            results (dict): International comparison results
            title (str, optional): Report title
            
        Returns:
            Document: The created document
        """
        # Get language
        language = report_gen.language
        
        # Create new document
        if title is None:
            title = report_gen.get_text('international_comparison')
        doc = report_gen.create_new_report(title)
        
        # Add executive summary
        summary_text = results.get('summary', 'This report compares local performance against international benchmarks.')
        report_gen.add_executive_summary(summary_text)
        
        # Add methodology section
        methodology_text = (
            "This analysis compares local assessment results against international benchmarks. "
            "It identifies areas where performance meets, exceeds, or falls below international standards, "
            "and provides recommendations for improving performance in critical areas."
        )
        report_gen.add_methodology(methodology_text)
        
        # Add results section
        report_gen.add_section(report_gen.get_text('results'), level=1)
        
        # Add comparison table
        comparison_data = results.get('comparison_data')
        if comparison_data is not None:
            report_gen.add_table(
                comparison_data, 
                title=report_gen.get_text('international_comparison_table', 'Comparison with International Benchmarks')
            )
        
        # Add visualizations
        visualizations = results.get('visualizations', [])
        for viz in visualizations:
            if 'figure' in viz and 'title' in viz:
                report_gen.add_plotly_figure(viz['figure'], viz['title'])
        
        # Add performance categories
        report_gen.add_section(report_gen.get_text('performance_categories', 'Performance Categories'), level=2)
        
        # Critical areas
        critical_areas = results.get('critical_areas', [])
        if critical_areas:
            report_gen.add_section(report_gen.get_text('critical_areas', 'Critical Areas'), level=3)
            for area in critical_areas:
                task = area.get('task', '')
                percentage = area.get('percentage', 0)
                report_gen.add_bullet_point(
                    f"{task}: {report_gen.format_number(percentage)}% " + 
                    report_gen.get_text('of_benchmark', 'of benchmark')
                )
        
        # Areas meeting or exceeding benchmarks
        meeting_areas = results.get('meeting_areas', [])
        if meeting_areas:
            report_gen.add_section(report_gen.get_text('meeting_areas', 'Areas Meeting or Exceeding Benchmarks'), level=3)
            for area in meeting_areas:
                task = area.get('task', '')
                percentage = area.get('percentage', 0)
                report_gen.add_bullet_point(
                    f"{task}: {report_gen.format_number(percentage)}% " + 
                    report_gen.get_text('of_benchmark', 'of benchmark')
                )
        
        # Add interpretations
        report_gen.add_section(report_gen.get_text('interpretation'), level=1)
        interpretations = results.get('interpretations', [])
        for interp in interpretations:
            text = interp.get('text', '')
            report_gen.add_paragraph(text)
        
        # Add recommendations
        report_gen.add_section(report_gen.get_text('recommendations'), level=1)
        
        # Critical area recommendations
        critical_recommendations = results.get('critical_recommendations', [])
        if critical_recommendations:
            report_gen.add_section(report_gen.get_text('critical_recommendations', 'Recommendations for Critical Areas'), level=2)
            for rec in critical_recommendations:
                report_gen.add_bullet_point(rec)
        
        # General recommendations
        general_recommendations = results.get('general_recommendations', [])
        if general_recommendations:
            report_gen.add_section(report_gen.get_text('general_recommendations', 'General Recommendations'), level=2)
            for rec in general_recommendations:
                report_gen.add_bullet_point(rec)
        
        # Set up headers and footers
        report_gen.setup_headers_and_footers(title=title)
        
        return doc


# Example usage:
# 
# # Create report generator
# report_generator = WordReportGenerator(language="en")
# 
# # Create statistical report
# results = {
#     "summary": "This report provides a statistical overview of the EGRA assessment results.",
#     "statistics": df.describe(),
#     "visualizations": [
#         {"figure": fig1, "title": "Distribution of Reading Scores"}
#     ],
#     "interpretations": [
#         {"indicator": "Reading Fluency", "text": "Students show varied performance in reading fluency", "score": 65}
#     ],
#     "recommendations": [
#         {"text": "Focus on improving fluency through regular practice", "indicator": "Reading Fluency", "score": 65}
#     ]
# }
# 
# # Use template to create report
# doc = StatisticalReportTemplate.create_report(report_generator, results)
# 
# # Save report
# report_generator.save("statistical_report.docx")