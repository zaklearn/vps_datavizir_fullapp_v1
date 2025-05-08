import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
from config import translations, egra_columns, egma_columns

# Define international benchmarks for EGRA and EGMA variables
# These values are based on international research and standards
# Sources: RTI International, USAID, World Bank, UNESCO
international_benchmarks = {
    # EGRA benchmarks (based on Grade 2-3 expectations)
    "clpm": {"standard": 40, "description": "Correct Letters Per Minute"},
    "phoneme": {"standard": 8, "description": "Phoneme Awareness (out of 10)"},
    "sound_word": {"standard": 18, "description": "Correct Words Read (out of 20)"},
    "cwpm": {"standard": 45, "description": "Correct Words Per Minute"},
    "listening": {"standard": 4, "description": "Listening Comprehension (out of 5)"},
    "orf": {"standard": 45, "description": "Oral Reading Fluency"},
    "comprehension": {"standard": 4, "description": "Reading Comprehension (out of 5)"},
    
    # EGMA benchmarks (based on Grade 2-3 expectations)
    "number_id": {"standard": 20, "description": "Number Identification (out of 20)"},
    "discrimin": {"standard": 8, "description": "Number Discrimination (out of 10)"},
    "missing_number": {"standard": 8, "description": "Missing Number (out of 10)"},
    "addition": {"standard": 16, "description": "Addition Problems (out of 20)"},
    "subtraction": {"standard": 14, "description": "Subtraction Problems (out of 20)"},
    "problems": {"standard": 4, "description": "Word Problems (out of 5)"}
}

def show_international_comparison(df, language):
    """
    Compares student performance against international benchmarks.
    
    Args:
        df (pandas.DataFrame): The data to analyze
        language (str): Selected language for UI elements
    """
    t = translations[language]  # Get translations for selected language
    
    st.markdown(f"""
    ### {t.get("title_international_comparison", "International Standards Comparison")}
    
    🔍 **{t.get("international_intro", "Objective: Compare local performance against international benchmarks to identify improvement areas.")}**
    """)
    
    # Get available assessment columns that have international benchmarks
    available_columns = [col for col in international_benchmarks.keys() if col in df.columns]
    
    if not available_columns:
        st.error(t.get("no_benchmark_columns", "No assessment columns with international benchmarks were found in the data."))
        return
    
    # Allow users to select columns for analysis
    st.subheader(t.get("select_variables", "📋 Select Variables for Comparison"))
    
    # Create two columns for selection
    col1, col2 = st.columns(2)
    
    # Split available columns into two groups for better UI organization
    first_half = available_columns[:len(available_columns)//2 + len(available_columns)%2]
    second_half = available_columns[len(available_columns)//2 + len(available_columns)%2:]
    
    with col1:
        selected_first = st.multiselect(
            t.get("variables_left", "Variables (left):"),
            options=first_half,
            default=first_half,
            format_func=lambda x: t["columns_of_interest"].get(x, international_benchmarks[x]["description"])
        )
    
    with col2:
        selected_second = st.multiselect(
            t.get("variables_right", "Variables (right):"),
            options=second_half,
            default=second_half,
            format_func=lambda x: t["columns_of_interest"].get(x, international_benchmarks[x]["description"])
        )
    
    selected_columns = selected_first + selected_second
    
    if selected_columns:
        try:
            # Calculate local mean scores
            local_means = df[selected_columns].mean().round(2)
            
            # Get international benchmarks for selected columns
            benchmarks = {col: international_benchmarks[col]["standard"] for col in selected_columns}
            
            # Calculate gaps between local means and benchmarks
            gaps = pd.Series({col: local_means[col] - benchmarks[col] for col in selected_columns})
            
            # Calculate percentage of benchmark achieved
            percentage_achieved = pd.Series({
                col: (local_means[col] / benchmarks[col] * 100).round(1) for col in selected_columns
            })
            
            # Prepare data for display
            comparison_data = pd.DataFrame({
                "variable": selected_columns,
                "local_mean": local_means.values,
                "benchmark": [benchmarks[col] for col in selected_columns],
                "gap": gaps.values,
                "percentage": percentage_achieved.values
            })
            
            # Add translated column names for display
            comparison_data["variable_name"] = comparison_data["variable"].apply(
                lambda x: t["columns_of_interest"].get(x, international_benchmarks[x]["description"])
            )
            
            # Order by gap (worst performing first)
            comparison_data = comparison_data.sort_values("gap", ascending=True)
            
            # Display comparison table
            st.subheader(t.get("comparison_table", "📊 Comparison with International Benchmarks"))
            
            # Format display table
            display_df = comparison_data.copy()
            display_df.columns = [
                "code",  # Hidden column with variable code
                t.get("local_mean", "Local Mean"),
                t.get("benchmark", "Benchmark"),
                t.get("gap", "Gap"),
                t.get("percentage", "% of Benchmark"),
                t.get("variable", "Assessment Variable")
            ]
            
            # Display the table without the code column
            st.dataframe(display_df[display_df.columns[1:]])
            
            # Visualization of comparison
            st.subheader(t.get("comparison_chart", "📈 Visualization of Comparison"))
            
            # Create bar chart comparing local means with benchmarks
            # Prepare data for chart
            chart_data = []
            for _, row in comparison_data.iterrows():
                # Add local mean
                chart_data.append({
                    "variable": row["variable_name"],
                    "score": row["local_mean"],
                    "type": t.get("local_mean", "Local Mean"),
                    "percentage": row["percentage"]
                })
                
                # Add benchmark
                chart_data.append({
                    "variable": row["variable_name"],
                    "score": row["benchmark"],
                    "type": t.get("benchmark", "Benchmark"),
                    "percentage": 100
                })
            
            chart_df = pd.DataFrame(chart_data)
            
            # Create bar chart
            fig = px.bar(
                chart_df,
                x="variable",
                y="score",
                color="type",
                barmode="group",
                title=t.get("comparison_chart_title", "Local Performance vs. International Benchmarks"),
                labels={
                    "variable": t.get("assessment_variable", "Assessment Variable"),
                    "score": t.get("score", "Score"),
                    "type": t.get("value_type", "Value Type")
                },
                color_discrete_map={
                    t.get("local_mean", "Local Mean"): "#3498DB",  # Blue
                    t.get("benchmark", "Benchmark"): "#F39C12"   # Orange
                }
            )
            
            # Update layout
            fig.update_layout(
                xaxis_tickangle=-45,
                legend_title=t.get("value_type", "Value Type"),
                height=600
            )
            
            st.plotly_chart(fig, use_container_width=True)
            
            # Visualization of percentage achieved
            st.subheader(t.get("percentage_chart", "📊 Percentage of Benchmark Achieved"))
            
            # Create percentage chart
            percentage_df = comparison_data.copy()
            percentage_df["achievement_level"] = percentage_df["percentage"].apply(
                lambda x: t.get("critical", "Critical") if x < 70 else (
                    t.get("concerning", "Concerning") if x < 85 else (
                        t.get("approaching", "Approaching") if x < 100 else t.get("meeting", "Meeting")
                    )
                )
            )
            
            percentage_fig = px.bar(
                percentage_df,
                x="variable_name",
                y="percentage",
                color="achievement_level",
                title=t.get("percentage_chart_title", "Percentage of International Benchmark Achieved"),
                labels={
                    "variable_name": t.get("assessment_variable", "Assessment Variable"),
                    "percentage": t.get("percentage_of_benchmark", "% of Benchmark"),
                    "achievement_level": t.get("achievement_level", "Achievement Level")
                },
                color_discrete_map={
                    t.get("critical", "Critical"): "#E74C3C",  # Red
                    t.get("concerning", "Concerning"): "#F39C12",  # Orange
                    t.get("approaching", "Approaching"): "#F1C40F",  # Yellow
                    t.get("meeting", "Meeting"): "#2ECC71"  # Green
                }
            )
            
            # Add reference line at 100%
            percentage_fig.add_hline(
                y=100, 
                line_dash="dash", 
                line_color="black",
                annotation_text=t.get("benchmark_line", "Benchmark")
            )
            
            # Add reference lines for achievement levels
            percentage_fig.add_hline(y=85, line_dash="dot", line_color="#F1C40F")  # Approaching (Yellow)
            percentage_fig.add_hline(y=70, line_dash="dot", line_color="#F39C12")  # Concerning (Orange)
            
            # Update layout
            percentage_fig.update_layout(
                xaxis_tickangle=-45,
                legend_title=t.get("achievement_level", "Achievement Level"),
                height=600
            )
            
            st.plotly_chart(percentage_fig, use_container_width=True)
            
            # Analysis of results
            st.subheader(t.get("results_analysis", "🔍 Analysis of Results"))
            
            # Group variables by achievement level
            critical_vars = percentage_df[percentage_df["achievement_level"] == t.get("critical", "Critical")]
            concerning_vars = percentage_df[percentage_df["achievement_level"] == t.get("concerning", "Concerning")]
            approaching_vars = percentage_df[percentage_df["achievement_level"] == t.get("approaching", "Approaching")]
            meeting_vars = percentage_df[percentage_df["achievement_level"] == t.get("meeting", "Meeting")]
            
            # Display variables by achievement level
            if not critical_vars.empty:
                st.markdown(f"**{t.get('critical_areas', 'Critical Areas')}** (<70% of benchmark)")
                for _, row in critical_vars.iterrows():
                    st.markdown(f"- {row['variable_name']}: {row['percentage']}% {t.get('of_benchmark', 'of benchmark')} ({row['gap']:.2f} {t.get('points_below', 'points below')})")
            
            if not concerning_vars.empty:
                st.markdown(f"**{t.get('concerning_areas', 'Concerning Areas')}** (70-85% of benchmark)")
                for _, row in concerning_vars.iterrows():
                    st.markdown(f"- {row['variable_name']}: {row['percentage']}% {t.get('of_benchmark', 'of benchmark')} ({abs(row['gap']):.2f} {t.get('points_below', 'points below')})")
            
            if not approaching_vars.empty:
                st.markdown(f"**{t.get('approaching_areas', 'Approaching Benchmark')}** (85-100% of benchmark)")
                for _, row in approaching_vars.iterrows():
                    st.markdown(f"- {row['variable_name']}: {row['percentage']}% {t.get('of_benchmark', 'of benchmark')} ({abs(row['gap']):.2f} {t.get('points_below', 'points below')})")
            
            if not meeting_vars.empty:
                st.markdown(f"**{t.get('meeting_areas', 'Meeting or Exceeding Benchmark')}** (≥100% of benchmark)")
                for _, row in meeting_vars.iterrows():
                    if row['gap'] > 0:
                        st.markdown(f"- {row['variable_name']}: {row['percentage']}% {t.get('of_benchmark', 'of benchmark')} ({row['gap']:.2f} {t.get('points_above', 'points above')})")
                    else:
                        st.markdown(f"- {row['variable_name']}: {row['percentage']}% {t.get('of_benchmark', 'of benchmark')} ({t.get('at_benchmark', 'at benchmark')})")
            
            # Policy implications and recommendations
            st.subheader(t.get("policy_implications", "Policy Implications and Recommendations"))
            
            # Reading skills analysis
            reading_vars = [var for var in selected_columns if var in egra_columns]
            reading_percentage = percentage_df[percentage_df["variable"].isin(reading_vars)]["percentage"].mean() if reading_vars else None
            
            # Math skills analysis
            math_vars = [var for var in selected_columns if var in egma_columns]
            math_percentage = percentage_df[percentage_df["variable"].isin(math_vars)]["percentage"].mean() if math_vars else None
            
            # Overall analysis
            overall_percentage = percentage_df["percentage"].mean()
            
            # Display overall analysis
            st.markdown(f"**{t.get('overall_performance', 'Overall Performance Analysis')}**")
            st.markdown(f"{t.get('average_achievement', 'Average achievement across all skills')}: **{overall_percentage:.1f}%** {t.get('of_benchmark', 'of benchmark')}")
            
            if reading_percentage is not None:
                st.markdown(f"{t.get('reading_average', 'Reading skills average')}: **{reading_percentage:.1f}%** {t.get('of_benchmark', 'of benchmark')}")
            
            if math_percentage is not None:
                st.markdown(f"{t.get('math_average', 'Math skills average')}: **{math_percentage:.1f}%** {t.get('of_benchmark', 'of benchmark')}")
            
            # Generate recommendations based on results
            st.markdown(f"**{t.get('recommendations', 'Recommendations')}**")
            
            # Areas with critical gaps
            if not critical_vars.empty:
                st.markdown(f"1. **{t.get('critical_recommendation', 'For Critical Areas')}**:")
                
                # Check if the critical areas are primarily in reading or math
                critical_reading = [var for var in critical_vars["variable"] if var in egra_columns]
                critical_math = [var for var in critical_vars["variable"] if var in egma_columns]
                
                if critical_reading:
                    reading_areas = [t["columns_of_interest"].get(col, international_benchmarks[col]["description"]) for col in critical_reading]
                    st.markdown(f"- {t.get('reading_intervention', 'Implement intensive reading intervention programs focused on')}: {', '.join(reading_areas)}")
                
                if critical_math:
                    math_areas = [t["columns_of_interest"].get(col, international_benchmarks[col]["description"]) for col in critical_math]
                    st.markdown(f"- {t.get('math_intervention', 'Develop targeted math intervention strategies for')}: {', '.join(math_areas)}")
                
                st.markdown(f"- {t.get('teacher_training', 'Provide specialized teacher training in effective instruction for these critical areas')}")
                st.markdown(f"- {t.get('additional_time', 'Allocate additional instructional time for these foundational skills')}")
                st.markdown(f"- {t.get('progress_monitoring', 'Implement frequent progress monitoring to track improvement')}")
            
            # Areas with concerning gaps
            if not concerning_vars.empty:
                st.markdown(f"2. **{t.get('concerning_recommendation', 'For Concerning Areas')}**:")
                st.markdown(f"- {t.get('targeted_support', 'Provide targeted support through small group instruction')}")
                st.markdown(f"- {t.get('instructional_materials', 'Review and enhance instructional materials and methods')}")
                st.markdown(f"- {t.get('regular_assessment', 'Conduct regular formative assessments to track progress')}")
            
            # Areas approaching benchmark
            if not approaching_vars.empty:
                st.markdown(f"3. **{t.get('approaching_recommendation', 'For Areas Approaching Benchmark')}**:")
                st.markdown(f"- {t.get('maintain_instruction', 'Maintain current instructional approaches with minor enhancements')}")
                st.markdown(f"- {t.get('continue_monitoring', 'Continue monitoring progress toward benchmark achievement')}")
            
            # Areas meeting or exceeding benchmark
            if not meeting_vars.empty:
                st.markdown(f"4. **{t.get('meeting_recommendation', 'For Areas Meeting Benchmark')}**:")
                st.markdown(f"- {t.get('identify_practices', 'Identify effective practices that led to success in these areas')}")
                st.markdown(f"- {t.get('apply_lessons', 'Apply lessons learned to areas still below benchmark')}")
                st.markdown(f"- {t.get('maintain_excellence', 'Set extended goals to maintain excellence')}")
            
            # Systemic recommendations
            st.markdown(f"5. **{t.get('systemic_recommendation', 'Systemic Recommendations')}**:")
            st.markdown(f"- {t.get('curriculum_alignment', 'Ensure curriculum alignment with international standards')}")
            st.markdown(f"- {t.get('professional_development', 'Invest in ongoing professional development for teachers')}")
            st.markdown(f"- {t.get('resource_allocation', 'Allocate resources based on identified performance gaps')}")
            st.markdown(f"- {t.get('community_involvement', 'Engage parents and communities in supporting student learning')}")
            
            # Export options
            col1, col2 = st.columns(2)
            
            # CSV Export
            with col1:
                csv = comparison_data.to_csv(index=False).encode('utf-8-sig')
                st.download_button(
                    t.get("export_international_csv", "📥 Download CSV"),
                    csv,
                    "international_comparison.csv",
                    "text/csv",
                    key='download-international-csv'
                )
            
            # Word Export
            with col2:
                if st.button(t.get("export_international_word", "📄 Export to Word")):
                    try:
                        doc = create_international_comparison_word_report(
                            comparison_data, fig, percentage_fig, reading_percentage, math_percentage, overall_percentage, t
                        )
                        
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                            doc.save(tmp.name)
                            with open(tmp.name, 'rb') as f:
                                docx = f.read()
                            st.download_button(
                                t.get("download_international_word", "📥 Download Word Report"),
                                docx,
                                "international_comparison.docx",
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        os.unlink(tmp.name)
                    except Exception as e:
                        st.error(f"Error creating Word report: {str(e)}")
        
        except Exception as e:
            st.error(f"Error in international comparison analysis: {str(e)}")
    
    else:
        st.warning(t.get("warning_select_variable", "Please select at least one variable to analyze."))

def create_international_comparison_word_report(comparison_data, fig1, fig2, reading_percentage, math_percentage, overall_percentage, t):
    """
    Creates a Word report with international comparison analysis.
    
    Args:
        comparison_data (pandas.DataFrame): Comparison data
        fig1 (plotly.graph_objs._figure.Figure): Comparison bar chart
        fig2 (plotly.graph_objs._figure.Figure): Percentage achievement chart
        reading_percentage (float): Average reading percentage achievement
        math_percentage (float): Average math percentage achievement
        overall_percentage (float): Overall percentage achievement
        t (dict): Translation dictionary
        
    Returns:
        docx.Document: Word document with the report
    """
    doc = Document()
    
    # Title
    doc.add_heading(t.get("title_international_comparison", "International Standards Comparison"), level=1)
    
    # Introduction
    doc.add_paragraph(t.get("international_intro", "Objective: Compare local performance against international benchmarks to identify improvement areas."))
    
    # Comparison table
    doc.add_heading(t.get("comparison_table", "Comparison with International Benchmarks"), level=2)
    
    # Create table
    table = doc.add_table(rows=len(comparison_data) + 1, cols=5)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = t.get("variable", "Assessment Variable")
    header_cells[1].text = t.get("local_mean", "Local Mean")
    header_cells[2].text = t.get("benchmark", "Benchmark")
    header_cells[3].text = t.get("gap", "Gap")
    header_cells[4].text = t.get("percentage", "% of Benchmark")
    
    # Add data rows
    for i, (_, row) in enumerate(comparison_data.iterrows(), 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = row["variable_name"]
        row_cells[1].text = f"{row['local_mean']:.2f}"
        row_cells[2].text = f"{row['benchmark']:.2f}"
        row_cells[3].text = f"{row['gap']:.2f}"
        row_cells[4].text = f"{row['percentage']:.1f}%"
    
    # Visualizations
    doc.add_heading(t.get("visualizations", "Visualizations"), level=2)
    
    # Save and add charts
    with tempfile.TemporaryDirectory() as tmp_dir:
        # Comparison bar chart
        img1_path = os.path.join(tmp_dir, "comparison_chart.png")
        fig1.write_image(img1_path, width=900, height=500)
        doc.add_picture(img1_path, width=Inches(6))
        doc.add_paragraph()
        
        # Percentage achievement chart
        img2_path = os.path.join(tmp_dir, "percentage_chart.png")
        fig2.write_image(img2_path, width=900, height=500)
        doc.add_picture(img2_path, width=Inches(6))
        doc.add_paragraph()
    
    # Analysis of results
    doc.add_heading(t.get("results_analysis", "Analysis of Results"), level=2)
    
    # Group variables by achievement level
    critical_vars = comparison_data[comparison_data["percentage"] < 70]
    concerning_vars = comparison_data[(comparison_data["percentage"] >= 70) & (comparison_data["percentage"] < 85)]
    approaching_vars = comparison_data[(comparison_data["percentage"] >= 85) & (comparison_data["percentage"] < 100)]
    meeting_vars = comparison_data[comparison_data["percentage"] >= 100]
    
    # Display variables by achievement level
    if not critical_vars.empty:
        p = doc.add_paragraph()
        p.add_run(f"{t.get('critical_areas', 'Critical Areas')} (<70% of benchmark)").bold = True
        
        for _, row in critical_vars.iterrows():
            doc.add_paragraph(f"{row['variable_name']}: {row['percentage']}% {t.get('of_benchmark', 'of benchmark')} ({row['gap']:.2f} {t.get('points_below', 'points below')})", style='List Bullet')
    
    if not concerning_vars.empty:
        p = doc.add_paragraph()
        p.add_run(f"{t.get('concerning_areas', 'Concerning Areas')} (70-85% of benchmark)").bold = True
        
        for _, row in concerning_vars.iterrows():
            doc.add_paragraph(f"{row['variable_name']}: {row['percentage']}% {t.get('of_benchmark', 'of benchmark')} ({abs(row['gap']):.2f} {t.get('points_below', 'points below')})", style='List Bullet')
    
    if not approaching_vars.empty:
        p = doc.add_paragraph()
        p.add_run(f"{t.get('approaching_areas', 'Approaching Benchmark')} (85-100% of benchmark)").bold = True
        
        for _, row in approaching_vars.iterrows():
            doc.add_paragraph(f"{row['variable_name']}: {row['percentage']}% {t.get('of_benchmark', 'of benchmark')} ({abs(row['gap']):.2f} {t.get('points_below', 'points below')})", style='List Bullet')
    
    if not meeting_vars.empty:
        p = doc.add_paragraph()
        p.add_run(f"{t.get('meeting_areas', 'Meeting or Exceeding Benchmark')} (≥100% of benchmark)").bold = True
        
        for _, row in meeting_vars.iterrows():
            if row['gap'] > 0:
                doc.add_paragraph(f"{row['variable_name']}: {row['percentage']}% {t.get('of_benchmark', 'of benchmark')} ({row['gap']:.2f} {t.get('points_above', 'points above')})", style='List Bullet')
            else:
                doc.add_paragraph(f"{row['variable_name']}: {row['percentage']}% {t.get('of_benchmark', 'of benchmark')} ({t.get('at_benchmark', 'at benchmark')})", style='List Bullet')
    
    # Overall analysis
    doc.add_heading(t.get("overall_performance", "Overall Performance Analysis"), level=2)
    doc.add_paragraph(f"{t.get('average_achievement', 'Average achievement across all skills')}: {overall_percentage:.1f}% {t.get('of_benchmark', 'of benchmark')}")
    
    if reading_percentage is not None:
        doc.add_paragraph(f"{t.get('reading_average', 'Reading skills average')}: {reading_percentage:.1f}% {t.get('of_benchmark', 'of benchmark')}")
    
    if math_percentage is not None:
        doc.add_paragraph(f"{t.get('math_average', 'Math skills average')}: {math_percentage:.1f}% {t.get('of_benchmark', 'of benchmark')}")
    
    # Recommendations
    doc.add_heading(t.get("recommendations", "Recommendations"), level=2)
    
    # Areas with critical gaps
    if not critical_vars.empty:
        p = doc.add_paragraph()
        p.add_run(f"{t.get('critical_recommendation', 'For Critical Areas')}:").bold = True
        
        # Check if the critical areas are primarily in reading or math
        critical_reading = [var for var in critical_vars["variable"] if var in egra_columns]
        critical_math = [var for var in critical_vars["variable"] if var in egma_columns]
        
        if critical_reading:
            reading_areas = [t["columns_of_interest"].get(col, international_benchmarks[col]["description"]) for col in critical_reading]
            doc.add_paragraph(f"{t.get('reading_intervention', 'Implement intensive reading intervention programs focused on')}: {', '.join(reading_areas)}", style='List Bullet')
        
        if critical_math:
            math_areas = [t["columns_of_interest"].get(col, international_benchmarks[col]["description"]) for col in critical_math]
            doc.add_paragraph(f"{t.get('math_intervention', 'Develop targeted math intervention strategies for')}: {', '.join(math_areas)}", style='List Bullet')
        
        doc.add_paragraph(t.get('teacher_training', 'Provide specialized teacher training in effective instruction for these critical areas'), style='List Bullet')
        doc.add_paragraph(t.get('additional_time', 'Allocate additional instructional time for these foundational skills'), style='List Bullet')
        doc.add_paragraph(t.get('progress_monitoring', 'Implement frequent progress monitoring to track improvement'), style='List Bullet')
    
    # Areas with concerning gaps
    if not concerning_vars.empty:
        p = doc.add_paragraph()
        p.add_run(f"{t.get('concerning_recommendation', 'For Concerning Areas')}:").bold = True
        
        doc.add_paragraph(t.get('targeted_support', 'Provide targeted support through small group instruction'), style='List Bullet')
        doc.add_paragraph(t.get('instructional_materials', 'Review and enhance instructional materials and methods'), style='List Bullet')
        doc.add_paragraph(t.get('regular_assessment', 'Conduct regular formative assessments to track progress'), style='List Bullet')
    
    # Areas approaching benchmark
    if not approaching_vars.empty:
        p = doc.add_paragraph()
        p.add_run(f"{t.get('approaching_recommendation', 'For Areas Approaching Benchmark')}:").bold = True
        
        doc.add_paragraph(t.get('maintain_instruction', 'Maintain current instructional approaches with minor enhancements'), style='List Bullet')
        doc.add_paragraph(t.get('continue_monitoring', 'Continue monitoring progress toward benchmark achievement'), style='List Bullet')
    
    # Areas meeting or exceeding benchmark
    if not meeting_vars.empty:
        p = doc.add_paragraph()
        p.add_run(f"{t.get('meeting_recommendation', 'For Areas Meeting Benchmark')}:").bold = True
        
        doc.add_paragraph(t.get('identify_practices', 'Identify effective practices that led to success in these areas'), style='List Bullet')
        doc.add_paragraph(t.get('apply_lessons', 'Apply lessons learned to areas still below benchmark'), style='List Bullet')
        doc.add_paragraph(t.get('maintain_excellence', 'Set extended goals to maintain excellence'), style='List Bullet')
    
    # Systemic recommendations
    p = doc.add_paragraph()
    p.add_run(f"{t.get('systemic_recommendation', 'Systemic Recommendations')}:").bold = True
    
    doc.add_paragraph(t.get('curriculum_alignment', 'Ensure curriculum alignment with international standards'), style='List Bullet')
    doc.add_paragraph(t.get('professional_development', 'Invest in ongoing professional development for teachers'), style='List Bullet')
    doc.add_paragraph(t.get('resource_allocation', 'Allocate resources based on identified performance gaps'), style='List Bullet')
    doc.add_paragraph(t.get('community_involvement', 'Engage parents and communities in supporting student learning'), style='List Bullet')
    
    # Information about benchmarks
    doc.add_heading(t.get("about_benchmarks", "About International Benchmarks"), level=2)
    doc.add_paragraph(t.get("benchmark_info", """
    The international benchmarks used in this analysis are based on research and standards from multiple sources including RTI International, USAID, World Bank, and UNESCO. These benchmarks represent achievement levels that have been associated with successful educational outcomes in various international contexts.
    
    These benchmarks should be interpreted as goals to work toward rather than absolute standards, as educational contexts can vary significantly across countries and regions. They provide valuable reference points for understanding local performance in a global context.
    """))
    
    # Methodology note
    doc.add_paragraph(t.get("methodology_note", """
    Methodology Note: This analysis compares mean scores against international benchmarks. It is important to also consider score distributions and the proportion of students meeting benchmarks, which may provide additional insights beyond mean performance.
    """), style='Normal')
    
    # Footer with date
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = t.get('report_date', 'Report generated on: ') + pd.Timestamp.now().strftime('%Y-%m-%d')
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc