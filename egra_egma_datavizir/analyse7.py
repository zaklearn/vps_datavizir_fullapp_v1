import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import scipy.stats as stats
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
from config import translations, egra_columns, egma_columns

def show_performance_school(df, language):
    """
    Analyzes and displays school performance comparisons for EGRA and EGMA variables.
    
    Args:
        df (pandas.DataFrame): The data to analyze
        language (str): Selected language for UI elements
    """
    t = translations[language]  # Get translations for selected language
    
    st.markdown(f"""
    ### {t.get("title_performance", "School Performance Analysis")}
    
    🔍 **{t.get("performance_intro", "Objective: Compare performance across different schools and identify significant differences.")}**
    """)
    
    # Check if school column exists
    if "school" not in df.columns:
        st.error(t.get("no_school_column", "Error: No 'school' column found in the data."))
        return
    
    # Get available assessment columns
    available_egra = [col for col in egra_columns if col in df.columns]
    available_egma = [col for col in egma_columns if col in df.columns]
    available_columns = available_egra + available_egma
    
    if not available_columns:
        st.error(t.get("no_assessment_columns", "No assessment columns found in the data."))
        return
    
    # Get unique schools
    schools = df["school"].unique()
    if len(schools) < 2:
        st.warning(t.get("insufficient_schools", "Warning: At least two schools are needed for comparison. Only {n} found.").format(n=len(schools)))
        if len(schools) == 0:
            return
    
    # Allow users to select columns for analysis
    st.subheader(t.get("select_variables", "📋 Select Variables for Analysis"))
    
    # Create two columns for selection
    col1, col2 = st.columns(2)
    
    with col1:
        selected_egra = st.multiselect(
            t.get("egra_variables", "EGRA Variables:"),
            options=available_egra,
            default=available_egra[:min(3, len(available_egra))],  # Default select up to 3 EGRA variables
            format_func=lambda x: t["columns_of_interest"].get(x, x)
        )
    
    with col2:
        selected_egma = st.multiselect(
            t.get("egma_variables", "EGMA Variables:"),
            options=available_egma,
            default=available_egma[:min(3, len(available_egma))],  # Default select up to 3 EGMA variables
            format_func=lambda x: t["columns_of_interest"].get(x, x)
        )
    
    selected_columns = selected_egra + selected_egma
    
    if selected_columns:
        try:
            # Calculate mean scores by school
            mean_scores_by_school = df.groupby("school")[selected_columns].mean().round(2)
            
            # Calculate sample sizes by school for reference
            sample_sizes = df.groupby("school").size().rename(t.get("sample_size", "Sample Size"))
            
            # Combine with mean scores for display
            performance_table = pd.concat([mean_scores_by_school, sample_sizes], axis=1)
            
            # Display the performance table
            st.subheader(t.get("performance_results", "📊 Mean Scores by School"))
            st.dataframe(performance_table, use_container_width=True)
            
            # Identify highest and lowest performing schools for each variable
            st.subheader(t.get("performance_highlights", "🏆 Performance Highlights"))
            
            highlight_data = []
            
            for col in selected_columns:
                # Get translated column name
                col_name = t["columns_of_interest"].get(col, col)
                
                # Find highest and lowest performing schools
                highest_school = mean_scores_by_school[col].idxmax()
                highest_score = mean_scores_by_school.loc[highest_school, col]
                
                lowest_school = mean_scores_by_school[col].idxmin()
                lowest_score = mean_scores_by_school.loc[lowest_school, col]
                
                # Store in highlight data
                highlight_data.append({
                    "variable": col_name,
                    "highest_school": highest_school,
                    "highest_score": highest_score,
                    "lowest_school": lowest_school,
                    "lowest_score": lowest_score,
                    "range": highest_score - lowest_score
                })
            
            # Convert to DataFrame and display
            highlight_df = pd.DataFrame(highlight_data)
            highlight_df.columns = [
                t.get("variable", "Variable"),
                t.get("highest_school", "Highest School"),
                t.get("highest_score", "Highest Score"),
                t.get("lowest_school", "Lowest School"),
                t.get("lowest_score", "Lowest Score"),
                t.get("score_range", "Score Range")
            ]
            
            st.dataframe(highlight_df, use_container_width=True)
            
            # Statistical significance testing
            st.subheader(t.get("statistical_testing", "📊 Statistical Significance Testing"))
            
            st.markdown(t.get("anova_explanation", """
            The analysis below uses the Kruskal-Wallis H-test, a non-parametric method for comparing scores from more than two schools.
            A p-value < 0.05 indicates statistically significant differences between at least some schools.
            """))
            
            # Run Kruskal-Wallis test for each variable
            test_results = []
            
            for col in selected_columns:
                col_name = t["columns_of_interest"].get(col, col)
                
                # Group data by school for this variable
                groups = [df[df["school"] == school][col].dropna() for school in schools]
                groups = [g for g in groups if len(g) > 0]  # Remove empty groups
                
                # Perform Kruskal-Wallis test if we have at least two groups with data
                if len(groups) >= 2:
                    try:
                        h_stat, p_value = stats.kruskal(*groups)
                        
                        test_results.append({
                            "variable": col_name,
                            "h_statistic": h_stat,
                            "p_value": p_value,
                            "significant": p_value < 0.05
                        })
                    except Exception as e:
                        # Handle errors in statistical testing
                        test_results.append({
                            "variable": col_name,
                            "h_statistic": None,
                            "p_value": None,
                            "significant": None,
                            "error": str(e)
                        })
            
            # Display test results if any tests were performed
            if test_results:
                test_df = pd.DataFrame(test_results)
                
                # Format the display DataFrame
                display_df = test_df.copy()
                display_df.columns = [
                    t.get("variable", "Variable"),
                    t.get("h_statistic", "H Statistic"),
                    t.get("p_value", "p-value"),
                    t.get("significant", "Significant Difference"),
                    *([t.get("error", "Error")] if "error" in test_df.columns else [])
                ]
                
                # Format p-values and significance for display
                if "p_value" in test_df.columns:
                    display_df[t.get("p_value", "p-value")] = display_df[t.get("p_value", "p-value")].apply(
                        lambda x: f"{x:.4f}" if x is not None else "N/A"
                    )
                
                if "significant" in test_df.columns:
                    display_df[t.get("significant", "Significant Difference")] = display_df[t.get("significant", "Significant Difference")].apply(
                        lambda x: t.get("significant_yes", "Yes") if x else t.get("significant_no", "No") if x is not None else "N/A"
                    )
                
                st.dataframe(display_df, use_container_width=True)
            
            # Visualization of distributions by school
            st.subheader(t.get("distribution_title", "📈 Score Distributions by School"))
            
            # Allow selection of distribution plot type
            plot_type = st.radio(
                t.get("plot_type", "Select plot type:"),
                [t.get("boxplot", "Box Plots"), t.get("violinplot", "Violin Plots")],
                horizontal=True
            )
            
            # Create distribution plots for each selected variable
            for column in selected_columns:
                col_name = t["columns_of_interest"].get(column, column)
                
                if plot_type == t.get("boxplot", "Box Plots"):
                    fig = px.box(
                        df,
                        x="school",
                        y=column,
                        color="school",
                        title=f"{col_name} - {t.get('by_school', 'by School')}",
                        labels={"school": t.get("school", "School"), column: col_name}
                    )
                else:  # Violin plots
                    fig = px.violin(
                        df,
                        x="school",
                        y=column,
                        color="school",
                        box=True,  # Include box plot inside violin
                        title=f"{col_name} - {t.get('by_school', 'by School')}",
                        labels={"school": t.get("school", "School"), column: col_name}
                    )
                
                # Update layout
                fig.update_layout(
                    xaxis_title=t.get("school", "School"),
                    yaxis_title=col_name,
                    legend_title=t.get("school", "School"),
                    height=500
                )
                
                st.plotly_chart(fig, use_container_width=True)
            
            # Export options
            col1, col2 = st.columns(2)
            
            # CSV Export
            with col1:
                # Combine data for CSV export
                export_data = {
                    "mean_scores": mean_scores_by_school.reset_index(),
                    "highlights": highlight_df,
                    "statistical_tests": display_df if test_results else None
                }
                
                # Save to a BytesIO object
                csv_buffer = pd.ExcelWriter("school_performance.xlsx", engine='xlsxwriter')
                export_data["mean_scores"].to_excel(csv_buffer, sheet_name='Mean Scores', index=False)
                export_data["highlights"].to_excel(csv_buffer, sheet_name='Highlights', index=False)
                
                if export_data["statistical_tests"] is not None:
                    export_data["statistical_tests"].to_excel(csv_buffer, sheet_name='Statistical Tests', index=False)
                
                csv_buffer.close()
                
                # Create a download button for CSV
                with open("school_performance.xlsx", "rb") as f:
                    excel_data = f.read()
                
                st.download_button(
                    t.get("export_performance_excel", "📥 Download Excel"),
                    excel_data,
                    "school_performance_analysis.xlsx",
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key='download-performance-excel'
                )
                
                # Clean up temp file
                os.remove("school_performance.xlsx")
            
            # Word Export
            with col2:
                if st.button(t.get("export_performance_word", "📄 Export to Word")):
                    try:
                        doc = create_performance_word_report(
                            df, mean_scores_by_school, highlight_df, 
                            test_results, selected_columns, t
                        )
                        
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                            doc.save(tmp.name)
                            with open(tmp.name, 'rb') as f:
                                docx = f.read()
                            st.download_button(
                                t.get("download_performance_word", "📥 Download Word Report"),
                                docx,
                                "school_performance_analysis.docx",
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        os.unlink(tmp.name)
                    except Exception as e:
                        st.error(f"Error creating Word report: {str(e)}")
        
        except Exception as e:
            st.error(f"Error in school performance analysis: {str(e)}")
    
    else:
        st.warning(t.get("warning_select_variable", "Please select at least one variable to analyze."))

def create_performance_word_report(df, mean_scores, highlight_df, test_results, selected_columns, t):
    """
    Creates a Word report with school performance analysis.
    
    Args:
        df (pandas.DataFrame): Original data
        mean_scores (pandas.DataFrame): Mean scores by school
        highlight_df (pandas.DataFrame): Performance highlights
        test_results (list): Statistical test results
        selected_columns (list): Selected columns for analysis
        t (dict): Translation dictionary
        
    Returns:
        docx.Document: Word document with the report
    """
    doc = Document()
    
    # Title
    doc.add_heading(t.get("title_performance", "School Performance Analysis"), level=1)
    
    # Introduction
    doc.add_paragraph(t.get("performance_intro", "Objective: Compare performance across different schools and identify significant differences."))
    
    # Mean scores by school
    doc.add_heading(t.get("performance_results", "Mean Scores by School"), level=2)
    
    # Reset index to make school a column
    mean_scores_table = mean_scores.reset_index()
    
    # Create table
    table = doc.add_table(rows=len(mean_scores_table) + 1, cols=len(mean_scores_table.columns))
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = t.get("school", "School")
    
    for i, col in enumerate(mean_scores_table.columns[1:], 1):
        header_cells[i].text = t["columns_of_interest"].get(col, col)
    
    # Apply header formatting
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
    
    # Add data rows
    for i, (_, row) in enumerate(mean_scores_table.iterrows(), 1):
        row_cells = table.rows[i].cells
        
        for j, value in enumerate(row):
            row_cells[j].text = str(value)
    
    # Performance highlights
    doc.add_heading(t.get("performance_highlights", "Performance Highlights"), level=2)
    
    # Create highlights table
    highlight_table = doc.add_table(rows=len(highlight_df) + 1, cols=len(highlight_df.columns))
    highlight_table.style = 'Table Grid'
    
    # Add headers
    header_cells = highlight_table.rows[0].cells
    for i, col in enumerate(highlight_df.columns):
        header_cells[i].text = col
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # Add data rows
    for i, (_, row) in enumerate(highlight_df.iterrows(), 1):
        row_cells = highlight_table.rows[i].cells
        
        for j, value in enumerate(row):
            row_cells[j].text = str(value)
    
    # Statistical significance testing
    doc.add_heading(t.get("statistical_testing", "Statistical Significance Testing"), level=2)
    
    doc.add_paragraph(t.get("anova_explanation", """
    The analysis below uses the Kruskal-Wallis H-test, a non-parametric method for comparing scores from more than two schools.
    A p-value < 0.05 indicates statistically significant differences between at least some schools.
    """))
    
    # Create test results table if results exist
    if test_results:
        # Convert to DataFrame for easier processing
        test_df = pd.DataFrame(test_results)
        
        # Create table
        test_table = doc.add_table(rows=len(test_df) + 1, cols=4)  # Assuming 4 columns: variable, h_stat, p_value, significant
        test_table.style = 'Table Grid'
        
        # Add headers
        header_cells = test_table.rows[0].cells
        header_cells[0].text = t.get("variable", "Variable")
        header_cells[1].text = t.get("h_statistic", "H Statistic")
        header_cells[2].text = t.get("p_value", "p-value")
        header_cells[3].text = t.get("significant", "Significant Difference")
        
        # Apply header formatting
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
        
        # Add data rows
        for i, row in enumerate(test_results, 1):
            row_cells = test_table.rows[i].cells
            
            row_cells[0].text = row["variable"]
            row_cells[1].text = f"{row['h_statistic']:.4f}" if row['h_statistic'] is not None else "N/A"
            row_cells[2].text = f"{row['p_value']:.4f}" if row['p_value'] is not None else "N/A"
            
            if row['significant'] is not None:
                row_cells[3].text = t.get("significant_yes", "Yes") if row['significant'] else t.get("significant_no", "No")
            else:
                row_cells[3].text = "N/A"
    
    # Distribution plots for each variable
    doc.add_heading(t.get("distribution_title", "Score Distributions by School"), level=2)
    
    with tempfile.TemporaryDirectory() as tmp_dir:
        for column in selected_columns:
            col_name = t["columns_of_interest"].get(column, column)
            doc.add_heading(col_name, level=3)
            
            # Create box plot
            fig = px.box(
                df,
                x="school",
                y=column,
                color="school",
                title=f"{col_name} - {t.get('by_school', 'by School')}",
                labels={"school": t.get("school", "School"), column: col_name}
            )
            
            # Save plot
            img_path = os.path.join(tmp_dir, f"{column}_boxplot.png")
            fig.write_image(img_path, width=800, height=500)
            
            # Add plot to document
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph()
    
    # Summary and recommendations
    doc.add_heading(t.get("summary_recommendations", "Summary and Recommendations"), level=2)
    
    # Identify variables with significant differences
    if test_results:
        sig_vars = [r["variable"] for r in test_results if r.get("significant")]
        
        if sig_vars:
            doc.add_paragraph(t.get("significant_differences_found", "Significant differences between schools were found in the following areas:"))
            for var in sig_vars:
                doc.add_paragraph(var, style='List Bullet')
            
            doc.add_paragraph(t.get("recommendations_sig", """
            Recommendations:
            - Investigate instructional practices in high-performing schools for these skills
            - Consider targeted professional development in lower-performing schools
            - Examine resource allocation and support systems across schools
            - Implement peer learning opportunities between schools
            """))
        else:
            doc.add_paragraph(t.get("no_significant_differences", """
            No statistically significant differences were found between schools. This suggests:
            - Relatively consistent instructional quality across schools
            - Similar student outcomes regardless of school attended
            - Successful standardization of educational practices
            """))
    
    # Find variables with largest disparities
    if not highlight_df.empty:
        large_gaps = highlight_df.sort_values(t.get("score_range", "Score Range"), ascending=False)
        
        if not large_gaps.empty:
            top_gap_var = large_gaps.iloc[0][t.get("variable", "Variable")]
            top_gap_high = large_gaps.iloc[0][t.get("highest_school", "Highest School")]
            top_gap_low = large_gaps.iloc[0][t.get("lowest_school", "Lowest School")]
            
            doc.add_paragraph(t.get("largest_performance_gap", "The largest performance gap was found in {}. {} had the highest scores while {} had the lowest.").format(
                top_gap_var, top_gap_high, top_gap_low
            ))
    
    # Footer with date
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = t.get('report_date', 'Report generated on: ') + pd.Timestamp.now().strftime('%Y-%m-%d')
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc