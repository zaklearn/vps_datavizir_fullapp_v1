# word_report_template.py
from docx import Document
from docx.shared import Inches, RGBColor
import io
from modules import language
from config import STANDARDS

def export_student_report(student_info, scores_egra, scores_egma, figures, rule_message, llm_message):
    """
    Generates a Word report with hybrid interpretation.
    
    Args:
        student_info (dict): Student info (name, ID, school, etc.)
        scores_egra (pd.DataFrame): EGRA scores table
        scores_egma (pd.DataFrame): EGMA scores table
        figures (list): Charts [fig_egra, fig_egma]
        rule_message (str): Systematic message
        llm_message (str): Enriched message

    Returns:
        BytesIO: Word document ready to download
    """
    doc = Document()

    try:
        # Header
        doc.add_heading(f"{language.t('student_report')} - {student_info.get('full_name', 'N/A')}", 0)
        
        # Student information - formatted as a table
        student_table = doc.add_table(rows=2, cols=3)
        student_table.style = 'Table Grid'  # Add borders
        
        # Headers
        headers = [language.t("personal_info"), language.t("school"), language.t("class")]
        for i, header in enumerate(headers):
            student_table.rows[0].cells[i].text = header
            if student_table.rows[0].cells[i].paragraphs and student_table.rows[0].cells[i].paragraphs[0].runs:
                student_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        # Values (with safer access using .get())
        student_table.rows[1].cells[0].text = (
            f"{language.t('name')}: {student_info.get('full_name', 'N/A')}\n"
            f"{language.t('age')}: {student_info.get('age', 'N/A')} {language.t('years')}\n"
            f"{language.t('gender')}: {student_info.get('gender', 'N/A')}\n"
            f"{language.t('language')}: {student_info.get('language', 'N/A')}\n"
            f"ID: {student_info.get('id', 'N/A')}"
        )
        student_table.rows[1].cells[1].text = student_info.get('school', 'N/A')
        student_table.rows[1].cells[2].text = student_info.get('group', 'N/A')
        
        doc.add_paragraph()

        # EGRA Section
        doc.add_heading(language.t("egra_results"), level=1)
        
        if scores_egra is not None and not scores_egra.empty:
            headers = list(scores_egra.columns)
            table = doc.add_table(rows=len(scores_egra)+1, cols=len(headers))
            table.style = 'Table Grid'  # Add borders
            
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = str(header)
                if table.rows[0].cells[i].paragraphs and table.rows[0].cells[i].paragraphs[0].runs:
                    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
                    
            for idx, row in scores_egra.iterrows():
                for j, val in enumerate(row):
                    table.rows[idx+1].cells[j].text = str(val)
                    
                    # Add colors for status cells
                    if j == 3:  # Assuming status is the 4th column (index 3)
                        cell = table.rows[idx+1].cells[j]
                        if "Mastery" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                        elif "Developing" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = RGBColor(255, 140, 0)  # Orange
                        elif "Emerging" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red

            if figures and len(figures) > 0:
                doc.add_paragraph()
                try:
                    img_stream = io.BytesIO(figures[0].to_image(format="png"))
                    doc.add_picture(img_stream, width=Inches(6))
                except Exception as e:
                    doc.add_paragraph(f"{language.t('error_generating_chart')}: {str(e)}")
        else:
            doc.add_paragraph(language.t("no_data_available"))

        # EGMA Section
        doc.add_heading(language.t("egma_results"), level=1)
        
        if scores_egma is not None and not scores_egma.empty:
            headers = list(scores_egma.columns)
            table = doc.add_table(rows=len(scores_egma)+1, cols=len(headers))
            table.style = 'Table Grid'  # Add borders
            
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = str(header)
                if table.rows[0].cells[i].paragraphs and table.rows[0].cells[i].paragraphs[0].runs:
                    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
                    
            for idx, row in scores_egma.iterrows():
                for j, val in enumerate(row):
                    table.rows[idx+1].cells[j].text = str(val)
                    
                    # Add colors for status cells
                    if j == 3:  # Assuming status is the 4th column (index 3)
                        cell = table.rows[idx+1].cells[j]
                        if "Mastery" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                        elif "Developing" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = RGBColor(255, 140, 0)  # Orange
                        elif "Emerging" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red

            if figures and len(figures) > 1:
                doc.add_paragraph()
                try:
                    img_stream = io.BytesIO(figures[1].to_image(format="png"))
                    doc.add_picture(img_stream, width=Inches(6))
                except Exception as e:
                    doc.add_paragraph(f"{language.t('error_generating_chart')}: {str(e)}")
        else:
            doc.add_paragraph(language.t("no_data_available"))

        # Interpretation Section - Always include if available
        if rule_message or llm_message:
            doc.add_heading(language.t("interpretation_recs"), level=1)
            
            if rule_message:
                doc.add_paragraph(f"{language.t('systematic_analysis')}:")
                doc.add_paragraph(rule_message)
            
            if llm_message:
                doc.add_paragraph(f"{language.t('enriched_synthesis')}:")
                doc.add_paragraph(llm_message)

    except Exception as e:
        # Add error section if something fails
        doc.add_heading(language.t("error_generating_report"), level=1)
        doc.add_paragraph(f"{language.t('error_occurred')}: {str(e)}")

    # Save document to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer