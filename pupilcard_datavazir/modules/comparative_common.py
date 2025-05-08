"""
Module comparative_common.py - Common utility functions for comparative modules.

This module provides shared functionality for the class_comparison.py and
school_comparison.py modules, including data processing, visualization, 
and report export.
"""

import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
import io
import traceback
from docx import Document
from docx.shared import Inches
from modules import language
from config import STANDARDS, EGRA_INDICATORS, EGMA_INDICATORS

def get_labels():
    """
    Get a dictionary of translated labels for all indicators
    
    Returns:
        dict: Dictionary of indicator labels
    """
    try:
        labels = {}
        for key in STANDARDS.keys():
            labels[key] = language.get_label(key)
        return labels
    except Exception as e:
        print(f"Error getting labels: {str(e)}")
        # Return empty dict as fallback
        return {}

def get_status(score, category):
    """
    Determine the status based on a score and category against standards
    
    Args:
        score (float): Score value
        category (str): Category key
        
    Returns:
        str: Status (mastery, developing, emerging)
    """
    try:
        if category not in STANDARDS:
            return "N/A"
            
        # Handle numeric conversion for safety
        try:
            score = float(score)
        except (ValueError, TypeError):
            return "N/A"
            
        if score >= STANDARDS[category]["Mastery"]:
            return language.t("mastery")
        elif isinstance(STANDARDS[category]["Developing"], tuple) and STANDARDS[category]["Developing"][0] <= score <= STANDARDS[category]["Developing"][1]:
            return language.t("developing")
        elif isinstance(STANDARDS[category]["Emerging"], tuple) and STANDARDS[category]["Emerging"][0] <= score <= STANDARDS[category]["Emerging"][1]:
            return language.t("emerging")
        else:
            return "N/A"
    except Exception as e:
        print(f"Error getting status: {str(e)}")
        return "Error"

def calculate_class_comparison(df, selected_school):
    """
    Calculate comparison metrics for classes within a school
    
    Args:
        df (DataFrame): Complete data
        selected_school (str): School name
        
    Returns:
        tuple: (egra_table, egma_table, class_means)
    """
    try:
        # Filter data for the selected school
        school_data = df[df['school'] == selected_school]
        
        if school_data.empty:
            return None, None, None
            
        # Calculate means by class
        class_means = school_data.groupby('group')[EGRA_INDICATORS + EGMA_INDICATORS].mean()
        labels = get_labels()  # Get labels when needed
        
        def create_comparison_table(indicators, means):
            """
            Create a comparison table for a set of indicators
            
            Args:
                indicators (list): List of indicator keys
                means (DataFrame): Means by group
                
            Returns:
                DataFrame: Formatted comparison table
            """
            try:
                rows = []
                for indicator in indicators:
                    row = {language.t('indicator'): labels[indicator]}
                    standard = STANDARDS[indicator]['Mastery']
                    for class_name in means.index:
                        score = means.loc[class_name, indicator]
                        row[f"{language.t('class')} {class_name}"] = f"{score:.1f}"
                    row[language.t('standard')] = standard
                    row[language.t('status')] = get_status(means[indicator].mean(), indicator)
                    rows.append(row)
                return pd.DataFrame(rows)
            except Exception as e:
                print(f"Error creating comparison table: {str(e)}")
                # Return empty dataframe instead of None to avoid downstream errors
                return pd.DataFrame()
            
        egra_table = create_comparison_table(EGRA_INDICATORS, class_means)
        egma_table = create_comparison_table(EGMA_INDICATORS, class_means)
        return egra_table, egma_table, class_means
    except Exception as e:
        print(f"Error calculating class comparison: {str(e)}")
        print(traceback.format_exc())
        return None, None, None

def plot_comparison(class_means, indicators, title):
    """
    Create a bar chart comparing classes for specified indicators
    
    Args:
        class_means (DataFrame): Means by class
        indicators (list): List of indicator keys
        title (str): Chart title
        
    Returns:
        Figure: Plotly figure object
    """
    try:
        fig = go.Figure()
        labels = get_labels()  # Get labels when needed
        
        if class_means is None or class_means.empty:
            # Return an empty figure with an error message
            fig.add_annotation(
                text="No data available",
                xref="paper", yref="paper",
                x=0.5, y=0.5,
                showarrow=False
            )
            return fig
        
        # Generate a list of colors from the viridis palette
        num_classes = len(class_means.index)
        colors = px.colors.sequential.Viridis  # List of viridis colors
        colors = colors[:num_classes]  # Adjust the color list if necessary

        # Add bars for each class with colors
        for ind, class_name in enumerate(class_means.index):
            # Error checking for each indicator
            y_values = []
            x_values = []
            text_values = []
            
            for indicator in indicators:
                try:
                    # Check if the indicator exists in the data
                    if indicator in class_means.columns:
                        value = class_means.loc[class_name, indicator]
                        # Check if the value is valid
                        if pd.notna(value):
                            y_values.append(value)
                            x_values.append(labels.get(indicator, indicator))  # Fallback to key if label missing
                            text_values.append(f"{value:.1f}")
                except Exception as e:
                    print(f"Error processing indicator {indicator} for class {class_name}: {str(e)}")
            
            if y_values:  # Only add trace if we have valid values
                fig.add_trace(go.Bar(
                    name=f'{language.t("class")} {class_name}',
                    x=x_values,
                    y=y_values,
                    marker=dict(color=colors[ind % len(colors)]),  # Use modulo to handle more classes than colors
                    text=text_values,
                    textposition='outside',
                ))

        # Add trace for standards
        standard_x = []
        standard_y = []
        
        for indicator in indicators:
            try:
                if indicator in STANDARDS:
                    standard_x.append(labels.get(indicator, indicator))
                    standard_y.append(STANDARDS[indicator]['Mastery'])
            except Exception as e:
                print(f"Error adding standard for indicator {indicator}: {str(e)}")
        
        if standard_x:  # Only add standards trace if we have valid values
            fig.add_trace(go.Scatter(
                name=language.t('standard'),
                x=standard_x,
                y=standard_y,
                mode='lines+markers',
                line=dict(color='red', dash='dash', width=2),
                marker=dict(symbol='diamond', size=10),
            ))

        # Update layout
        fig.update_layout(
            title=title,
            xaxis_title=language.t('indicator'),
            yaxis_title=language.t('score'),
            hovermode='x unified',
            xaxis={'tickangle': -45},
            barmode='group',
            showlegend=True,
            legend=dict(
                yanchor="top",
                y=0.99,
                xanchor="left",
                x=0.01
            )
        )
        return fig
    except Exception as e:
        print(f"Error plotting comparison: {str(e)}")
        print(traceback.format_exc())
        # Return basic empty figure instead of None
        fig = go.Figure()
        fig.add_annotation(
            text=f"Error generating chart: {str(e)}",
            xref="paper", yref="paper",
            x=0.5, y=0.5,
            showarrow=False
        )
        return fig

def calculate_school_comparison(df, selected_schools):
    """
    Calculate comparison metrics between schools
    
    Args:
        df (DataFrame): Complete data
        selected_schools (list): List of school names
        
    Returns:
        tuple: (egra_comparison, egma_comparison)
    """
    try:
        egra_indicators = EGRA_INDICATORS
        egma_indicators = EGMA_INDICATORS
        
        # Filter data for selected schools
        filtered_df = df[df['school'].isin(selected_schools)]
        
        if filtered_df.empty:
            return None, None
            
        # Calculate means by school
        school_means = filtered_df.groupby('school')[egra_indicators + egma_indicators].mean()
        labels = get_labels()  # Get labels when needed
        
        def create_school_comparison_table(indicators):
            """
            Create a comparison table for schools
            
            Args:
                indicators (list): List of indicator keys
                
            Returns:
                DataFrame: Formatted comparison table
            """
            try:
                comparison_data = []
                for school in selected_schools:
                    # Skip if school not in means (e.g., no data)
                    if school not in school_means.index:
                        continue
                        
                    row = {language.t('school'): school}
                    for indicator in indicators:
                        score = school_means.loc[school, indicator]
                        row[labels[indicator]] = f"{score:.1f}"
                        row[f'{labels[indicator]} ({language.t("standard")})'] = STANDARDS[indicator]['Mastery']
                        row[f'{labels[indicator]} ({language.t("status")})'] = get_status(score, indicator)
                    comparison_data.append(row)
                return pd.DataFrame(comparison_data)
            except Exception as e:
                print(f"Error creating school comparison table: {str(e)}")
                return pd.DataFrame()
            
        egra_comparison = create_school_comparison_table(egra_indicators)
        egma_comparison = create_school_comparison_table(egma_indicators)
        return egra_comparison, egma_comparison
    except Exception as e:
        print(f"Error calculating school comparison: {str(e)}")
        print(traceback.format_exc())
        return None, None

def export_comparison_to_word(report_title, egra_table, egma_table, figures, summary_text=None, llm_summary=None):
    """
    Exports comparison results for Class and School modules to Word.
    
    Args:
        report_title (str): The report title
        egra_table (DataFrame): The EGRA results table
        egma_table (DataFrame): The EGMA results table
        figures (list): List containing figures (EGRA, EGMA)
        summary_text (str, optional): Systematic summary text
        llm_summary (str, optional): LLM-generated interpretation
    
    Returns:
        BytesIO: A buffer containing the Word document
    """
    doc = Document()
    
    try:
        # Report title
        doc.add_heading(report_title, 0)
        
        # --- EGRA Section ---
        doc.add_heading(language.t("egra_results"), level=1)
        
        if egra_table is not None and not egra_table.empty:
            headers = list(egra_table.columns)
            table_word = doc.add_table(rows=len(egra_table) + 1, cols=len(headers))
            table_word.style = 'Table Grid'  # Add borders
            
            # Add headers
            for i, header in enumerate(headers):
                table_word.rows[0].cells[i].text = str(header)
                
            # Add data rows    
            for idx, row in egra_table.iterrows():
                for col_idx, value in enumerate(row):
                    table_word.rows[idx + 1].cells[col_idx].text = str(value)
        else:
            doc.add_paragraph(language.t("no_data_available"))
        
        # Add charts
        doc.add_paragraph()
        if figures and len(figures) > 0 and figures[0] is not None:
            try:
                img_bytes = figures[0].to_image(format="png")
                img_stream = io.BytesIO(img_bytes)
                doc.add_picture(img_stream, width=Inches(6))
            except Exception as e:
                doc.add_paragraph(f"{language.t('error_generating_chart')}: {str(e)}")
        
        # --- EGMA Section ---
        doc.add_heading(language.t("egma_results"), level=1)
        
        if egma_table is not None and not egma_table.empty:
            headers = list(egma_table.columns)
            table_word2 = doc.add_table(rows=len(egma_table) + 1, cols=len(headers))
            table_word2.style = 'Table Grid'  # Add borders
            
            # Add headers
            for i, header in enumerate(headers):
                table_word2.rows[0].cells[i].text = str(header)
                
            # Add data rows
            for idx, row in egma_table.iterrows():
                for col_idx, value in enumerate(row):
                    table_word2.rows[idx + 1].cells[col_idx].text = str(value)
        else:
            doc.add_paragraph(language.t("no_data_available"))
        
        # Add chart
        doc.add_paragraph()
        if figures and len(figures) > 1 and figures[1] is not None:
            try:
                img_bytes = figures[1].to_image(format="png")
                img_stream = io.BytesIO(img_bytes)
                doc.add_picture(img_stream, width=Inches(6))
            except Exception as e:
                doc.add_paragraph(f"{language.t('error_generating_chart')}: {str(e)}")
        
        # Add interpretation if available
        if summary_text or llm_summary:
            doc.add_heading(language.t("interpretation_recs"), level=1)
            
            if summary_text:
                doc.add_paragraph(language.t("systematic_summary") + ":")
                doc.add_paragraph(summary_text)
            
            if llm_summary:
                doc.add_paragraph(language.t("llm_interpretation") + ":")
                doc.add_paragraph(llm_summary)
    
        # Save document to buffer
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer
        
    except Exception as e:
        # In case of error, create a minimal document with error message
        print(f"Error generating comparison report: {str(e)}")
        print(traceback.format_exc())
        
        error_doc = Document()
        error_doc.add_heading(language.t("error_generating_report"), 0)
        error_doc.add_paragraph(f"{language.t('error_occurred')}: {str(e)}")
        error_buffer = io.BytesIO()
        error_doc.save(error_buffer)
        error_buffer.seek(0)
        return error_buffer