"""
Module pupil_report_card.py - Analysis and visualization of student performance.

This module allows analysis and visualization of student performance on EGRA and EGMA tests.
It offers two navigation modes:
1. Complete student list view
2. Search for a specific student

The module handles student profile display, performance charts,
and report export in Word format.
"""

# ===== IMPORTS AND CONFIGURATION =====
import streamlit as st
import pandas as pd
import plotly.express as px
import io
import random
import traceback
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

# Import from centralized configuration
from config import STANDARDS, EGRA_INDICATORS, EGMA_INDICATORS, TRANSLATIONS
from modules import interpretation
from modules.word_report_template import export_student_report
from modules import language

# ===== UTILITY FUNCTIONS =====

def get_status(score, category):
    """
    Determines status (Mastery, Developing, Emerging) based on score and category
    
    Args:
        score: Score obtained by the student
        category: Evaluation category
        
    Returns:
        str: Student status (Mastery, Developing, Emerging, N/A)
    """
    try:
        # Check that the category exists
        if category not in STANDARDS:
            return "Unknown category"
        
        # Handling NaN or None values
        if pd.isna(score) or score is None:
            return "N/A"
        
        # Converting score to number if it's a string
        if isinstance(score, str):
            try:
                score = float(score)
            except ValueError:
                return "N/A"
        
        # Ensure score is a float for reliable comparisons
        score = float(score)
        
        # Checking score for Mastery
        if score >= STANDARDS[category]["Mastery"]:
            return "Mastery"
        
        # Checking score for Developing
        developing_range = STANDARDS[category]["Developing"]
        if isinstance(developing_range, tuple) and developing_range[0] <= score <= developing_range[1]:
            return "Developing"
        elif not isinstance(developing_range, tuple) and score == developing_range:
            return "Developing"
        
        # Checking score for Emerging
        emerging_range = STANDARDS[category]["Emerging"]
        if isinstance(emerging_range, tuple) and emerging_range[0] <= score <= emerging_range[1]:
            return "Emerging"
        elif not isinstance(emerging_range, tuple) and score == emerging_range:
            return "Emerging"
        
        # If no condition is met
        return "N/A"
    except Exception as e:
        st.error(f"Error determining status: {str(e)}")
        return "Error"


def style_dataframe(df, status_column=None):
    """
    Applies style to DataFrame for display in Streamlit
    
    Args:
        df: DataFrame to style
        status_column: Name of the column containing status (if applicable)
    
    Returns:
        Styler object
    """
    try:
        # Protection against empty DataFrames
        if df.empty:
            return df.style
        
        # Create a style function that colors cells based on status
        def highlight_status(val):
            if val == "Mastery":
                return 'background-color: rgba(0, 128, 0, 0.2); color: green; font-weight: bold'
            elif val == "Developing":
                return 'background-color: rgba(255, 165, 0, 0.2); color: orange; font-weight: bold'
            elif val == "Emerging":
                return 'background-color: rgba(255, 0, 0, 0.2); color: red; font-weight: bold'
            else:
                return ''

        # Apply styles to DataFrame
        if status_column and status_column in df.columns:
            return df.style.apply(
                lambda row: [highlight_status(val) if col == status_column else '' for col, val in row.items()],
                axis=1
            ).set_properties(**{
                'text-align': 'center',
                'border': '1px solid lightgrey',
                'padding': '5px'
            }).set_table_styles([
                {'selector': 'th', 'props': [('font-weight', 'bold'), ('text-align', 'center'), ('background-color', '#f0f2f6')]},
            ])
        else:
            return df.style.set_properties(**{
                'text-align': 'center',
                'border': '1px solid lightgrey',
                'padding': '5px'
            }).set_table_styles([
                {'selector': 'th', 'props': [('font-weight', 'bold'), ('text-align', 'center'), ('background-color', '#f0f2f6')]},
            ])
    except Exception as e:
        st.error(f"Error styling DataFrame: {str(e)}")
        # Return unstyled dataframe as fallback
        return df.style

# ===== INTERFACE AND EXPORT FUNCTIONS =====

def apply_table_style(table):
    """
    Applies style to Word table to improve readability
    
    Args:
        table: Word table to style
    """
    try:
        # Protection against empty or invalid tables
        if not table or not hasattr(table, 'rows') or len(table.rows) == 0:
            return
        
        # Header style
        for cell in table.rows[0].cells:
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Borders and alignment for all cells
        for row in table.rows:
            for cell in row.cells:
                if cell.paragraphs:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Add style to statuses if text exists
                    if cell.paragraphs[0].runs:
                        if "Mastery" in cell.text:
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Green
                        elif "Developing" in cell.text:
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 140, 0)  # Orange
                        elif "Emerging" in cell.text:
                            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red
    except Exception as e:
        # Silently handle formatting errors - don't crash the report
        print(f"Error applying table style: {str(e)}")


def export_to_word(scores_egra, scores_egma, student_info, figures, rule_message=None, llm_message=None):
    """
    Exports results to a Word document
    
    Args:
        scores_egra: DataFrame of EGRA scores
        scores_egma: DataFrame of EGMA scores
        student_info: Dictionary with student information
        figures: List of Plotly figures
        rule_message: Systematic analysis message (optional)
        llm_message: LLM-generated interpretation (optional)
        
    Returns:
        BytesIO: Word document in binary format
    """
    try:
        # Parameter verification
        if scores_egra is None or scores_egma is None or scores_egra.empty or scores_egma.empty:
            raise ValueError(f"{language.t('egra_egma_required')}")
        
        if student_info is None:
            raise ValueError(f"{language.t('student_info_required')}")
        
        if figures is None or len(figures) < 2:
            raise ValueError(f"{language.t('charts_required')}")
        
        doc = Document()
        
        # Title
        doc.add_heading(f"{language.t('report_for')} {student_info.get('full_name', 'N/A')}", 0)
        
        # Student information
        student_table = doc.add_table(rows=2, cols=3)
        student_table.style = 'Table Grid'  # Add borders
        
        # Headers
        headers = [language.t("personal_info"), language.t("school"), language.t("class")]
        for i, header in enumerate(headers):
            student_table.rows[0].cells[i].text = header
            if student_table.rows[0].cells[i].paragraphs and student_table.rows[0].cells[i].paragraphs[0].runs:
                student_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        # Values
        student_table.rows[1].cells[0].text = (
            f"{language.t('name')}: {student_info.get('full_name', 'N/A')}\n"
            f"{language.t('age')}: {student_info.get('age', 'N/A')} {language.t('years')}\n"
            f"{language.t('gender')}: {student_info.get('gender', 'N/A')}\n"
            f"{language.t('language')}: {student_info.get('language', 'N/A')}\n"
            f"ID: {student_info.get('id', 'N/A')}"
        )
        student_table.rows[1].cells[1].text = student_info.get('school', 'N/A')
        student_table.rows[1].cells[2].text = student_info.get('group', 'N/A')
        
        doc.add_paragraph("")
        
        # EGRA Section
        doc.add_heading(language.t("egra_results"), level=1)
        
        # EGRA Table
        try:
            egra_table = doc.add_table(rows=len(scores_egra)+1, cols=4)
            egra_table.style = 'Table Grid'  # Add borders
            headers = [language.t("indicator"), language.t("student_score"), language.t("standard_score"), language.t("status")]
            for i, header in enumerate(headers):
                egra_table.rows[0].cells[i].text = header
            
            for idx, row in enumerate(scores_egra.values, start=1):
                for col_idx, value in enumerate(row):
                    egra_table.rows[idx].cells[col_idx].text = str(value)
            
            apply_table_style(egra_table)
        except Exception as e:
            doc.add_paragraph(f"{language.t('error_creating_table')}: {str(e)}")
        
        # EGRA Chart
        doc.add_paragraph()
        try:
            if len(figures) > 0 and figures[0] is not None:
                img_bytes = figures[0].to_image(format="png")
                img_stream = io.BytesIO(img_bytes)
                doc.add_picture(img_stream, width=Inches(6))
        except Exception as e:
            # Error handling when generating image
            doc.add_paragraph(f"{language.t('error_generating_chart')}: {str(e)}")
        
        # EGMA Section
        doc.add_heading(language.t("egma_results"), level=1)
        
        # EGMA Table
        try:
            egma_table = doc.add_table(rows=len(scores_egma)+1, cols=4)
            egma_table.style = 'Table Grid'  # Add borders
            for i, header in enumerate(headers):
                egma_table.rows[0].cells[i].text = header
            
            for idx, row in enumerate(scores_egma.values, start=1):
                for col_idx, value in enumerate(row):
                    egma_table.rows[idx].cells[col_idx].text = str(value)
            
            apply_table_style(egma_table)
        except Exception as e:
            doc.add_paragraph(f"{language.t('error_creating_table')}: {str(e)}")
        
        # EGMA Chart
        doc.add_paragraph()
        try:
            if len(figures) > 1 and figures[1] is not None:  # Additional check
                img_bytes = figures[1].to_image(format="png")
                img_stream = io.BytesIO(img_bytes)
                doc.add_picture(img_stream, width=Inches(6))
        except Exception as e:
            # Error handling when generating image
            doc.add_paragraph(f"{language.t('error_generating_chart')}: {str(e)}")
                
        # Add interpretation section if available
        if rule_message or llm_message:
            doc.add_heading(language.t("interpretation_recs"), level=1)
            
            if rule_message:
                doc.add_paragraph(f"{language.t('systematic_analysis')}:")
                doc.add_paragraph(rule_message)
            
            if llm_message:
                doc.add_paragraph(f"{language.t('enriched_synthesis')}:")
                doc.add_paragraph(llm_message)
        
        # Save document to buffer
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer
    
    except Exception as e:
        # In case of error, create a minimal document with error message
        error_doc = Document()
        error_doc.add_heading(language.t("error_generating_report"), 0)
        error_doc.add_paragraph(f"{language.t('error_occurred')}: {str(e)}")
        error_buffer = io.BytesIO()
        error_doc.save(error_buffer)
        error_buffer.seek(0)
        return error_buffer


def display_student_profile(df, student_data):
    """
    Displays complete student profile
    
    Args:
        df: Complete DataFrame with all data
        student_data: Pandas Series with selected student data
    
    Returns:
        tuple: (figures, scores_egra, scores_egma, student_info, rule_message, llm_message)
    """
    try:
        # Data validity check
        if student_data is None or len(student_data) == 0:
            st.error(f"{language.t('no_data_available')}.")
            return None, None, None, None, None, None
        
        # Basic student information (with type handling)
        try:
            gender_value = "Boy" if student_data.get('st_gender') == "M" else "Girl"
            student_info = {
                'full_name': f"{student_data.get('last_name', '')} {student_data.get('first_name', '')}",
                'age': int(student_data.get('st_age', 0)),
                'gender': gender_value,
                'language': student_data.get('language', ''),
                'id': int(float(student_data.get('pupil_id', 0))),
                'school': student_data.get('school', ''),
                'group': student_data.get('group', '')
            }
        except (ValueError, TypeError) as e:
            st.error(f"{language.t('error_processing_data')}: {e}")
            st.write(f"{language.t('problematic_values')}: {dict(student_data)}")
            return None, None, None, None, None, None
        
        st.markdown(f"## 🎓 {language.t('profile_for')} {student_info['full_name']} (ID: {student_info['id']})")
        
        # Display basic information
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric(language.t("age"), f"{student_info['age']} {language.t('years')}")
        with col2:
            st.metric(language.t("gender"), student_info['gender'])
        with col3:
            st.metric(language.t("language"), student_info['language'])
        with col4:
            st.metric(language.t("class"), student_info['group'])
        
        st.markdown(f"### 📊 {language.t('student_perf')}")
        
        figures = []  # To store charts
        scores_egra = None
        scores_egma = None
        rule_message = None
        llm_message = None
        
        # Define status column name 
        status_col = "Status"
        
        for title, indicators in [
            (f"📖 {language.t('egra_results')}", EGRA_INDICATORS),
            (f"🔢 {language.t('egma_results')}", EGMA_INDICATORS)
        ]:
            col1, col2 = st.columns([1, 1])
            with col1:
                st.markdown(f"#### {title}")
                
                # Data validation to avoid errors
                valid_indicators = []
                for ind in indicators:
                    # Check if indicator exists in data and is not NaN
                    if ind in student_data and not pd.isna(student_data[ind]):
                        valid_indicators.append(ind)
                
                if not valid_indicators:
                    st.warning(f"⚠️ {language.t('no_data')} {title}")
                    # Add empty chart for EGRA if needed
                    if "EGRA" in title:
                        figures.append(px.bar(title=f"{language.t('no_data')} {title}"))
                    continue
                
                # Create scores DataFrame with type handling
                try:
                    scores = pd.DataFrame({
                        "Indicator": [language.get_label(col) for col in valid_indicators],
                        "Student Score": [float(student_data[col]) if not pd.isna(student_data[col]) else 0 for col in valid_indicators],
                        "Standard Score": [float(STANDARDS[col]["Mastery"]) for col in valid_indicators],
                        status_col: [get_status(student_data[col], col) for col in valid_indicators]
                    })
                    
                    if "EGRA" in title:
                        scores_egra = scores
                    else:
                        scores_egma = scores
                    
                    # Apply style to dataframe and display as a static table (no scrollbars)
                    styled_scores = style_dataframe(scores, status_col)
                    # FIX: Don't call .style on the Styler object again
                    st.table(styled_scores.hide(axis="index"))
                    
                except Exception as e:
                    st.error(f"{language.t('error_creating_table')}: {e}")
                    if "EGRA" in title:
                        figures.append(px.bar(title=f"{language.t('error')}: {e}"))
                    continue
            
            with col2:
                try:
                    # Create chart
                    fig = px.bar(scores, x="Indicator", y=["Student Score", "Standard Score"], 
                              color_discrete_sequence=px.colors.sequential.Viridis,
                              barmode="group", title=f"{title} {language.t('comparison')}")
                    fig.update_layout(xaxis_tickangle=-45)
                    st.plotly_chart(fig, use_container_width=True)
                    figures.append(fig)
                except Exception as e:
                    st.error(f"{language.t('error_creating_chart')}: {e}")
                    # Add empty chart to maintain index
                    figures.append(px.bar(title=f"{language.t('error')}: {e}"))
        
        # Add interpretation section if we have both EGRA and EGMA data
        if scores_egra is not None and not scores_egra.empty and scores_egma is not None and not scores_egma.empty:
            try:
                # Generate interpretation
                rule_message = interpretation.generate_rule_based_interpretation(scores_egra, scores_egma)
                prompt = interpretation.create_llm_prompt(rule_message, scores_egra, scores_egma)
                llm_message = interpretation.generate_llm_message(prompt)
                
                # Display interpretation
                st.markdown(f"### 🧠 {language.t('interpretation')}")
                st.markdown(f"**{language.t('systematic_analysis')} :**")
                st.info(rule_message)
                st.markdown(f"**{language.t('enriched_synthesis')} :**")
                st.success(llm_message)
            except Exception as e:
                st.error(f"Error generating interpretation: {str(e)}")
                rule_message = None
                llm_message = None
            
            # Display single download button for report
            st.markdown(f"### 💾 {language.t('export_report')}")
            
            if st.button(f"📄 {language.t('export_word')}"):
                try:
                    buffer = export_student_report(student_info, scores_egra, scores_egma, figures, rule_message, llm_message)
                    
                    # Create meaningful filename
                    now = datetime.now().strftime("%Y%m%d_%H%M%S")
                    safe_name = "".join([c if c.isalnum() else "_" for c in student_info['full_name']])
                    filename = f"report_{safe_name}_{now}.docx"
                    
                    st.download_button(
                        label=f"📄 {language.t('download')}",
                        data=buffer,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    st.success(f"✅ {language.t('success_report')}")
                except Exception as e:
                    st.error(f"❌ {language.t('error_report')}: {str(e)}")
        
        return figures, scores_egra, scores_egma, student_info, rule_message, llm_message
    
    except Exception as e:
        st.error(f"Unexpected error in display_student_profile: {str(e)}")
        st.error(traceback.format_exc())
        return None, None, None, None, None, None

# ===== MAIN FUNCTION =====

def main():
    """
    Main application function
    """
    try:
        # Session variable initialization
        if 'selected_student_id' not in st.session_state:
            st.session_state.selected_student_id = None
        
        # Initialize tracking variables to detect data changes
        if 'previous_file_name' not in st.session_state:
            st.session_state.previous_file_name = None
        
        # Sidebar for data loading
        with st.sidebar:
            st.header(f"📂 {language.t('data_source')}")
            use_demo = st.checkbox(language.t("use_demo"))
            uploaded_file = st.file_uploader(language.t("upload_excel"), type=["xlsx"])

        if use_demo:
            try:
                # Generate demo data
                demo_data = generate_demo_data()
                df = pd.DataFrame(demo_data)
                st.success("✅ Demo data loaded successfully!")
            except Exception as e:
                st.error(f"Error generating demo data: {str(e)}")
                return
        elif uploaded_file:
            try:
                df = pd.read_excel(uploaded_file)
                st.success("✅ File loaded successfully!")
            except Exception as e:
                st.error(f"❌ Error loading file: {str(e)}")
                return
        else:
            st.info(f"👈 {language.t('choose_data')}.")
            return

        # Navigation
        st.sidebar.header(f"📌 {language.t('navigation')}")
        action = st.sidebar.radio(language.t("what_to_do"), [
            language.t("view_list"),
            language.t("search_student")
        ])

        # Check required columns
        pupil_columns = ["school", "group", "pupil_id", "last_name", "first_name", "language", "st_gender", "st_age"]
        score_columns = list(STANDARDS.keys())
        
        # More robust column check
        df_columns = list(df.columns)
        missing_columns = [col for col in pupil_columns if col not in df_columns]
        
        if missing_columns:
            st.error(f"❌ {language.t('missing_columns')}: {', '.join(missing_columns)}")
            st.info(f"{language.t('required_columns')}: " + ", ".join(pupil_columns))
            st.info(f"{language.t('available_columns')}: " + ", ".join(df_columns))
            return

        # Data preparation with error handling
        try:
            # Convert numeric types if needed
            for col in ["pupil_id", "st_age"] + [c for c in score_columns if c in df.columns]:
                try:
                    df[col] = pd.to_numeric(df[col], errors='coerce')
                except Exception:
                    pass  # Ignore columns that can't be converted
            
            # Get student information columns
            pupil_info = df[pupil_columns].copy()
        except Exception as e:
            st.error(f"❌ {language.t('error_data_prep')}: {e}")
            return
        
        # MODE 1: COMPLETE STUDENT LIST
        if action == language.t("view_list"):
            st.header(f"📋 {language.t('complete_list')}")
            
            # Display return element if a student is selected
            if st.session_state.selected_student_id is not None:
                if st.button(language.t("return_list")):
                    st.session_state.selected_student_id = None
                    st.rerun()
                
                try:
                    student_matches = df[df["pupil_id"] == st.session_state.selected_student_id]
                    if student_matches.empty:
                        st.error(f"{language.t('no_student')} {st.session_state.selected_student_id}")
                        st.session_state.selected_student_id = None
                        st.rerun()
                    
                    student_row = student_matches.iloc[0]
                    figures, scores_egra, scores_egma, student_info, rule_message, llm_message = display_student_profile(df, student_row)
                    
                except Exception as e:
                    st.error(f"❌ {language.t('error_displaying_profile')}: {e}")
                    st.error(traceback.format_exc())
                    st.session_state.selected_student_id = None
            
            else:
                st.markdown(f"### {language.t('student_list')}")
                
                for _, row in pupil_info.iterrows():
                    col1, col2, col3, col4, col5, col6, col7 = st.columns([2, 1, 2, 2, 1, 1, 2])
                    
                    with col1:
                        st.text(f"{row['last_name']} {row['first_name']}")
                    with col2:
                        st.text(str(int(float(row['pupil_id']))))
                    with col3:
                        st.text(row['school'])
                    with col4:
                        st.text(row['group'])
                    with col5:
                        st.text(f"{int(row['st_age'])} {language.t('years')}")
                    with col6:
                        st.text("Boy" if row['st_gender'] == "M" else "Girl")
                    with col7:
                        if st.button(language.t("view_profile"), key=f"btn_{row['pupil_id']}"):
                            st.session_state.selected_student_id = int(float(row['pupil_id']))
                            st.session_state.report_downloaded = False
                            st.rerun()
        # MODE 2: SEARCH FOR A STUDENT
        else:  # language.t("search_student")
            st.header(f"🔍 {language.t('search_title')}")
            
            col1, col2, col3 = st.columns(3)
            
            try:
                with col1:
                    school_options = [""] + sorted(list(pupil_info["school"].unique()))
                    school_selected = st.selectbox(language.t("school"), school_options)
                
                if school_selected:
                    filtered_df = pupil_info[pupil_info["school"] == school_selected]
                    
                    with col2:
                        class_options = [""] + sorted(list(filtered_df["group"].unique()))
                        class_selected = st.selectbox(language.t("class"), class_options)
                    
                    if class_selected:
                        filtered_df = filtered_df[filtered_df["group"] == class_selected]
                        
                        with col3:
                            # Create student list for this class
                            student_options = [""]
                            for _, row in filtered_df.iterrows():
                                try:
                                    pupil_id = int(float(row['pupil_id']))
                                    student_options.append(f"{row['last_name']} {row['first_name']} (ID: {pupil_id})")
                                except (ValueError, TypeError):
                                    continue  # Ignore invalid IDs
                            
                            if len(student_options) <= 1:  # If list is empty after ""
                                st.warning(f"⚠️ {language.t('no_students_in_class')}.")
                                return
                            
                            student_selected = st.selectbox(language.t("student"), student_options)
                        
                        if student_selected and student_selected != "":
                            # Extract student ID from selected string
                            try:
                                student_id_str = student_selected.split("ID: ")[1].replace(")", "")
                                student_id = int(float(student_id_str))
                                
                                # Get student data
                                student_row = df[df["pupil_id"] == student_id].iloc[0]
                                
                                # Display student profile
                                figures, scores_egra, scores_egma, student_info, rule_message, llm_message = display_student_profile(df, student_row)
                                
                            except Exception as e:
                                st.error(f"❌ {language.t('error_retrieving_data')}: {str(e)}")
            except Exception as e:
                st.error(f"❌ {language.t('unexpected_error')}: {str(e)}")
                st.error(traceback.format_exc())
    
    except Exception as e:
        st.error(f"Critical application error: {str(e)}")
        st.error(traceback.format_exc())
        st.info("Please refresh the page and try again. If the problem persists, contact support.")


def generate_demo_data():
    """
    Generate demo data for testing
    
    Returns:
        dict: Dictionary with demo data
    """
    school_names = [
        "Greenwood International School", "Maplewood Academy", "Oxford Preparatory School",
        "Global Leaders Institute", "Harborview High School"
    ]
    
    class_names = ["3A", "3B", "4A", "4B", "5A"]
    
    last_names = ["Smith", "Johnson", "Brown", "Garcia", "Williams", "Lee", "Kim", "Patel", "Schneider", "Ivanov"]
    first_names = ["Emma", "Liam", "Olivia", "Noah", "Sophia", "Lucas", "Mia", "Ethan", "Isabella", "Nathan"]
    
    demo_data = {
        "school": [random.choice(school_names) for _ in range(25)],
        "group": [random.choice(class_names) for _ in range(25)],
        "pupil_id": list(range(1, 26)),
        "last_name": [random.choice(last_names) for _ in range(25)],
        "first_name": [random.choice(first_names) for _ in range(25)],
        "language": ["English", "French", "Spanish", "German", "Mandarin"] * 5,
        "st_gender": [random.choice(["M", "F"]) for _ in range(25)],
        "st_age": [random.randint(6, 12) for _ in range(25)],
        "clpm": [random.randint(10, 60) for _ in range(25)],
        "phoneme": [random.randint(10, 50) for _ in range(25)],
        "sound_word": [random.randint(30, 80) for _ in range(25)],
        "cwpm": [random.randint(5, 35) for _ in range(25)],
        "listening": [random.randint(40, 90) for _ in range(25)],
        "orf": [random.randint(20, 70) for _ in range(25)],
        "comprehension": [random.randint(30, 85) for _ in range(25)],
        "number_id": [random.randint(10, 50) for _ in range(25)],
        "discrimin": [random.randint(20, 60) for _ in range(25)],
        "missing_number": [random.randint(15, 55) for _ in range(25)],
        "addition": [random.randint(5, 25) for _ in range(25)],
        "subtraction": [random.randint(5, 25) for _ in range(25)],
        "problems": [random.randint(10, 50) for _ in range(25)]
    }
    
    return demo_data


if __name__ == '__main__':
    main()